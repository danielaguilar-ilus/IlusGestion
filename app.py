import base64
import io
import json
import os
import re
import secrets
import smtplib
import threading
import time
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from functools import wraps

from flask import (Flask, Response, flash, g, jsonify, make_response, redirect,
                   render_template, request, send_file, session, url_for)
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

from config import MAX_BULTOS, MYSQL_CONFIG, ERP_CONFIG, EMAIL_CONFIG, CLOUDINARY_CONFIG
try:
    from config import ANTHROPIC_API_KEY as _ANTHROPIC_KEY_CFG
except ImportError:
    _ANTHROPIC_KEY_CFG = ""

# ── Motor ERP unificado (cubicador, asignar, retiros, mantenciones, etc.) ──
import erp_engine
_ERP = erp_engine.init_engine(
    base_url=ERP_CONFIG.get("api_url", "https://lab.random.cl/ilus"),
    token=ERP_CONFIG.get("api_token", ""),
    doc_ttl=90,
    ent_ttl=300,
    timeout=6,
    retries=2,
)

def _get_ai_key():
    """Resuelve la API key de Anthropic: env var > config.py"""
    return os.environ.get("ANTHROPIC_API_KEY") or _ANTHROPIC_KEY_CFG or ""

app = Flask(__name__)
app.secret_key = "ilus-etiquetas-2026"


@app.template_filter("hm")
def _jinja_hm(value):
    if value is None:
        return ""
    return str(value)[:5]

# ══════════════════════════════════════════════════════════════
#  PLAYWRIGHT BROWSER POOL — instancia única reutilizada
#  Evita el overhead de launch/close (~700ms) por cada PDF.
#  Thread-safe: un Lock protege el acceso concurrente.
# ══════════════════════════════════════════════════════════════
_pw_lock     = threading.Lock()
_pw_ctx      = None   # sync_playwright() context manager
_pw_browser  = None   # Browser instance reutilizado


def _pw_browser_get():
    """Devuelve el browser compartido; lo lanza si aún no existe o murió."""
    global _pw_ctx, _pw_browser
    with _pw_lock:
        # Verificar si el browser sigue vivo
        if _pw_browser is not None:
            try:
                _ = _pw_browser.contexts  # ping liviano
                return _pw_browser
            except Exception:
                # Murió — limpiar
                try:
                    _pw_ctx.__exit__(None, None, None)
                except Exception:
                    pass
                _pw_ctx = _pw_browser = None

        # Lanzar nuevo browser
        from playwright.sync_api import sync_playwright
        _pw_ctx     = sync_playwright()
        pw          = _pw_ctx.__enter__()
        _pw_browser = pw.chromium.launch(
            args=["--no-sandbox", "--disable-dev-shm-usage",
                  "--disable-gpu", "--disable-extensions"]
        )
        return _pw_browser


def _pw_pdf(html: str, *, width: str = None, height: str = None,
            page_format: str = None, margin: dict = None,
            wait_fn: str = None, wait_timeout: int = 5000) -> bytes:
    """
    Genera PDF con el browser pool compartido.
    - width/height → tamaño personalizado (etiquetas)
    - page_format  → 'A4', 'Letter', etc.
    - wait_fn      → JS expression string para page.wait_for_function()
    - margin       → dict top/right/bottom/left en mm ('0mm')
    """
    browser = _pw_browser_get()
    page    = browser.new_page()
    try:
        page.set_content(html, wait_until="domcontentloaded")
        if wait_fn:
            try:
                page.wait_for_function(wait_fn, timeout=wait_timeout)
            except Exception:
                pass   # timeout: seguimos con lo que hay
        pdf_kw = dict(print_background=True)
        mrg = margin or {"top": "0mm", "right": "0mm",
                         "bottom": "0mm", "left": "0mm"}
        if width and height:
            pdf_kw.update(width=width, height=height, margin=mrg)
        elif page_format:
            pdf_kw.update(format=page_format, margin=mrg)
        return page.pdf(**pdf_kw)
    finally:
        page.close()

@app.template_filter('from_json')
def from_json_filter(value):
    """Parsea un string JSON almacenado en DB; devuelve lista/dict o []."""
    if not value:
        return []
    try:
        return json.loads(value)
    except Exception:
        return []


@app.template_filter('to_chile')
def to_chile_filter(value):
    """
    Convierte un datetime UTC (o naive asumido UTC) a hora local Chile.
    Útil para mostrar timestamps de la BD (que están en UTC) en hora local.
    Chile: UTC-4 invierno, UTC-3 verano (DST). Para simplicidad usamos UTC-3
    (que es el offset estándar continental los meses no-DST y la mayor parte
     del año tras la unificación horaria).
    """
    if not value:
        return value
    try:
        from datetime import timezone, timedelta
        # Si es naive, asumimos que está en UTC (que es como MySQL guarda por
        # default en servidores cloud como Railway).
        if hasattr(value, 'tzinfo'):
            if value.tzinfo is None:
                value = value.replace(tzinfo=timezone.utc)
            return value.astimezone(timezone(timedelta(hours=-3)))
    except Exception:
        pass
    return value

@app.template_filter('fkg')
def fkg_filter(value, decimals=1):
    """Formato kg estilo chileno: 2926.0 → '2.926,0' (punto miles, coma decimal, 1 decimal)"""
    try:
        n = float(value)
        formatted = f"{n:.{decimals}f}"          # "2926.0"
        int_part, dec_part = formatted.split('.')
        neg = int_part.startswith('-')
        int_abs = int_part.lstrip('-')
        # Agregar punto como separador de miles
        with_sep = ""
        for i, ch in enumerate(reversed(int_abs)):
            if i > 0 and i % 3 == 0:
                with_sep = '.' + with_sep
            with_sep = ch + with_sep
        return ('-' if neg else '') + with_sep + ',' + dec_part
    except Exception:
        return '—'

@app.template_filter('fvol')
def fvol_filter(value):
    """Formato volumen cm³ estilo chileno: 5673482 → '5.673.482' (punto miles, sin decimal)"""
    try:
        n = int(round(float(value)))
        s = str(abs(n))
        with_sep = ""
        for i, ch in enumerate(reversed(s)):
            if i > 0 and i % 3 == 0:
                with_sep = '.' + with_sep
            with_sep = ch + with_sep
        return ('-' if n < 0 else '') + with_sep
    except Exception:
        return '—'

BASE_DIR        = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER   = os.path.join(BASE_DIR, "static", "uploads")
COLABS_FOLDER   = os.path.join(BASE_DIR, "static", "uploads", "colaboradores")
ERP_TABLE_PRODUCTS = ERP_CONFIG.get("table_products", "MAEPR")
ALLOWED_EXT     = {"png", "jpg", "jpeg", "webp", "gif"}
MAX_PHOTOS      = 2

AUTH_TABLE     = MYSQL_CONFIG.get("users_table",    "app_users")
PRODUCTS_TABLE = MYSQL_CONFIG.get("products_table", "app_products")
BULTOS_TABLE   = MYSQL_CONFIG.get("bultos_table",   "app_bultos")
PHOTOS_TABLE   = MYSQL_CONFIG.get("photos_table",   "app_photos")
ERP_TABLE      = MYSQL_CONFIG["table"]

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")

DEFAULT_USERS = (
    ("daniel.aguilar@sphs.cl", "Daniel Aguilar", "superadmin", "19109364Daniel"),
)

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(COLABS_FOLDER, exist_ok=True)

# ─────────────────────────────────────────────
#  Cloudinary — almacenamiento de fotos en la nube
# ─────────────────────────────────────────────
_CLD_READY = False
_cloudinary_uploader = None
try:
    import cloudinary as _cld_module
    import cloudinary.uploader as _cloudinary_uploader
    if CLOUDINARY_CONFIG.get("cloud_name") and CLOUDINARY_CONFIG.get("api_key"):
        _cld_module.config(
            cloud_name = CLOUDINARY_CONFIG["cloud_name"],
            api_key    = CLOUDINARY_CONFIG["api_key"],
            api_secret = CLOUDINARY_CONFIG["api_secret"],
            secure     = True,
        )
        _CLD_READY = True
        print("[ILUS] Cloudinary configurado —", CLOUDINARY_CONFIG["cloud_name"])
    else:
        print("[ILUS] Cloudinary sin credenciales — fotos locales.")
except Exception as _cld_err:
    print(f"[ILUS] Cloudinary no disponible: {_cld_err} — fotos locales.")


def _cloud_upload(file_obj, public_id: str, folder: str = "ilus") -> str:
    """Sube a Cloudinary y devuelve la URL segura. Lanza excepción si falla."""
    result = _cloudinary_uploader.upload(
        file_obj,
        public_id     = public_id,
        folder        = folder,
        overwrite     = True,
        resource_type = "image",
    )
    return result["secure_url"]


def _cloud_delete(url_or_filename: str) -> None:
    """Elimina de Cloudinary si es URL; del disco si es nombre local."""
    if url_or_filename.startswith("http"):
        try:
            match = re.search(r"/upload/(?:v\d+/)?(.+)\.[^.]+$", url_or_filename)
            if match and _cloudinary_uploader:
                _cloudinary_uploader.destroy(match.group(1))
        except Exception as exc:
            print(f"[ILUS] Cloudinary delete error: {exc}")
    else:
        path = os.path.join(UPLOAD_FOLDER, url_or_filename)
        if os.path.exists(path):
            os.remove(path)


# ─────────────────────────────────────────────
#  DB helpers
# ─────────────────────────────────────────────

def get_erp_conn():
    """
    Conexión de SOLO LECTURA al ERP externo (cloud.random.cl).
    Devuelve None si no se puede conectar (no interrumpe la app).
    NUNCA usar para escribir.
    """
    import pymysql, pymysql.cursors
    connect_timeout = ERP_CONFIG.get("connect_timeout", 10)
    read_timeout    = ERP_CONFIG.get("read_timeout", 15)
    try:
        return pymysql.connect(
            host=ERP_CONFIG["host"],
            port=ERP_CONFIG["port"],
            user=ERP_CONFIG["user"],
            password=ERP_CONFIG["password"],
            database=ERP_CONFIG["database"],
            connect_timeout=connect_timeout,
            read_timeout=read_timeout,
            write_timeout=read_timeout,
            charset="utf8mb4",
            cursorclass=pymysql.cursors.DictCursor,
            autocommit=True,
        )
    except Exception:
        try:
            return pymysql.connect(
                host=ERP_CONFIG["host"],
                port=ERP_CONFIG["port"],
                user=ERP_CONFIG["user"],
                password=ERP_CONFIG["password"],
                database=ERP_CONFIG["database"],
                connect_timeout=connect_timeout,
                read_timeout=read_timeout,
                write_timeout=read_timeout,
                cursorclass=pymysql.cursors.DictCursor,
                autocommit=True,
            )
        except Exception:
            return None


def get_mysql():
    """Abre una conexión directa (usada solo por init_db y utilidades internas)."""
    import pymysql, pymysql.cursors
    return pymysql.connect(
        host=MYSQL_CONFIG["host"],
        port=MYSQL_CONFIG["port"],
        user=MYSQL_CONFIG["user"],
        password=MYSQL_CONFIG["password"],
        database=MYSQL_CONFIG["database"],
        connect_timeout=MYSQL_CONFIG.get("connect_timeout", 15),
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
        autocommit=False,
    )


# ── Pool de conexiones ──────────────────────────────────────────
# Mantiene conexiones TCP abiertas y las reutiliza entre requests.
# Elimina el costo de ~250-400 ms por apertura en cada clic.
_db_pool = None
_db_pool_lock = threading.Lock()

def _get_pool():
    global _db_pool
    if _db_pool is not None:
        return _db_pool
    with _db_pool_lock:
        if _db_pool is not None:          # doble chequeo post-lock
            return _db_pool
        try:
            import pymysql, pymysql.cursors
            from dbutils.pooled_db import PooledDB
            _db_pool = PooledDB(
                creator        = pymysql,
                mincached      = 1,           # conexiones siempre listas
                maxcached      = 5,           # máximo en pool inactivo
                maxconnections = 10,          # total permitido
                blocking       = False,       # falla rápido si no hay conexión libre
                ping           = 1,           # verifica conexión antes de entregar (reconecta si está muerta)
                host           = MYSQL_CONFIG["host"],
                port           = MYSQL_CONFIG["port"],
                user           = MYSQL_CONFIG["user"],
                password       = MYSQL_CONFIG["password"],
                database       = MYSQL_CONFIG["database"],
                connect_timeout= MYSQL_CONFIG.get("connect_timeout", 15),
                read_timeout   = 20,
                write_timeout  = 20,
                charset        = "utf8mb4",
                cursorclass    = pymysql.cursors.DictCursor,
                autocommit     = False,
            )
            print("[ILUS] Pool de conexiones MySQL activo (DBUtils).")
        except ImportError:
            # Si DBUtils no está instalado, degrada a conexión directa
            print("[ILUS][WARN] DBUtils no disponible — sin pool de conexiones.")
            _db_pool = None
    return _db_pool


def get_db():
    """Devuelve una conexión del pool (o una directa si el pool no está disponible).
    Se reutiliza durante todo el request — una sola apertura TCP por página."""
    if "_db" not in g:
        pool = _get_pool()
        g._db = pool.connection() if pool else get_mysql()
    return g._db


# ════════════════════════════════════════════════════════════════════════
# RANDOM ERP — SQL Server READ-ONLY (Cuatro capas de seguridad)
# ════════════════════════════════════════════════════════════════════════
#
# Conexión directa a SQL Server del ERP Random (cloud.random.cl:8058).
# El usuario `usr_sport` tiene rol db_owner por ahora — mientras Random
# nos crea un usuario db_datareader puro, este código IMPONE read-only
# desde la aplicación con 4 capas:
#
#   Capa 1: WHITELIST — _random_sql_query() solo acepta SELECT/WITH
#   Capa 2: BLACKLIST — bloquea palabras peligrosas (INSERT, DROP, etc.)
#   Capa 3: PARAMETRIZACIÓN — pymssql con %s, imposible SQL injection
#   Capa 4: AUTOCOMMIT OFF — sin commit explícito, escrituras no persisten
#
# Único punto de entrada para Random ERP. Cualquier código que necesite
# leer del ERP DEBE pasar por _random_sql_query() o _random_sql_one().
# ════════════════════════════════════════════════════════════════════════

RANDOM_SQL_CFG = {
    "server":   os.environ.get("RANDOM_SQL_HOST", "cloud.random.cl"),
    "port":     int(os.environ.get("RANDOM_SQL_PORT", "8058")),
    "user":     os.environ.get("RANDOM_SQL_USER", "").strip(),
    "password": os.environ.get("RANDOM_SQL_PASS", "").strip(),
    "database": os.environ.get("RANDOM_SQL_DB",   "rd095bd01"),
}

_random_pool      = None
_random_pool_lock = threading.Lock()

# Capa 2: tokens prohibidos (case-insensitive). Si alguno aparece en el SQL,
# la query es rechazada antes de tocar la BD.
_RANDOM_FORBIDDEN_TOKENS = (
    "INSERT", "UPDATE", "DELETE", "DROP", "ALTER", "TRUNCATE",
    "EXEC", "EXECUTE", "MERGE", "GRANT", "REVOKE", "CREATE",
    "BACKUP", "RESTORE", "SHUTDOWN", "OPENROWSET", "OPENQUERY",
    "BULK", "DBCC", "KILL", "RECONFIGURE",
    "INTO ", "; ", "/*", "*/",
    "XP_CMDSHELL", "SP_CONFIGURE", "SP_EXECUTESQL",
)


def _random_sql_pool():
    """Pool perezoso de conexiones a Random SQL Server."""
    global _random_pool
    if _random_pool is not None:
        return _random_pool
    if not RANDOM_SQL_CFG["user"] or not RANDOM_SQL_CFG["password"]:
        return None  # No configurado — endpoints lo manejan devolviendo lista vacía
    with _random_pool_lock:
        if _random_pool is not None:
            return _random_pool
        try:
            import pymssql
            from dbutils.pooled_db import PooledDB
            _random_pool = PooledDB(
                creator        = pymssql,
                mincached      = 1,
                maxcached      = 3,
                maxconnections = 5,
                blocking       = True,
                ping           = 1,
                server         = RANDOM_SQL_CFG["server"],
                port           = RANDOM_SQL_CFG["port"],
                user           = RANDOM_SQL_CFG["user"],
                password       = RANDOM_SQL_CFG["password"],
                database       = RANDOM_SQL_CFG["database"],
                login_timeout  = 10,
                timeout        = 25,
                autocommit     = False,   # Capa 4: sin commit, escrituras NO persisten
                as_dict        = True,
            )
            print(f"[RANDOM SQL] Pool activo → {RANDOM_SQL_CFG['server']}:{RANDOM_SQL_CFG['port']}/{RANDOM_SQL_CFG['database']}")
        except ImportError:
            print("[RANDOM SQL][WARN] pymssql no instalado — instalar con: pip install pymssql")
            _random_pool = None
        except Exception as e:
            print(f"[RANDOM SQL][ERROR] No se pudo crear pool: {e}")
            _random_pool = None
    return _random_pool


def _random_sql_validate(sql: str) -> None:
    """
    Capa 1 + 2 de seguridad. Lanza PermissionError si la query no es segura.
    Se ejecuta SIEMPRE antes de tocar la BD.
    """
    if not sql or not isinstance(sql, str):
        raise PermissionError("Random SQL: query vacía o no es string")
    sql_clean = sql.strip()
    sql_upper = sql_clean.upper()
    # Capa 1: solo SELECT (también permitimos CTE con WITH)
    first_token = sql_upper.split(None, 1)[0] if sql_upper else ""
    if first_token not in ("SELECT", "WITH"):
        raise PermissionError(f"Random SQL: solo SELECT permitido (recibido: '{first_token}')")
    # Capa 2: blacklist de tokens en cualquier lugar
    for tok in _RANDOM_FORBIDDEN_TOKENS:
        if tok in sql_upper:
            raise PermissionError(f"Random SQL: token prohibido en query: '{tok.strip()}'")


def _random_sql_query(sql: str, params=None, max_rows: int = 500):
    """
    ÚNICO punto de entrada permitido para Random ERP SQL Server.
    Devuelve lista de dicts o None si no hay configuración / error de conexión.
    Lanza PermissionError si el SQL viola las reglas de seguridad.

    NUNCA construir SQL con f-strings o concatenación. SIEMPRE usar params.
    Ejemplo correcto:
        _random_sql_query("SELECT TOP 10 * FROM MAEEN WHERE RTEN LIKE %s", (f"{rut}%",))
    """
    _random_sql_validate(sql)   # Capas 1 y 2 — falla antes de conectar
    pool = _random_sql_pool()
    if pool is None:
        return None
    conn = None
    try:
        conn = pool.connection()
        with conn.cursor(as_dict=True) as cur:
            cur.execute(sql, params or ())
            rows = cur.fetchmany(max_rows) if max_rows else cur.fetchall()
        # Capa 4: NO llamamos conn.commit() — cualquier escritura accidental
        # se descarta al cerrar la conexión.
        return [dict(r) for r in rows]
    except PermissionError:
        raise
    except Exception as e:
        print(f"[RANDOM SQL][QUERY ERROR] {e}  ::  SQL={sql[:120]}  ::  params={params}")
        return None
    finally:
        if conn is not None:
            try: conn.close()
            except Exception: pass


def _random_sql_one(sql: str, params=None):
    """Versión que devuelve solo la primera fila o None."""
    rows = _random_sql_query(sql, params, max_rows=1)
    return rows[0] if rows else None


def mysql_fetchone(query, params=None):
    with get_db().cursor() as cur:
        cur.execute(query, params or ())
        return cur.fetchone()


def mysql_fetchall(query, params=None):
    with get_db().cursor() as cur:
        cur.execute(query, params or ())
        return cur.fetchall()


def mysql_execute(query, params=None):
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(query, params or ())
    conn.commit()


# ─────────────────────────────────────────────
#  Schema init
# ─────────────────────────────────────────────

def ensure_erp_table_index():
    """Añade índice en etiquetas.SKU si no existe — mejora velocidad de búsqueda."""
    try:
        conn = get_mysql()
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COUNT(*) AS cnt FROM information_schema.statistics
                WHERE table_schema = DATABASE()
                  AND table_name   = %s
                  AND index_name   = 'idx_sku'
            """, (ERP_TABLE,))
            if cur.fetchone()["cnt"] == 0:
                cur.execute(f"CREATE INDEX idx_sku ON `{ERP_TABLE}` (SKU(90))")
                conn.commit()
        conn.close()
    except Exception as e:
        print(f"[INFO] No se pudo crear índice en {ERP_TABLE}: {e}")


def sync_erp_table(product, bultos, conn):
    """
    Sincroniza los datos de app_products + app_bultos en la tabla compartida `etiquetas`.
    Usa INSERT ... ON DUPLICATE KEY UPDATE para manejar crear y editar.
    `Codigo` es la PK (INT) de esa tabla.
    """
    try:
        codigo_int = int(product.get("codigo") or 0)
        if not codigo_int:
            return

        pv_calc = lambda l, a, al: round(float(l or 0) * float(a or 0) * float(al or 0) / 4000.0, 4)

        sorted_b = sorted(bultos, key=lambda b: int(b.get("bulto_num", 0)))
        peso_total = round(sum(float(b.get("peso", 0)) for b in sorted_b), 2)
        pv_total   = round(sum(pv_calc(b.get("largo",0), b.get("ancho",0), b.get("alto",0)) for b in sorted_b), 4)

        data = {
            "Codigo":                   codigo_int,
            "SKU":                      product.get("sku", ""),
            "Nombre":                   product.get("nombre", ""),
            "Estado":                   product.get("estado", "Confirmado"),
            "Bultos":                   len(sorted_b),
            "Peso Total":               peso_total,
            "Peso Volumetrico total":   pv_total,
        }

        # Por bulto (posiciones 1–27)
        for i in range(1, 28):
            suf_dim = "" if i == 1 else str(i)            # "Largo ( cm )", "Largo ( cm )2"...
            suf_pv  = "" if i == 1 else f" {i}"           # "Peso Volumetrico", "Peso Volumetrico 2"...
            b = sorted_b[i - 1] if i <= len(sorted_b) else None
            if b:
                l = float(b.get("largo", 0))
                a = float(b.get("ancho", 0))
                al = float(b.get("alto", 0))
                p  = float(b.get("peso", 0))
                data[f"Largo ( cm ){suf_dim}"] = l
                data[f"Ancho ( cm ){suf_dim}"] = a
                data[f"Alto ( cm ){suf_dim}"]  = al
                data[f"Peso (kg){suf_dim}"]    = p
                data[f"Peso Volumetrico{suf_pv}"] = pv_calc(l, a, al)
            else:
                data[f"Largo ( cm ){suf_dim}"] = 0
                data[f"Ancho ( cm ){suf_dim}"] = 0
                data[f"Alto ( cm ){suf_dim}"]  = 0
                data[f"Peso (kg){suf_dim}"]    = 0
                data[f"Peso Volumetrico{suf_pv}"] = 0

        cols   = list(data.keys())
        c_sql  = ", ".join(f"`{c}`" for c in cols)
        p_sql  = ", ".join(["%s"] * len(cols))
        upd    = ", ".join(f"`{c}`=%s" for c in cols if c != "Codigo")
        sql    = f"INSERT INTO `{ERP_TABLE}` ({c_sql}) VALUES ({p_sql}) ON DUPLICATE KEY UPDATE {upd}"

        ins_vals = [data[c] for c in cols]
        upd_vals = [data[c] for c in cols if c != "Codigo"]

        with conn.cursor() as cur:
            cur.execute(sql, ins_vals + upd_vals)
    except Exception as exc:
        print(f"[WARN] sync_erp_table: {exc}")


def delete_from_erp_table(codigo, conn):
    """Elimina el registro de la tabla etiquetas por Codigo (PK INT)."""
    try:
        codigo_int = int(codigo or 0)
        if not codigo_int:
            return
        with conn.cursor() as cur:
            cur.execute(f"DELETE FROM `{ERP_TABLE}` WHERE `Codigo`=%s", (codigo_int,))
    except Exception as exc:
        print(f"[WARN] delete_from_erp_table: {exc}")


def init_mysql_schema():
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{AUTH_TABLE}` (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    username     VARCHAR(190) NOT NULL UNIQUE,
                    nombre       VARCHAR(190) NOT NULL,
                    password_hash VARCHAR(255) NOT NULL,
                    phone        VARCHAR(40)  DEFAULT NULL,
                    role         VARCHAR(20)  NOT NULL DEFAULT 'editor',
                    active       TINYINT(1)   NOT NULL DEFAULT 1,
                    created_at   TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    updated_at   TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            try:
                cur.execute(f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN phone VARCHAR(40) DEFAULT NULL AFTER password_hash")
            except Exception:
                pass
            # Datos de calidad opcionales del usuario (perfil) — todos NULL-able
            for col_sql in [
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN foto_url   VARCHAR(500) DEFAULT NULL COMMENT 'URL de foto de perfil (Cloudinary)'",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN rut        VARCHAR(20)  DEFAULT NULL",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN cargo      VARCHAR(120) DEFAULT NULL",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN genero     VARCHAR(20)  DEFAULT NULL COMMENT 'masculino|femenino|otro|prefiero_no_decir'",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN direccion  VARCHAR(300) DEFAULT NULL",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN comuna     VARCHAR(100) DEFAULT NULL",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN ciudad     VARCHAR(100) DEFAULT NULL",
                f"ALTER TABLE `{AUTH_TABLE}` ADD COLUMN fecha_nac  DATE         DEFAULT NULL",
            ]:
                try:
                    cur.execute(col_sql)
                except Exception:
                    pass
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PRODUCTS_TABLE}` (
                    id         INT AUTO_INCREMENT PRIMARY KEY,
                    sku        VARCHAR(120) NOT NULL UNIQUE,
                    nombre     VARCHAR(255) NOT NULL,
                    estado     VARCHAR(40)  NOT NULL DEFAULT 'Pendiente',
                    codigo     VARCHAR(120) NOT NULL DEFAULT '',
                    erp_sync   TINYINT(1)   NOT NULL DEFAULT 0,
                    created_by VARCHAR(190) DEFAULT NULL,
                    updated_by VARCHAR(190) DEFAULT NULL,
                    created_at TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{BULTOS_TABLE}` (
                    id         INT AUTO_INCREMENT PRIMARY KEY,
                    product_id INT  NOT NULL,
                    bulto_num  INT  NOT NULL,
                    largo      DECIMAL(12,2) NOT NULL DEFAULT 0,
                    ancho      DECIMAL(12,2) NOT NULL DEFAULT 0,
                    alto       DECIMAL(12,2) NOT NULL DEFAULT 0,
                    peso       DECIMAL(12,2) NOT NULL DEFAULT 0,
                    UNIQUE KEY uniq_product_bulto (product_id, bulto_num),
                    CONSTRAINT fk_bulto_product
                        FOREIGN KEY (product_id) REFERENCES `{PRODUCTS_TABLE}`(id)
                        ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PHOTOS_TABLE}` (
                    id         INT AUTO_INCREMENT PRIMARY KEY,
                    product_id INT          NOT NULL,
                    filename   VARCHAR(255) NOT NULL,
                    orden      INT          NOT NULL DEFAULT 1,
                    created_at TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    CONSTRAINT fk_photo_product
                        FOREIGN KEY (product_id) REFERENCES `{PRODUCTS_TABLE}`(id)
                        ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            for username, nombre, role, password in DEFAULT_USERS:
                cur.execute(f"SELECT id FROM `{AUTH_TABLE}` WHERE username=%s", (username,))
                existing = cur.fetchone()
                if existing:
                    cur.execute(
                        f"UPDATE `{AUTH_TABLE}` SET nombre=%s, role=%s, active=1 WHERE id=%s",
                        (nombre, role, existing["id"]),
                    )
                else:
                    cur.execute(
                        f"INSERT INTO `{AUTH_TABLE}` (username,nombre,password_hash,role,active) VALUES (%s,%s,%s,%s,1)",
                        (username, nombre, generate_password_hash(password), role),
                    )
        conn.commit()
    finally:
        conn.close()


def init_resets_table():
    """Crea la tabla de tokens de recuperación de contraseña."""
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"""
            CREATE TABLE IF NOT EXISTS `{RESETS_TABLE}` (
                id         INT AUTO_INCREMENT PRIMARY KEY,
                user_id    INT NOT NULL,
                token      VARCHAR(120) NOT NULL UNIQUE,
                expires_at DATETIME NOT NULL,
                used       TINYINT(1) NOT NULL DEFAULT 0,
                created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                INDEX idx_token (token),
                INDEX idx_user  (user_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """)
    conn.commit()


def init_transporte_tables():
    """Crea tablas del módulo Transporte y Distribución."""
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_commitments (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    tido            VARCHAR(5)   NOT NULL,
                    nudo            VARCHAR(15)  NOT NULL,
                    endo            VARCHAR(20)  NOT NULL DEFAULT '',
                    fecha_emision   DATE,
                    fecha_entrega   DATE,
                    cliente_nombre  VARCHAR(200),
                    cliente_rut     VARCHAR(20),
                    comuna          VARCHAR(100),
                    direccion       VARCHAR(300),
                    telefono        VARCHAR(50),
                    email           VARCHAR(150),
                    valor_neto      DECIMAL(14,2) DEFAULT 0,
                    valor_bruto     DECIMAL(14,2) DEFAULT 0,
                    costo_zz        DECIMAL(10,2) DEFAULT 0,
                    tiene_saldo     TINYINT(1) DEFAULT 1,
                    guia_numero     VARCHAR(20),
                    estado          VARCHAR(50) DEFAULT 'Pendiente',
                    clasificacion   ENUM('despacho','retiro','instalacion','mantencion','garantia') DEFAULT 'despacho',
                    fecha_agenda    DATE,
                    notas           TEXT,
                    erp_synced_at   DATETIME,
                    created_by      VARCHAR(190),
                    updated_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    UNIQUE KEY uq_doc (tido, nudo),
                    INDEX idx_saldo    (tiene_saldo),
                    INDEX idx_estado   (estado),
                    INDEX idx_fecha    (fecha_emision),
                    INDEX idx_clasif   (clasificacion)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migrar columna clasificacion si ya existía con ENUM antiguo
            try:
                cur.execute("""ALTER TABLE transport_commitments
                    MODIFY COLUMN clasificacion
                    ENUM('despacho','retiro','instalacion','mantencion','garantia')
                    DEFAULT 'despacho'""")
            except Exception:
                pass
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_commitment_lines (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    commitment_id   INT NOT NULL,
                    koprct          VARCHAR(30) NOT NULL,
                    nokopr          VARCHAR(300),
                    cantidad        DECIMAL(12,3) DEFAULT 0,
                    cant_despachada DECIMAL(12,3) DEFAULT 0,
                    saldo           DECIMAL(12,3) DEFAULT 0,
                    bodega          VARCHAR(10),
                    peso_unitario   DECIMAL(10,3) DEFAULT 0,
                    volumen_unitario DECIMAL(14,2) DEFAULT 0,
                    FOREIGN KEY (commitment_id)
                        REFERENCES transport_commitments(id) ON DELETE CASCADE,
                    INDEX idx_comm (commitment_id),
                    INDEX idx_koprct (koprct),
                    INDEX idx_comm_saldo (commitment_id, saldo)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migración: agregar índices si la tabla ya existía
            for _idx in [
                "ALTER TABLE transport_commitment_lines ADD INDEX idx_koprct (koprct)",
                "ALTER TABLE transport_commitment_lines ADD INDEX idx_comm_saldo (commitment_id, saldo)",
            ]:
                try: cur.execute(_idx)
                except Exception: pass
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_manifests (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    correlativo     VARCHAR(20) UNIQUE,
                    fecha           DATE NOT NULL,
                    courier         VARCHAR(80) NOT NULL,
                    estado          ENUM('En preparación','En curso','Cerrado','Entregado completo')
                                    DEFAULT 'En preparación',
                    total_items     INT DEFAULT 0,
                    peso_total      DECIMAL(10,3) DEFAULT 0,
                    vol_total       DECIMAL(14,2) DEFAULT 0,
                    peso_pred_total DECIMAL(10,3) DEFAULT 0,
                    costo_total     DECIMAL(12,2) DEFAULT 0,
                    notas           TEXT,
                    created_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_manifest_items (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    manifest_id     INT NOT NULL,
                    commitment_id   INT NOT NULL,
                    orden           INT DEFAULT 0,
                    estado_entrega  ENUM(
                        'En preparación','Entregado a transporte',
                        'En ruta','Entregado','Entrega fallida','Devolución'
                    ) DEFAULT 'En preparación',
                    added_at        DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (manifest_id)
                        REFERENCES transport_manifests(id) ON DELETE CASCADE,
                    FOREIGN KEY (commitment_id)
                        REFERENCES transport_commitments(id),
                    UNIQUE KEY uq_item (manifest_id, commitment_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_logs (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    entity_type ENUM('commitment','manifest','manifest_item') NOT NULL,
                    entity_id   INT NOT NULL,
                    accion      VARCHAR(80) NOT NULL,
                    detalle     TEXT,
                    usuario     VARCHAR(190),
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_entity (entity_type, entity_id),
                    INDEX idx_user   (usuario)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── COURIERS ──────────────────────────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_couriers (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    nombre          VARCHAR(120) NOT NULL,
                    rut             VARCHAR(20),
                    contacto        VARCHAR(120),
                    telefono        VARCHAR(50),
                    email           VARCHAR(150),
                    tipo            ENUM('nacional','regional','local','internacional') DEFAULT 'nacional',
                    activo          TINYINT(1) DEFAULT 1,
                    notas           TEXT,
                    logo_url        VARCHAR(400),
                    peso_max_bulto  DECIMAL(10,2) DEFAULT 0   COMMENT 'kg máx por bulto (0=sin límite)',
                    peso_max_guia   DECIMAL(10,2) DEFAULT 0   COMMENT 'kg máx por guía (0=sin límite)',
                    vol_max_bulto   DECIMAL(12,2) DEFAULT 0   COMMENT 'cm³ máx por bulto (0=sin límite)',
                    factor_vol      DECIMAL(10,4) DEFAULT 5000 COMMENT 'divisor peso volumétrico (cm³/kg)',
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_courier_tarifas (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    courier_id      INT NOT NULL,
                    zona            VARCHAR(80) NOT NULL DEFAULT 'General',
                    peso_desde      DECIMAL(10,3) NOT NULL DEFAULT 0,
                    peso_hasta      DECIMAL(10,3) NOT NULL DEFAULT 0   COMMENT '0=sin tope',
                    precio_base     DECIMAL(10,2) NOT NULL DEFAULT 0,
                    precio_kg_extra DECIMAL(10,2) DEFAULT 0   COMMENT 'precio por kg sobre peso_desde',
                    moneda          CHAR(3) DEFAULT 'CLP',
                    activo          TINYINT(1) DEFAULT 1,
                    FOREIGN KEY (courier_id)
                        REFERENCES transport_couriers(id) ON DELETE CASCADE,
                    INDEX idx_courier (courier_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── COMMUNE-BASED PRICING & CONTRACTS ────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_courier_comunas (
                    id            INT AUTO_INCREMENT PRIMARY KEY,
                    courier_id    INT NOT NULL,
                    codigo        VARCHAR(20),
                    sucursal      VARCHAR(120),
                    comuna        VARCHAR(120) NOT NULL,
                    zona          VARCHAR(80),
                    region        VARCHAR(80),
                    dias_transito VARCHAR(20),
                    dias_entrega  VARCHAR(50),
                    precios_json  MEDIUMTEXT,
                    FOREIGN KEY (courier_id) REFERENCES transport_couriers(id) ON DELETE CASCADE,
                    INDEX idx_courier_comuna (courier_id, comuna),
                    UNIQUE KEY uq_courier_comuna (courier_id, comuna(100))
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS transport_courier_contratos (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    courier_id   INT NOT NULL,
                    nombre       VARCHAR(200),
                    descripcion  TEXT,
                    archivo_url  VARCHAR(400),
                    tipo         ENUM('contrato','tarifario','acuerdo','otro') DEFAULT 'contrato',
                    vigente      TINYINT(1) DEFAULT 1,
                    fecha_inicio DATE,
                    fecha_fin    DATE,
                    subido_por   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (courier_id) REFERENCES transport_couriers(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migrations for transport_couriers additional fields
            for _mig in [
                "ALTER TABLE transport_couriers ADD COLUMN website VARCHAR(200)",
                "ALTER TABLE transport_couriers ADD COLUMN direccion VARCHAR(300)",
                "ALTER TABLE transport_couriers ADD COLUMN logo_square_url VARCHAR(400)",
                "ALTER TABLE transport_couriers ADD COLUMN logo_label_url VARCHAR(400)",
                "ALTER TABLE transport_couriers ADD COLUMN nombre_fantasia VARCHAR(120)",
                "ALTER TABLE transport_couriers ADD COLUMN giro VARCHAR(150)",
                "ALTER TABLE transport_couriers ADD COLUMN contacto_cargo VARCHAR(120)",
                "ALTER TABLE transport_couriers ADD COLUMN renovacion_automatica TINYINT(1) DEFAULT 0",
            ]:
                try: cur.execute(_mig)
                except Exception: pass
            # ── DEFAULT COURIERS (solo si la tabla está vacía) ─────────
            cur.execute("SELECT COUNT(*) AS n FROM transport_couriers")
            row = cur.fetchone()
            if (row or {}).get('n', 1) == 0:
                _defaults = [
                    ('FedEx',                'internacional'),
                    ('Transportes Melling',  'nacional'),
                    ('Transporte Felca',     'regional'),
                    ('Envíame',              'nacional'),
                ]
                for _nom, _tipo in _defaults:
                    cur.execute(
                        "INSERT INTO transport_couriers (nombre, tipo) VALUES (%s, %s)",
                        (_nom, _tipo)
                    )
        conn.commit()
    finally:
        conn.close()


PICKUP_REQUESTS_TABLE = "pickup_requests"
PICKUP_PACKAGES_TABLE = "pickup_packages"
PICKUP_PROPOSALS_TABLE = "pickup_proposals"
PICKUP_LOGS_TABLE = "pickup_logs"
PICKUP_ATTACHMENTS_TABLE = "pickup_attachments"
PICKUP_SIGNATURES_TABLE = "pickup_signatures"
PICKUP_SETTINGS_TABLE = "pickup_settings"
PICKUP_TEMPLATES_TABLE = "pickup_templates"


def init_pickup_tables():
    """Crea el modulo de solicitudes de retiro en bodega."""
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_REQUESTS_TABLE}` (
                    id                         INT AUTO_INCREMENT PRIMARY KEY,
                    code                       VARCHAR(30) UNIQUE,
                    document_type              VARCHAR(30) NOT NULL,
                    document_number            VARCHAR(60) NOT NULL,
                    customer_name              VARCHAR(200) NOT NULL,
                    customer_rut               VARCHAR(30) NOT NULL,
                    contact_name               VARCHAR(160) NOT NULL,
                    contact_email              VARCHAR(180) NOT NULL,
                    contact_phone              VARCHAR(60) NOT NULL,
                    pickup_person_name         VARCHAR(160) NOT NULL,
                    pickup_person_rut          VARCHAR(30) NOT NULL,
                    pickup_person_phone        VARCHAR(60) NOT NULL,
                    pickup_person_relation     VARCHAR(40) NOT NULL,
                    requested_date             DATE NOT NULL,
                    requested_time_from        TIME NOT NULL,
                    requested_time_to          TIME NOT NULL,
                    proposed_date              DATE NULL,
                    proposed_time_from         TIME NULL,
                    proposed_time_to           TIME NULL,
                    confirmed_date             DATE NULL,
                    confirmed_time_from        TIME NULL,
                    confirmed_time_to          TIME NULL,
                    status                     VARCHAR(40) NOT NULL DEFAULT 'solicitud_recibida',
                    information_quality_score  INT DEFAULT 0,
                    risk_score                 INT DEFAULT 0,
                    total_packages             INT DEFAULT 0,
                    total_weight_kg            DECIMAL(12,3) DEFAULT 0,
                    total_volumetric_weight    DECIMAL(12,3) DEFAULT 0,
                    total_volume_m3            DECIMAL(12,4) DEFAULT 0,
                    invoice_total_amount       DECIMAL(14,2) DEFAULT 0,
                    observations               TEXT,
                    internal_notes             TEXT,
                    public_token               VARCHAR(160) NOT NULL,
                    signature_status           VARCHAR(30) DEFAULT 'pendiente',
                    created_ip                 VARCHAR(80),
                    created_user_agent         VARCHAR(300),
                    created_at                 DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at                 DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    closed_at                  DATETIME NULL,
                    INDEX idx_pickup_status (status),
                    INDEX idx_pickup_date (requested_date),
                    INDEX idx_pickup_doc (document_type, document_number),
                    INDEX idx_pickup_token (public_token)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_PACKAGES_TABLE}` (
                    id                  INT AUTO_INCREMENT PRIMARY KEY,
                    request_id          INT NOT NULL,
                    package_number      INT NOT NULL,
                    length_cm           DECIMAL(10,2) DEFAULT 0,
                    width_cm            DECIMAL(10,2) DEFAULT 0,
                    height_cm           DECIMAL(10,2) DEFAULT 0,
                    weight_kg           DECIMAL(10,3) DEFAULT 0,
                    volumetric_weight   DECIMAL(10,3) DEFAULT 0,
                    volume_m3           DECIMAL(12,4) DEFAULT 0,
                    FOREIGN KEY (request_id) REFERENCES `{PICKUP_REQUESTS_TABLE}`(id) ON DELETE CASCADE,
                    INDEX idx_pickup_pkg (request_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_PROPOSALS_TABLE}` (
                    id             INT AUTO_INCREMENT PRIMARY KEY,
                    request_id     INT NOT NULL,
                    proposed_by    VARCHAR(30) NOT NULL,
                    date           DATE NOT NULL,
                    time_from      TIME NOT NULL,
                    time_to        TIME NOT NULL,
                    message        TEXT,
                    reason         VARCHAR(200),
                    status         VARCHAR(30) DEFAULT 'pending',
                    token          VARCHAR(160),
                    created_at     DATETIME DEFAULT CURRENT_TIMESTAMP,
                    answered_at    DATETIME NULL,
                    FOREIGN KEY (request_id) REFERENCES `{PICKUP_REQUESTS_TABLE}`(id) ON DELETE CASCADE,
                    INDEX idx_pickup_prop (request_id, status)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_LOGS_TABLE}` (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    request_id  INT NOT NULL,
                    actor_type  VARCHAR(30) NOT NULL,
                    actor_name  VARCHAR(180),
                    action      VARCHAR(80) NOT NULL,
                    old_status  VARCHAR(40),
                    new_status  VARCHAR(40),
                    notes       TEXT,
                    ip          VARCHAR(80),
                    user_agent  VARCHAR(300),
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (request_id) REFERENCES `{PICKUP_REQUESTS_TABLE}`(id) ON DELETE CASCADE,
                    INDEX idx_pickup_log (request_id, created_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_ATTACHMENTS_TABLE}` (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    request_id      INT NOT NULL,
                    filename        VARCHAR(260) NOT NULL,
                    original_name   VARCHAR(260),
                    mime_type       VARCHAR(120),
                    uploaded_by     VARCHAR(80) DEFAULT 'cliente',
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (request_id) REFERENCES `{PICKUP_REQUESTS_TABLE}`(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_SIGNATURES_TABLE}` (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    request_id      INT NOT NULL,
                    signer_name     VARCHAR(180) NOT NULL,
                    signer_rut      VARCHAR(30) NOT NULL,
                    accepted_terms  TINYINT(1) DEFAULT 1,
                    ip              VARCHAR(80),
                    user_agent      VARCHAR(300),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (request_id) REFERENCES `{PICKUP_REQUESTS_TABLE}`(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_SETTINGS_TABLE}` (
                    id              INT PRIMARY KEY DEFAULT 1,
                    warehouse_name  VARCHAR(160) DEFAULT 'Bodega ILUS Quilicura',
                    warehouse_addr  VARCHAR(260) DEFAULT 'Av. Presidente Eduardo Frei Montalva 9770, Bod 30, Quilicura.',
                    maps_url        VARCHAR(500) DEFAULT 'https://www.google.com/maps/search/?api=1&query=Av.%20Presidente%20Eduardo%20Frei%20Montalva%209770%20Bod%2030%20Quilicura',
                    open_time       TIME DEFAULT '09:00:00',
                    close_time      TIME DEFAULT '17:30:00',
                    work_days       VARCHAR(30) DEFAULT '1,2,3,4,5',
                    holidays        TEXT,
                    alert_enabled   TINYINT(1) DEFAULT 0,
                    alert_title     VARCHAR(160) DEFAULT 'Aviso importante',
                    alert_message   TEXT,
                    hero_image_1    VARCHAR(260),
                    hero_image_2    VARCHAR(260),
                    hero_image_3    VARCHAR(260),
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PICKUP_TEMPLATES_TABLE}` (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    code        VARCHAR(60) UNIQUE,
                    title       VARCHAR(180) NOT NULL,
                    body        TEXT NOT NULL,
                    channel     VARCHAR(30) DEFAULT 'email_whatsapp',
                    active      TINYINT(1) DEFAULT 1
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Tabla de bloqueos de retiros: días u horas no disponibles
            cur.execute("""
                CREATE TABLE IF NOT EXISTS pickup_blocks (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    fecha        DATE NOT NULL,
                    hora_inicio  TIME NULL
                                 COMMENT 'NULL = bloquea día completo',
                    hora_fin     TIME NULL
                                 COMMENT 'NULL = bloquea hasta el cierre',
                    motivo       VARCHAR(200) DEFAULT '',
                    created_by   VARCHAR(190) DEFAULT NULL,
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_fecha (fecha)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migración: capacidad por franja (peso/volumen/cupos) + colación + step
            for _mig in [
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN slot_minutes INT DEFAULT 60 COMMENT 'Duración de cada franja en minutos'",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN slot_step_minutes INT DEFAULT 30 COMMENT 'Cada cuántos min se ofrece un nuevo bloque'",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN lunch_start TIME DEFAULT '13:00:00' COMMENT 'Inicio colación'",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN lunch_end TIME DEFAULT '14:00:00' COMMENT 'Fin colación'",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN max_picks_per_slot INT DEFAULT 5",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN max_kg_per_slot DECIMAL(10,2) DEFAULT 500.00",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN max_m3_per_slot DECIMAL(10,3) DEFAULT 5.000",
                f"ALTER TABLE `{PICKUP_SETTINGS_TABLE}` ADD COLUMN max_picks_per_day INT DEFAULT 30",
            ]:
                try: cur.execute(_mig)
                except Exception: pass

            # Migración: workflow de validación de documentación (proceso interno)
            # Permite que ILUS valide los datos del cliente antes de proponer fecha.
            for _mig in [
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN doc_validation_status "
                f"ENUM('pendiente','en_revision','ok','incompleto') DEFAULT 'pendiente' "
                f"COMMENT 'Estado de validación de la documentación del cliente'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN doc_validated_at DATETIME NULL "
                f"COMMENT 'Fecha en que se validó la documentación'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN doc_validated_by VARCHAR(190) NULL "
                f"COMMENT 'Usuario que validó la documentación'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN doc_erp_data MEDIUMTEXT NULL "
                f"COMMENT 'JSON con snapshot del documento del ERP (header + lineas)'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN doc_validation_notes TEXT NULL "
                f"COMMENT 'Notas internas del validador'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN peso_real_kg DECIMAL(10,2) DEFAULT NULL "
                f"COMMENT 'Peso real total calculado desde catálogo'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN peso_vol_kg DECIMAL(10,2) DEFAULT NULL "
                f"COMMENT 'Peso volumétrico total calculado desde catálogo'",
                f"ALTER TABLE `{PICKUP_REQUESTS_TABLE}` ADD COLUMN tiempo_estimado_min INT DEFAULT NULL "
                f"COMMENT 'Tiempo estimado de retiro en minutos (basado en bultos)'",
            ]:
                try: cur.execute(_mig)
                except Exception: pass
            cur.execute(f"INSERT IGNORE INTO `{PICKUP_SETTINGS_TABLE}` (id) VALUES (1)")
            cur.execute(
                f"""UPDATE `{PICKUP_SETTINGS_TABLE}`
                    SET warehouse_addr='Av. Presidente Eduardo Frei Montalva 9770, Bod 30, Quilicura.',
                        maps_url='https://www.google.com/maps/search/?api=1&query=Av.%20Presidente%20Eduardo%20Frei%20Montalva%209770%20Bod%2030%20Quilicura'
                    WHERE id=1
                      AND (warehouse_addr IS NULL OR warehouse_addr='' OR warehouse_addr='Bodega principal ILUS, Quilicura')"""
            )
            defaults = [
                ("cierre_anticipado", "Bodega cierra anticipadamente", "Hoy la bodega cerrara anticipadamente. Si tu retiro se ve afectado, te propondremos un nuevo horario."),
                ("sin_disponibilidad", "Sin disponibilidad", "No tenemos disponibilidad para el horario solicitado. Te enviamos una propuesta alternativa para confirmar."),
                ("info_incompleta", "Falta informacion", "Necesitamos completar informacion del documento, contacto o persona autorizada para poder avanzar con tu retiro."),
                ("validar_identidad", "Validacion de identidad", "Para proteger tu pedido, necesitamos validar la identidad de la persona autorizada para retirar."),
                ("recordatorio", "Recordatorio de retiro", "Recuerda presentarte con documento de identidad y la autorizacion correspondiente si retira un tercero."),
            ]
            for code, title, body in defaults:
                cur.execute(
                    f"""INSERT IGNORE INTO `{PICKUP_TEMPLATES_TABLE}` (code,title,body)
                        VALUES (%s,%s,%s)""",
                    (code, title, body),
                )
        conn.commit()
    finally:
        conn.close()


def init_db():
    """Inicializa el esquema MySQL. Sin SQLite — todo va a MySQL."""
    init_mysql_schema()
    init_hrm_tables()
    init_eval_tables()
    init_resets_table()
    init_transporte_tables()
    init_pickup_tables()
    ensure_erp_table_index()


# ─────────────────────────────────────────────
#  Helpers de negocio
# ─────────────────────────────────────────────

def calc_pv(largo, ancho, alto):
    return round(float(largo or 0) * float(ancho or 0) * float(alto or 0) / 4000.0, 2)


def enrich(raw_bultos):
    return [{**dict(b), "peso_vol": calc_pv(b["largo"], b["ancho"], b["alto"])} for b in raw_bultos]


def to_f(value):
    try:
        return float(str(value or "0").replace(",", "."))
    except (TypeError, ValueError):
        return 0.0


def current_username():
    """Devuelve el nombre (no el correo) del usuario activo para registrar quién actuó."""
    return g.user["nombre"] if getattr(g, "user", None) else None


# ══════════════════════════════════════════════════════════════════════
#  PERMISOS — sistema unificado
#
#  Dos fuentes de verdad combinadas:
#    1) Matriz dinámica `rol_permisos` (admin edita en /admin/roles)
#       → fuente principal cuando el rol tiene filas en esa tabla
#    2) Fallback legacy hardcoded
#       → roles del sistema que aún no han sido editados en la matriz
#
#  Las "keys" de g.permissions son lo que leen templates y decoradores:
#    view / edit / create / delete / print
#    etiquetas / retiros / mantenciones / transporte / cubicador
#    comunicaciones / admin / ajustes / hrm / superadmin
#
#  Caché en proceso: `_ROLE_PERMS_CACHE[role] = dict`
#  Se invalida cuando admin guarda la matriz → `invalidate_role_cache(role)`
# ══════════════════════════════════════════════════════════════════════

PERMS_KEYS = (
    "view", "edit", "print", "create", "delete",
    "admin", "superadmin", "ajustes", "hrm",
    "cubicador", "transporte", "mantenciones",
    "etiquetas", "retiros", "comunicaciones",
)

_ROLE_PERMS_CACHE = {}   # in-process cache, busted por admin_roles_matrix_save


def _empty_perms():
    return {k: False for k in PERMS_KEYS}


def _legacy_permission_set(role):
    """Fallback hardcoded para roles del sistema NO presentes en rol_permisos."""
    base = _empty_perms()
    if role == "superadmin":
        return {k: True for k in PERMS_KEYS}
    if role == "admin":
        return {**base,
                "view": True, "edit": True, "print": True,
                "create": True, "delete": True, "admin": True, "hrm": True,
                "cubicador": True, "transporte": True, "mantenciones": True,
                "ajustes": True,
                "etiquetas": True, "retiros": True, "comunicaciones": True}
    if role == "ejecutivo":
        # Histórico: sólo mantenciones (alguna versión vieja)
        return {**base, "mantenciones": True, "view": True,
                "edit": True, "create": True, "print": True}
    if role == "editor":
        return {**base, "view": True, "edit": True, "print": True,
                "create": True, "hrm": True, "cubicador": True, "transporte": True,
                "etiquetas": True, "retiros": True}
    if role == "lector":
        return {**base, "view": True, "etiquetas": True, "retiros": True}
    if role == "vendedor":
        return {**base, "view": True, "cubicador": True, "etiquetas": True}
    return base


def _build_perms_from_matrix(role):
    """Compila g.permissions desde la matriz rol_permisos. Llama get_role_permissions()."""
    base = _empty_perms()
    try:
        matrix = get_role_permissions(role)   # {modulo: {accion: bool}}
    except Exception:
        return _legacy_permission_set(role)

    eti = matrix.get("etiquetas", {})
    ret = matrix.get("retiros", {})
    man = matrix.get("mantenciones", {})
    tra = matrix.get("transporte", {})
    com = matrix.get("comunicaciones", {})
    adm = matrix.get("admin", {})

    # Flags de módulo (gates de sidebar)
    base["etiquetas"]      = bool(eti.get("ver"))
    base["retiros"]        = bool(ret.get("ver"))
    base["mantenciones"]   = bool(man.get("ver"))
    base["transporte"]     = bool(tra.get("ver"))
    base["cubicador"]      = bool(tra.get("cubicador"))
    base["comunicaciones"] = bool(com.get("ver"))

    # Acciones legacy — mapeadas desde Etiquetas (campo "view" controla el primer producto)
    base["view"]    = bool(eti.get("ver") or ret.get("ver") or man.get("ver")
                            or tra.get("ver") or com.get("ver"))
    base["create"]  = bool(eti.get("crear")    or man.get("crear"))
    base["edit"]    = bool(eti.get("editar")   or man.get("editar"))
    base["delete"]  = bool(eti.get("eliminar") or man.get("eliminar"))
    base["print"]   = bool(eti.get("imprimir"))

    # Admin / ajustes
    base["admin"]   = bool(adm.get("usuarios") or adm.get("roles"))
    base["ajustes"] = bool(adm.get("ajustes")  or adm.get("usuarios") or adm.get("roles")
                            or adm.get("marketing") or adm.get("login_imagenes"))

    # HRM no está en la matriz actual → fallback: lo permitimos si tiene admin general
    base["hrm"]        = bool(adm.get("usuarios") or adm.get("roles"))
    base["superadmin"] = False
    return base


def _role_has_matrix_rows(role):
    """True si el rol tiene filas en rol_permisos (== matriz fue editada al menos 1 vez)."""
    try:
        row = mysql_fetchone(
            "SELECT COUNT(*) AS n FROM rol_permisos WHERE rol_slug=%s", (role,)
        )
        return bool(row and (row.get("n") or 0) > 0)
    except Exception:
        return False


def permission_set(role):
    """Devuelve dict de permisos para un rol. Con caché en proceso."""
    if role is None:
        return _empty_perms()
    if role == "superadmin":
        return {k: True for k in PERMS_KEYS}
    # Cache check
    cached = _ROLE_PERMS_CACHE.get(role)
    if cached is not None:
        return cached
    # Si la matriz tiene filas para este rol → usar matriz dinámica
    # Si no → fallback legacy hardcoded
    if _role_has_matrix_rows(role):
        perms = _build_perms_from_matrix(role)
    else:
        perms = _legacy_permission_set(role)
    _ROLE_PERMS_CACHE[role] = perms
    return perms


def invalidate_role_cache(role=None):
    """Borra el caché de permisos de un rol (o todos si role=None).
    Debe llamarse después de guardar la matriz de un rol."""
    if role:
        _ROLE_PERMS_CACHE.pop(role, None)
    else:
        _ROLE_PERMS_CACHE.clear()


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT


# ══════════════════════════════════════════════════════════════════════
#  VALIDACIONES DE DATOS — para uso con datos reales del cliente
# ══════════════════════════════════════════════════════════════════════

def normalizar_rut(rut):
    """Normaliza RUT chileno: '12.345.678-9' → '123456789'.
    Devuelve None si vacío. Conserva el dígito verificador como último char."""
    if not rut:
        return None
    s = str(rut).strip().upper()
    if not s:
        return None
    # Quitar puntos, guiones, espacios
    return s.replace(".", "").replace("-", "").replace(" ", "")


def formatear_rut(rut_normalizado):
    """Da formato chileno con puntos y guión: '123456789' → '12.345.678-9'."""
    rut = normalizar_rut(rut_normalizado)
    if not rut or len(rut) < 2:
        return rut or ""
    num, dv = rut[:-1], rut[-1]
    # Insertar puntos cada 3 dígitos
    num_rev = num[::-1]
    chunks = [num_rev[i:i+3] for i in range(0, len(num_rev), 3)]
    num_fmt = ".".join(chunks)[::-1]
    return f"{num_fmt}-{dv}"


def validar_rut(rut):
    """Valida RUT chileno con dígito verificador (algoritmo módulo 11).
    Acepta cualquier formato (con/sin puntos, con/sin guión).
    Devuelve (True, rut_normalizado) o (False, mensaje_error)."""
    rut_norm = normalizar_rut(rut)
    if not rut_norm:
        return False, "RUT vacío"
    if len(rut_norm) < 2:
        return False, "RUT muy corto"
    if len(rut_norm) > 12:
        return False, "RUT muy largo"
    num, dv = rut_norm[:-1], rut_norm[-1]
    if not num.isdigit():
        return False, "RUT contiene caracteres no numéricos"
    if dv not in "0123456789K":
        return False, "Dígito verificador inválido"

    # Calcular DV esperado con módulo 11
    suma = 0
    multiplicador = 2
    for d in reversed(num):
        suma += int(d) * multiplicador
        multiplicador = 2 if multiplicador == 7 else multiplicador + 1
    resto = suma % 11
    dv_esperado = 11 - resto
    if dv_esperado == 11:
        dv_esperado = "0"
    elif dv_esperado == 10:
        dv_esperado = "K"
    else:
        dv_esperado = str(dv_esperado)

    if dv != dv_esperado:
        return False, f"Dígito verificador inválido (esperado: {dv_esperado})"
    return True, rut_norm


_EMAIL_RE = re.compile(r"^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$")


def validar_email(email):
    """Valida formato básico de email. Devuelve (True, email_normalizado) o (False, error)."""
    if not email:
        return True, None  # email es opcional
    e = str(email).strip().lower()
    if not e:
        return True, None
    if len(e) > 200:
        return False, "Email muy largo (máx 200 caracteres)"
    if not _EMAIL_RE.match(e):
        return False, "Formato de email inválido"
    return True, e


def normalizar_telefono(tel):
    """Normaliza teléfono chileno: quita espacios, paréntesis, etc.
    Acepta '+56 9 1234 5678', '9 1234 5678', '912345678', etc.
    Devuelve número limpio con prefijo +56 si aplica."""
    if not tel:
        return None
    s = str(tel).strip()
    if not s:
        return None
    # Mantener solo dígitos y el '+' inicial
    s_clean = "".join(c for c in s if c.isdigit() or c == "+")
    if not s_clean:
        return None
    # Si no empieza con +, agregar +56 si parece chileno
    if not s_clean.startswith("+"):
        # 9XXXXXXXX (móvil chileno: 9 dígitos sin código país)
        if len(s_clean) == 9 and s_clean.startswith("9"):
            return "+56" + s_clean
        if len(s_clean) == 8:  # fijo Santiago 8 dígitos
            return "+562" + s_clean
        return s_clean  # devolver tal cual si no matchea
    return s_clean


def delete_photo_file(filename):
    """Elimina foto local o de Cloudinary según el contenido de filename."""
    _cloud_delete(filename)


def _photo_src(filename, subfolder="uploads"):
    """URL de la foto: directa si es Cloudinary, estática local si no."""
    if not filename:
        return ""
    if filename.startswith("http"):
        return filename
    return url_for("static", filename=f"{subfolder}/{filename}")


# ─────────────────────────────────────────────
#  Auth helpers
# ─────────────────────────────────────────────

def get_auth_user_by_id(user_id):
    return mysql_fetchone(
        f"SELECT id,username,nombre,password_hash,phone,role,active,foto_url,rut,cargo,genero,direccion,comuna,ciudad,fecha_nac FROM `{AUTH_TABLE}` WHERE id=%s",
        (user_id,),
    )


def get_auth_user_by_username(username):
    return mysql_fetchone(
        f"SELECT id,username,nombre,password_hash,phone,role,active,foto_url,rut,cargo,genero,direccion,comuna,ciudad,fecha_nac FROM `{AUTH_TABLE}` WHERE username=%s",
        (username,),
    )


def load_current_user():
    """
    Carga el usuario actual y sus permisos.

    Estrategia híbrida (perf + freshness):
    — Cachea user en session con TTL=10s.
    — Mientras el TTL es válido, NO toca BD (~ahorra 10-40 ms por request).
    — Cuando expira, re-lee de BD (1 query indexada).
    — `permission_set(role)` usa caché en proceso por rol (invalidado por
      `admin_roles_matrix_save`), así que cambios de matriz aplican
      a más tardar 10s después.
    — Cambios de rol del usuario aplican al expirar el TTL o tras logout.
    """
    g.user = None
    g.permissions = permission_set(None)
    user_id = session.get("user_id")
    if not user_id:
        return

    # ── Intento 1: cache de session si aún fresco (TTL 10s) ─────────
    cached = session.get("_uc")
    if cached and cached.get("id") == user_id:
        cached_ts = cached.get("ts", 0)
        if (time.time() - cached_ts) < 10:
            g.user = cached
            g.permissions = permission_set(cached["role"])
            return
        # TTL expirado → cae a fetch fresh abajo

    # ── Intento 2: consulta BD ───────────────────────────────────────
    try:
        user = get_auth_user_by_id(user_id)
    except Exception as exc:
        session.clear()
        flash(f"No fue posible validar la sesion: {exc}", "danger")
        return

    if not user or not user["active"]:
        session.clear()
        return

    user = dict(user)
    g.user = user
    g.permissions = permission_set(user["role"])

    # Renovar cache con timestamp
    session["_uc"] = {
        "id":       user["id"],
        "username": user["username"],
        "nombre":   user["nombre"],
        "role":     user["role"],
        "active":   user["active"],
        "ts":       time.time(),
    }


def login_required(view):
    @wraps(view)
    def wrapped(*a, **kw):
        if not g.user:
            flash("Inicia sesion para continuar.", "warning")
            return redirect(url_for("login", next=request.path))
        return view(*a, **kw)
    return wrapped


def require_permission(permission):
    def decorator(view):
        @wraps(view)
        def wrapped(*a, **kw):
            if not g.user:
                flash("Inicia sesion para continuar.", "warning")
                return redirect(url_for("login", next=request.path))
            # ★ superadmin SIEMPRE pasa todos los checks de permiso
            # (es la cuenta dueña del sistema; no debe quedar bloqueada
            # de ningún módulo aunque sea nuevo y no esté en su matriz)
            if g.permissions.get("superadmin"):
                return view(*a, **kw)
            if not g.permissions.get(permission):
                flash("No tienes permisos para realizar esta accion.", "danger")
                return redirect(url_for("index"))
            return view(*a, **kw)
        return wrapped
    return decorator


@app.before_request
def before_request():
    load_current_user()


@app.teardown_appcontext
def close_db(exc=None):
    db = g.pop("_db", None)
    if db is not None:
        try:
            db.close()
        except Exception:
            pass


# ─────────────────────────────────────────────
#  Product queries
# ─────────────────────────────────────────────

# ── Caché simple del listado de productos ──────────────────────
# Evita re-ejecutar el UNION ALL + 4 JOINs en cada recarga de inicio.
# Se invalida automáticamente al guardar/editar cualquier producto.
_listing_cache: dict = {}
_listing_cache_lock = threading.Lock()
_LISTING_TTL = 45   # segundos

def _invalidate_listing_cache():
    """Llama esto cada vez que se crea, edita o elimina un producto."""
    with _listing_cache_lock:
        _listing_cache.clear()


def get_product_listing(search_query=""):
    """
    Devuelve TODOS los productos visibles:
    - Productos del ERP (con o sin etiqueta en app_products)
    - Productos creados directamente en la app que NO están en el ERP
    """
    cache_key = search_query.strip().lower()
    now = time.time()
    with _listing_cache_lock:
        entry = _listing_cache.get(cache_key)
        if entry and (now - entry[1]) < _LISTING_TTL:
            return entry[0]
    erp_params = []
    app_params = []
    if search_query:
        like_u = f"%{search_query.upper()}%"
        like_p = f"%{search_query}%"
        erp_where = "WHERE UPPER(TRIM(e.`SKU`)) LIKE %s OR e.`Nombre` LIKE %s"
        app_where = "AND (p.sku LIKE %s OR p.nombre LIKE %s)"
        erp_params = [like_u, like_p]
        app_params = [like_u, like_p]
    else:
        erp_where = ""
        app_where = ""

    sql = f"""
        (
          SELECT
              UPPER(TRIM(e.`SKU`))                     AS sku,
              TRIM(COALESCE(e.`Nombre`,  ''))           AS nombre,
              COALESCE(p.estado, 'Pendiente')           AS estado,
              COALESCE(p.codigo, '')                    AS codigo,
              p.id                                      AS app_product_id,
              p.created_by,
              p.updated_by,
              COALESCE(COUNT(DISTINCT b.id),  0)        AS total_bultos,
              COALESCE(SUM(b.peso), 0)                  AS peso_total,
              ROUND(COALESCE(SUM((b.largo*b.ancho*b.alto)/4000),0),2) AS pv_total,
              COALESCE(COUNT(DISTINCT ph.id), 0)        AS total_fotos
          FROM `{ERP_TABLE}` e
          LEFT JOIN `{PRODUCTS_TABLE}` p  ON p.sku = UPPER(TRIM(e.`SKU`))
          LEFT JOIN `{BULTOS_TABLE}`   b  ON b.product_id = p.id
          LEFT JOIN `{PHOTOS_TABLE}`   ph ON ph.product_id = p.id
          {erp_where}
          GROUP BY UPPER(TRIM(e.`SKU`)), TRIM(COALESCE(e.`Nombre`,'')),
                   p.id, p.estado, p.codigo, p.created_by, p.updated_by
        )
        UNION ALL
        (
          SELECT
              p.sku                                     AS sku,
              p.nombre                                  AS nombre,
              p.estado                                  AS estado,
              p.codigo                                  AS codigo,
              p.id                                      AS app_product_id,
              p.created_by,
              p.updated_by,
              COALESCE(COUNT(DISTINCT b.id),  0)        AS total_bultos,
              COALESCE(SUM(b.peso), 0)                  AS peso_total,
              ROUND(COALESCE(SUM((b.largo*b.ancho*b.alto)/4000),0),2) AS pv_total,
              COALESCE(COUNT(DISTINCT ph.id), 0)        AS total_fotos
          FROM `{PRODUCTS_TABLE}` p
          LEFT JOIN `{BULTOS_TABLE}`   b  ON b.product_id = p.id
          LEFT JOIN `{PHOTOS_TABLE}`   ph ON ph.product_id = p.id
          LEFT JOIN `{ERP_TABLE}`      e  ON p.sku = UPPER(TRIM(e.`SKU`))
          WHERE e.`SKU` IS NULL
          {app_where}
          GROUP BY p.id, p.sku, p.nombre, p.estado, p.codigo, p.created_by, p.updated_by
        )
        ORDER BY nombre, sku
    """
    result = mysql_fetchall(sql, erp_params + app_params)
    with _listing_cache_lock:
        _listing_cache[cache_key] = (result, time.time())
    return result


def get_erp_product_by_sku(sku):
    return mysql_fetchone(
        f"""SELECT UPPER(TRIM(`SKU`)) AS sku,
                   TRIM(COALESCE(`Nombre`,'')) AS nombre,
                   TRIM(COALESCE(`Estado`,'Pendiente')) AS estado,
                   TRIM(COALESCE(`Codigo`,'')) AS codigo
            FROM `{ERP_TABLE}` WHERE UPPER(TRIM(`SKU`))=%s LIMIT 1""",
        (sku.strip().upper(),),
    )


def get_full(product_id):
    product = mysql_fetchone(f"SELECT * FROM `{PRODUCTS_TABLE}` WHERE id=%s", (product_id,))
    if not product:
        return None, [], []
    bultos = mysql_fetchall(
        f"SELECT * FROM `{BULTOS_TABLE}` WHERE product_id=%s ORDER BY bulto_num", (product_id,)
    )
    photos = mysql_fetchall(
        f"SELECT * FROM `{PHOTOS_TABLE}` WHERE product_id=%s ORDER BY orden", (product_id,)
    )
    return product, enrich(bultos), photos


def ensure_product_record_from_erp(sku):
    normalized = sku.strip().upper()
    existing = mysql_fetchone(f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE sku=%s", (normalized,))
    if existing:
        return existing["id"]
    erp = get_erp_product_by_sku(normalized)
    if not erp:
        return None
    codigo = next_codigo()
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            f"""INSERT INTO `{PRODUCTS_TABLE}` (sku,nombre,estado,codigo,erp_sync,created_by,updated_by)
                VALUES (%s,%s,%s,%s,1,%s,%s)""",
            (erp["sku"], erp["nombre"], erp["estado"], codigo,
             current_username(), current_username()),
        )
        pid = cur.lastrowid
    conn.commit()
    return pid


def save_bultos_mysql(conn, product_id, form):
    with conn.cursor() as cur:
        for idx in range(1, MAX_BULTOS + 1):
            largo = to_f(form.get(f"largo_{idx}"))
            ancho = to_f(form.get(f"ancho_{idx}"))
            alto  = to_f(form.get(f"alto_{idx}"))
            peso  = to_f(form.get(f"peso_{idx}"))
            if largo > 0 or ancho > 0 or alto > 0 or peso > 0:
                cur.execute(
                    f"""INSERT INTO `{BULTOS_TABLE}` (product_id,bulto_num,largo,ancho,alto,peso)
                        VALUES (%s,%s,%s,%s,%s,%s)""",
                    (product_id, idx, largo, ancho, alto, peso),
                )


def next_codigo():
    """Devuelve el siguiente código de impresión autoincremental (001, 002, ...)."""
    row = mysql_fetchone(
        f"SELECT MAX(CAST(codigo AS UNSIGNED)) AS max_c "
        f"FROM `{PRODUCTS_TABLE}` WHERE codigo REGEXP '^[0-9]+$'"
    )
    nxt = (int(row["max_c"]) + 1) if row and row["max_c"] else 1
    return str(nxt).zfill(3)


def validate_bultos_form(form):
    """Devuelve lista de errores si los bultos no son válidos."""
    errors = []
    has_any = False
    for idx in range(1, MAX_BULTOS + 1):
        largo = to_f(form.get(f"largo_{idx}"))
        ancho = to_f(form.get(f"ancho_{idx}"))
        alto  = to_f(form.get(f"alto_{idx}"))
        peso  = to_f(form.get(f"peso_{idx}"))
        vals  = [largo, ancho, alto, peso]
        if any(v > 0 for v in vals):
            has_any = True
            if not all(v > 0 for v in vals):
                errors.append(f"Bulto {idx}: todos los valores (largo, ancho, alto, peso) deben ser mayores a 0.")
    if not has_any:
        errors.append("Debes agregar al menos un bulto con medidas y peso completos.")
    return errors


# ─────────────────────────────────────────────
#  PDF builder — Playwright (pixel-perfect al HTML preview)
# ─────────────────────────────────────────────

def _logo_data_url():
    """Logo PNG como data:image URL — se incrusta en el HTML standalone sin depender del servidor."""
    for fname in ("Logo.png", "logo.png", "LOGO.png"):
        path = os.path.join(BASE_DIR, "static", fname)
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                return f"data:image/png;base64,{b64}"
            except Exception:
                pass
    return ""


def _label_format(fmt):
    formats = {
        "150x100": {"key": "150x100", "w": "150mm", "h": "100mm", "label": "150 x 100 mm"},
        "100x50": {"key": "100x50", "w": "100mm", "h": "50mm", "label": "100 x 50 mm"},
    }
    return formats.get(fmt, formats["150x100"])


def _label_bulto_data(bulto):
    data = dict(bulto)
    largo = float(data.get("largo_cm") or data.get("largo") or 0)
    ancho = float(data.get("ancho_cm") or data.get("ancho") or 0)
    alto = float(data.get("alto_cm") or data.get("alto") or 0)
    kg = float(data.get("kg") or data.get("peso") or data.get("peso_bruto") or 0)
    data.update({
        "largo": largo,
        "ancho": ancho,
        "alto": alto,
        "largo_cm": largo,
        "ancho_cm": ancho,
        "alto_cm": alto,
        "kg": kg,
        "peso": kg,
        "peso_vol": calc_pv(largo, ancho, alto),
    })
    return data


def build_labels_pdf(product, label_bultos, total_bultos, fmt="150x100"):
    """
    Genera un PDF con una o mas etiquetas usando Playwright (Chromium headless).
    El resultado es pixel-perfect idéntico al HTML preview porque usa
    el mismo template label_standalone.html y el mismo @media print CSS.
    Si Playwright no está instalado, lanza ImportError con instrucción de instalación.
    """
    fecha        = datetime.now().strftime("%d-%m-%Y %H:%M")
    label_format = _label_format(fmt)
    enriched     = [_label_bulto_data(b) for b in label_bultos]

    html = render_template(
        "label_standalone.html",
        product      = product,
        bultos       = enriched,
        total_bultos = total_bultos,
        fecha        = fecha,
        qty_per_bulto= {},
        logo_url     = _logo_data_url(),
        fmt          = label_format["key"],
        label_format = label_format,
    )

    # Usa el browser pool — sin overhead de launch/close
    return _pw_pdf(
        html,
        width  = label_format["w"],
        height = label_format["h"],
        wait_fn = (
            "() => {"
            "  const codes = Array.from(document.querySelectorAll('.barcode'));"
            "  return codes.length > 0 && codes.every(c => c.dataset.rendered === '1');"
            "}"
        ),
    )


def build_label_pdf(product, bulto, total_bultos, fmt="150x100"):
    return build_labels_pdf(product, [bulto], total_bultos, fmt)


# ── Mantener la antigua firma por compatibilidad ── (ya no usa reportlab)
def _build_label_pdf_legacy(product, bulto, total_bultos):
    """Fallback ReportLab — ya no se usa, conservado como referencia."""
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.graphics.barcode import code128

    W, H    = 150 * mm, 70 * mm
    HALF    = W / 2                 # 75 mm por cara
    HDR_H   = 14 * mm
    META_H  =  8 * mm
    NAME_H  =  9 * mm
    FTR_H   =  5 * mm
    hdr_y   = H - HDR_H
    meta_y  = hdr_y - META_H
    name_y  = meta_y - NAME_H
    body_y  = FTR_H
    body_h  = name_y - body_y      # ≈ 34mm

    BLACK = colors.black
    WHITE = colors.white
    GRAY  = colors.HexColor("#aaaaaa")
    DGRAY = colors.HexColor("#2c2c2c")

    buf   = io.BytesIO()
    c     = canvas.Canvas(buf, pagesize=(W, H))
    fecha = datetime.now().strftime("%d-%m-%Y %H:%M")

    # ── Intenta cargar logo (convertido a blanco via Pillow si está disponible) ──
    logo_img_path = None
    try:
        from PIL import Image as PILImage
        import tempfile
        src = os.path.join(app.static_folder, "Logo.png")
        if os.path.exists(src):
            img = PILImage.open(src).convert("RGBA")
            r, g, b, a = img.split()
            white_bg = PILImage.new("RGBA", img.size, (0, 0, 0, 255))
            # Pintar píxeles no-transparentes de blanco
            white_layer = PILImage.new("RGBA", img.size, (255, 255, 255, 255))
            white_bg.paste(white_layer, mask=a)
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            white_bg.save(tmp.name)
            logo_img_path = tmp.name
    except Exception:
        logo_img_path = None

    def draw_half(ox):
        # ══════════════════════════════════════════
        #  HEADER  (negro, 14mm)
        # ══════════════════════════════════════════
        c.setFillColor(BLACK)
        c.rect(ox, hdr_y, HALF, HDR_H, fill=1, stroke=0)

        # Proporciones: SHS izq 40% | ILUS der 60%
        shs_w  = HALF * 0.40
        ilus_w = HALF * 0.60
        shs_cx = ox + shs_w / 2
        ilus_x = ox + shs_w          # borde derecho de SHS = borde izq de ILUS
        ilus_cx = ilus_x + ilus_w / 2

        # Separador interno
        c.setStrokeColor(DGRAY)
        c.setLineWidth(0.6)
        c.line(ilus_x, hdr_y + 1*mm, ilus_x, hdr_y + HDR_H - 1*mm)

        # — Sports Health Solutions (izquierda) —
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 8)
        c.drawCentredString(shs_cx, hdr_y + 9.2*mm, "SPORTS")
        c.drawCentredString(shs_cx, hdr_y + 5.8*mm, "HEALTH")
        c.setFillColor(GRAY)
        c.setFont("Helvetica-Oblique", 4)
        c.drawCentredString(shs_cx, hdr_y + 2.2*mm, "S O L U T I O N S")

        # — Logo ILUS (derecha) —
        if logo_img_path:
            try:
                lh = HDR_H * 0.78
                lw = ilus_w * 0.82
                c.drawImage(logo_img_path,
                            ilus_cx - lw / 2,
                            hdr_y + (HDR_H - lh) / 2,
                            width=lw, height=lh,
                            preserveAspectRatio=True, mask='auto')
            except Exception:
                _draw_ilus_text(ilus_cx)
        else:
            _draw_ilus_text(ilus_cx)

    def _draw_ilus_text(cx):
        """Fallback: texto ILUS. en blanco cuando no hay imagen."""
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 18)
        iw = pdfmetrics.stringWidth("ILUS.", "Helvetica-Bold", 18)
        # cx es el centro de la zona ILUS dentro de la cara actual
        c.drawCentredString(cx, hdr_y + 4*mm, "ILUS.")

    def draw_body(ox):
        # ══════════════════════════════════════════
        #  META BAR  (blanco, 8mm)
        # ══════════════════════════════════════════
        c.setFillColor(WHITE)
        c.rect(ox, meta_y, HALF, META_H, fill=1, stroke=0)
        c.setStrokeColor(BLACK)
        c.setLineWidth(0.8)
        c.line(ox, meta_y,          ox + HALF, meta_y)
        c.line(ox, meta_y + META_H, ox + HALF, meta_y + META_H)

        ty = meta_y + META_H / 2 - 1.4*mm

        # Ícono caja (aproximación con rectángulo)
        cx_icon = ox + 5*mm
        c.setFillColor(BLACK)
        c.setLineWidth(0.6)
        c.rect(cx_icon - 1.4*mm, ty + 0.3*mm, 2.8*mm, 2.8*mm, fill=0, stroke=1)

        c.setFont("Helvetica-Bold", 6.5)
        c.setFillColor(BLACK)
        c.drawString(ox + 7.5*mm, ty,
                     f"BULTO {bulto['bulto_num']} DE {total_bultos}")

        # Separador |
        sep_x = ox + HALF * 0.52
        c.setFont("Helvetica", 10)
        c.setFillColor(GRAY)
        c.drawCentredString(sep_x, ty - 0.5*mm, "|")

        c.setFont("Helvetica-Bold", 6.5)
        c.setFillColor(BLACK)
        c.drawString(sep_x + 2.5*mm, ty, "PESO BRUTO")
        c.setFont("Helvetica-Bold", 7)
        c.drawRightString(ox + HALF - 2*mm, ty,
                          f"{float(bulto['peso']):.0f} KG")

        # ══════════════════════════════════════════
        #  NOMBRE PRODUCTO  (blanco, 9mm)
        # ══════════════════════════════════════════
        c.setFillColor(WHITE)
        c.rect(ox, name_y, HALF, NAME_H, fill=1, stroke=0)
        c.setStrokeColor(BLACK)
        c.setLineWidth(0.8)
        c.line(ox, name_y, ox + HALF, name_y)

        nombre = (product["nombre"] or "").upper()
        avail  = HALF - 5*mm
        fs_n   = 7.5
        while pdfmetrics.stringWidth(nombre, "Helvetica-Bold", fs_n) > avail and fs_n > 4.5:
            fs_n -= 0.15
        c.setFont("Helvetica-Bold", fs_n)
        c.setFillColor(BLACK)
        # Si cabe en una línea
        if pdfmetrics.stringWidth(nombre, "Helvetica-Bold", fs_n) <= avail:
            c.drawString(ox + 2.5*mm, name_y + NAME_H/2 - 1.3*mm, nombre)
        else:
            # Partir en dos líneas
            mid = len(nombre) // 2
            cut = nombre.rfind(" ", 0, mid + 10) or mid
            c.setFont("Helvetica-Bold", max(fs_n - 0.3, 5.0))
            c.drawString(ox + 2.5*mm, name_y + NAME_H - 3.8*mm, nombre[:cut].strip())
            c.drawString(ox + 2.5*mm, name_y + 1.8*mm,          nombre[cut:].strip())

        # ══════════════════════════════════════════
        #  BODY  (dims izq 50% | código der 50%)
        # ══════════════════════════════════════════
        dims_w   = HALF * 0.50
        dims_end = ox + dims_w
        bc_x     = dims_end + 0.5*mm          # inicio columna código
        bc_w     = HALF - dims_w - 0.5*mm     # ancho disponible código

        # Separador vertical dims | código
        c.setStrokeColor(BLACK)
        c.setLineWidth(0.8)
        c.line(dims_end, body_y, dims_end, name_y)

        # — Dimensiones (4 filas centradas verticalmente) —
        rows = [
            ("↔",  "LARGO:",     f"{float(bulto['largo']):.0f} cm"),
            ("↕",  "ALTO:",      f"{float(bulto['alto']):.0f} cm"),
            ("⊞",  "ANCHO:",     f"{float(bulto['ancho']):.0f} cm"),
            ("🔒", "PESO VOL.:", f"{float(bulto['peso_vol']):.2f} Vol"),
        ]
        row_h  = body_h / len(rows)
        lbl_x  = ox + 3*mm      # inicio etiqueta (3mm desde borde izq)
        val_x  = dims_end - 1*mm  # valores anclados a la derecha

        for i, (ico, lbl, val) in enumerate(rows):
            ry = name_y - (i + 0.62) * row_h

            # Etiqueta
            c.setFont("Helvetica-Bold", 5.5)
            c.setFillColor(BLACK)
            c.drawString(lbl_x, ry, lbl)

            # Valor (ancla derecha)
            c.setFont("Helvetica-Bold", 6.5)
            c.drawRightString(val_x, ry, val)

        # — Barcode CODE128 —
        bc_val = (product["sku"] or "000")[:30]

        # Reservar espacio: barcode + SKU box + fecha
        sku_h    = body_h * 0.30     # caja SKU negra
        fecha_h  = 3   * mm          # línea de fecha
        bc_h     = body_h - sku_h - fecha_h - 1.5*mm  # altura barras

        bw_pt = 1.0 if len(bc_val) > 20 else (0.85 if len(bc_val) > 12 else 1.0)
        bc_y  = body_y + sku_h + fecha_h + 0.5*mm

        try:
            bar = code128.Code128(bc_val, barWidth=bw_pt,
                                  barHeight=bc_h, humanReadable=False, quiet=False)
            while bar.width > bc_w - 1*mm and bw_pt > 0.35:
                bw_pt -= 0.04
                bar = code128.Code128(bc_val, barWidth=bw_pt,
                                      barHeight=bc_h, humanReadable=False, quiet=False)
            bc_cx = bc_x + bc_w / 2
            bar.drawOn(c, bc_cx - bar.width / 2, bc_y)
        except Exception:
            pass

        # — Caja negra CÓDIGO / SKU —
        sku_y = body_y + fecha_h
        c.setFillColor(BLACK)
        c.rect(bc_x, sku_y, bc_w, sku_h, fill=1, stroke=0)

        c.setFont("Helvetica-Bold", 3.8)
        c.setFillColor(GRAY)
        c.drawCentredString(bc_x + bc_w / 2, sku_y + sku_h - 2.2*mm, "CÓDIGO / SKU")

        sku_val = product["sku"]
        fs_s    = 9.5
        while pdfmetrics.stringWidth(sku_val, "Helvetica-Bold", fs_s) > bc_w - 2*mm and fs_s > 5:
            fs_s -= 0.3
        c.setFont("Helvetica-Bold", fs_s)
        c.setFillColor(WHITE)
        c.drawCentredString(bc_x + bc_w / 2, sku_y + sku_h * 0.20, sku_val)

        # — Fecha bajo el SKU box (horizontal, no interfiere con el barcode) —
        c.setFont("Helvetica", 3.5)
        c.setFillColor(GRAY)
        c.drawCentredString(bc_x + bc_w / 2, body_y + 0.6*mm, fecha)

    # ── Dibuja las dos caras ──────────────────────────────────────────
    draw_half(0)
    draw_body(0)
    draw_half(HALF)
    draw_body(HALF)

    # ── Borde exterior ────────────────────────────────────────────────
    c.setStrokeColor(BLACK)
    c.setLineWidth(0.8)
    c.rect(0, 0, W, H, fill=0, stroke=1)

    # ── Línea de doblez central (punteada) ───────────────────────────
    c.setStrokeColor(colors.HexColor("#999999"))
    c.setLineWidth(0.6)
    c.setDash([2.5, 2.5])
    c.line(HALF, 0, HALF, H)
    c.setDash([])

    # ── Footer: franjas diagonales B&W  (-45°, igual que el HTML) ────
    c.setFillColor(BLACK)
    c.rect(0, 0, W, FTR_H, fill=1, stroke=0)
    c.setFillColor(WHITE)
    sg, sw = 8*mm, 4*mm
    for i in range(-2, int(W / sg) + 4):
        x0 = i * sg
        p  = c.beginPath()
        p.moveTo(x0, 0);            p.lineTo(x0 + sw, 0)
        p.lineTo(x0 + sw + FTR_H, FTR_H); p.lineTo(x0 + FTR_H, FTR_H)
        p.close()
        c.drawPath(p, fill=1, stroke=0)
    c.setStrokeColor(BLACK)
    c.setLineWidth(0.8)
    c.line(0, FTR_H, W, FTR_H)

    c.save()
    buf.seek(0)

    # Limpiar archivo temporal del logo blanco
    if logo_img_path:
        try:
            os.unlink(logo_img_path)
        except Exception:
            pass

    return buf.read()   # fin de _build_label_pdf_legacy


# ─────────────────────────────────────────────
#  Context processor
# ─────────────────────────────────────────────

@app.context_processor
def inject_globals():
    return {
        "has_logo":    os.path.exists(os.path.join(app.static_folder, "Logo.png")),
        "current_user": g.user,
        "permissions":  g.permissions,
        "role_label":   {
                            "superadmin": "Super Administrador",
                            "admin":      "Administrador",
                            "ejecutivo":  "Ejecutivo",
                            "editor":     "Editor",
                            "lector":     "Lector",
                            "vendedor":   "Vendedor",
                        }.get(g.user["role"] if g.user else "", "Usuario"),
        "photo_src":    _photo_src,
    }


# ─────────────────────────────────────────────
#  API — Búsqueda en ERP externo (SOLO LECTURA)
# ─────────────────────────────────────────────

@app.route("/api/sku-lookup")
@login_required
def sku_lookup():
    """
    Búsqueda exacta por SKU — usada por el escáner de código de barras.
    Responde en <50ms sin recargar la página.
    Devuelve: {status, pid?, sku, nombre, prepare_url?}
    """
    sku = request.args.get("sku", "").strip().upper()
    if not sku:
        return jsonify({"status": "not_found", "sku": ""})

    # ── 1. Buscar en app_products (nuestra BD) ───────────────────────
    hit = mysql_fetchone(
        f"SELECT id, nombre, sku FROM `{PRODUCTS_TABLE}` WHERE UPPER(TRIM(sku)) = %s",
        (sku,)
    )
    if hit:
        return jsonify({
            "status":     "found_app",
            "pid":        hit["id"],
            "sku":        hit["sku"],
            "nombre":     (hit["nombre"] or "").strip(),
            "detail_url": url_for("product_detail", pid=hit["id"]),
            "edit_url":   url_for("edit_product",   pid=hit["id"]),
        })

    # ── 2. Buscar en tabla ERP local ─────────────────────────────────
    erp_hit = None
    try:
        erp_hit = mysql_fetchone(
            f"SELECT UPPER(TRIM(`SKU`)) AS sku, TRIM(COALESCE(`Nombre`,'')) AS nombre "
            f"FROM `{ERP_TABLE}` WHERE UPPER(TRIM(`SKU`)) = %s",
            (sku,)
        )
    except Exception:
        pass

    if erp_hit:
        return jsonify({
            "status":      "found_erp",
            "sku":         erp_hit["sku"],
            "nombre":      (erp_hit["nombre"] or "").strip(),
            "prepare_url": url_for("prepare_product_from_erp", sku=erp_hit["sku"]),
        })

    # ── 3. REST API ERP — búsqueda exacta por código ─────────────────
    try:
        TOKEN = ERP_CONFIG.get("api_token", "")
        body  = _erp_get(
            "/productos",
            {"kopr": sku, "empresa": "01", "fields": "KOPR,NOKOPR", "visible": "true"},
            TOKEN, timeout=6,
        )
        items = body.get("data") or []
        if items:
            p = items[0]
            p_sku  = (p.get("KOPR") or "").strip().upper()
            p_name = (p.get("NOKOPR") or "").strip()
            if p_sku:
                return jsonify({
                    "status":      "found_erp",
                    "sku":         p_sku,
                    "nombre":      p_name,
                    "prepare_url": url_for("prepare_product_from_erp", sku=p_sku),
                })
    except Exception:
        pass

    return jsonify({"status": "not_found", "sku": sku})


@app.route("/api/product-search")
@login_required
def product_search():
    """
    Typeahead unificado para el formulario de nuevo producto.
    Fuentes: app_products → tabla ERP local → ERP REST API (fallback).
    Devuelve JSON: [{sku, nombre, source, already_exists}]
    """
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify([])

    like_u  = f"%{q.upper()}%"
    like_p  = f"%{q}%"
    results = []
    seen    = set()

    # ── 1. app_products ──────────────────────────────────────────────
    try:
        rows = mysql_fetchall(
            f"""SELECT sku, nombre FROM `{PRODUCTS_TABLE}`
                WHERE UPPER(sku) LIKE %s OR UPPER(nombre) LIKE %s
                ORDER BY
                  CASE WHEN UPPER(sku) = %s THEN 0 ELSE 1 END,
                  nombre
                LIMIT 10""",
            (like_u, like_u, q.upper()),
        )
        for r in rows:
            sku = (r.get("sku") or "").strip().upper()
            if sku and sku not in seen:
                results.append({"sku": sku, "nombre": (r.get("nombre") or "").strip(),
                                 "source": "app", "already_exists": True})
                seen.add(sku)
    except Exception:
        pass

    # ── 2. Tabla ERP local ───────────────────────────────────────────
    try:
        rows2 = mysql_fetchall(
            f"""SELECT UPPER(TRIM(`SKU`)) AS sku,
                       TRIM(COALESCE(`Nombre`,'')) AS nombre
                FROM `{ERP_TABLE}`
                WHERE UPPER(TRIM(`SKU`)) LIKE %s OR UPPER(`Nombre`) LIKE %s
                ORDER BY
                  CASE WHEN UPPER(TRIM(`SKU`)) = %s THEN 0 ELSE 1 END,
                  `Nombre`
                LIMIT 20""",
            (like_u, like_u, q.upper()),
        )
        existing_skus = {r["sku"] for r in results}
        for r in rows2:
            sku = (r.get("sku") or "").strip().upper()
            if not sku or sku in seen:
                continue
            already = sku in existing_skus
            results.append({"sku": sku, "nombre": (r.get("nombre") or "").strip(),
                             "source": "erp", "already_exists": already})
            seen.add(sku)
    except Exception:
        pass

    # ── 3. REST API ERP /productos — fuente en tiempo real ───────────
    if len(q) >= 2:
        try:
            TOKEN = ERP_CONFIG.get("api_token", "")
            body  = _erp_get(
                "/productos",
                {
                    "search":  q,            # busca en código Y descripción
                    "empresa": "01",
                    "fields":  "KOPR,NOKOPR",
                    "visible": "true",
                    "venta":   "true",       # solo productos de venta (excluye servicios)
                },
                TOKEN, timeout=6,
            )
            existing_skus = {r["sku"] for r in results}
            for p in (body.get("data") or []):
                p_sku  = (p.get("KOPR") or "").strip().upper()
                p_name = (p.get("NOKOPR") or "").strip()
                if not p_sku or p_sku in seen:
                    continue
                already = p_sku in existing_skus
                results.append({
                    "sku":           p_sku,
                    "nombre":        p_name,
                    "source":        "erp-api",
                    "already_exists": already,
                })
                seen.add(p_sku)
                if len(results) >= 25:
                    break
        except Exception:
            pass   # si la API no responde, igual devolvemos lo local

    # ── Verificación final: consulta batch a app_products para todos los SKUs ──
    # Esto corrige casos donde la búsqueda por nombre falló (acentos/case) pero
    # el SKU SÍ existe en la BD (evita mostrar "No registrado" cuando ya está).
    if results:
        try:
            all_skus = list({r["sku"] for r in results})
            ph = ",".join(["%s"] * len(all_skus))
            verified = mysql_fetchall(
                f"SELECT UPPER(TRIM(sku)) AS sku FROM `{PRODUCTS_TABLE}` WHERE UPPER(TRIM(sku)) IN ({ph})",
                tuple(all_skus),
            )
            verified_set = {r["sku"] for r in (verified or [])}
            for r in results:
                r["already_exists"] = r["sku"] in verified_set
        except Exception:
            pass  # si falla la verificación, dejamos los valores anteriores

    # Ordenar: exactos primero, no-existentes primero, luego nombre
    results.sort(key=lambda x: (
        1 if x["already_exists"] else 0,
        0 if x["sku"] == q.upper() else 1,
        x["nombre"].lower()
    ))

    return jsonify(results[:20])


# ─────────────────────────────────────────────
#  Auth routes
# ─────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    if g.user:
        return redirect(url_for("index"))
    next_url = request.args.get("next") or request.form.get("next") or url_for("index")
    imgs = _login_images_active()
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")
        try:
            user = get_auth_user_by_username(username)
        except Exception as exc:
            flash(f"No fue posible conectar: {exc}", "danger")
            return render_template("login.html", next_url=next_url, username=username, login_images=imgs)
        if not user or not user["active"] or not check_password_hash(user["password_hash"], password):
            flash("Usuario o contraseña incorrectos.", "danger")
            return render_template("login.html", next_url=next_url, username=username, login_images=imgs)
        session.clear()
        session["user_id"] = user["id"]
        flash(f"Bienvenido, {user['nombre']}.", "success")
        return redirect(next_url)
    # Anti-cache para forzar al navegador a recargar diseño actualizado
    resp = make_response(render_template("login.html", next_url=next_url, username="", login_images=imgs))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/logout", methods=["POST"])
@login_required
def logout():
    session.clear()
    flash("Sesion cerrada.", "success")
    return redirect(url_for("login"))


# ─────────────────────────────────────────────
#  Plantilla maestra ILUS para todos los correos
# ─────────────────────────────────────────────

def _ilus_email_html(
    titulo: str,
    subtitulo: str = "",
    saludo: str = "",
    parrafos: list = None,          # lista de strings HTML
    btn_primario_txt: str = "",
    btn_primario_url: str = "",
    btn_secundario_txt: str = "",
    btn_secundario_url: str = "",
    info_lineas: list = None,       # lista de (icono, clave, valor)
) -> str:
    """
    Genera HTML de correo con el diseño oficial ILUS:
    Header negro con logo → banda oscura título/subtítulo → cuerpo blanco → botones → footer negro.
    """
    # ── Logo y empresa desde config (base64 o URL) ──────────────────────────
    try:
        cc       = _get_client_cfg()
        logo_src = cc.get("logo_url") or ""
        company  = cc.get("company_name") or "ILUS Sport &amp; Health"
    except Exception:
        logo_src = ""
        company  = "ILUS Sport &amp; Health"
    if not logo_src:
        logo_src = "https://ilusfitness.com/cdn/shop/files/Logo_ILUS_Fitness_Blanco_equipamiento_para_gimnasios.png"

    # ── Info box ─────────────────────────────────────────────────────────────
    info_html = ""
    if info_lineas:
        rows_html = "".join(
            f'<p style="margin:0 0 4px;font-size:13px;color:#333">'
            f'<strong>{k}:</strong> {v}</p>'
            for _, k, v in info_lineas
        )
        info_html = (
            f'<div style="background:#f8f8f8;border-left:4px solid #DC143C;'
            f'padding:15px 18px;margin:20px 0;border-radius:6px">'
            f'{rows_html}</div>'
        )

    # ── Párrafos ─────────────────────────────────────────────────────────────
    saludo_html = (
        f'<p style="font-size:14px;color:#111827;line-height:1.6;margin:0 0 14px">'
        f'Hola <strong>{saludo}</strong>,</p>'
    ) if saludo else ""

    body_html = "".join(
        f'<p style="font-size:14px;color:#111827;line-height:1.6;margin:0 0 14px">{p}</p>'
        for p in (parrafos or [])
    )

    # ── Botones ──────────────────────────────────────────────────────────────
    btn1 = ""
    if btn_primario_txt and btn_primario_url:
        btn1 = (
            f'<a href="{btn_primario_url}" style="display:inline-block;background:#DC143C;'
            f'color:#ffffff;padding:12px 25px;text-decoration:none;font-size:13px;'
            f'border-radius:5px;margin:5px;font-weight:bold">{btn_primario_txt}</a>'
        )
    btn2 = ""
    if btn_secundario_txt and btn_secundario_url:
        btn2 = (
            f'<a href="{btn_secundario_url}" style="display:inline-block;background:#000;'
            f'color:#ffffff;padding:12px 25px;text-decoration:none;font-size:13px;'
            f'border-radius:5px;margin:5px;font-weight:bold">{btn_secundario_txt}</a>'
        )
    btns_html = (
        f'<div style="padding:0 30px 30px;text-align:center">{btn1}{btn2}</div>'
    ) if (btn1 or btn2) else ""

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{titulo}</title>
</head>
<body style="margin:0;padding:0;background:#f1f2f4;font-family:Arial,Helvetica,sans-serif;color:#111827">

<div style="max-width:580px;margin:28px auto;background:#ffffff;border-radius:10px;
            overflow:hidden;box-shadow:0 6px 20px rgba(15,23,42,.10)">

  <!-- HEADER: negro + logo -->
  <div style="background:#000;padding:24px 28px;text-align:center">
    <img src="{logo_src}" alt="{company}"
         style="height:48px;display:block;margin:0 auto;max-width:230px;width:auto;object-fit:contain">
  </div>

  <!-- TÍTULO: banda oscura -->
  <div style="background:#111;color:#fff;text-align:center;padding:28px 32px;border-top:1px solid #202020">
    <h1 style="margin:0;font-size:22px;line-height:1.25;font-weight:800">{titulo}</h1>
    {f'<p style="margin-top:8px;font-size:13px;line-height:1.45;color:#f3f4f6;margin-bottom:0">{subtitulo}</p>' if subtitulo else ''}
  </div>

  <!-- CONTENIDO -->
  <div style="padding:32px 30px">
    {saludo_html}
    {body_html}
    {info_html}
  </div>

  <!-- BOTONES -->
  {btns_html}

  <!-- FOOTER: negro -->
  <div style="background:#000;padding:24px 32px;text-align:center">
    <div style="color:#DC143C;font-size:13px;font-weight:700;text-transform:uppercase">{company}</div>
    <div style="color:#9ca3af;font-size:11px;margin-top:5px">
      Equipamiento profesional para alto rendimiento
    </div>
    <div style="margin-top:14px;font-size:11px;line-height:1.6;color:#6b7280">
      Este correo fue generado automáticamente.<br>
      Para soporte, utiliza nuestros canales oficiales.
    </div>
  </div>

</div>
</body>
</html>"""


def _smtp_ipv4(host: str) -> str:
    """Devuelve una IP v4 del host SMTP para evitar fallas por IPv6 en algunos hosting."""
    import socket

    infos = socket.getaddrinfo(host, None, socket.AF_INET, socket.SOCK_STREAM)
    if not infos:
        return host
    return infos[0][4][0]


class _IPv4SMTP(smtplib.SMTP):
    def _get_socket(self, host, port, timeout):
        import socket

        ip = _smtp_ipv4(host)
        return socket.create_connection((ip, port), timeout, self.source_address)


class _IPv4SMTP_SSL(smtplib.SMTP_SSL):
    def _get_socket(self, host, port, timeout):
        import socket

        ip = _smtp_ipv4(host)
        sock = socket.create_connection((ip, port), timeout, self.source_address)
        return self.context.wrap_socket(sock, server_hostname=self._host)


def _open_smtp_client(host: str, port: int, secure: bool, timeout: int = 15, context=None):
    """Abre SMTP forzando IPv4, pero conserva el hostname para TLS/SNI."""
    if secure:
        return _IPv4SMTP_SSL(host, port, context=context, timeout=timeout)
    return _IPv4SMTP(host, port, timeout=timeout)


def _get_resend_cfg() -> dict:
    """
    Devuelve configuración de Resend API.
    Resend es necesario en hosting cloud (Railway, Heroku) porque Gmail/SMTP
    típicamente bloquea o limita conexiones desde rangos IP de proveedores cloud.
    Prioridad: env vars (Railway) → BD (configurado vía front).
    """
    # 1. Env var (Railway / Docker / local .env)
    api_key = os.environ.get("RESEND_API_KEY", "").strip()
    if api_key:
        return {
            "api_key":   api_key,
            "from_addr": os.environ.get("RESEND_FROM_ADDR", "onboarding@resend.dev").strip(),
            "_source":   "env",
        }
    # 2. BD
    try:
        row = mysql_fetchone("SELECT * FROM comm_resend_config ORDER BY id DESC LIMIT 1")
        if row and row.get("api_key"):
            return {
                "api_key":   row["api_key"],
                "from_addr": row.get("from_addr") or "onboarding@resend.dev",
                "_source":   "db",
            }
    except Exception:
        pass
    return {"api_key": "", "from_addr": "", "_source": ""}


def _send_via_resend(to, subject: str, html: str, from_addr: str = None) -> bool:
    """
    Envía email vía API HTTPS de Resend (no usa puertos SMTP).
    Funciona desde cualquier IP, incluyendo cloud hosting (Railway, Heroku, AWS).

    En caso de fallo, deja info en `g._last_resend_error` (dict con message/http_code/raw_body).
    """
    import urllib.request as _ur
    import urllib.error as _ue

    cfg = _get_resend_cfg()
    if not cfg.get("api_key"):
        g._last_resend_error = {
            "name": "no_config",
            "message": "RESEND_API_KEY no está configurada (env var o BD).",
            "http_code": 0, "status_code": 0, "raw_body": "",
        }
        return False

    sender = (from_addr or cfg.get("from_addr") or "onboarding@resend.dev").strip()
    recipients = [to] if isinstance(to, str) else list(to)
    payload = json.dumps({
        "from": sender,
        "to": recipients,
        "subject": subject,
        "html": html,
    }).encode("utf-8")

    req = _ur.Request(
        "https://api.resend.com/emails",
        data=payload,
        headers={
            "Authorization": f"Bearer {cfg['api_key']}",
            "Content-Type":  "application/json",
        },
        method="POST",
    )
    try:
        with _ur.urlopen(req, timeout=20) as resp:
            print(f"[ILUS][RESEND] Email enviado a {recipients} (HTTP {resp.status}, source={cfg.get('_source')})")
            g._last_resend_error = None
            return True
    except _ue.HTTPError as e:
        body = ""
        try:
            body = e.read().decode("utf-8", errors="replace")
        except Exception:
            pass
        try:
            err_data = json.loads(body) if body else {}
        except Exception:
            err_data = {}
        g._last_resend_error = {
            "name":        err_data.get("name", ""),
            "message":     err_data.get("message", str(e)),
            "http_code":   e.code,
            "status_code": e.code,
            "raw_body":    body[:600],
        }
        print(f"[ILUS][RESEND] HTTP {e.code} — {body[:200]}")
        return False
    except Exception as e:
        g._last_resend_error = {
            "name":      "exception",
            "message":   str(e),
            "http_code": 0, "status_code": 0,
            "raw_body":  "",
        }
        print(f"[ILUS][RESEND] Excepción: {e}")
        return False


def _email_log(destinatario, asunto, evento, estado, error_msg=None, metadata=None):
    """Registra cada intento de envío de email para trazabilidad."""
    try:
        actor = ""
        try: actor = current_username() if g.user else "sistema"
        except Exception: actor = "sistema"
        mysql_execute(
            "INSERT INTO email_log (destinatario,asunto,evento,canal,estado,error_msg,actor,metadata) "
            "VALUES (%s,%s,%s,'email',%s,%s,%s,%s)",
            (str(destinatario)[:300], str(asunto)[:500], evento or "manual",
             estado, (error_msg or "")[:1000], actor[:190],
             json.dumps(metadata or {}, ensure_ascii=False)[:1500])
        )
    except Exception as exc:
        print(f"[EMAIL LOG] {exc}")


def _send_ilus_email(to_addr: str, subject: str, html_body: str, *, evento: str = None, **kwargs) -> bool:
    """
    Envía un correo HTML usando la configuración SMTP dinámica.
    Prioridad: Resend API (Railway) → SMTP con env vars → SMTP BD → SMTP config.py
    Loguea automáticamente el resultado en email_log.
    """
    # _send_ilus_email_inner hace el envío real; lo wrappeamos
    sent = False
    err  = None
    try:
        sent = _send_ilus_email_real(to_addr, subject, html_body, **kwargs)
    except Exception as exc:
        err = str(exc)[:1000]
    # Log
    try:
        _email_log(to_addr, subject, evento, 'enviado' if sent else 'fallido',
                   error_msg=err)
    except Exception: pass
    return sent


def _send_ilus_email_real(to_addr: str, subject: str, html_body: str) -> bool:
    """
    Implementación real con fallback inteligente:
      1) Resend API (HTTPS) — funciona desde cloud hosting (Railway, Heroku)
      2) SMTP — fallback si Resend no está configurado o falla

    Esto resuelve el bug "el email funciona en local pero no en Railway":
    Gmail bloquea/limita conexiones SMTP desde IPs cloud, pero acepta emails
    enviados vía Resend porque pasan por sus propios MTAs ya whitelisted.
    """
    # ── 1. Intento Resend primero (si está configurado) ─────────────────
    resend_cfg = _get_resend_cfg()
    if resend_cfg.get("api_key"):
        # Si el remitente Resend está configurado, lo usamos. Si no, dejamos el default
        # de Resend (onboarding@resend.dev) ya manejado por _send_via_resend.
        from_for_resend = None
        try:
            smtp_cfg = _get_smtp_cfg()
            from_name = smtp_cfg.get("from_name", "ILUS Sport & Health")
            from_addr = smtp_cfg.get("from_addr") or resend_cfg.get("from_addr")
            if from_addr:
                from_for_resend = f"{from_name} <{from_addr}>"
        except Exception:
            pass

        if _send_via_resend(to_addr, subject, html_body, from_addr=from_for_resend):
            return True
        # Resend falló — guardar error legible y caer a SMTP
        err = getattr(g, "_last_resend_error", None) or {}
        try:
            g._last_email_error = f"Resend: {err.get('message','fallo')} (HTTP {err.get('http_code',0)})"
        except Exception:
            pass
        print(f"[ILUS][EMAIL] Resend falló, intento SMTP como fallback…")

    # ── 2. Fallback a SMTP ──────────────────────────────────────────────
    try:
        cfg = _get_smtp_cfg()
    except Exception:
        cfg = dict(EMAIL_CONFIG)

    from_name = cfg.get("from_name", "ILUS Sport & Health")
    from_addr_cfg = cfg.get("from_addr") or cfg.get("smtp_user", "")

    # ── Envío vía SMTP (único método; configurable desde el front) ──────
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"]    = f"{from_name} <{from_addr_cfg}>"
    msg["To"]      = to_addr
    msg.attach(MIMEText(html_body, "html", "utf-8"))

    host      = cfg["smtp_host"]
    port      = int(cfg.get("smtp_port", 587))
    secure    = bool(cfg.get("secure"))
    user      = cfg["smtp_user"]
    passwd    = cfg.get("smtp_pass", "")
    from_addr = cfg.get("from_addr") or user

    def _try_send(p, sec, timeout=25):
        if sec:
            with _open_smtp_client(host, p, True, timeout=timeout) as srv:
                srv.login(user, passwd)
                srv.sendmail(from_addr, [to_addr], msg.as_string())
        else:
            with _open_smtp_client(host, p, False, timeout=timeout) as srv:
                srv.ehlo()
                srv.starttls()
                srv.login(user, passwd)
                srv.sendmail(from_addr, [to_addr], msg.as_string())

    # Intentos: puerto configurado → puerto alternativo automático
    attempts = [(port, secure)]
    if port == 587 and not secure:
        attempts.append((465, True))   # fallback SSL si STARTTLS falla
    elif port == 465 and secure:
        attempts.append((587, False))  # fallback STARTTLS si SSL falla

    last_exc = None
    for p, sec in attempts:
        try:
            _try_send(p, sec)
            print(f"[ILUS][EMAIL] Enviado a {to_addr} via :{p} (source={cfg.get('_source','?')})")
            return True
        except Exception as exc:
            last_exc = exc
            print(f"[ILUS][EMAIL] Intento :{p} falló — {exc}")

    try:
        g._last_email_error = str(last_exc)
    except Exception:
        pass
    print(f"[ILUS][EMAIL] Todos los intentos fallaron para {to_addr}: {last_exc}")
    return False


def _password_strength_errors(password: str) -> list[str]:
    """Politica minima para claves creadas desde enlaces publicos."""
    errors = []
    if len(password or "") < 12:
        errors.append("La contraseña debe tener al menos 12 caracteres.")
    if not re.search(r"[a-z]", password or ""):
        errors.append("Incluye al menos una letra minuscula.")
    if not re.search(r"[A-Z]", password or ""):
        errors.append("Incluye al menos una letra mayuscula.")
    if not re.search(r"\d", password or ""):
        errors.append("Incluye al menos un numero.")
    if not re.search(r"[^A-Za-z0-9]", password or ""):
        errors.append("Incluye al menos un simbolo.")
    return errors


def _issue_password_token(user_id: int, minutes: int = 60) -> tuple[str, datetime]:
    """Crea un token de un solo uso e invalida enlaces anteriores del usuario."""
    token = secrets.token_urlsafe(64)
    expires = datetime.utcnow() + timedelta(minutes=minutes)
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            f"UPDATE `{RESETS_TABLE}` SET used=1 WHERE user_id=%s AND used=0",
            (user_id,),
        )
        cur.execute(
            f"INSERT INTO `{RESETS_TABLE}` (user_id, token, expires_at) VALUES (%s,%s,%s)",
            (user_id, token, expires),
        )
    conn.commit()
    return token, expires


def _send_password_access_email(
    to_addr: str,
    to_name: str,
    action_url: str,
    *,
    actor_name: str = "ILUS",
    mode: str = "reset",
    minutes: int = 60,
) -> bool:
    """Envia correo ILUS para crear o cambiar clave mediante token seguro."""
    # Validación SMTP previa antes de enviar
    try:
        diag = _smtp_connection_diagnose(_get_smtp_cfg())
        if not diag.get("ok"):
            detail = diag.get("message") or "No se pudo validar la conexion SMTP."
            suggestions = diag.get("suggestions") or []
            if suggestions:
                detail += " " + " ".join(suggestions[:3])
            g._last_email_error = detail
            return False
    except Exception as exc:
        g._last_email_error = f"No se pudo validar SMTP antes de enviar: {exc}"
        return False

    is_setup = mode == "setup"
    titulo = "Crear contraseña de acceso" if is_setup else "Cambio de contraseña solicitado"
    subject = "ILUS - Crea tu contraseña de acceso" if is_setup else "ILUS - Cambio seguro de contraseña"
    button = "Crear mi contraseña" if is_setup else "Cambiar contraseña"
    intro = (
        f"<strong>{actor_name}</strong> creo una cuenta para ti en el sistema ILUS."
        if is_setup
        else f"<strong>{actor_name}</strong> solicito un cambio de contraseña para tu cuenta ILUS."
    )
    html_body = _ilus_email_html(
        titulo=titulo,
        subtitulo="Acceso seguro con enlace de un solo uso",
        saludo=f"Hola, {to_name}",
        parrafos=[
            intro,
            "Por seguridad, la contraseña no se envia ni se escribe manualmente. "
            "Debes definirla desde el boton de este correo.",
            f"Este enlace vence en <strong>{minutes} minutos</strong>, solo puede usarse una vez "
            "y reemplaza cualquier enlace anterior.",
            f'Si no esperabas este correo, ignoralo o avisa al administrador.<br>'
            f'<span style="font-size:11px;color:#777">Enlace directo: '
            f'<a href="{action_url}" style="color:#CC0000">{action_url}</a></span>',
        ],
        btn_primario_txt=button,
        btn_primario_url=action_url,
        info_lineas=[
            ("", "Cuenta", to_addr),
            ("", "Solicitado por", actor_name),
        ],
    )
    sent = _send_ilus_email(to_addr, subject, html_body)
    if not sent and not getattr(g, "_last_email_error", ""):
        g._last_email_error = "El servidor SMTP rechazo el envio. Revisa la configuracion en Comunicaciones."
    return sent


def _portal_login_url():
    return url_for("login", _external=True)


def _send_access_notification_email(to_addr: str, to_name: str, login_url: str, *, actor_name: str = "ILUS") -> bool:
    """Notifica que el acceso al portal fue habilitado, sin compartir la clave."""
    html_body = _ilus_email_html(
        titulo="Acceso al portal ILUS habilitado",
        subtitulo="Ya puedes ingresar a la aplicacion",
        saludo=f"Hola, {to_name}",
        parrafos=[
            f"<strong>{actor_name}</strong> habilito tus credenciales de acceso al portal ILUS.",
            "Por seguridad, este correo no incluye tu contraseña. Usa la clave entregada por el administrador "
            "o solicita un cambio de contraseña desde el portal cuando el correo este disponible.",
            "Presiona el boton para ir a la aplicacion.",
        ],
        btn_primario_txt="Ingresar al portal",
        btn_primario_url=login_url,
        info_lineas=[
            ("", "Cuenta", to_addr),
            ("", "Portal", login_url),
        ],
    )
    sent = _send_ilus_email(to_addr, "ILUS - Acceso al portal habilitado", html_body)
    if not sent and not getattr(g, "_last_email_error", ""):
        g._last_email_error = "No se pudo enviar la notificacion de acceso por email."
    return sent


def _notify_user_access(username: str, nombre: str, phone: str = "", *,
                        mode: str = "manual", action_url: str = "",
                        email_purpose: str = "invite") -> dict:
    """
    Envia notificaciones email/WhatsApp sobre acceso o creacion/cambio de clave.

    email_purpose:
      - 'invite' (default): cuenta nueva → email "Crea tu contraseña" (mode=setup, 24h)
      - 'change': cambio solicitado por admin → email "Cambio seguro de contraseña" (mode=reset, 60min)
                  Gmail los trata como mensajes distintos (subject distinto), evita threading
                  con la invitación previa y el usuario distingue claramente cada acción.
    """
    actor = g.user["nombre"] if getattr(g, "user", None) else "ILUS"
    login_url = _portal_login_url()
    result = {"email": None, "whatsapp": None, "errors": []}

    try:
        if mode == "token" and action_url:
            if email_purpose == "change":
                # Email de cambio de contraseña — subject y CTA distintos a los de invitación
                result["email"] = _send_password_access_email(
                    username, nombre, action_url,
                    actor_name=actor, mode="reset", minutes=60
                )
            else:
                result["email"] = _send_invitation_email(username, nombre, action_url, actor)
        else:
            result["email"] = _send_access_notification_email(username, nombre, login_url, actor_name=actor)
        if not result["email"]:
            result["errors"].append(getattr(g, "_last_email_error", "") or "No se pudo enviar email.")
    except Exception as exc:
        result["email"] = False
        result["errors"].append(f"Email: {exc}")

    if phone:
        try:
            wa_cfg = _get_wa_cfg()
            if wa_cfg.get("account_sid") and wa_cfg.get("auth_token") and wa_cfg.get("from_number"):
                if mode == "token" and action_url and email_purpose == "change":
                    body = (
                        f"Hola {nombre}, se solicito un *cambio de contrasena* para tu cuenta ILUS.\n\n"
                        f"Define tu nueva contrasena aqui: {action_url}\n"
                        f"El enlace es de un solo uso y vence en 60 minutos.\n\n"
                        f"Si no esperabas esta solicitud, ignora este mensaje."
                    )
                elif mode == "token" and action_url:
                    body = (
                        f"Hola {nombre}, ILUS habilito tu acceso al portal.\n\n"
                        f"Crea tu contrasena aqui: {action_url}\n"
                        f"El enlace es de un solo uso y vence pronto.\n\n"
                        f"Portal: {login_url}"
                    )
                else:
                    body = (
                        f"Hola {nombre}, ILUS habilito tus credenciales de acceso al portal.\n\n"
                        f"Ingresa aqui: {login_url}\n"
                        f"Por seguridad, este mensaje no incluye tu contrasena."
                    )
                sid = _send_whatsapp(wa_cfg["account_sid"], wa_cfg["auth_token"], wa_cfg["from_number"], phone, body)
                result["whatsapp"] = sid
            else:
                result["whatsapp"] = False
                result["errors"].append("WhatsApp no esta configurado.")
        except Exception as exc:
            result["whatsapp"] = False
            result["errors"].append(f"WhatsApp: {exc}")
    return result


def _access_notification_flash(result: dict, *, token_mode: bool = False) -> tuple[str, str]:
    """Construye un mensaje corto para el admin segun los canales que respondieron."""
    channels = []
    if result.get("email"):
        channels.append("email")
    if result.get("whatsapp"):
        channels.append("WhatsApp")
    if channels:
        action = "enlace seguro enviado" if token_mode else "notificacion de acceso enviada"
        return f"{action.capitalize()} por {', '.join(channels)}.", "success"
    errors = [str(e) for e in result.get("errors", []) if e]
    detail = " ".join(errors[:2]) or "Revisa SMTP o WhatsApp en Comunicaciones."
    return f"Usuario guardado, pero no salio la notificacion. {detail}", "warning"


# ─────────────────────────────────────────────
#  Recuperación de contraseña
# ─────────────────────────────────────────────

def _send_recovery_email(to_addr: str, to_name: str, reset_url: str) -> bool:
    """Envía el correo HTML de recuperación con diseño ILUS unificado."""
    html_body = _ilus_email_html(
        titulo          = "Recuperar contraseña",
        subtitulo       = "Sistema de Gestión ILUS Sport &amp; Health",
        saludo          = f"Hola, {to_name}",
        parrafos        = [
            "Recibimos una solicitud para restablecer la contraseña de tu cuenta en el sistema ILUS.",
            "Haz clic en el botón a continuación para crear una nueva contraseña. "
            "Este enlace es válido por <strong>60 minutos</strong>.",
            f'Si no solicitaste este cambio, puedes ignorar este correo — '
            f'tu contraseña seguirá siendo la misma.<br>'
            f'<span style="font-size:11px;color:#bbb">O copia: '
            f'<a href="{reset_url}" style="color:#CC0000">{reset_url}</a></span>',
        ],
        btn_primario_txt = "Restablecer contraseña",
        btn_primario_url = reset_url,
    )
    return _send_ilus_email(to_addr, "Recuperar contraseña — ILUS Sport & Health", html_body)


def _send_recovery_email(to_addr: str, to_name: str, reset_url: str) -> bool:
    """Version vigente: cambio de clave con token seguro."""
    return _send_password_access_email(
        to_addr, to_name, reset_url, actor_name="ILUS", mode="reset", minutes=60
    )


def _send_invitation_email(to_addr: str, to_name: str, set_url: str, creator_name: str = "ILUS") -> bool:
    """Version vigente: alta de usuario con token seguro."""
    return _send_password_access_email(
        to_addr, to_name, set_url, actor_name=creator_name, mode="setup", minutes=1440
    )


@app.route("/auth/olvidar-contrasena", methods=["GET", "POST"])
def forgot_password():
    if g.user:
        return redirect(url_for("index"))

    if request.method == "POST":
        email_input = request.form.get("email", "").strip().lower()

        # Siempre mostrar el mismo mensaje (no revelar si el email existe)
        flash("Si el correo está registrado, recibirás instrucciones en breve.", "info")

        try:
            if EMAIL_RE.match(email_input):
                user = get_auth_user_by_username(email_input)
                if user and user.get("active"):
                    token, _expires = _issue_password_token(user["id"], minutes=60)

                    reset_url = url_for("reset_password", token=token, _external=True)
                    _send_recovery_email(user["username"], user["nombre"], reset_url)
        except Exception as _e:
            print(f"[ILUS][RESET] Error en flujo de recuperación: {_e}")

        return redirect(url_for("forgot_password"))

    return render_template("forgot_password.html")


@app.route("/auth/restablecer/<token>", methods=["GET", "POST"])
def reset_password(token):
    if g.user:
        return redirect(url_for("index"))

    conn = get_db()
    row  = None
    with conn.cursor() as cur:
        cur.execute(
            f"""SELECT r.*, u.nombre, u.username
                FROM `{RESETS_TABLE}` r
                JOIN `{AUTH_TABLE}` u ON u.id = r.user_id
                WHERE r.token=%s AND r.used=0 AND r.expires_at > UTC_TIMESTAMP()""",
            (token,)
        )
        row = cur.fetchone()

    if not row:
        flash("El enlace no es válido o ha expirado.", "danger")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        pw1 = request.form.get("password", "")
        pw2 = request.form.get("password2", "")
        if pw1 != pw2:
            flash("Las contraseñas no coinciden.", "danger")
            return render_template("reset_password.html", token=token, nombre=row["nombre"])

        strength_errors = _password_strength_errors(pw1)
        if strength_errors:
            flash(" ".join(strength_errors), "danger")
            return render_template("reset_password.html", token=token, nombre=row["nombre"])

        new_hash = generate_password_hash(pw1)
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE `{AUTH_TABLE}` SET password_hash=%s WHERE id=%s",
                (new_hash, row["user_id"])
            )
            cur.execute(
                f"UPDATE `{RESETS_TABLE}` SET used=1 WHERE token=%s",
                (token,)
            )
        conn.commit()

        flash("Contraseña actualizada. Ahora puedes iniciar sesión.", "success")
        return redirect(url_for("login"))

    return render_template("reset_password.html", token=token, nombre=row["nombre"])


# ── PERFIL DE USUARIO (autogestión) ─────────────────────────────────────
# Cualquier usuario logueado puede editar SU PROPIO perfil. No es admin.
# Seguridad: el id se toma de g.user (sesión), nunca del request.

def _normalize_phone_chile(raw: str) -> str:
    """Normaliza teléfono Chile: deja solo dígitos y +; máximo 16 chars."""
    if not raw:
        return ""
    cleaned = re.sub(r"[^\d+]", "", raw.strip())
    return cleaned[:16]


@app.route("/mi-cuenta")
@login_required
def mi_cuenta():
    """Página de perfil — accesible para cualquier usuario logueado."""
    u = mysql_fetchone(
        f"SELECT id,username,nombre,phone,role,active,created_at,"
        f"foto_url,rut,cargo,genero,direccion,comuna,ciudad,fecha_nac "
        f"FROM `{AUTH_TABLE}` WHERE id=%s",
        (g.user["id"],)
    )
    if not u:
        return redirect(url_for("logout"))
    return render_template("mi_cuenta.html", usuario=dict(u))


@app.route("/mi-cuenta/datos", methods=["POST"])
@login_required
def mi_cuenta_datos():
    """
    Actualiza datos del perfil del usuario actual.
    Campos editables: nombre, phone, rut, cargo, genero, direccion, comuna, ciudad, fecha_nac.
    NO editables: username, role (eso es identidad/seguridad).
    """
    d = request.get_json(silent=True) or {}
    nombre = (d.get("nombre") or "").strip()[:190]
    phone  = _normalize_phone_chile(d.get("phone") or "")
    if len(nombre) < 2:
        return jsonify({"error": "El nombre debe tener al menos 2 caracteres"}), 400
    if phone and not re.match(r"^\+?\d{8,16}$", phone):
        return jsonify({"error": "Teléfono inválido. Usa formato +56912345678"}), 400

    rut       = (d.get("rut") or "").strip()[:20] or None
    cargo     = (d.get("cargo") or "").strip()[:120] or None
    genero    = (d.get("genero") or "").strip()[:20] or None
    direccion = (d.get("direccion") or "").strip()[:300] or None
    comuna    = (d.get("comuna") or "").strip()[:100] or None
    ciudad    = (d.get("ciudad") or "").strip()[:100] or None
    fecha_nac = (d.get("fecha_nac") or "").strip() or None
    if genero and genero not in ("masculino","femenino","otro","prefiero_no_decir"):
        genero = None

    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE `{AUTH_TABLE}` SET nombre=%s, phone=%s, rut=%s, cargo=%s, "
                f"genero=%s, direccion=%s, comuna=%s, ciudad=%s, fecha_nac=%s WHERE id=%s",
                (nombre, phone or None, rut, cargo, genero, direccion, comuna, ciudad,
                 fecha_nac, g.user["id"])
            )
        conn.commit()
        return jsonify({"ok": True, "nombre": nombre})
    except Exception as e:
        return jsonify({"error": f"No se pudo guardar: {e}"}), 500


@app.route("/mi-cuenta/foto", methods=["POST"])
@login_required
def mi_cuenta_foto():
    """
    Sube la foto de perfil del usuario a Cloudinary y guarda la URL en BD.
    Acepta multipart/form-data con campo 'foto' (jpg/png/webp, máx 5MB).
    """
    f = request.files.get("foto")
    if not f or not f.filename:
        return jsonify({"error": "No se recibió archivo"}), 400
    ext = (f.filename.rsplit(".", 1)[-1] or "").lower()
    if ext not in {"jpg","jpeg","png","webp","gif"}:
        return jsonify({"error": "Formato no permitido. Usa JPG, PNG o WEBP"}), 400
    # Tamaño máximo 5MB
    f.stream.seek(0, 2)
    size = f.stream.tell()
    f.stream.seek(0)
    if size > 5 * 1024 * 1024:
        return jsonify({"error": "Archivo demasiado grande (máx 5MB)"}), 400

    try:
        import cloudinary, cloudinary.uploader
        # Subir a Cloudinary en folder específico de avatares
        result = cloudinary.uploader.upload(
            f,
            folder=f"ilus/avatars",
            public_id=f"user_{g.user['id']}",
            overwrite=True,
            resource_type="image",
            transformation=[
                {"width": 400, "height": 400, "crop": "fill", "gravity": "face"},
                {"quality": "auto", "fetch_format": "auto"},
            ],
        )
        url = result.get("secure_url")
        if not url:
            return jsonify({"error": "Error al subir a Cloudinary"}), 500

        conn = get_db()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    f"UPDATE `{AUTH_TABLE}` SET foto_url=%s WHERE id=%s",
                    (url, g.user["id"])
                )
            conn.commit()
        finally:
            conn.close()

        return jsonify({"ok": True, "url": url})
    except Exception as e:
        return jsonify({"error": f"Error al procesar imagen: {e}"}), 500


@app.route("/mi-cuenta/foto", methods=["DELETE"])
@login_required
def mi_cuenta_foto_eliminar():
    """Elimina la foto de perfil (deja foto_url=NULL, vuelve a iniciales)."""
    try:
        conn = get_db()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    f"UPDATE `{AUTH_TABLE}` SET foto_url=NULL WHERE id=%s",
                    (g.user["id"],)
                )
            conn.commit()
        finally:
            conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/mi-cuenta/password", methods=["POST"])
@login_required
def mi_cuenta_password():
    """
    Cambia la contraseña del usuario actual.
    Seguridad:
      - Requiere contraseña ACTUAL para validar identidad.
      - Aplica política de fortaleza (_password_strength_errors).
      - La nueva no puede ser igual a la actual.
      - No se loguea ningún valor de contraseña.
    """
    d = request.get_json(silent=True) or {}
    current = d.get("current") or ""
    new1    = d.get("new") or ""
    new2    = d.get("confirm") or ""

    if not current or not new1 or not new2:
        return jsonify({"error": "Completa los 3 campos"}), 400

    # Recargar hash desde BD (no confiar en sesión cacheada)
    row = mysql_fetchone(
        f"SELECT id,password_hash FROM `{AUTH_TABLE}` WHERE id=%s",
        (g.user["id"],)
    )
    if not row or not check_password_hash(row["password_hash"], current):
        return jsonify({"error": "La contraseña actual es incorrecta"}), 400

    if new1 != new2:
        return jsonify({"error": "La nueva contraseña y su confirmación no coinciden"}), 400

    if check_password_hash(row["password_hash"], new1):
        return jsonify({"error": "La nueva contraseña debe ser distinta a la actual"}), 400

    errs = _password_strength_errors(new1)
    if errs:
        return jsonify({"error": " ".join(errs)}), 400

    new_hash = generate_password_hash(new1)
    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE `{AUTH_TABLE}` SET password_hash=%s WHERE id=%s",
                (new_hash, g.user["id"])
            )
        conn.commit()
        return jsonify({"ok": True, "message": "Contraseña actualizada correctamente"})
    except Exception:
        return jsonify({"error": "No se pudo guardar"}), 500


# ─────────────────────────────────────────────
#  Productos — listado
# ─────────────────────────────────────────────

@app.route("/")
@login_required
def index():
    q     = request.args.get("q", "").strip()
    exact = request.args.get("exact") == "1"

    # ── Búsqueda exacta por escáner de código de barras ──────────
    # Si viene con exact=1 (del scanner físico), busca el SKU exacto.
    # Si encuentra exactamente un producto con ficha → va directo a él.
    if exact and q:
        sku_norm = q.upper()
        hit = mysql_fetchone(
            f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE sku=%s", (sku_norm,)
        )
        if hit:
            return redirect(url_for("product_detail", pid=hit["id"]))
        # No tiene ficha aún → intenta crear desde ERP y redirige al índice filtrado
        flash(f"SKU {sku_norm} encontrado en ERP pero sin ficha. Usa 'Preparar' para crearlo.", "info")
        return redirect(url_for("index", q=sku_norm))

    products = get_product_listing(q)

    # Cobertura fotográfica — se calcula en Python para evitar problemas
    # de tipos (Decimal vs int) al comparar desde MySQL
    _fotos = [int(p.get("total_fotos") or 0) for p in products]
    foto0  = sum(1 for f in _fotos if f == 0)
    foto1  = sum(1 for f in _fotos if f == 1)
    foto2  = sum(1 for f in _fotos if f >= 2)

    return render_template("index.html", products=products, q=q,
                           foto0=foto0, foto1=foto1, foto2=foto2)


@app.route("/products/refresh-cache", methods=["POST"])
@login_required
def refresh_listing_cache():
    """Limpia el caché del listado de productos para ver datos frescos."""
    _invalidate_listing_cache()
    return redirect(url_for("index"))


# ─────────────────────────────────────────────
#  Productos — quick view JSON (para modal)
# ─────────────────────────────────────────────

@app.route("/products/<int:pid>/quick")
@login_required
def product_quick(pid):
    product, bultos, photos = get_full(pid)
    if not product:
        return jsonify({"error": "not found"}), 404
    photo_urls = [_photo_src(ph["filename"]) for ph in photos]
    return jsonify({
        "id":          product["id"],
        "sku":         product["sku"],
        "nombre":      product["nombre"],
        "estado":      product["estado"],
        "codigo":      product["codigo"] or "",
        "created_by":  product["created_by"] or "-",
        "updated_by":  product["updated_by"] or "-",
        "peso_total":  round(sum(float(b["peso"])     for b in bultos), 2),
        "pv_total":    round(sum(float(b["peso_vol"]) for b in bultos), 2),
        "total_bultos": len(bultos),
        "bultos": [
            {
                "num":      b["bulto_num"],
                "largo":    float(b["largo"]),
                "ancho":    float(b["ancho"]),
                "alto":     float(b["alto"]),
                "peso":     float(b["peso"]),
                "peso_vol": float(b["peso_vol"]),
            }
            for b in bultos
        ],
        "photos": photo_urls,
    })


# ─────────────────────────────────────────────
#  Productos — CRUD
# ─────────────────────────────────────────────

@app.route("/products/from-erp/<path:sku>")
@require_permission("create")
def prepare_product_from_erp(sku):
    pid = ensure_product_record_from_erp(sku)
    if not pid:
        flash("No se encontro ese producto en el ERP.", "danger")
        return redirect(url_for("index"))
    flash("Producto preparado desde el ERP.", "success")
    return redirect(url_for("edit_product", pid=pid))


@app.route("/products/new", methods=["GET", "POST"])
@require_permission("create")
def new_product():
    # Pre-calcular el código que se asignará (para mostrarlo en el form)
    auto_codigo = next_codigo()

    if request.method == "POST":
        sku    = request.form.get("sku",    "").strip().upper()
        nombre = request.form.get("nombre", "").strip()
        estado = request.form.get("estado", "Confirmado").strip()
        # El código NO viene del formulario: se genera server-side
        codigo = next_codigo()

        errors = []
        if not sku:
            errors.append("El SKU es requerido.")
        if not nombre:
            errors.append("El nombre es requerido.")
        if mysql_fetchone(f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE sku=%s", (sku,)):
            errors.append(f"El SKU <b>{sku}</b> ya existe.")
        errors += validate_bultos_form(request.form)

        if errors:
            return render_template("product_form.html", errors=errors, product=None,
                                   fd=request.form, max_b=MAX_BULTOS, auto_codigo=codigo)

        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(
                f"""INSERT INTO `{PRODUCTS_TABLE}` (sku,nombre,estado,codigo,created_by,updated_by)
                    VALUES (%s,%s,%s,%s,%s,%s)""",
                (sku, nombre, estado, codigo, current_username(), current_username()),
            )
            pid = cur.lastrowid
            save_bultos_mysql(conn, pid, request.form)
        conn.commit()

        # Sincronizar tabla compartida etiquetas
        _, bultos_sync, _ = get_full(pid)
        p_sync = {"sku": sku, "nombre": nombre, "estado": estado, "codigo": codigo}
        sync_erp_table(p_sync, bultos_sync, conn)
        conn.commit()

        _invalidate_listing_cache()
        flash(f"Producto <b>{sku}</b> creado con código <b>{codigo}</b>.", "success")
        return redirect(url_for("product_detail", pid=pid))

    return render_template("product_form.html", errors=[], product=None, fd={},
                           max_b=MAX_BULTOS, auto_codigo=auto_codigo)


@app.route("/products/<int:pid>/edit", methods=["GET", "POST"])
@require_permission("edit")
def edit_product(pid):
    product, bultos, photos = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))

    if request.method == "POST":
        sku    = request.form.get("sku",    "").strip().upper()
        nombre = request.form.get("nombre", "").strip()
        estado = request.form.get("estado", "Pendiente").strip()
        # El código NO se puede cambiar: siempre se conserva el valor actual
        codigo = product["codigo"]

        errors = []
        if not sku:
            errors.append("El SKU es requerido.")
        if not nombre:
            errors.append("El nombre es requerido.")
        if mysql_fetchone(f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE sku=%s AND id<>%s", (sku, pid)):
            errors.append(f"El SKU <b>{sku}</b> ya está en uso.")
        errors += validate_bultos_form(request.form)

        if errors:
            return render_template("product_form.html", errors=errors, product=product,
                                   bultos=bultos, fd=request.form, max_b=MAX_BULTOS, photos=photos)

        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(
                f"""UPDATE `{PRODUCTS_TABLE}`
                    SET sku=%s, nombre=%s, estado=%s, updated_by=%s
                    WHERE id=%s""",
                (sku, nombre, estado, current_username(), pid),
            )
            cur.execute(f"DELETE FROM `{BULTOS_TABLE}` WHERE product_id=%s", (pid,))
            save_bultos_mysql(conn, pid, request.form)
        conn.commit()

        # Sincronizar tabla compartida etiquetas
        _, bultos_sync, _ = get_full(pid)
        p_sync = {"sku": sku, "nombre": nombre, "estado": estado, "codigo": codigo}
        sync_erp_table(p_sync, bultos_sync, conn)
        conn.commit()

        _invalidate_listing_cache()
        flash(f"Producto <b>{sku}</b> actualizado.", "success")
        return redirect(url_for("product_detail", pid=pid))

    return render_template("product_form.html", errors=[], product=product,
                           bultos=bultos, photos=photos, fd={}, max_b=MAX_BULTOS)


@app.route("/products/<int:pid>")
@login_required
def product_detail(pid):
    product, bultos, photos = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))
    peso_total = round(sum(float(b["peso"])     for b in bultos), 2)
    pv_total   = round(sum(float(b["peso_vol"]) for b in bultos), 2)
    return render_template("product_detail.html", product=product, bultos=bultos,
                           photos=photos, peso_total=peso_total, pv_total=pv_total,
                           total_bultos=len(bultos), can_add_photo=(len(photos) < MAX_PHOTOS))


@app.route("/products/<int:pid>/delete", methods=["POST"])
@require_permission("delete")
def delete_product(pid):
    product, _, photos = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))

    # Verificar confirmación escribiendo el SKU
    confirm = request.form.get("confirm_sku", "").strip().upper()
    if confirm != product["sku"]:
        flash("Confirmación incorrecta. El producto NO fue eliminado.", "danger")
        return redirect(url_for("product_detail", pid=pid))

    for photo in photos:
        delete_photo_file(photo["filename"])

    conn = get_db()
    # Eliminar de la tabla compartida etiquetas ANTES de borrar en app_products
    delete_from_erp_table(product.get("codigo"), conn)
    with conn.cursor() as cur:
        cur.execute(f"DELETE FROM `{PRODUCTS_TABLE}` WHERE id=%s", (pid,))
    conn.commit()

    _invalidate_listing_cache()
    flash(f"Producto <b>{product['sku']}</b> eliminado permanentemente.", "warning")
    return redirect(url_for("index"))


# ─────────────────────────────────────────────
#  Fotos
# ─────────────────────────────────────────────

@app.route("/products/<int:pid>/photos", methods=["POST"])
@require_permission("edit")
def upload_photo(pid):
    product, _, photos = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))
    if len(photos) >= MAX_PHOTOS:
        flash(f"Maximo {MAX_PHOTOS} fotos por producto.", "warning")
        return redirect(url_for("product_detail", pid=pid))

    file = request.files.get("photo")
    if not file or not file.filename:
        flash("No se selecciono archivo.", "warning")
        return redirect(url_for("product_detail", pid=pid))
    if not allowed_file(file.filename):
        flash("Formato no permitido. Usa JPG, PNG, WEBP o GIF.", "danger")
        return redirect(url_for("product_detail", pid=pid))

    ext       = file.filename.rsplit(".", 1)[1].lower()
    ts        = int(datetime.now().timestamp())
    if _CLD_READY:
        try:
            filename = _cloud_upload(file, public_id=f"p{pid}_{ts}", folder="ilus/products")
            print(f"[ILUS] Foto subida a Cloudinary: {filename}")
        except Exception as exc:
            print(f"[ILUS] Cloudinary upload error: {exc}")
            flash(f"Error al subir la foto a la nube: {exc}", "danger")
            return redirect(url_for("product_detail", pid=pid))
    else:
        filename = f"p{pid}_{ts}.{ext}"
        try:
            file.save(os.path.join(UPLOAD_FOLDER, filename))
            print(f"[ILUS] Foto guardada localmente: {filename}")
        except Exception as exc:
            print(f"[ILUS] Error guardando foto local: {exc}")
            flash(f"Error al guardar la foto: {exc}", "danger")
            return redirect(url_for("product_detail", pid=pid))

    # Guardamos URL completa (Cloudinary) o nombre local
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            f"INSERT INTO `{PHOTOS_TABLE}` (product_id,filename,orden) VALUES (%s,%s,%s)",
            (pid, filename, len(photos) + 1),
        )
    conn.commit()

    flash("Foto agregada correctamente.", "success")
    return redirect(url_for("product_detail", pid=pid))


@app.route("/products/<int:pid>/photos/<int:photo_id>/delete", methods=["POST"])
@require_permission("delete")
def delete_photo(pid, photo_id):
    product, _, _ = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))

    # Verificar confirmación por SKU
    confirm = request.form.get("confirm_sku", "").strip().upper()
    if confirm != product["sku"]:
        flash("Confirmación incorrecta. La foto NO fue eliminada.", "danger")
        return redirect(url_for("product_detail", pid=pid))

    photo = mysql_fetchone(
        f"SELECT * FROM `{PHOTOS_TABLE}` WHERE id=%s AND product_id=%s", (photo_id, pid)
    )
    if photo:
        delete_photo_file(photo["filename"])
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(f"DELETE FROM `{PHOTOS_TABLE}` WHERE id=%s", (photo_id,))
        conn.commit()
        flash("Foto eliminada.", "success")
    else:
        flash("Foto no encontrada.", "danger")
    return redirect(url_for("product_detail", pid=pid))


# ─────────────────────────────────────────────
#  Impresión — HTML
# ─────────────────────────────────────────────

def _render_labels_view(pid, template_name):
    only = request.args.get("bulto", type=int)
    fmt = request.args.get("fmt", "150x100")
    label_format = _label_format(fmt)
    product, bultos, _ = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))

    enriched_bultos = [_label_bulto_data(b) for b in bultos]
    valid = [b for b in enriched_bultos if b["largo"] > 0 and b["ancho"] > 0 and b["alto"] > 0]
    if not valid:
        flash("Ningun bulto tiene medidas completas.", "danger")
        return redirect(url_for("product_detail", pid=pid))
    if only:
        valid = [b for b in valid if int(b["bulto_num"]) == only]

    # Cantidades por bulto (vienen del formulario de print_setup)
    qty_per_bulto = {}
    for b in valid:
        qty = request.args.get(f"qty_{b['bulto_num']}", 1, type=int)
        qty_per_bulto[int(b["bulto_num"])] = max(1, min(qty, 10))

    fecha = datetime.now().strftime("%d-%m-%Y %H:%M")
    response = make_response(render_template(template_name, product=product, bultos=valid,
                                             total_bultos=len(bultos), fecha=fecha,
                                             qty_per_bulto=qty_per_bulto,
                                             fmt=label_format["key"],
                                             label_format=label_format,
                                             logo_url=_logo_data_url()))
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    return response


@app.route("/products/<int:pid>/labels")
@require_permission("print")
def print_labels(pid):
    return _render_labels_view(pid, "print_labels.html")


@app.route("/print/labels")
@require_permission("print")
def print_labels_alt():
    pid = request.args.get("pid", type=int)
    if not pid:
        flash("Selecciona un producto para imprimir etiquetas.", "warning")
        return redirect(url_for("index"))
    return _render_labels_view(pid, "print_labels.html")


@app.route("/products/<int:pid>/print-setup")
@require_permission("print")
def print_setup(pid):
    """Pantalla de configuración de impresión: copias por bulto."""
    product, bultos, _ = get_full(pid)
    if not product:
        flash("Producto no encontrado.", "danger")
        return redirect(url_for("index"))

    valid = [b for b in bultos if float(b["largo"]) > 0 and float(b["ancho"]) > 0 and float(b["alto"]) > 0]
    if not valid:
        flash("Ningun bulto tiene medidas completas.", "danger")
        return redirect(url_for("product_detail", pid=pid))

    # Si viene un bulto específico, filtrar
    only = request.args.get("bulto", type=int)
    if only:
        valid = [b for b in valid if int(b["bulto_num"]) == only]

    peso_total = round(sum(float(b["peso"]) for b in valid), 2)
    return render_template("print_setup.html", product=product, bultos=valid,
                           total_bultos=len(bultos), peso_total=peso_total,
                           only_bulto=only)


@app.route("/products/<int:pid>/download-all-pdf")
@require_permission("print")
def download_all_pdf(pid):
    """Descarga las etiquetas seleccionadas como un PDF multipagina."""
    return _labels_pdf_response(pid, force_download=True)


@app.route("/products/<int:pid>/labels-preview.pdf")
@require_permission("print")
def labels_pdf_preview(pid):
    """PDF inline para previsualizacion en modal."""
    return _labels_pdf_response(pid, force_download=request.args.get("download") == "1")


def _labels_pdf_response(pid, force_download=False):
    product, bultos, _ = get_full(pid)
    if not product:
        return "Producto no encontrado", 404

    valid = [_label_bulto_data(b) for b in bultos]
    valid = [b for b in valid if b["largo"] > 0 and b["ancho"] > 0 and b["alto"] > 0]
    only = request.args.get("bulto", type=int)
    if only:
        valid = [b for b in valid if int(b["bulto_num"]) == only]

    if not valid:
        return "Sin bultos con medidas completas", 404

    selected = []
    for b in valid:
        qty = request.args.get(f"qty_{b['bulto_num']}", 1, type=int)
        qty = max(1, min(qty, 10))
        for _ in range(qty):
            selected.append(dict(b))

    try:
        pdf_bytes = build_labels_pdf(product, selected, len(bultos), request.args.get("fmt", "150x100"))
    except Exception as exc:
        return f"Error generando PDF: {exc}", 500

    fecha = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"ILUS_{product['sku']}_{fecha}.pdf"
    disposition = "attachment" if force_download else "inline"
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        },
    )


# ─────────────────────────────────────────────
#  Impresión — PDF (150 × 50 mm)
# ─────────────────────────────────────────────

@app.route("/products/<int:pid>/bulto/<int:bnum>/pdf")
@require_permission("print")
def download_pdf(pid, bnum):
    product, bultos, _ = get_full(pid)
    if not product:
        return "Producto no encontrado", 404
    bulto = next((b for b in bultos if int(b["bulto_num"]) == bnum), None)
    if not bulto or float(bulto["largo"]) == 0:
        return "Bulto sin medidas completas", 404

    try:
        pdf_bytes = build_label_pdf(product, bulto, len(bultos), request.args.get("fmt", "150x100"))
    except Exception as exc:
        return f"Error generando PDF: {exc}", 500

    filename = f"ILUS_{product['sku']}_B{bnum:02d}.pdf"
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────
#  Exportar Excel
# ─────────────────────────────────────────────

@app.route("/products/export/excel")
@login_required
def export_excel():
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        flash("Instala openpyxl: pip install openpyxl", "danger")
        return redirect(url_for("index"))

    # ── Datos ───────────────────────────────────────────────────────────
    products = mysql_fetchall(
        f"""SELECT p.id, p.sku, p.nombre, p.estado, p.codigo,
                   p.created_by, p.updated_by, p.created_at, p.updated_at
            FROM `{PRODUCTS_TABLE}` p ORDER BY p.sku"""
    )
    bultos = mysql_fetchall(
        f"""SELECT b.product_id, b.bulto_num, b.largo, b.ancho, b.alto, b.peso,
                   ROUND((b.largo*b.ancho*b.alto)/4000, 2) AS peso_vol
            FROM `{BULTOS_TABLE}` b ORDER BY b.product_id, b.bulto_num"""
    )
    # índice de bultos por product_id
    bultos_by_pid = {}
    for b in bultos:
        bultos_by_pid.setdefault(int(b["product_id"]), []).append(b)

    wb = openpyxl.Workbook()

    # ── Estilos ─────────────────────────────────────────────────────────
    RED     = "CC0000"
    BLACK   = "111111"
    WHITE   = "FFFFFF"
    LGRAY   = "F5F5F5"
    hdr_font  = Font(name="Calibri", bold=True, color=WHITE, size=11)
    hdr_fill  = PatternFill("solid", fgColor=BLACK)
    red_fill  = PatternFill("solid", fgColor=RED)
    alt_fill  = PatternFill("solid", fgColor=LGRAY)
    center    = Alignment(horizontal="center", vertical="center")
    left      = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    thin      = Side(style="thin", color="DDDDDD")
    border    = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_hdr(cell, red=False):
        cell.font      = hdr_font
        cell.fill      = red_fill if red else hdr_fill
        cell.alignment = center
        cell.border    = border

    def style_cell(cell, align=None, bold=False, color=None):
        cell.font      = Font(name="Calibri", bold=bold, color=color or "000000", size=10)
        cell.alignment = align or left
        cell.border    = border

    # ══════════════════════════════════════════
    #  Hoja 1 — Resumen productos
    # ══════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Productos"
    ws1.freeze_panes = "A2"

    hdrs1 = ["SKU", "Nombre", "Estado", "Código", "Bultos",
             "Peso Total (kg)", "Peso Vol Total", "Creado por", "Actualizado por"]
    col_w1 = [16, 45, 14, 10, 8, 16, 15, 22, 22]
    for i, (h, w) in enumerate(zip(hdrs1, col_w1), 1):
        c = ws1.cell(row=1, column=i, value=h)
        style_hdr(c, red=(i <= 1))
        ws1.column_dimensions[c.column_letter].width = w
    ws1.row_dimensions[1].height = 20

    for row_i, p in enumerate(products, 2):
        blist = bultos_by_pid.get(int(p["id"]), [])
        ptotal = sum(float(b["peso"])     for b in blist)
        pvtot  = sum(float(b["peso_vol"]) for b in blist)
        fill = alt_fill if row_i % 2 == 0 else PatternFill()
        vals = [p["sku"], p["nombre"], p["estado"], p["codigo"],
                len(blist), round(ptotal, 2), round(pvtot, 2),
                p["created_by"] or "", p["updated_by"] or ""]
        for col_i, v in enumerate(vals, 1):
            c = ws1.cell(row=row_i, column=col_i, value=v)
            c.fill   = fill
            c.border = border
            c.font   = Font(name="Calibri", size=10,
                            bold=(col_i == 1),
                            color=RED if col_i == 1 else "000000")
            c.alignment = center if col_i in (1, 3, 4, 5, 6, 7) else left
        ws1.row_dimensions[row_i].height = 16

    # ══════════════════════════════════════════
    #  Hoja 2 — Detalle bultos
    # ══════════════════════════════════════════
    ws2 = wb.create_sheet("Bultos")
    ws2.freeze_panes = "A2"

    hdrs2 = ["SKU", "Nombre", "Código", "Bulto N°",
             "Largo (cm)", "Ancho (cm)", "Alto (cm)", "Peso (kg)", "Peso Vol"]
    col_w2 = [16, 45, 10, 9, 11, 11, 11, 11, 11]
    for i, (h, w) in enumerate(zip(hdrs2, col_w2), 1):
        c = ws2.cell(row=1, column=i, value=h)
        style_hdr(c, red=(i <= 1))
        ws2.column_dimensions[c.column_letter].width = w
    ws2.row_dimensions[1].height = 20

    row_i = 2
    for p in products:
        blist = bultos_by_pid.get(int(p["id"]), [])
        for b in blist:
            fill = alt_fill if row_i % 2 == 0 else PatternFill()
            vals = [p["sku"], p["nombre"], p["codigo"], int(b["bulto_num"]),
                    float(b["largo"]), float(b["ancho"]), float(b["alto"]),
                    float(b["peso"]), float(b["peso_vol"])]
            for col_i, v in enumerate(vals, 1):
                c = ws2.cell(row=row_i, column=col_i, value=v)
                c.fill   = fill
                c.border = border
                c.font   = Font(name="Calibri", size=10,
                                bold=(col_i == 1),
                                color=RED if col_i == 1 else "000000")
                c.alignment = center if col_i != 2 else left
            ws2.row_dimensions[row_i].height = 16
            row_i += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fecha = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"ILUS_Etiquetas_{fecha}.xlsx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────
#  Usuarios (solo admin)
# ─────────────────────────────────────────────

@app.route("/admin/users")
@require_permission("admin")
def users_index():
    users = mysql_fetchall(f"SELECT * FROM `{AUTH_TABLE}` ORDER BY nombre, username")
    return render_template("users.html", users=users)


@app.route("/admin/users/new", methods=["GET", "POST"])
@require_permission("admin")
def new_user():
    if request.method == "POST":
        username    = request.form.get("username", "").strip().lower()
        nombre      = request.form.get("nombre",   "").strip()
        phone       = request.form.get("phone",    "").strip()
        role        = request.form.get("role",      "editor")
        active      = 1 if request.form.get("active") == "1" else 0
        send_invite = request.form.get("send_invite") == "1"
        wa_number   = request.form.get("wa_number", "").strip()
        manual_password = request.form.get("manual_password", "")
        is_superadmin = bool(g.permissions.get("superadmin")) if getattr(g, "permissions", None) else False

        errors = []
        if not username:
            errors.append("El correo es requerido.")
        elif not EMAIL_RE.match(username):
            errors.append("El correo no tiene un formato válido.")
        if not nombre:
            errors.append("El nombre y apellido son requeridos.")
        if phone and not re.match(r"^\+?[0-9\s\-]{7,20}$", phone):
            errors.append("El telefono WhatsApp debe usar formato internacional, por ejemplo +56912345678.")
        if not _es_rol_valido(role):
            errors.append("Rol no valido.")
        elif role == "superadmin" and not is_superadmin:
            errors.append("Solo un superadministrador puede asignar el rol superadmin.")
        if manual_password and not is_superadmin:
            errors.append("Solo un superadministrador puede definir clave manual.")
        if manual_password and len(manual_password) < 8:
            errors.append("La clave manual debe tener al menos 8 caracteres.")
        if get_auth_user_by_username(username):
            errors.append("Ese correo ya está registrado.")

        if errors:
            return render_template("user_form.html", errors=errors, user=None, fd=request.form,
                                   roles=_get_roles_disponibles())

        # Crear usuario con contraseña placeholder (bloqueada hasta que use el enlace)
        placeholder_hash = generate_password_hash(manual_password or secrets.token_hex(32))
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(
                f"INSERT INTO `{AUTH_TABLE}` (username,nombre,password_hash,phone,role,active) VALUES (%s,%s,%s,%s,%s,%s)",
                (username, nombre, placeholder_hash, phone or None, role, active),
            )
        conn.commit()

        # Generar token de invitación (usa la misma tabla de resets, válido 24h)
        if send_invite and not manual_password:
            try:
                new_uid   = mysql_fetchone(f"SELECT id FROM `{AUTH_TABLE}` WHERE username=%s", (username,))["id"]
                token, _expires = _issue_password_token(new_uid, minutes=1440)
                set_url   = url_for("reset_password", token=token, _external=True)
                result = _notify_user_access(username, nombre, wa_number or phone, mode="token", action_url=set_url)
                msg, level = _access_notification_flash(result, token_mode=True)
                flash(f"Usuario creado. {msg}", level)
            except Exception as _ie:
                flash(f"Usuario creado, pero fallo el envio de invitacion: {_ie}", "warning")
        else:
            if manual_password:
                result = _notify_user_access(username, nombre, wa_number or phone, mode="manual")
                msg, level = _access_notification_flash(result, token_mode=False)
                flash(f"Usuario creado con clave manual. {msg}", level)
            else:
                flash("Usuario creado. Recuerda establecer la contrasena o enviar la invitacion despues.", "success")

        if manual_password:
            flash("Clave manual asignada por superadmin.", "info")
        return redirect(url_for("users_index"))

    return render_template("user_form.html", errors=[], user=None, fd={},
                           roles=_get_roles_disponibles())


@app.route("/admin/users/<int:user_id>/edit", methods=["GET", "POST"])
@require_permission("admin")
def edit_user(user_id):
    user = get_auth_user_by_id(user_id)
    if not user:
        flash("Usuario no encontrado.", "danger")
        return redirect(url_for("users_index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        nombre   = request.form.get("nombre",   "").strip()
        phone    = request.form.get("phone",    "").strip()
        manual_password = request.form.get("manual_password", "")
        is_superadmin = bool(g.permissions.get("superadmin")) if getattr(g, "permissions", None) else False
        role     = request.form.get("role",     "editor")
        active   = 1 if request.form.get("active") == "1" else 0

        errors = []
        if not username:
            errors.append("El correo es requerido.")
        elif not EMAIL_RE.match(username):
            errors.append("El correo no tiene un formato válido.")
        if not nombre:
            errors.append("El nombre y apellido son requeridos.")
        if phone and not re.match(r"^\+?[0-9\s\-]{7,20}$", phone):
            errors.append("El telefono WhatsApp debe usar formato internacional, por ejemplo +56912345678.")
        if not _es_rol_valido(role):
            errors.append("Rol no valido.")
        elif role == "superadmin" and not is_superadmin:
            errors.append("Solo un superadministrador puede asignar el rol superadmin.")
        if manual_password and not is_superadmin:
            errors.append("Solo un superadministrador puede definir clave manual.")
        if manual_password and len(manual_password) < 8:
            errors.append("La clave manual debe tener al menos 8 caracteres.")
        if mysql_fetchone(
            f"SELECT id FROM `{AUTH_TABLE}` WHERE username=%s AND id<>%s", (username, user_id)
        ):
            errors.append("Ese correo ya está en uso.")

        if errors:
            return render_template("user_form.html", errors=errors, user=user, fd=request.form,
                                   roles=_get_roles_disponibles())

        conn = get_db()
        with conn.cursor() as cur:
            if manual_password:
                cur.execute(
                    f"UPDATE `{AUTH_TABLE}` SET username=%s,nombre=%s,password_hash=%s,phone=%s,role=%s,active=%s WHERE id=%s",
                    (username, nombre, generate_password_hash(manual_password), phone or None, role, active, user_id),
                )
            else:
                cur.execute(
                    f"UPDATE `{AUTH_TABLE}` SET username=%s,nombre=%s,phone=%s,role=%s,active=%s WHERE id=%s",
                    (username, nombre, phone or None, role, active, user_id),
                )
        conn.commit()

        # Invalida caché de session si se editó el usuario actual
        if g.user and g.user["id"] == user_id:
            session.pop("_uc", None)

        if manual_password:
            result = _notify_user_access(username, nombre, phone, mode="manual")
            msg, level = _access_notification_flash(result, token_mode=False)
            flash(f"Usuario actualizado con clave manual. {msg}", level)
        else:
            flash("Usuario actualizado.", "success")
        return redirect(url_for("users_index"))

    return render_template("user_form.html", errors=[], user=user, fd={},
                           roles=_get_roles_disponibles())


@app.route("/admin/users/<int:user_id>/delete", methods=["POST"])
@require_permission("admin")
def delete_user(user_id):
    if g.user and g.user["id"] == user_id:
        flash("No puedes eliminar tu propio usuario mientras estás conectado.", "danger")
        return redirect(url_for("users_index"))

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"DELETE FROM `{AUTH_TABLE}` WHERE id=%s", (user_id,))
    conn.commit()

    flash("Usuario eliminado.", "warning")
    return redirect(url_for("users_index"))


@app.route("/admin/users/<int:user_id>/invite", methods=["POST"])
@require_permission("admin")
def invite_user(user_id):
    """Envía (o reenvía) el correo de invitación para crear contraseña."""
    user = get_auth_user_by_id(user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    try:
        token, _expires = _issue_password_token(user_id, minutes=1440)
        set_url = url_for("reset_password", token=token, _external=True)
        result = _notify_user_access(user["username"], user["nombre"], user.get("phone") or "", mode="token", action_url=set_url)
        if result.get("email") or result.get("whatsapp"):
            msg, _level = _access_notification_flash(result, token_mode=True)
            return jsonify({"ok": True, "message": msg})
        detalle = " ".join(result.get("errors") or []) or "Revisa la configuracion SMTP o WhatsApp en Comunicaciones."
        return jsonify({"error": f"No se pudo enviar la invitacion. {detalle}"}), 500
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ══════════════════════════════════════════════════════════════
#  MÓDULO: ROLES
# ══════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════
#  ROLES DINÁMICOS — matriz módulo × acción editable
# ══════════════════════════════════════════════════════════════════════

PERMISSIONS_MATRIX = {
    "etiquetas":      {"label":"Etiquetas",      "icon":"bi-tags",
                       "acciones":["ver","crear","editar","eliminar","imprimir"]},
    "mantenciones":   {"label":"Mantenciones",   "icon":"bi-wrench-adjustable",
                       "acciones":["ver","crear","editar","eliminar","contratos","ai","reportes","repuestos"]},
    "retiros":        {"label":"Retiros",        "icon":"bi-box-arrow-up-right",
                       "acciones":["ver","gestionar","monitor","marketing"]},
    "transporte":     {"label":"Transporte",     "icon":"bi-truck",
                       "acciones":["ver","cubicador","asignar","manifiestos","couriers"]},
    "comunicaciones": {"label":"Comunicaciones", "icon":"bi-chat-dots",
                       "acciones":["ver","configurar","enviar","plantillas"]},
    "admin":          {"label":"Administración", "icon":"bi-gear-wide-connected",
                       "acciones":["ajustes","usuarios","roles","marketing","login_imagenes"]},
}


def _get_roles_disponibles():
    """
    Devuelve la lista de roles activos desde roles_dinamicos.
    Se usa para poblar dropdowns de creación/edición de usuario.
    Si la tabla aún no existe (primer arranque), devuelve los roles del sistema hardcoded.
    """
    try:
        rows = mysql_fetchall(
            "SELECT slug,nombre,color,is_system FROM roles_dinamicos WHERE activo=1 "
            "ORDER BY is_system DESC, nombre"
        )
        if rows:
            return [dict(r) for r in rows]
    except Exception:
        pass
    # Fallback si la tabla aún no fue inicializada
    return [
        {"slug":"superadmin", "nombre":"Super Administrador", "color":"#dc2626", "is_system":1},
        {"slug":"admin",      "nombre":"Administrador",      "color":"#2563eb", "is_system":1},
        {"slug":"editor",     "nombre":"Editor",             "color":"#ea580c", "is_system":1},
        {"slug":"lector",     "nombre":"Solo lectura",       "color":"#6b7280", "is_system":1},
    ]


def _es_rol_valido(slug):
    """True si el slug existe en roles_dinamicos activos. Solo superadmin puede asignar 'superadmin'."""
    if not slug:
        return False
    try:
        row = mysql_fetchone(
            "SELECT 1 FROM roles_dinamicos WHERE slug=%s AND activo=1 LIMIT 1", (slug,)
        )
        return bool(row)
    except Exception:
        return slug in {"superadmin","admin","editor","lector","vendedor","ejecutivo"}


def get_role_permissions(slug):
    """Devuelve dict {modulo:{accion:bool}} para un rol dado."""
    if slug == "superadmin":
        return {m: {a: True for a in cfg["acciones"]} for m, cfg in PERMISSIONS_MATRIX.items()}
    rows = mysql_fetchall(
        "SELECT modulo,accion,permitido FROM rol_permisos WHERE rol_slug=%s", (slug,)
    )
    perms = {m: {a: False for a in cfg["acciones"]} for m, cfg in PERMISSIONS_MATRIX.items()}
    for r in rows:
        if r["modulo"] in perms and r["accion"] in perms[r["modulo"]]:
            perms[r["modulo"]][r["accion"]] = bool(r["permitido"])
    return perms


def has_role_permission(slug, modulo, accion):
    """Verifica si un rol dinámico tiene un permiso específico."""
    if slug == "superadmin":
        return True
    row = mysql_fetchone(
        "SELECT permitido FROM rol_permisos WHERE rol_slug=%s AND modulo=%s AND accion=%s",
        (slug, modulo, accion)
    )
    return bool(row and row.get("permitido"))


@app.route("/admin/roles")
@require_permission("admin")
def admin_roles():
    """Vista UNIFICADA: usuarios + lista de roles dinámicos + matriz editable."""
    users = mysql_fetchall(
        f"SELECT id,username,nombre,role,active FROM `{AUTH_TABLE}` ORDER BY nombre", ()
    )
    roles = mysql_fetchall(
        "SELECT * FROM roles_dinamicos WHERE activo=1 ORDER BY is_system DESC, nombre"
    )
    perms_by_role = {}
    for r in roles:
        perms_by_role[r["slug"]] = get_role_permissions(r["slug"])
    return render_template("admin_roles.html",
        users=[dict(u) for u in users],
        roles=[dict(r) for r in roles],
        matrix=PERMISSIONS_MATRIX,
        perms=perms_by_role,
    )


@app.route("/admin/roles/<int:uid>", methods=["PUT"])
@require_permission("admin")
def admin_roles_update(uid):
    """Asigna rol a un usuario (acepta cualquier rol activo registrado)."""
    d = request.get_json(silent=True) or {}
    role = d.get("role", "")
    valid_rows = mysql_fetchall("SELECT slug FROM roles_dinamicos WHERE activo=1")
    valid = {r["slug"] for r in valid_rows} | {"superadmin"}
    if role not in valid:
        return jsonify({"error":"Rol no válido"}), 400
    if g.user and g.user["id"] == uid and not g.permissions.get("superadmin"):
        return jsonify({"error":"No puedes cambiar tu propio rol"}), 403
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE `{AUTH_TABLE}` SET role=%s WHERE id=%s", (role, uid))
        conn.commit()
        return jsonify({"ok":True})
    finally:
        conn.close()


@app.route("/admin/roles/matrix")
@require_permission("admin")
def admin_roles_matrix():
    """Alias legacy → redirige a la vista unificada."""
    return redirect(url_for("admin_roles") + "#matriz")


@app.route("/admin/roles/matrix", methods=["POST"])
@require_permission("admin")
def admin_roles_matrix_save():
    """Guarda la matriz completa (form: rol_slug.modulo.accion=on)."""
    rol_slug = (request.form.get("rol_slug") or "").strip()
    if not rol_slug:
        return jsonify({"error":"Rol no especificado"}), 400
    if rol_slug == "superadmin":
        return jsonify({"error":"superadmin tiene acceso total y no es editable"}), 400
    permisos_recibidos = set()
    for k,v in request.form.items():
        if k.startswith("p_"):
            try:
                _, modulo, accion = k.split(".",2)
                permisos_recibidos.add((modulo, accion))
            except ValueError: continue
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Borra y reinsertar (más simple)
            cur.execute("DELETE FROM rol_permisos WHERE rol_slug=%s", (rol_slug,))
            for modulo, cfg in PERMISSIONS_MATRIX.items():
                for accion in cfg["acciones"]:
                    permitido = 1 if (modulo, accion) in permisos_recibidos else 0
                    cur.execute(
                        "INSERT INTO rol_permisos (rol_slug,modulo,accion,permitido) "
                        "VALUES (%s,%s,%s,%s)",
                        (rol_slug, modulo, accion, permitido)
                    )
        conn.commit()
        # CRÍTICO: invalidar caché para que el cambio aplique en próximo request
        invalidate_role_cache(rol_slug)
        return jsonify({"ok":True, "rol":rol_slug,
                        "permisos":len(permisos_recibidos)})
    finally:
        conn.close()


@app.route("/admin/roles/nuevo", methods=["POST"])
@require_permission("admin")
def admin_rol_crear():
    """Crea un nuevo rol dinámico."""
    nombre = (request.form.get("nombre") or "").strip()
    color  = (request.form.get("color") or "#6b7280").strip()[:20]
    descripcion = (request.form.get("descripcion") or "").strip()[:300]
    if not nombre:
        flash("Nombre del rol obligatorio.", "danger")
        return redirect(url_for("admin_roles_matrix"))
    slug = re.sub(r"[^a-z0-9_]","", nombre.lower().replace(" ","_"))[:60]
    if not slug:
        flash("Nombre inválido.", "danger")
        return redirect(url_for("admin_roles_matrix"))
    try:
        mysql_execute(
            "INSERT INTO roles_dinamicos (slug,nombre,descripcion,color,is_system) VALUES (%s,%s,%s,%s,0)",
            (slug, nombre, descripcion, color)
        )
        flash(f"Rol \"{nombre}\" creado. Configura sus permisos.", "success")
    except Exception as exc:
        flash(f"Error al crear rol: {exc}", "danger")
    return redirect(url_for("admin_roles_matrix"))


@app.route("/admin/roles/<slug>/eliminar", methods=["POST"])
@require_permission("admin")
def admin_rol_eliminar(slug):
    """Elimina un rol dinámico (solo no-sistema)."""
    row = mysql_fetchone("SELECT is_system FROM roles_dinamicos WHERE slug=%s", (slug,))
    if not row:
        flash("Rol no encontrado.", "danger")
        return redirect(url_for("admin_roles_matrix"))
    if row.get("is_system"):
        flash("No se pueden eliminar roles del sistema.", "danger")
        return redirect(url_for("admin_roles_matrix"))
    # Reasignar usuarios a 'lector'
    mysql_execute(f"UPDATE `{AUTH_TABLE}` SET role='lector' WHERE role=%s", (slug,))
    mysql_execute("DELETE FROM rol_permisos WHERE rol_slug=%s", (slug,))
    mysql_execute("DELETE FROM roles_dinamicos WHERE slug=%s", (slug,))
    flash(f"Rol {slug} eliminado.", "success")
    return redirect(url_for("admin_roles_matrix"))


# ══════════════════════════════════════════════════════════════
#  MÓDULO: HRM — COLABORADORES
# ══════════════════════════════════════════════════════════════

@app.route("/admin/users/<int:user_id>/password-link", methods=["POST"])
@require_permission("admin")
def user_password_link(user_id):
    """Envia un enlace seguro para que el usuario cambie su propia clave."""
    user = get_auth_user_by_id(user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    try:
        token, _expires = _issue_password_token(user_id, minutes=60)
        reset_url = url_for("reset_password", token=token, _external=True)
        # email_purpose='change' → email "Cambio seguro de contraseña" (no se mezcla con invitación previa)
        result = _notify_user_access(
            user["username"], user["nombre"], user.get("phone") or "",
            mode="token", action_url=reset_url, email_purpose="change"
        )
        if result.get("email") or result.get("whatsapp"):
            msg, _level = _access_notification_flash(result, token_mode=True)
            return jsonify({"ok": True, "message": msg})
        detalle = " ".join(result.get("errors") or []) or "Revisa la configuracion SMTP o WhatsApp en Comunicaciones."
        return jsonify({"error": f"No se pudo enviar el enlace seguro. {detalle}"}), 500
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ══════════════════════════════════════════════════════════════════════
#  IMÁGENES DEL LOGIN — carrusel hasta 5 fotos
# ══════════════════════════════════════════════════════════════════════
LOGIN_IMAGES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                "static", "uploads", "login")
os.makedirs(LOGIN_IMAGES_DIR, exist_ok=True)


def _login_images_active():
    try:
        rows = mysql_fetchall(
            "SELECT id,archivo_path,titulo,subtitulo,orden FROM login_images "
            "WHERE activa=1 ORDER BY orden ASC, id ASC LIMIT 5"
        )
        return [dict(r) for r in rows]
    except Exception:
        return []


@app.route("/admin/login-imagenes", methods=["GET"])
@require_permission("admin")
def admin_login_imagenes():
    rows = mysql_fetchall(
        "SELECT * FROM login_images ORDER BY orden ASC, id ASC"
    )
    return render_template("admin/login_imagenes.html", imagenes=[dict(r) for r in rows])


@app.route("/admin/login-imagenes", methods=["POST"])
@require_permission("admin")
def admin_login_imagenes_subir():
    """Sube UNA O VARIAS imágenes al carrusel del login (máx 5 activas).

    Soporta multi-upload (input name="imagenes" multiple) y también el
    legacy input name="imagen" (single) para compatibilidad. Valida cada
    archivo con _validate_uploaded_image() y respeta el límite de 5
    activas — si se exceden, las restantes se cargan como inactivas.
    """
    # Soporta multi-upload (imagenes[]) y legacy single (imagen)
    files = request.files.getlist("imagenes")
    legacy = request.files.get("imagen")
    if legacy and legacy.filename and legacy not in files:
        files.append(legacy)
    if not files or all(not f.filename for f in files):
        flash("Selecciona al menos una imagen.", "warning")
        return redirect(url_for("admin_login_imagenes"))

    titulo_base    = (request.form.get("titulo") or "").strip()[:200]
    subtitulo_base = (request.form.get("subtitulo") or "").strip()[:300]
    want_active    = 1 if request.form.get("activa") else 0

    # Determinar orden inicial
    row = mysql_fetchone("SELECT COALESCE(MAX(orden),0) AS mx FROM login_images")
    nx = (row or {}).get("mx", 0) + 1

    # Contar activas actuales
    activas_row = mysql_fetchone("SELECT COUNT(*) AS n FROM login_images WHERE activa=1")
    activas = (activas_row or {}).get("n", 0)

    saved = 0
    errors: list[str] = []
    for i, f in enumerate(files):
        if not f or not f.filename:
            continue
        ext, err = _validate_uploaded_image(f, label=f.filename)
        if err:
            errors.append(err)
            continue
        fname = f"login_{int(time.time())}_{i}_{secure_filename(f.filename)}"
        fpath = os.path.join(LOGIN_IMAGES_DIR, fname)
        try:
            f.save(fpath)
        except Exception as e:
            errors.append(f"Error guardando {f.filename}: {e}")
            continue
        # Si excede 5 activas, marcar como inactiva
        activa = want_active
        if activa and activas >= 5:
            activa = 0
        else:
            activas += activa
        try:
            mysql_execute(
                "INSERT INTO login_images (archivo_path,titulo,subtitulo,orden,activa,created_by) "
                "VALUES (%s,%s,%s,%s,%s,%s)",
                (
                    f"uploads/login/{fname}",
                    titulo_base or None,
                    subtitulo_base or None,
                    nx + saved,
                    activa,
                    current_username(),
                )
            )
            saved += 1
        except Exception as e:
            errors.append(f"Error BD {f.filename}: {e}")
    if saved:
        flash(f"✅ {saved} imagen(es) agregadas al carrusel.", "success")
    for err in errors[:3]:
        flash(err, "warning")
    return redirect(url_for("admin_login_imagenes"))


@app.route("/admin/login-imagenes/<int:iid>", methods=["POST"])
@require_permission("admin")
def admin_login_imagen_update(iid):
    """Actualiza atributos (titulo, subtitulo, orden, activa) de una imagen."""
    action = request.form.get("action") or ""
    if action == "delete":
        row = mysql_fetchone("SELECT archivo_path FROM login_images WHERE id=%s", (iid,))
        if row:
            try:
                fp = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", row["archivo_path"])
                if os.path.exists(fp): os.remove(fp)
            except Exception: pass
        mysql_execute("DELETE FROM login_images WHERE id=%s", (iid,))
        flash("Imagen eliminada.", "success")
    else:
        titulo    = (request.form.get("titulo") or "").strip()[:200]
        subtitulo = (request.form.get("subtitulo") or "").strip()[:300]
        orden     = int(request.form.get("orden") or 0)
        activa    = 1 if request.form.get("activa") else 0
        # Limitar a 5 activas
        if activa:
            ac = mysql_fetchone(
                "SELECT COUNT(*) AS n FROM login_images WHERE activa=1 AND id <> %s", (iid,)
            )
            if ac and ac.get("n",0) >= 5:
                activa = 0
                flash("Ya hay 5 imágenes activas. Desactiva alguna para activar esta.", "warning")
        mysql_execute(
            "UPDATE login_images SET titulo=%s, subtitulo=%s, orden=%s, activa=%s WHERE id=%s",
            (titulo, subtitulo, orden, activa, iid)
        )
        flash("Imagen actualizada.", "success")
    return redirect(url_for("admin_login_imagenes"))


# ═════════════════════════════════════════════════════════════════════
# CARRUSEL PÚBLICO DE RETIROS (multi-imagen desde BD)
# Mismo patrón que login_images: subir, actualizar, eliminar.
# ═════════════════════════════════════════════════════════════════════

RETIROS_IMAGES_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "static", "uploads", "retiros"
)
os.makedirs(RETIROS_IMAGES_DIR, exist_ok=True)


def _retiros_carousel_active():
    """Devuelve lista de imágenes activas ordenadas para el carrusel público."""
    try:
        rows = mysql_fetchall(
            "SELECT id, archivo_path, titulo, subtitulo FROM retiros_carousel "
            "WHERE activa=1 ORDER BY orden ASC, id ASC"
        )
        return [dict(r) for r in (rows or [])]
    except Exception:
        return []


# Helper de validación reutilizable (seguridad: extensión + tamaño + MIME)
_ALLOWED_IMAGE_EXTS = {"png", "jpg", "jpeg", "webp"}
_MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 8 MB


def _validate_uploaded_image(f, label="imagen"):
    """Valida un FileStorage de imagen. Retorna (ext_ok, error_msg)."""
    if not f or not f.filename:
        return None, f"Selecciona un archivo ({label})."
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in _ALLOWED_IMAGE_EXTS:
        return None, f"Formato no permitido en {label}. Usa PNG/JPG/WEBP."
    # Tamaño aproximado (lee header sin cargar todo)
    f.stream.seek(0, 2)
    size = f.stream.tell()
    f.stream.seek(0)
    if size > _MAX_IMAGE_BYTES:
        return None, f"{label}: archivo muy grande ({size//(1024*1024)} MB). Máx 8 MB."
    if size == 0:
        return None, f"{label}: archivo vacío."
    return ext, None


@app.route("/admin/retiros-carousel", methods=["GET"])
@require_permission("marketing")
def admin_retiros_carousel():
    """Listado del carrusel público de retiros (admin)."""
    rows = mysql_fetchall(
        "SELECT * FROM retiros_carousel ORDER BY orden ASC, id ASC"
    )
    return render_template(
        "admin/retiros_carousel.html",
        imagenes=[dict(r) for r in (rows or [])],
    )


@app.route("/admin/retiros-carousel/subir", methods=["POST"])
@require_permission("marketing")
def admin_retiros_carousel_subir():
    """Sube UNA O VARIAS imágenes al carrusel público de retiros.

    Soporta multi-upload (input name="imagenes" multiple). Valida cada
    archivo por extensión y tamaño máx 8MB.
    """
    files = request.files.getlist("imagenes")
    if not files or all(not f.filename for f in files):
        flash("Selecciona al menos una imagen.", "warning")
        return redirect(url_for("admin_retiros_carousel"))

    # Determinar orden inicial = max actual + 1
    row = mysql_fetchone("SELECT COALESCE(MAX(orden),0) AS mx FROM retiros_carousel")
    nx = (row or {}).get("mx", 0) + 1

    titulo_base = (request.form.get("titulo") or "").strip()[:200]
    subtitulo_base = (request.form.get("subtitulo") or "").strip()[:300]
    activa = 1 if request.form.get("activa") else 0

    saved = 0
    errors = []
    for f in files:
        if not f or not f.filename:
            continue
        ext, err = _validate_uploaded_image(f, label=f.filename)
        if err:
            errors.append(err)
            continue
        fname = f"retiros_{int(time.time())}_{saved}_{secure_filename(f.filename)}"
        fpath = os.path.join(RETIROS_IMAGES_DIR, fname)
        try:
            f.save(fpath)
        except Exception as e:
            errors.append(f"Error guardando {f.filename}: {e}")
            continue
        try:
            mysql_execute(
                "INSERT INTO retiros_carousel "
                "(archivo_path,titulo,subtitulo,orden,activa,created_by) "
                "VALUES (%s,%s,%s,%s,%s,%s)",
                (
                    f"uploads/retiros/{fname}",
                    titulo_base or None,
                    subtitulo_base or None,
                    nx + saved,
                    activa,
                    current_username(),
                )
            )
            saved += 1
        except Exception as e:
            errors.append(f"Error BD {f.filename}: {e}")
    if saved:
        flash(f"✅ {saved} imagen(es) agregadas al carrusel de retiros.", "success")
    for err in errors[:3]:
        flash(err, "warning")
    return redirect(url_for("admin_retiros_carousel"))


# ═════════════════════════════════════════════════════════════════════
# BANNER PÚBLICO DE RETIROS (avisos)
# Anuncios temporales que se muestran arriba del hero en /retiros/solicitar
# (ej. "Cierre anticipado por inventario", "Feriado regional", etc.)
# ═════════════════════════════════════════════════════════════════════

def _retiros_announcements_active():
    """Devuelve lista de avisos vigentes para el banner público."""
    try:
        rows = mysql_fetchall(
            "SELECT id, titulo, mensaje, tipo, icon FROM retiros_announcements "
            "WHERE activa=1 "
            "  AND (fecha_desde IS NULL OR fecha_desde <= NOW()) "
            "  AND (fecha_hasta IS NULL OR fecha_hasta >= NOW()) "
            "ORDER BY orden ASC, id DESC"
        )
        return [dict(r) for r in (rows or [])]
    except Exception:
        return []


@app.route("/admin/retiros-avisos", methods=["GET"])
@require_permission("marketing")
def admin_retiros_avisos():
    rows = mysql_fetchall(
        "SELECT * FROM retiros_announcements ORDER BY activa DESC, orden ASC, id DESC"
    )
    return render_template(
        "admin/retiros_avisos.html",
        avisos=[dict(r) for r in (rows or [])],
    )


@app.route("/admin/retiros-avisos/nuevo", methods=["POST"])
@require_permission("marketing")
def admin_retiros_avisos_nuevo():
    titulo = (request.form.get("titulo") or "").strip()[:200]
    if not titulo:
        flash("El título es obligatorio.", "warning")
        return redirect(url_for("admin_retiros_avisos"))
    mensaje      = (request.form.get("mensaje") or "").strip()
    tipo         = (request.form.get("tipo") or "info").strip().lower()
    if tipo not in ("info", "warning", "danger", "success"):
        tipo = "info"
    icon         = (request.form.get("icon") or "info-circle").strip()[:40]
    fecha_desde  = (request.form.get("fecha_desde") or "").strip() or None
    fecha_hasta  = (request.form.get("fecha_hasta") or "").strip() or None
    activa       = 1 if request.form.get("activa") else 0
    orden        = int(request.form.get("orden") or 0)
    mysql_execute(
        "INSERT INTO retiros_announcements "
        "(titulo,mensaje,tipo,icon,fecha_desde,fecha_hasta,activa,orden,created_by) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)",
        (titulo, mensaje, tipo, icon, fecha_desde, fecha_hasta, activa, orden, current_username())
    )
    flash("Aviso publicado.", "success")
    return redirect(url_for("admin_retiros_avisos"))


@app.route("/admin/retiros-avisos/<int:aid>", methods=["POST"])
@require_permission("marketing")
def admin_retiros_avisos_update(aid):
    action = request.form.get("action") or ""
    if action == "delete":
        mysql_execute("DELETE FROM retiros_announcements WHERE id=%s", (aid,))
        flash("Aviso eliminado.", "success")
        return redirect(url_for("admin_retiros_avisos"))
    titulo = (request.form.get("titulo") or "").strip()[:200]
    mensaje = (request.form.get("mensaje") or "").strip()
    tipo = (request.form.get("tipo") or "info").strip().lower()
    if tipo not in ("info", "warning", "danger", "success"):
        tipo = "info"
    icon = (request.form.get("icon") or "info-circle").strip()[:40]
    fecha_desde = (request.form.get("fecha_desde") or "").strip() or None
    fecha_hasta = (request.form.get("fecha_hasta") or "").strip() or None
    activa = 1 if request.form.get("activa") else 0
    orden = int(request.form.get("orden") or 0)
    mysql_execute(
        "UPDATE retiros_announcements "
        "SET titulo=%s, mensaje=%s, tipo=%s, icon=%s, "
        "fecha_desde=%s, fecha_hasta=%s, activa=%s, orden=%s "
        "WHERE id=%s",
        (titulo, mensaje, tipo, icon, fecha_desde, fecha_hasta, activa, orden, aid)
    )
    flash("Aviso actualizado.", "success")
    return redirect(url_for("admin_retiros_avisos"))


@app.route("/admin/retiros-carousel/<int:iid>", methods=["POST"])
@require_permission("marketing")
def admin_retiros_carousel_update(iid):
    """Actualiza o elimina una imagen del carrusel."""
    action = request.form.get("action") or ""
    if action == "delete":
        row = mysql_fetchone("SELECT archivo_path FROM retiros_carousel WHERE id=%s", (iid,))
        if row:
            try:
                fp = os.path.join(
                    os.path.dirname(os.path.abspath(__file__)),
                    "static", row["archivo_path"]
                )
                if os.path.exists(fp):
                    os.remove(fp)
            except Exception:
                pass
        mysql_execute("DELETE FROM retiros_carousel WHERE id=%s", (iid,))
        flash("Imagen eliminada del carrusel.", "success")
    else:
        titulo    = (request.form.get("titulo") or "").strip()[:200]
        subtitulo = (request.form.get("subtitulo") or "").strip()[:300]
        orden     = int(request.form.get("orden") or 0)
        activa    = 1 if request.form.get("activa") else 0
        mysql_execute(
            "UPDATE retiros_carousel "
            "SET titulo=%s, subtitulo=%s, orden=%s, activa=%s WHERE id=%s",
            (titulo or None, subtitulo or None, orden, activa, iid)
        )
        flash("Imagen actualizada.", "success")
    return redirect(url_for("admin_retiros_carousel"))


HRM_AREAS_TABLE  = "hrm_areas"
HRM_CARGOS_TABLE = "hrm_cargos"
HRM_COLAB_TABLE  = "hrm_colaboradores"
PREG_GEN_TABLE   = "eval_preguntas_genericas"

CHILE_REGIONES = [
    "Arica y Parinacota",
    "Tarapacá",
    "Antofagasta",
    "Atacama",
    "Coquimbo",
    "Valparaíso",
    "Metropolitana de Santiago",
    "O'Higgins",
    "Maule",
    "Ñuble",
    "Biobío",
    "La Araucanía",
    "Los Ríos",
    "Los Lagos",
    "Aysén",
    "Magallanes",
]

RESETS_TABLE = "app_password_resets"

GENEROS = {
    "masculino":        "Masculino",
    "femenino":         "Femenino",
    "otro":             "Otro",
    "no_especificado":  "Prefiero no indicar",
}

ESTADOS_COLAB = {
    "activo":   "Activo",
    "inactivo": "Inactivo",
    "licencia": "Con licencia",
}


def init_hrm_tables():
    """Crea las tablas del módulo HRM si no existen."""
    os.makedirs(COLABS_FOLDER, exist_ok=True)
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{HRM_AREAS_TABLE}` (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    nombre      VARCHAR(120) NOT NULL UNIQUE,
                    descripcion TEXT,
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{HRM_CARGOS_TABLE}` (
                    id         INT AUTO_INCREMENT PRIMARY KEY,
                    nombre     VARCHAR(120) NOT NULL,
                    descriptor TEXT,
                    area_id    INT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    CONSTRAINT fk_cargo_area FOREIGN KEY (area_id)
                        REFERENCES `{HRM_AREAS_TABLE}`(id) ON DELETE SET NULL
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{HRM_COLAB_TABLE}` (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    nombre_completo VARCHAR(200) NOT NULL,
                    rut             VARCHAR(20),
                    email           VARCHAR(190),
                    telefono        VARCHAR(30),
                    direccion       TEXT,
                    genero          ENUM('masculino','femenino','otro','no_especificado')
                                    DEFAULT 'no_especificado',
                    cargo_id        INT,
                    area_id         INT,
                    foto_filename   VARCHAR(255),
                    fecha_ingreso   DATE,
                    estado          ENUM('activo','inactivo','licencia') DEFAULT 'activo',
                    created_by      VARCHAR(190),
                    updated_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    CONSTRAINT fk_colab_cargo FOREIGN KEY (cargo_id)
                        REFERENCES `{HRM_CARGOS_TABLE}`(id) ON DELETE SET NULL,
                    CONSTRAINT fk_colab_area FOREIGN KEY (area_id)
                        REFERENCES `{HRM_AREAS_TABLE}`(id) ON DELETE SET NULL
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PREG_GEN_TABLE}` (
                    id             INT AUTO_INCREMENT PRIMARY KEY,
                    seccion        ENUM('tecnica','operativa','conductual','cumplimiento') NOT NULL,
                    texto          TEXT NOT NULL,
                    tipo_respuesta ENUM('escala_1_5','texto_libre','multiple','si_no','porcentaje')
                                   DEFAULT 'escala_1_5',
                    opciones       JSON,
                    es_obligatoria TINYINT(1) DEFAULT 1,
                    activa         TINYINT(1) DEFAULT 1,
                    created_by     VARCHAR(190),
                    updated_by     VARCHAR(190),
                    created_at     DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at     DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
        conn.commit()
        # ── Migración segura: agregar columnas región/comuna si no existen ──
        for col, definition in [
            ("region",  "VARCHAR(100)"),
            ("comuna",  "VARCHAR(100)"),
        ]:
            try:
                with conn.cursor() as cur:
                    cur.execute(f"""
                        ALTER TABLE `{HRM_COLAB_TABLE}`
                        ADD COLUMN IF NOT EXISTS `{col}` {definition}
                    """)
                conn.commit()
            except Exception:
                pass  # MySQL <8 no soporta IF NOT EXISTS en ALTER; ignorar si ya existe
    finally:
        conn.close()


def _get_or_create_area(nombre_raw, conn):
    """Devuelve area_id. Si no existe, la crea."""
    nombre = nombre_raw.strip()
    if not nombre:
        return None
    row = mysql_fetchone(f"SELECT id FROM `{HRM_AREAS_TABLE}` WHERE nombre=%s", (nombre,))
    if row:
        return row["id"]
    with conn.cursor() as cur:
        cur.execute(f"INSERT INTO `{HRM_AREAS_TABLE}` (nombre) VALUES (%s)", (nombre,))
        return cur.lastrowid


def _get_or_create_cargo(nombre_raw, descriptor_raw, area_id, conn):
    """Devuelve cargo_id. Si no existe, lo crea."""
    nombre = nombre_raw.strip()
    if not nombre:
        return None
    row = mysql_fetchone(
        f"SELECT id FROM `{HRM_CARGOS_TABLE}` WHERE nombre=%s AND (area_id=%s OR area_id IS NULL)",
        (nombre, area_id),
    )
    if row:
        return row["id"]
    desc = (descriptor_raw or "").strip() or None
    with conn.cursor() as cur:
        cur.execute(
            f"INSERT INTO `{HRM_CARGOS_TABLE}` (nombre, descriptor, area_id) VALUES (%s,%s,%s)",
            (nombre, desc, area_id),
        )
        return cur.lastrowid


def _save_colab_foto(file, colab_id):
    """Sube la foto a Cloudinary (o disco local) y devuelve filename/URL, o None."""
    if not file or not file.filename:
        return None
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXT:
        return None
    if _CLD_READY:
        try:
            return _cloud_upload(file, public_id=f"colab_{colab_id}", folder="ilus/colabs")
        except Exception as exc:
            print(f"[ILUS] Cloudinary colab upload error: {exc}")
            return None
    else:
        fname = f"colab_{colab_id}.{ext}"
        file.save(os.path.join(COLABS_FOLDER, fname))
        return fname


# ── Listado colaboradores ─────────────────────────────────────

@app.route("/colaboradores/")
@login_required
def colab_index():
    if not g.permissions.get("hrm") and not g.permissions.get("view"):
        flash("Sin permisos.", "danger")
        return redirect(url_for("index"))
    rows = mysql_fetchall(f"""
        SELECT c.*,
               a.nombre AS area_nombre,
               ca.nombre AS cargo_nombre
        FROM `{HRM_COLAB_TABLE}` c
        LEFT JOIN `{HRM_AREAS_TABLE}` a  ON a.id  = c.area_id
        LEFT JOIN `{HRM_CARGOS_TABLE}` ca ON ca.id = c.cargo_id
        ORDER BY c.nombre_completo ASC
    """)
    return render_template("colaboradores/index.html",
                           colaboradores=rows, estados=ESTADOS_COLAB)


# ── Nueva ficha ───────────────────────────────────────────────

@app.route("/colaboradores/nuevo", methods=["GET", "POST"])
@require_permission("hrm")
def colab_nuevo():
    areas  = mysql_fetchall(f"SELECT * FROM `{HRM_AREAS_TABLE}` ORDER BY nombre")
    cargos = mysql_fetchall(f"SELECT c.*, a.nombre AS area_nombre FROM `{HRM_CARGOS_TABLE}` c LEFT JOIN `{HRM_AREAS_TABLE}` a ON a.id=c.area_id ORDER BY c.nombre")

    if request.method == "POST":
        nombre     = request.form.get("nombre_completo", "").strip()
        if not nombre:
            flash("El nombre completo es requerido.", "danger")
            return render_template("colaboradores/form.html",
                                   colab=None, fd=request.form,
                                   areas=areas, cargos=cargos,
                                   generos=GENEROS, estados=ESTADOS_COLAB,
                                   regiones=CHILE_REGIONES)

        conn = get_db()

        # Área: id existente O nombre libre → auto-crear
        area_id = request.form.get("area_id", "").strip()
        area_nueva = request.form.get("area_nueva", "").strip()
        if area_nueva:
            area_id = _get_or_create_area(area_nueva, conn)
        elif area_id:
            area_id = int(area_id)
        else:
            area_id = None

        # Cargo: id existente O nombre libre → auto-crear
        cargo_id = request.form.get("cargo_id", "").strip()
        cargo_nuevo  = request.form.get("cargo_nuevo", "").strip()
        cargo_desc   = request.form.get("cargo_descriptor", "").strip()
        if cargo_nuevo:
            cargo_id = _get_or_create_cargo(cargo_nuevo, cargo_desc, area_id, conn)
        elif cargo_id:
            cargo_id = int(cargo_id)
        else:
            cargo_id = None

        fecha_ing = request.form.get("fecha_ingreso") or None
        region_val = request.form.get("region", "").strip() or None
        comuna_val = request.form.get("comuna", "").strip() or None

        with conn.cursor() as cur:
            cur.execute(f"""
                INSERT INTO `{HRM_COLAB_TABLE}`
                    (nombre_completo, rut, email, telefono, direccion,
                     region, comuna, genero, cargo_id, area_id,
                     fecha_ingreso, estado, created_by)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            """, (
                nombre,
                request.form.get("rut", "").strip() or None,
                request.form.get("email", "").strip() or None,
                request.form.get("telefono", "").strip() or None,
                request.form.get("direccion", "").strip() or None,
                region_val, comuna_val,
                request.form.get("genero", "no_especificado"),
                cargo_id, area_id, fecha_ing,
                request.form.get("estado", "activo"),
                current_username(),
            ))
            cid = cur.lastrowid

        # Foto (después de obtener el id)
        foto = request.files.get("foto")
        fname = _save_colab_foto(foto, cid)
        if fname:
            with conn.cursor() as cur:
                cur.execute(f"UPDATE `{HRM_COLAB_TABLE}` SET foto_filename=%s WHERE id=%s",
                            (fname, cid))
        conn.commit()

        flash(f"Colaborador {nombre} creado correctamente.", "success")
        return redirect(url_for("colab_ficha", cid=cid))

    return render_template("colaboradores/form.html",
                           colab=None, fd={},
                           areas=areas, cargos=cargos,
                           generos=GENEROS, estados=ESTADOS_COLAB,
                           regiones=CHILE_REGIONES)


# ── Ficha de colaborador ──────────────────────────────────────

@app.route("/colaboradores/<int:cid>/")
@login_required
def colab_ficha(cid):
    row = mysql_fetchone(f"""
        SELECT c.*,
               a.nombre  AS area_nombre,
               ca.nombre AS cargo_nombre,
               ca.descriptor AS cargo_descriptor
        FROM `{HRM_COLAB_TABLE}` c
        LEFT JOIN `{HRM_AREAS_TABLE}`  a  ON a.id  = c.area_id
        LEFT JOIN `{HRM_CARGOS_TABLE}` ca ON ca.id = c.cargo_id
        WHERE c.id=%s
    """, (cid,))
    if not row:
        flash("Colaborador no encontrado.", "danger")
        return redirect(url_for("colab_index"))
    return render_template("colaboradores/ficha.html",
                           c=row, generos=GENEROS, estados=ESTADOS_COLAB)


# ── Editar colaborador ────────────────────────────────────────

@app.route("/colaboradores/<int:cid>/editar", methods=["GET", "POST"])
@require_permission("hrm")
def colab_editar(cid):
    colab  = mysql_fetchone(f"SELECT * FROM `{HRM_COLAB_TABLE}` WHERE id=%s", (cid,))
    if not colab:
        flash("Colaborador no encontrado.", "danger")
        return redirect(url_for("colab_index"))
    areas  = mysql_fetchall(f"SELECT * FROM `{HRM_AREAS_TABLE}` ORDER BY nombre")
    cargos = mysql_fetchall(f"SELECT c.*, a.nombre AS area_nombre FROM `{HRM_CARGOS_TABLE}` c LEFT JOIN `{HRM_AREAS_TABLE}` a ON a.id=c.area_id ORDER BY c.nombre")

    if request.method == "POST":
        nombre = request.form.get("nombre_completo", "").strip()
        if not nombre:
            flash("El nombre es requerido.", "danger")
            return render_template("colaboradores/form.html",
                                   colab=colab, fd=request.form,
                                   areas=areas, cargos=cargos,
                                   generos=GENEROS, estados=ESTADOS_COLAB,
                                   regiones=CHILE_REGIONES)

        conn = get_db()

        area_id = request.form.get("area_id", "").strip()
        area_nueva = request.form.get("area_nueva", "").strip()
        if area_nueva:
            area_id = _get_or_create_area(area_nueva, conn)
        elif area_id:
            area_id = int(area_id)
        else:
            area_id = None

        cargo_id = request.form.get("cargo_id", "").strip()
        cargo_nuevo = request.form.get("cargo_nuevo", "").strip()
        cargo_desc  = request.form.get("cargo_descriptor", "").strip()
        if cargo_nuevo:
            cargo_id = _get_or_create_cargo(cargo_nuevo, cargo_desc, area_id, conn)
        elif cargo_id:
            cargo_id = int(cargo_id)
        else:
            cargo_id = None

        fecha_ing  = request.form.get("fecha_ingreso") or None
        region_val = request.form.get("region", "").strip() or None
        comuna_val = request.form.get("comuna", "").strip() or None

        with conn.cursor() as cur:
            cur.execute(f"""
                UPDATE `{HRM_COLAB_TABLE}` SET
                    nombre_completo=%s, rut=%s, email=%s, telefono=%s,
                    direccion=%s, region=%s, comuna=%s,
                    genero=%s, cargo_id=%s, area_id=%s,
                    fecha_ingreso=%s, estado=%s, updated_by=%s
                WHERE id=%s
            """, (
                nombre,
                request.form.get("rut", "").strip() or None,
                request.form.get("email", "").strip() or None,
                request.form.get("telefono", "").strip() or None,
                request.form.get("direccion", "").strip() or None,
                region_val, comuna_val,
                request.form.get("genero", "no_especificado"),
                cargo_id, area_id, fecha_ing,
                request.form.get("estado", "activo"),
                current_username(), cid,
            ))

        # Foto nueva (opcional)
        foto = request.files.get("foto")
        fname = _save_colab_foto(foto, cid)
        if fname:
            with conn.cursor() as cur:
                cur.execute(f"UPDATE `{HRM_COLAB_TABLE}` SET foto_filename=%s WHERE id=%s",
                            (fname, cid))
        conn.commit()

        flash("Ficha actualizada.", "success")
        return redirect(url_for("colab_ficha", cid=cid))

    return render_template("colaboradores/form.html",
                           colab=colab, fd={},
                           areas=areas, cargos=cargos,
                           generos=GENEROS, estados=ESTADOS_COLAB,
                           regiones=CHILE_REGIONES)


# ── Eliminar colaborador ──────────────────────────────────────

@app.route("/colaboradores/<int:cid>/eliminar", methods=["POST"])
@require_permission("delete")
def colab_eliminar(cid):
    colab = mysql_fetchone(f"SELECT foto_filename, nombre_completo FROM `{HRM_COLAB_TABLE}` WHERE id=%s", (cid,))
    if not colab:
        flash("Colaborador no encontrado.", "danger")
        return redirect(url_for("colab_index"))
    # Eliminar foto si existe
    if colab.get("foto_filename"):
        try:
            os.remove(os.path.join(COLABS_FOLDER, colab["foto_filename"]))
        except Exception:
            pass
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"DELETE FROM `{HRM_COLAB_TABLE}` WHERE id=%s", (cid,))
    conn.commit()
    flash(f"Colaborador {colab['nombre_completo']} eliminado.", "warning")
    return redirect(url_for("colab_index"))


# ══════════════════════════════════════════════════════════════
#  MÓDULO: PREGUNTAS GENÉRICAS (solo superadmin)
# ══════════════════════════════════════════════════════════════

def _require_superadmin(view):
    @wraps(view)
    def wrapped(*a, **kw):
        if not g.user:
            flash("Inicia sesión para continuar.", "warning")
            return redirect(url_for("login", next=request.path))
        if not g.permissions.get("superadmin"):
            flash("Esta acción requiere permisos de Super Administrador.", "danger")
            return redirect(url_for("index"))
        return view(*a, **kw)
    return wrapped


@app.route("/admin/preguntas-genericas/")
@_require_superadmin
def preg_gen_index():
    rows = mysql_fetchall(f"""
        SELECT * FROM `{PREG_GEN_TABLE}`
        ORDER BY seccion, id ASC
    """)
    for r in rows:
        _parse_opciones(r)
    por_seccion = {s: [] for s in SECCIONES}
    for r in rows:
        sec = r.get("seccion", "tecnica")
        if sec in por_seccion:
            por_seccion[sec].append(r)
    return render_template("admin/preguntas_genericas.html",
                           por_seccion=por_seccion,
                           secciones=SECCIONES, tipos=TIPOS_RESPUESTA)


@app.route("/admin/preguntas-genericas/guardar", methods=["POST"])
@_require_superadmin
def preg_gen_guardar():
    pid_raw        = request.form.get("preg_id", "").strip()
    seccion        = request.form.get("seccion", "tecnica")
    texto          = request.form.get("texto", "").strip()
    tipo_respuesta = request.form.get("tipo_respuesta", "escala_1_5")
    es_obligatoria = 1 if request.form.get("es_obligatoria") else 0
    activa         = 1 if request.form.get("activa", "1") else 0
    opciones_raw   = request.form.get("opciones", "").strip()

    if not texto:
        return jsonify({"ok": False, "error": "El texto es requerido"}), 400
    if seccion not in SECCIONES:
        return jsonify({"ok": False, "error": "Sección inválida"}), 400

    opciones_json = None
    if tipo_respuesta == "multiple" and opciones_raw:
        items = [o.strip() for o in opciones_raw.split("\n") if o.strip()]
        opciones_json = json.dumps(items, ensure_ascii=False)

    conn = get_db()
    if pid_raw:
        # Editar
        with conn.cursor() as cur:
            cur.execute(f"""
                UPDATE `{PREG_GEN_TABLE}`
                SET seccion=%s, texto=%s, tipo_respuesta=%s,
                    opciones=%s, es_obligatoria=%s, activa=%s, updated_by=%s
                WHERE id=%s
            """, (seccion, texto, tipo_respuesta, opciones_json,
                  es_obligatoria, activa, current_username(), int(pid_raw)))
        conn.commit()
        p = mysql_fetchone(f"SELECT * FROM `{PREG_GEN_TABLE}` WHERE id=%s", (int(pid_raw),))
    else:
        # Crear
        with conn.cursor() as cur:
            cur.execute(f"""
                INSERT INTO `{PREG_GEN_TABLE}`
                    (seccion, texto, tipo_respuesta, opciones, es_obligatoria, activa, created_by)
                VALUES (%s,%s,%s,%s,%s,%s,%s)
            """, (seccion, texto, tipo_respuesta, opciones_json,
                  es_obligatoria, activa, current_username()))
            pid_new = cur.lastrowid
        conn.commit()
        p = mysql_fetchone(f"SELECT * FROM `{PREG_GEN_TABLE}` WHERE id=%s", (pid_new,))

    _parse_opciones(p)
    return jsonify({"ok": True, "pregunta": dict(p)})


@app.route("/admin/preguntas-genericas/<int:pid>/toggle", methods=["POST"])
@_require_superadmin
def preg_gen_toggle(pid):
    """Activa / desactiva una pregunta genérica."""
    p = mysql_fetchone(f"SELECT id, activa FROM `{PREG_GEN_TABLE}` WHERE id=%s", (pid,))
    if not p:
        return jsonify({"ok": False, "error": "No encontrada"}), 404
    nuevo = 0 if p["activa"] else 1
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"UPDATE `{PREG_GEN_TABLE}` SET activa=%s, updated_by=%s WHERE id=%s",
                    (nuevo, current_username(), pid))
    conn.commit()
    return jsonify({"ok": True, "activa": nuevo})


@app.route("/admin/preguntas-genericas/<int:pid>/eliminar", methods=["POST"])
@_require_superadmin
def preg_gen_eliminar(pid):
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"DELETE FROM `{PREG_GEN_TABLE}` WHERE id=%s", (pid,))
    conn.commit()
    return jsonify({"ok": True})




# ══════════════════════════════════════════════════════════════
#  MÓDULO: GESTIÓN DE EVALUACIONES
# ══════════════════════════════════════════════════════════════

EVAL_TABLE = "eval_evaluaciones"
PREG_TABLE = "eval_preguntas"

SECCIONES = {
    "tecnica":      {"label": "Evaluación Técnica",         "color": "#1a4a8a", "icon": "bi-tools"},
    "operativa":    {"label": "Evaluación Operativa",        "color": "#1a7a1a", "icon": "bi-gear-fill"},
    "conductual":   {"label": "Evaluación Conductual",       "color": "#7a4a00", "icon": "bi-person-check-fill"},
    "cumplimiento": {"label": "Cumplimiento de Procesos",    "color": "#6a006a", "icon": "bi-clipboard2-check-fill"},
}

TIPOS_RESPUESTA = {
    "escala_1_5":  "Escala 1 – 5",
    "texto_libre": "Texto libre",
    "multiple":    "Selección múltiple",
    "si_no":       "Sí / No",
    "porcentaje":  "Porcentaje (0–100)",
}

TIPOS_EVAL = {
    "diagnostica": "Evaluación Diagnóstica",
    "periodica":   "Evaluación Periódica",
    "especial":    "Evaluación Especial",
}


def init_eval_tables():
    """Crea las tablas del módulo de evaluaciones si no existen."""
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{EVAL_TABLE}` (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    nombre      VARCHAR(200) NOT NULL,
                    descripcion TEXT,
                    tipo        ENUM('diagnostica','periodica','especial') DEFAULT 'diagnostica',
                    estado      ENUM('borrador','publicada','archivada')   DEFAULT 'borrador',
                    created_by  VARCHAR(190),
                    updated_by  VARCHAR(190),
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at  DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute(f"""
                CREATE TABLE IF NOT EXISTS `{PREG_TABLE}` (
                    id             INT AUTO_INCREMENT PRIMARY KEY,
                    evaluacion_id  INT NOT NULL,
                    seccion        ENUM('tecnica','operativa','conductual','cumplimiento') NOT NULL,
                    orden          INT          DEFAULT 0,
                    texto          TEXT         NOT NULL,
                    tipo_respuesta ENUM('escala_1_5','texto_libre','multiple','si_no','porcentaje')
                                               DEFAULT 'escala_1_5',
                    opciones       JSON,
                    es_obligatoria TINYINT(1)   DEFAULT 1,
                    created_at     DATETIME     DEFAULT CURRENT_TIMESTAMP,
                    updated_at     DATETIME     DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    CONSTRAINT fk_preg_eval FOREIGN KEY (evaluacion_id)
                        REFERENCES `{EVAL_TABLE}`(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
        conn.commit()
    finally:
        conn.close()


# ── helpers internos ──────────────────────────────────────────

def _parse_opciones(p):
    """Convierte el campo opciones JSON→list en cada fila de pregunta."""
    if p.get("opciones") and isinstance(p["opciones"], str):
        try:
            p["opciones"] = json.loads(p["opciones"])
        except Exception:
            p["opciones"] = []
    elif not p.get("opciones"):
        p["opciones"] = []
    return p


def _preguntas_por_seccion(eid):
    rows = mysql_fetchall(
        f"SELECT * FROM `{PREG_TABLE}` WHERE evaluacion_id=%s ORDER BY seccion, orden ASC",
        (eid,),
    )
    por_seccion = {s: [] for s in SECCIONES}
    for p in rows:
        _parse_opciones(p)
        sec = p.get("seccion", "tecnica")
        if sec in por_seccion:
            por_seccion[sec].append(p)
    return por_seccion


# ── Listado ───────────────────────────────────────────────────

@app.route("/evaluaciones/")
@login_required
def eval_index():
    evals = mysql_fetchall(f"""
        SELECT e.*, COUNT(p.id) AS total_preguntas
        FROM `{EVAL_TABLE}` e
        LEFT JOIN `{PREG_TABLE}` p ON p.evaluacion_id = e.id
        GROUP BY e.id
        ORDER BY e.created_at DESC
    """)
    return render_template("evaluaciones/index.html",
                           evals=evals, tipos=TIPOS_EVAL)


# ── Crear evaluación ──────────────────────────────────────────

@app.route("/evaluaciones/nueva", methods=["GET", "POST"])
@require_permission("edit")
def eval_nueva():
    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        desc   = request.form.get("descripcion", "").strip()
        tipo   = request.form.get("tipo", "diagnostica")
        if not nombre:
            flash("El nombre es requerido.", "danger")
            return render_template("evaluaciones/form.html",
                                   ev=None, fd=request.form, tipos=TIPOS_EVAL)
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(f"""
                INSERT INTO `{EVAL_TABLE}` (nombre, descripcion, tipo, estado, created_by)
                VALUES (%s, %s, %s, 'borrador', %s)
            """, (nombre, desc, tipo, current_username()))
            eid = cur.lastrowid
        conn.commit()
        flash("Evaluación creada. Ahora agrega las preguntas.", "success")
        return redirect(url_for("eval_constructor", eid=eid))
    return render_template("evaluaciones/form.html",
                           ev=None, fd={}, tipos=TIPOS_EVAL)


# ── Editar metadatos ──────────────────────────────────────────

@app.route("/evaluaciones/<int:eid>/editar", methods=["GET", "POST"])
@require_permission("edit")
def eval_editar(eid):
    ev = mysql_fetchone(f"SELECT * FROM `{EVAL_TABLE}` WHERE id=%s", (eid,))
    if not ev:
        flash("Evaluación no encontrada.", "danger")
        return redirect(url_for("eval_index"))
    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        desc   = request.form.get("descripcion", "").strip()
        tipo   = request.form.get("tipo", "diagnostica")
        if not nombre:
            flash("El nombre es requerido.", "danger")
            return render_template("evaluaciones/form.html",
                                   ev=ev, fd=request.form, tipos=TIPOS_EVAL)
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(f"""
                UPDATE `{EVAL_TABLE}`
                SET nombre=%s, descripcion=%s, tipo=%s, updated_by=%s
                WHERE id=%s
            """, (nombre, desc, tipo, current_username(), eid))
        conn.commit()
        flash("Evaluación actualizada.", "success")
        return redirect(url_for("eval_constructor", eid=eid))
    return render_template("evaluaciones/form.html",
                           ev=ev, fd={}, tipos=TIPOS_EVAL)


# ── Constructor de preguntas ──────────────────────────────────

@app.route("/evaluaciones/<int:eid>/constructor")
@login_required
def eval_constructor(eid):
    ev = mysql_fetchone(f"SELECT * FROM `{EVAL_TABLE}` WHERE id=%s", (eid,))
    if not ev:
        flash("Evaluación no encontrada.", "danger")
        return redirect(url_for("eval_index"))
    por_seccion = _preguntas_por_seccion(eid)
    total = sum(len(v) for v in por_seccion.values())
    return render_template("evaluaciones/constructor.html",
                           ev=ev, por_seccion=por_seccion,
                           secciones=SECCIONES, tipos=TIPOS_RESPUESTA,
                           tipos_eval=TIPOS_EVAL, total_preguntas=total)


# ── API: Agregar pregunta ─────────────────────────────────────

@app.route("/api/evaluaciones/<int:eid>/preguntas", methods=["POST"])
@require_permission("edit")
def api_agregar_pregunta(eid):
    ev = mysql_fetchone(f"SELECT id, estado FROM `{EVAL_TABLE}` WHERE id=%s", (eid,))
    if not ev:
        return jsonify({"ok": False, "error": "Evaluación no encontrada"}), 404
    if ev["estado"] == "archivada":
        return jsonify({"ok": False, "error": "Evaluación archivada, no se puede modificar"}), 403

    seccion        = request.form.get("seccion", "tecnica")
    texto          = request.form.get("texto", "").strip()
    tipo_respuesta = request.form.get("tipo_respuesta", "escala_1_5")
    es_obligatoria = 1 if request.form.get("es_obligatoria") else 0
    opciones_raw   = request.form.get("opciones", "").strip()

    if not texto:
        return jsonify({"ok": False, "error": "El texto de la pregunta es requerido"}), 400
    if seccion not in SECCIONES:
        return jsonify({"ok": False, "error": "Sección inválida"}), 400

    max_ord = mysql_fetchone(
        f"SELECT COALESCE(MAX(orden),0) AS mo FROM `{PREG_TABLE}` WHERE evaluacion_id=%s AND seccion=%s",
        (eid, seccion),
    )
    nuevo_orden = int(max_ord["mo"]) + 1

    opciones_json = None
    if tipo_respuesta == "multiple" and opciones_raw:
        items = [o.strip() for o in opciones_raw.split("\n") if o.strip()]
        opciones_json = json.dumps(items, ensure_ascii=False)

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"""
            INSERT INTO `{PREG_TABLE}`
                (evaluacion_id, seccion, orden, texto, tipo_respuesta, opciones, es_obligatoria)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (eid, seccion, nuevo_orden, texto, tipo_respuesta, opciones_json, es_obligatoria))
        pid = cur.lastrowid
    conn.commit()

    p = mysql_fetchone(f"SELECT * FROM `{PREG_TABLE}` WHERE id=%s", (pid,))
    return jsonify({"ok": True, "pregunta": dict(_parse_opciones(p))})


# ── API: Editar pregunta ──────────────────────────────────────

@app.route("/api/preguntas/<int:pid>", methods=["POST"])
@require_permission("edit")
def api_editar_pregunta(pid):
    p = mysql_fetchone(f"""
        SELECT p.*, e.estado FROM `{PREG_TABLE}` p
        JOIN `{EVAL_TABLE}` e ON e.id = p.evaluacion_id
        WHERE p.id=%s
    """, (pid,))
    if not p:
        return jsonify({"ok": False, "error": "Pregunta no encontrada"}), 404
    if p["estado"] == "archivada":
        return jsonify({"ok": False, "error": "Evaluación archivada, no editable"}), 403

    texto          = request.form.get("texto", "").strip()
    tipo_respuesta = request.form.get("tipo_respuesta", p["tipo_respuesta"])
    es_obligatoria = 1 if request.form.get("es_obligatoria") else 0
    opciones_raw   = request.form.get("opciones", "").strip()

    if not texto:
        return jsonify({"ok": False, "error": "El texto es requerido"}), 400

    opciones_json = None
    if tipo_respuesta == "multiple" and opciones_raw:
        items = [o.strip() for o in opciones_raw.split("\n") if o.strip()]
        opciones_json = json.dumps(items, ensure_ascii=False)

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"""
            UPDATE `{PREG_TABLE}`
            SET texto=%s, tipo_respuesta=%s, opciones=%s, es_obligatoria=%s
            WHERE id=%s
        """, (texto, tipo_respuesta, opciones_json, es_obligatoria, pid))
    conn.commit()

    updated = mysql_fetchone(f"SELECT * FROM `{PREG_TABLE}` WHERE id=%s", (pid,))
    return jsonify({"ok": True, "pregunta": dict(_parse_opciones(updated))})


# ── API: Eliminar pregunta ────────────────────────────────────

@app.route("/api/preguntas/<int:pid>/eliminar", methods=["POST"])
@require_permission("edit")
def api_eliminar_pregunta(pid):
    p = mysql_fetchone(f"""
        SELECT p.*, e.estado FROM `{PREG_TABLE}` p
        JOIN `{EVAL_TABLE}` e ON e.id = p.evaluacion_id
        WHERE p.id=%s
    """, (pid,))
    if not p:
        return jsonify({"ok": False, "error": "Pregunta no encontrada"}), 404
    if p["estado"] == "archivada":
        return jsonify({"ok": False, "error": "Evaluación archivada"}), 403

    eid     = p["evaluacion_id"]
    seccion = p["seccion"]

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"DELETE FROM `{PREG_TABLE}` WHERE id=%s", (pid,))
        # reindexar posiciones en la sección
        cur.execute(f"""
            SELECT id FROM `{PREG_TABLE}`
            WHERE evaluacion_id=%s AND seccion=%s
            ORDER BY orden ASC
        """, (eid, seccion))
        for i, r in enumerate(cur.fetchall(), 1):
            cur.execute(f"UPDATE `{PREG_TABLE}` SET orden=%s WHERE id=%s", (i, r["id"]))
    conn.commit()
    return jsonify({"ok": True})


# ── API: Mover pregunta ↑ ↓ ──────────────────────────────────

@app.route("/api/preguntas/<int:pid>/mover", methods=["POST"])
@require_permission("edit")
def api_mover_pregunta(pid):
    direccion = request.form.get("direccion")   # "up" | "down"
    p = mysql_fetchone(f"SELECT * FROM `{PREG_TABLE}` WHERE id=%s", (pid,))
    if not p:
        return jsonify({"ok": False, "error": "No encontrada"}), 404

    eid, seccion, orden = p["evaluacion_id"], p["seccion"], p["orden"]

    if direccion == "up":
        otra = mysql_fetchone(f"""
            SELECT id, orden FROM `{PREG_TABLE}`
            WHERE evaluacion_id=%s AND seccion=%s AND orden<%s
            ORDER BY orden DESC LIMIT 1
        """, (eid, seccion, orden))
    else:
        otra = mysql_fetchone(f"""
            SELECT id, orden FROM `{PREG_TABLE}`
            WHERE evaluacion_id=%s AND seccion=%s AND orden>%s
            ORDER BY orden ASC LIMIT 1
        """, (eid, seccion, orden))

    if not otra:
        return jsonify({"ok": True, "moved": False})   # ya en el límite

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"UPDATE `{PREG_TABLE}` SET orden=%s WHERE id=%s", (otra["orden"], pid))
        cur.execute(f"UPDATE `{PREG_TABLE}` SET orden=%s WHERE id=%s", (orden, otra["id"]))
    conn.commit()
    return jsonify({"ok": True, "moved": True, "swapped_with": otra["id"]})


# ── API: Reordenar por drag-and-drop ─────────────────────────

@app.route("/api/evaluaciones/<int:eid>/reordenar", methods=["POST"])
@require_permission("edit")
def api_reordenar(eid):
    """Body JSON: {"seccion": "tecnica", "orden": [id1, id2, ...]}"""
    data    = request.get_json(force=True) or {}
    seccion = data.get("seccion", "")
    orden   = data.get("orden", [])

    if seccion not in SECCIONES:
        return jsonify({"ok": False, "error": "Sección inválida"}), 400

    conn = get_db()
    with conn.cursor() as cur:
        for i, preg_id in enumerate(orden, 1):
            cur.execute(f"""
                UPDATE `{PREG_TABLE}` SET orden=%s
                WHERE id=%s AND evaluacion_id=%s AND seccion=%s
            """, (i, preg_id, eid, seccion))
    conn.commit()
    return jsonify({"ok": True})


# ── API: Cambiar estado evaluación ───────────────────────────

@app.route("/api/evaluaciones/<int:eid>/estado", methods=["POST"])
@require_permission("edit")
def api_eval_estado(eid):
    ev = mysql_fetchone(f"SELECT * FROM `{EVAL_TABLE}` WHERE id=%s", (eid,))
    if not ev:
        return jsonify({"ok": False, "error": "No encontrada"}), 404

    nuevo = request.form.get("estado", "")
    transiciones = {
        "borrador":  ["publicada"],
        "publicada": ["borrador", "archivada"],
        "archivada": [],
    }
    if nuevo not in transiciones.get(ev["estado"], []):
        return jsonify({"ok": False,
                        "error": f"Transición inválida: {ev['estado']} → {nuevo}"}), 400

    if nuevo == "publicada":
        total = mysql_fetchone(
            f"SELECT COUNT(*) AS c FROM `{PREG_TABLE}` WHERE evaluacion_id=%s", (eid,)
        )
        if int(total["c"]) == 0:
            return jsonify({"ok": False,
                            "error": "Agrega al menos una pregunta antes de publicar"}), 400

    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"""
            UPDATE `{EVAL_TABLE}` SET estado=%s, updated_by=%s WHERE id=%s
        """, (nuevo, current_username(), eid))
    conn.commit()
    return jsonify({"ok": True, "estado": nuevo})


# ═══════════════════════════════════════════════════════════════
#  MÓDULO CUBICADOR
#  Busca documentos de venta en el ERP Random y cruza con
#  los datos de peso/volumen de la base de etiquetas.
#  Solo visible para rol: vendedor, admin, superadmin
# ═══════════════════════════════════════════════════════════════

TIPOS_DOC_CUBICADOR = [
    ("FCV", "Factura"),
    ("BLV", "Boleta"),
    ("GDV", "Guía de Despacho"),
    ("VD",  "Nota de Venta Directa"),
    ("WEB", "Nota de Venta Web"),
    ("COV", "Cotización"),
]

# VD y WEB usan TIDO=NVV en el ERP; el NUDO lleva el prefijo dentro (10 chars)
# NVI puede mapear a NVV también dependiendo del ERP
_ERP_TIDO_NUDO_MAP = {
    "VD":  ("NVV", lambda n: "VD"  + str(n).zfill(8)),
    "WEB": ("NVV", lambda n: "WEB" + str(n).zfill(7)),
}


def _normalize_phone_cl(raw: str) -> str:
    """
    Normaliza números de teléfono chilenos al formato +56XXXXXXXXX (12 chars).
    Casos soportados:
      +56936535760  →  +56936535760  (ya correcto)
      +569 3653 5760 → +56936535760  (espacios)
      56936535760    → +56936535760  (sin +)
      936535760      → +56936535760  (solo 9 dígitos, celular)
      9 3653 5760    → +56936535760  (empieza en 9 con espacios)
      236535760      → +56236535760  (fijo, 8 dígitos con código)
    """
    if not raw:
        return ""
    # Eliminar espacios, guiones, paréntesis, puntos
    p = re.sub(r'[\s\-\(\)\.]+', '', raw.strip())
    if not p:
        return ""
    # Ya tiene + al inicio
    if p.startswith('+'):
        digits = p[1:]
        # +56XXXXXXXXX → correcto
        if digits.startswith('56') and len(digits) == 11:
            return '+' + digits
        # +569XXXXXXXX (con 9 extra) → +56XXXXXXXXX
        if digits.startswith('569') and len(digits) == 12:
            return '+56' + digits[2:]
        # Cualquier otro +XX… → devolver tal cual
        return p
    # Empieza con 56 sin +
    if p.startswith('56') and len(p) == 11:
        return '+' + p
    # Empieza con 9 (celular nacional, 9 dígitos)
    if p.startswith('9') and len(p) == 9:
        return '+56' + p
    # Empieza con 9 y tiene 8 dígitos (sin el 9 de área)
    if p.startswith('9') and len(p) == 8:
        return '+569' + p
    # Número fijo de 8 dígitos (sin código ciudad)
    if p.isdigit() and len(p) == 8:
        return '+562' + p
    # Fallback: agregar +56 y entregar lo que hay
    return '+56' + p


# ── Mapa de códigos ERP (CIEN_región, CMEN_comuna) → nombre de comuna ──────
# CIEN: código de región de 3 dígitos (ej. "013"=RM, "005"=Valparaíso)
# CMEN: código de 3 chars (ej. "LOB"=Lo Barnechea, "VAL"=Valdivia/Valparaíso)
_CMEN_MAP: dict[str, dict[str, str]] = {
    # ── Región Metropolitana (013) — CÓDIGOS EXACTOS DEL ERP RANDOM ───
    # (verificados contra Sport and Health Solutions SPA Gestión Admin)
    "013": {
        "CEI":"Cerrillos",      "CER":"Cerro Navia",     "COL":"Colina",
        "CON":"Conchalí",       "CUR":"Curacaví",        "ELB":"El Bosque",
        "ELM":"El Monte",       "EST":"Estación Central","HUE":"Huechuraba",
        "IND":"Independencia",  "ISL":"Isla de Maipo",   "LAC":"La Cisterna",
        "LAF":"La Florida",     "LAG":"La Granja",       "LAP":"La Pintana",
        "LAR":"La Reina",       "LAM":"Lampa",           "LAS":"Las Condes",
        "LOB":"Lo Barnechea",   "LOE":"Lo Espejo",       "LOP":"Lo Prado",
        "MAC":"Macul",          "MAI":"Maipú",           "MAR":"María Pinto",
        "MEL":"Melipilla",      "PAD":"Padre Hurtado",   "PAI":"Paine",
        "PED":"Pedro Aguirre Cerda",
        "PEF":"Peñaflor",       "PEN":"Peñalolén",        "PEA":"Peñalolén",
        "PIR":"Pirque",         "PRO":"Providencia",     "PUD":"Pudahuel",
        "PUE":"Puente Alto",    "PUA":"Puente Alto",     "QUI":"Quilicura",
        "QNO":"Quinta Normal",  "REC":"Recoleta",        "REN":"Renca",
        "SBE":"San Bernardo",   "SJM":"San José de Maipo",
        "SMI":"San Miguel",     "SPE":"San Pedro",       "SJO":"San Joaquín",
        "SRA":"San Ramón",      "SAN":"Santiago",        "STG":"Santiago",
        "TAL":"Talagante",      "TIL":"Tiltil",          "VIT":"Vitacura",
        "ALH":"Alhué",          "BUI":"Buin",            "CTA":"Calera de Tango",
        "CAL":"Calera de Tango","NUN":"Ñuñoa",
    },
    # ── Valparaíso (005) ───────────────────────────────────────
    "005": {
        "VAL":"Valparaíso",  "VDM":"Viña del Mar", "CON":"Concón",      "QUI":"Quilpué",
        "VLA":"Villa Alemana","SAN":"San Antonio",  "QLL":"Quillota",    "LAC":"La Calera",
        "LAN":"Los Andes",   "SFE":"San Felipe",   "LIM":"Limache",     "OLM":"Olmué",
        "CAB":"Cabildo",     "LLI":"La Ligua",     "ZAP":"Zapallar",    "PAP":"Papudo",
        "QTE":"Quintero",    "PCU":"Puchuncaví",   "CAS":"Casablanca",  "SES":"San Esteban",
        "LLY":"Llaillay",    "PUT":"Putaendo",      "SMR":"Santa María", "ALG":"Algarrobo",
        "CTG":"Cartagena",   "SDO":"Santo Domingo", "EQU":"El Quisco",   "ETA":"El Tabo",
        "RIN":"Rinconada",   "CAL":"Calle Larga",  "JSF":"Juan Fernández","IPA":"Isla de Pascua",
    },
    # ── O'Higgins (006) ────────────────────────────────────────
    "006": {
        "RAN":"Rancagua",    "GRA":"Graneros",     "MOS":"Mostazal",    "COD":"Codegua",
        "OLI":"Olivar",      "COL":"Coltauco",     "DON":"Doñihue",     "REN":"Rengo",
        "REQ":"Requínoa",    "SFE":"San Fernando", "CHI":"Chimbarongo", "STA":"Santa Cruz",
        "NAN":"Nancagua",    "PAL":"Palmilla",      "PIC":"Pichilemu",   "LOL":"Lolol",
        "MAR":"Marchihue",   "PAR":"Paredones",    "SVC":"San Vicente", "LCA":"Las Cabras",
        "PEU":"Peumo",       "PID":"Pichidegua",   "MAL":"Malloa",      "MCL":"Machalí",
    },
    # ── Maule (007) ────────────────────────────────────────────
    "007": {
        "TAL":"Talca",       "CUR":"Curicó",        "LIN":"Linares",     "CON":"Constitución",
        "CAU":"Cauquenes",   "MOL":"Molina",        "TEN":"Teno",        "ROM":"Romeral",
        "HUA":"Hualañé",     "LIC":"Licantén",      "RAU":"Rauco",       "SCL":"San Clemente",
        "PEN":"Pencahue",    "MAU":"Maule",          "CUR":"Curepto",     "EMP":"Empedrado",
        "SJV":"San Javier",  "VLA":"Villa Alegre",  "YER":"Yerbas Buenas","COL":"Colbún",
        "LON":"Longaví",     "PAR":"Parral",         "RET":"Retiro",
    },
    # ── Biobío (008) ───────────────────────────────────────────
    "008": {
        "CON":"Concepción",  "TAL":"Talcahuano",   "HUA":"Hualpén",    "SAN":"San Pedro de la Paz",
        "COR":"Coronel",     "LOT":"Lota",          "TOM":"Tomé",        "PEN":"Penco",
        "CHI":"Chiguayante", "HUL":"Hualqui",       "SJU":"Santa Juana", "FLO":"Florida",
        "ARA":"Arauco",      "CAN":"Cañete",        "LEB":"Lebu",        "LOS":"Los Álamos",
        "CRN":"Curanilahue", "LAJ":"Laja",          "NAC":"Nacimiento",  "MUL":"Mulchén",
        "NEG":"Negrete",     "LAA":"Los Ángeles",   "YUM":"Yumbel",      "CAB":"Cabrero",
        "SRO":"San Rosendo", "NTL":"Nacimiento",
        "CHI":"Chillán",     "CHV":"Chillán Viejo", "BUL":"Bulnes",      "SCA":"San Carlos",
        "SFB":"San Fabián",  "SNN":"San Nicolás",   "NIH":"Ninhue",      "COE":"Coelemu",
        "PEM":"Pemuco",      "ELC":"El Carmen",     "PIN":"Pinto",       "COI":"Coihueco",
        "YUN":"Yungay",      "SIG":"San Ignacio",
    },
    # ── Araucanía (009) ────────────────────────────────────────
    "009": {
        "TEM":"Temuco",      "PDL":"Padre las Casas","VIL":"Villarrica", "PUC":"Pucón",
        "ANG":"Angol",       "VIC":"Victoria",       "LAU":"Lautaro",     "FRE":"Freire",
        "GOR":"Gorbea",      "LON":"Loncoche",       "CUR":"Curacautín", "LON":"Lonquimay",
        "MEL":"Melipeuco",   "CUN":"Cunco",          "VLC":"Vilcún",     "PER":"Perquenco",
        "GAL":"Galvarino",   "COL":"Collipulli",     "ERC":"Ercilla",    "PUR":"Purén",
        "TRA":"Traiguén",    "REN":"Renaico",        "PIT":"Pitrufquén", "TOL":"Toltén",
        "CAR":"Carahue",     "NEI":"Nueva Imperial", "CHO":"Cholchol",   "SAA":"Saavedra",
    },
    # ── Los Ríos (016) ─────────────────────────────────────────
    "016": {
        "VAL":"Valdivia",    "LUN":"La Unión",       "RBO":"Río Bueno",  "LRA":"Lago Ranco",
        "FUT":"Futrono",     "PAN":"Panguipulli",    "LLA":"Los Lagos",  "COR":"Corral",
        "MAR":"Mariquina",   "LAN":"Lanco",          "MAF":"Máfil",      "PAI":"Paillaco",
    },
    # ── Los Lagos (010) ────────────────────────────────────────
    "010": {
        "PMO":"Puerto Montt","PVA":"Puerto Varas",   "OSO":"Osorno",     "CAS":"Castro",
        "ANC":"Ancud",       "QUE":"Quellón",        "CAL":"Calbuco",    "MAU":"Maullín",
        "LMU":"Los Muermos", "FRU":"Frutillar",      "LLA":"Llanquihue", "PUR":"Purranque",
        "POC":"Puerto Octay","FRE":"Fresia",          "SPB":"San Pablo",  "PUY":"Puyehue",
        "RNE":"Río Negro",   "SJC":"San Juan de la Costa",
        "CHA":"Chaitén",     "FUL":"Futaleufú",      "PAL":"Palena",     "HUL":"Hualaihué",
    },
    # ── Aysén (011) ────────────────────────────────────────────
    "011": {
        "COY":"Coyhaique",   "PAY":"Puerto Aysén",   "CCH":"Chile Chico","COC":"Cochrane",
        "OHI":"O'Higgins",   "TOR":"Tortel",          "CIS":"Cisnes",     "LVE":"Lago Verde",
        "RIB":"Río Ibáñez",
    },
    # ── Magallanes (012) ───────────────────────────────────────
    "012": {
        "PUA":"Punta Arenas","PNA":"Puerto Natales", "POR":"Porvenir",   "PRI":"Primavera",
        "TIM":"Timaukel",    "LBL":"Laguna Blanca",  "RVE":"Río Verde",  "SGR":"San Gregorio",
        "CAH":"Cabo de Hornos",
    },
    # ── Tarapacá (001) ─────────────────────────────────────────
    "001": {
        "IQU":"Iquique",     "ALH":"Alto Hospicio",  "POZ":"Pozo Almonte","PIC":"Pica",
        "COL":"Colchane",    "CAM":"Camiña",          "HUA":"Huara",
    },
    # ── Arica y Parinacota (015) ───────────────────────────────
    "015": {
        "ARI":"Arica",       "CAM":"Camarones",      "PUT":"Putre",       "GLA":"General Lagos",
    },
    # ── Antofagasta (002) ──────────────────────────────────────
    "002": {
        "ANT":"Antofagasta", "CAL":"Calama",          "TOC":"Tocopilla",  "MEJ":"Mejillones",
        "TAL":"Taltal",       "SPA":"San Pedro de Atacama","OLL":"Ollagüe","MRE":"María Elena",
    },
    # ── Atacama (003) ──────────────────────────────────────────
    "003": {
        "COP":"Copiapó",     "CLD":"Caldera",         "CHA":"Chañaral",  "DIA":"Diego de Almagro",
        "VAL":"Vallenar",    "FRE":"Freirina",         "HUA":"Huasco",    "ALC":"Alto del Carmen",
        "TIA":"Tierra Amarilla",
    },
    # ── Coquimbo (004) ─────────────────────────────────────────
    "004": {
        "LSE":"La Serena",   "COQ":"Coquimbo",        "OVA":"Ovalle",    "ILL":"Illapel",
        "LVI":"Los Vilos",   "SAL":"Salamanca",        "CAN":"Canela",    "MPT":"Monte Patria",
        "PUN":"Punitaqui",   "VIC":"Vicuña",           "ANT":"Andacollo", "PAI":"Paihuano",
        "COM":"Combarbalá",  "LHG":"La Higuera",
    },
}


def _cmen_to_comuna(cien: str, cmen: str) -> str:
    """
    Convierte el código de región (CIEN) + código de comuna (CMEN) del ERP
    a nombre de comuna legible. Ej: ('016','VAL') → 'Valdivia'
    Si no hay match, retorna el CMEN tal cual (sirve como seed para autocomplete).
    """
    if not cmen:
        return ""
    region_map = _CMEN_MAP.get(str(cien).zfill(3), {})
    nombre = region_map.get(cmen.upper().strip())
    if nombre:
        return nombre
    # Fallback: buscar en todas las regiones (por si el CIEN viene mal)
    for rmap in _CMEN_MAP.values():
        if cmen.upper().strip() in rmap:
            return rmap[cmen.upper().strip()]
    return cmen   # retorna el código como seed


# Mapa CIEN → nombre de región
_REGION_NOMBRES: dict = {
    "001": "Tarapacá",          "002": "Antofagasta",         "003": "Atacama",
    "004": "Coquimbo",          "005": "Valparaíso",          "006": "O'Higgins",
    "007": "Maule",             "008": "Biobío",              "009": "Araucanía",
    "010": "Los Lagos",         "011": "Aysén",               "012": "Magallanes",
    "013": "Metropolitana",     "014": "Los Ríos",            "015": "Arica y Parinacota",
    "016": "Los Ríos",
}


def _nudo_variants(nudo_raw):
    """
    El ERP guarda NUDO como string de 10 chars con ceros a la izquierda.
    Devuelve lista de variantes a probar.

    Casos especiales:
    - NV WEB: prefijo "WEB" + 7 dígitos (ej. WEB0021756)
    - NV directa: prefijo "VD" + 8 dígitos (ej. VD00009344)
    - Otros docs: solo zfill a 10

    Si el NUDO ya tiene letras (prefijo), respetamos el prefijo y
    ajustamos solo el padding de los dígitos.
    """
    s = str(nudo_raw).strip()
    if not s:
        return []

    # Separar prefijo alfa (si existe) de la parte numérica
    import re as _re
    m = _re.match(r"^([A-Za-z]*)(\d+)$", s)
    if m:
        prefix, num = m.group(1).upper(), m.group(2)
        variants = [
            prefix + num.zfill(10 - len(prefix)),     # padding al total 10
            prefix + num.zfill(7),                     # WEB+7 dígitos (típico)
            prefix + num.zfill(8),                     # VD+8 dígitos
            prefix + num.zfill(6),
            prefix + num,                              # sin padding
            num.zfill(10),                             # sin prefijo, padded
            num,                                       # tal cual
        ]
    else:
        # NUDO puramente numérico (factura, boleta, etc.)
        variants = [s.zfill(10), s, s.zfill(8), s.zfill(7)]

    return list(dict.fromkeys(variants))


def _erp_get(path, params, token, timeout=10):
    """
    GET a la REST API del ERP usando urllib (sin dependencias externas).
    Retorna el JSON decodificado o lanza excepción.
    """
    import urllib.request as _urlreq
    import urllib.parse   as _urlparse
    import json           as _json_mod

    url = ERP_CONFIG.get("api_url", "https://lab.random.cl/ilus").rstrip("/") + path
    qs  = _urlparse.urlencode(params)
    req = _urlreq.Request(
        f"{url}?{qs}",
        headers={"Authorization": f"Bearer {token}"},
    )
    with _urlreq.urlopen(req, timeout=timeout) as resp:
        return _json_mod.loads(resp.read().decode("utf-8"))


# ── CACHE SIMPLE EN MEMORIA para búsquedas de documentos ERP ─────────
# Evita refetch del mismo documento en ventanas cortas (cuando el
# usuario edita campos del cubicador o cambia de pestaña).
_ERP_DOC_CACHE = {}        # { 'TIDO|NUDO': (timestamp, header, lineas) }
_ERP_DOC_CACHE_TTL = 90    # segundos
_ERP_ENT_CACHE = {}        # { 'RUT': (timestamp, entidad) }
_ERP_ENT_CACHE_TTL = 300   # las entidades cambian poco — 5 min

# Cache del catálogo de couriers + tarifas (cambia poco, ~5 min)
_COURIERS_CACHE = {"ts": 0, "data": None}
_COURIERS_CACHE_TTL = 300

def _get_couriers_cached():
    """Devuelve la lista de couriers activos con sus tarifas concatenadas.
    Cacheado 5 min para evitar re-ejecutar el JOIN+GROUP_CONCAT en cada cotización.
    """
    import time as _time
    now = _time.time()
    if _COURIERS_CACHE["data"] is not None and (now - _COURIERS_CACHE["ts"]) < _COURIERS_CACHE_TTL:
        return _COURIERS_CACHE["data"]
    rows = mysql_fetchall(
        """SELECT c.*,
               GROUP_CONCAT(
                   CONCAT_WS('|', t.zona, t.peso_desde, t.peso_hasta,
                             t.precio_base, t.precio_kg_extra)
                   ORDER BY t.zona, t.peso_desde SEPARATOR ';;'
               ) AS tarifas_raw
           FROM transport_couriers c
           LEFT JOIN transport_courier_tarifas t ON t.courier_id = c.id AND t.activo = 1
           WHERE c.activo = 1
           GROUP BY c.id
           ORDER BY c.nombre""",
        (),
    ) or []
    data = [dict(r) for r in rows]
    _COURIERS_CACHE["ts"] = now
    _COURIERS_CACHE["data"] = data
    return data

def _invalidate_couriers_cache():
    """Llamar después de editar couriers/tarifas para refrescar el cache."""
    _COURIERS_CACHE["ts"] = 0
    _COURIERS_CACHE["data"] = None


# ── Endpoint DEBUG: muestra el RAW del ERP para diagnosticar campos ──
@app.route("/api/erp/documento-raw", methods=["GET", "POST"])
@login_required
def erp_documento_raw():
    """Devuelve el JSON crudo del ERP Random para diagnosticar qué campos
    vienen en cada tipo de documento. Útil cuando un documento (ej. 10599
    de NV) no muestra cliente/dirección — aquí se ve si el problema es
    falta de campos o nombres distintos."""
    if not (g.permissions.get("admin") or g.permissions.get("superadmin")):
        return jsonify({"error":"Solo admin/superadmin"}), 403
    if request.method == "POST":
        d = request.get_json(silent=True) or {}
        tido = (d.get("tido") or "").strip().upper()
        nudo = (d.get("nudo") or "").strip()
    else:
        tido = (request.args.get("tido") or "").strip().upper()
        nudo = (request.args.get("nudo") or "").strip()
    if not tido or not nudo:
        return jsonify({"error":"tido y nudo son obligatorios"}), 400

    # Mapear VD/WEB → NVV con prefijo
    if tido in _ERP_TIDO_NUDO_MAP:
        erp_tido, nudo_fn = _ERP_TIDO_NUDO_MAP[tido]
        erp_nudo = nudo_fn(nudo)
    else:
        erp_tido = tido
        erp_nudo = nudo

    nudos = _nudo_variants(erp_nudo)
    TOKEN = ERP_CONFIG.get("api_token", "")
    raw_body = None
    used_nudo = ""
    for nv in nudos:
        try:
            raw_body = _erp_get(
                "/documentos/render",
                {"tido": erp_tido, "nudo": nv, "empresa": "01"},
                TOKEN, timeout=10,
            )
            data = raw_body.get("data") or []
            if data:
                used_nudo = nv
                break
        except Exception as exc:
            return jsonify({"error": f"ERP no respondió: {exc}"}), 503

    if not raw_body or not raw_body.get("data"):
        return jsonify({"error": "Documento no encontrado"}), 404

    raw_header = raw_body["data"][0].get("maeedo") or {}
    raw_lineas = raw_body["data"][0].get("maeddo") or []

    return jsonify({
        "tido_buscado":    tido,
        "tido_erp":        erp_tido,
        "nudo_buscado":    nudo,
        "nudo_usado":      used_nudo,
        "header_keys":     sorted(raw_header.keys()),
        "header":          raw_header,
        "n_lineas":        len(raw_lineas),
        "primera_linea":   raw_lineas[0] if raw_lineas else None,
    })


# ════════════════════════════════════════════════════════════════════
#  MOTOR ÚNICO DE BÚSQUEDA DE DOCUMENTOS ERP — REUTILIZABLE
# ════════════════════════════════════════════════════════════════════
# Este endpoint es el ÚNICO punto de entrada para buscar documentos del
# ERP Random desde cualquier módulo de la app:
#   - Cubicador / Asignar y Cotizar (POST /api/erp/documento)
#   - Mantenciones — repuestos
#   - Retiros — validación de documentación
#   - Módulos futuros (cobranzas, devoluciones, etc.)
#
# Devuelve estructura UNIFICADA con todos los datos del cliente extraídos
# desde HEADER + LÍNEAS + /entidades (con variantes RUT en paralelo).
# Códigos de comuna se resuelven automáticamente a nombres.
#
# Uso desde JS:
#   const r = await fetch('/api/erp/documento', {
#     method: 'POST',
#     headers: {'Content-Type':'application/json'},
#     body: JSON.stringify({ tido: 'FCV', nudo: '10599' })
#   });
#   const data = await r.json();
#   // data.hdr.cliente_nombre, data.hdr.email, data.hdr.comuna, etc.
#   // data.lineas[0].sku, data.lineas[0].nombre, etc.
#
# Tipos de documento soportados:
#   FCV (Factura), BLV (Boleta), VD (NV directa), WEB (NV web),
#   NVI (NV interna), GDV (Guía despacho), GDI (Guía int), GTI (Traspaso)
#
# El endpoint cachea respuestas 90 segundos para evitar refetch en
# llamadas repetidas con los mismos parámetros.
# ────────────────────────────────────────────────────────────────────

def _get_build_version():
    """Devuelve commit hash + timestamp del HEAD actual de git.
    Si git no está disponible (Railway puede no tenerlo), usa el mtime
    de app.py como fallback. Esto permite verificar visualmente desde
    el navegador si Railway ya desplegó el commit más reciente.
    """
    import subprocess
    try:
        sha = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=os.path.dirname(os.path.abspath(__file__)),
            stderr=subprocess.DEVNULL,
            timeout=2,
        ).decode().strip()
        ts = subprocess.check_output(
            ["git", "log", "-1", "--format=%ai", "HEAD"],
            cwd=os.path.dirname(os.path.abspath(__file__)),
            stderr=subprocess.DEVNULL,
            timeout=2,
        ).decode().strip()
        return sha, ts
    except Exception:
        try:
            mtime = os.path.getmtime(os.path.abspath(__file__))
            return "no-git", datetime.fromtimestamp(mtime).isoformat()
        except Exception:
            return "unknown", "unknown"


@app.route("/api/erp/peek", methods=["GET"])
@login_required
def erp_engine_peek():
    """Inspector simple del ERP: abrí en el navegador como
        /api/erp/peek?tido=FCV&nudo=10683
    y devuelve el JSON con:
      - lo que vio el motor (cliente_nombre, observaciones, comuna, etc.)
      - obs_source: de dónde extrajo la observación (o por qué no la tiene)
      - raw_header_keys: todos los campos del header del ERP
      - first_3_lines_obdo: OBDO de las primeras 3 líneas — para verificar
        si el ERP lo devuelve o no.
      - first_line_keys: todos los campos de la primera línea.

    Diseñado para ser leído desde el navegador sin DevTools.
    """
    if not g.permissions.get("cubicador"):
        return jsonify({"error": "Sin permiso"}), 403
    tido = (request.args.get("tido") or "FCV").strip().upper()
    nudo = (request.args.get("nudo") or "").strip()
    if not nudo:
        return jsonify({"error": "Falta parámetro 'nudo'. Ejemplo: /api/erp/peek?tido=FCV&nudo=10683"}), 400
    try:
        _ERP.invalidate_doc(tido, nudo)  # forzar refetch
        doc = _ERP.fetch_document(tido, nudo)
    except Exception as e:
        return jsonify({"error": str(e), "tido": tido, "nudo": nudo}), 500
    if not doc:
        return jsonify({"error": "Documento no encontrado", "tido": tido, "nudo": nudo}), 404

    # Inspeccionar las primeras 3 líneas crudas: ¿tienen OBDO?
    raw_lines = doc.get("lineas_raw") or []
    first_3 = []
    for ln in raw_lines[:3]:
        if not isinstance(ln, dict):
            continue
        first_3.append({
            "sku":     ln.get("KOPRCT") or ln.get("koprct"),
            "OBDO":    ln.get("OBDO") or ln.get("obdo") or "",
            "NOKOEN":  ln.get("NOKOEN") or ln.get("nokoen") or "",
            "DIEN":    ln.get("DIEN") or ln.get("dien") or "",
            "COMUNA":  ln.get("COMUNA") or ln.get("comuna") or "",
        })

    return jsonify({
        "tido": doc.get("tido"),
        "nudo": doc.get("nudo"),
        "encontrado": True,
        "extraido": {
            "cliente_nombre": doc.get("cliente_nombre"),
            "cliente_rut":    doc.get("cliente_rut"),
            "email":          doc.get("email"),
            "telefono":       doc.get("telefono"),
            "direccion":      doc.get("direccion"),
            "comuna":         doc.get("comuna"),
            "observaciones":  doc.get("observaciones"),
        },
        "diagnostico": doc.get("diagnostics"),
        "primeras_3_lineas": first_3,
        "header_keys_disponibles": sorted((doc.get("raw_sample") or {}).keys()),
        "primera_linea_keys_disponibles": sorted((doc.get("raw_linea_sample") or {}).keys()),
        # ★ NUEVO: contenido de MAEEDOOB (tabla de observaciones del doc)
        "maeedoob_campos": doc.get("raw_obs_sample", {}),
        "totales": {
            "neto":  doc.get("valor_neto"),
            "iva":   doc.get("valor_iva"),
            "bruto": doc.get("valor_bruto"),
        },
        "n_lineas": doc.get("n_lineas"),
    })


@app.route("/api/erp/health", methods=["GET"])
def erp_engine_health():
    """Health-check del motor ERP. Útil para verificar en producción si:
      - El módulo erp_engine está cargado
      - El cliente está inicializado
      - Hay token configurado
      - El caché está funcionando
      - Qué commit/versión está corriendo (vs lo que esperás)
    NO requiere login para diagnosticar problemas de despliegue sin sesión.
    """
    sha, build_ts = _get_build_version()
    try:
        c = erp_engine.get_client()
        with c._lock:
            doc_size = len(c._doc_cache)
            ent_size = len(c._ent_cache)
        return jsonify({
            "ok": True,
            "engine_loaded": True,
            "base_url": c.base_url,
            "has_token": bool(c.token),
            "token_length": len(c.token or ""),
            "doc_ttl": c.doc_ttl,
            "ent_ttl": c.ent_ttl,
            "timeout": c.timeout,
            "retries": c.retries,
            "doc_cache_size": doc_size,
            "ent_cache_size": ent_size,
            "build_commit": sha,
            "build_timestamp": build_ts,
            "server_time": datetime.now().isoformat(),
        })
    except Exception as e:
        return jsonify({
            "ok": False,
            "engine_loaded": False,
            "error": str(e),
            "build_commit": sha,
            "build_timestamp": build_ts,
        }), 500


@app.route("/api/erp/diagnose", methods=["POST"])
@login_required
def erp_engine_diagnose():
    """Diagnóstico profundo de una búsqueda en el ERP.

    POST JSON: { tido, nudo }
    Devuelve TODO lo que el motor intentó: variantes de NUDO probadas,
    variantes de RUT, latencia de cada llamada, fallback chain. Sin cruzar
    con BD local. Útil para depurar cuando un documento "no aparece".
    """
    if not g.permissions.get("cubicador"):
        return jsonify({"error": "Sin permiso"}), 403
    d = request.get_json(silent=True) or {}
    tido = (d.get("tido") or "").strip().upper()
    nudo = (d.get("nudo") or "").strip()
    if not tido or not nudo:
        return jsonify({"error": "tido y nudo son obligatorios"}), 400
    try:
        # Invalidar la cache del motor para forzar refetch real
        _ERP.invalidate_doc(tido, nudo)
        doc = _ERP.fetch_document(tido, nudo)
        if not doc:
            return jsonify({
                "ok": False,
                "error": "Documento no encontrado",
                "nudo_variants_tried": erp_engine.nudo_variants(
                    erp_engine.TIDO_NUDO_MAP[tido][1](nudo)
                    if tido in erp_engine.TIDO_NUDO_MAP else nudo
                ),
            }), 404
        return jsonify({
            "ok": True,
            "encontrado": True,
            "cliente_nombre": doc.get("cliente_nombre"),
            "cliente_rut": doc.get("cliente_rut"),
            "email": doc.get("email"),
            "telefono": doc.get("telefono"),
            "direccion": doc.get("direccion"),
            "comuna": doc.get("comuna"),
            "observaciones": doc.get("observaciones"),
            "n_lineas": doc.get("n_lineas"),
            "datos_completos": doc.get("datos_completos"),
            "diagnostics": doc.get("diagnostics"),
            "raw_sample": doc.get("raw_sample"),
            "raw_linea_sample": doc.get("raw_linea_sample"),
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/erp/documento", methods=["GET","POST"])
@login_required
def erp_documento_unificado():
    """Busca un documento ERP por TIDO+NUDO. Devuelve cabecera + líneas normalizadas.

    Query params o JSON: tido, nudo
    Response: {hdr: {nrazon, nruc, comuna, direccion, ...}, lineas: [{sku, nombre, cantidad, ...}]}

    Pensado como motor único de búsqueda para todas las pantallas que necesiten
    importar datos de un documento del ERP (factura, boleta, NV, cotización).
    """
    if request.method == "POST":
        d = request.get_json(silent=True) or {}
        tido = (d.get("tido") or "").strip().upper()
        nudo = (d.get("nudo") or "").strip()
    else:
        tido = (request.args.get("tido") or "").strip().upper()
        nudo = (request.args.get("nudo") or "").strip()
    if not tido or not nudo:
        return jsonify({"error":"tido y nudo son obligatorios"}), 400
    try:
        hdr, lineas = _cubicador_fetch(tido, nudo)
    except Exception as exc:
        return jsonify({"error": f"ERP no respondió: {exc}"}), 503
    if not hdr:
        return jsonify({"error":"Documento no encontrado en ERP", "tido":tido, "nudo":nudo}), 404

    # Como _cubicador_fetch YA hace todo el trabajo pesado (extracción
    # desde header + líneas + /entidades + resolución de comuna), aquí
    # solo pasamos esos campos al response. Mantiene compatibilidad
    # con código viejo (que usa razon_social/rut/etc) Y con código nuevo
    # (que usa cliente_nombre/cliente_rut/email/etc).
    h = {
        "tido":         tido,
        "nudo":         nudo,
        # Nombres "nuevos" (preferidos)
        "cliente_nombre": hdr.get("cliente_nombre", ""),
        "cliente_rut":    hdr.get("cliente_rut", ""),
        "email":          hdr.get("email", ""),
        "telefono":       hdr.get("telefono", ""),
        "direccion":      hdr.get("direccion", ""),
        "comuna":         hdr.get("comuna", ""),
        "observaciones":  hdr.get("observaciones", ""),
        "fecha":          hdr.get("fecha", ""),
        "valor_neto":     hdr.get("valor_neto", 0),
        "valor_bruto":    hdr.get("valor_bruto", 0),
        "valor_iva":      hdr.get("valor_iva", 0),
        # Aliases para compatibilidad con código antiguo
        "razon_social":   hdr.get("cliente_nombre", ""),
        "rut":            hdr.get("cliente_rut", ""),
        "obs":            hdr.get("observaciones", ""),
        # Diagnóstico para el frontend
        "raw_sample":     hdr.get("raw_sample", {}),
        "datos_completos": hdr.get("datos_completos", False),
        "all_fields":     hdr.get("all_fields", []),
    }

    # Normalizar líneas (formato compacto pero útil)
    out_lineas = []
    for ln in (lineas or []):
        out_lineas.append({
            "sku":         ln.get("sku", ""),
            "nombre":      ln.get("descripcion_erp") or ln.get("nombre_app", ""),
            "cantidad":    float(ln.get("cantidad") or 0),
            "peso_kg_u":   float(ln.get("peso_kg_u") or 0),
            "peso_vol_u":  float(ln.get("peso_vol_u") or 0),
            "vol_u":       float(ln.get("vol_u") or 0),
            "tiene_bultos": bool(ln.get("tiene_bultos")),
            "es_zz":       bool(ln.get("es_zz")),
            "vaneli":      float(ln.get("vaneli") or 0),
        })

    return jsonify({"hdr": h, "lineas": out_lineas})


def _cubicador_fetch(tido, nudo):
    """
    SHIM de compatibilidad sobre erp_engine.ERPClient.

    Delegamos toda la lógica ERP al motor unificado (erp_engine.py) y aquí
    solo hacemos:
      1. Llamar al motor → obtener documento + líneas crudas
      2. Cruzar SKUs con nuestra BD local en UNA query batch (peso/volumen/bultos)
      3. Devolver (header_dict, lineas_list) en el formato legacy que esperan
         los templates y los demás callers.

    El motor maneja: variantes NUDO, variantes RUT en paralelo, fallback a
    /entidades por nombre cuando RUT no resuelve, caché, retry, logger,
    resolución de comuna desde código, normalización de teléfono, etc.

    Cualquier módulo nuevo (retiros, mantenciones, etc.) debe llamar
    directamente a `erp_engine.get_client().fetch_document(tido, nudo)` y
    no a esta función — esto es solo para mantener la API legacy.
    """
    # 1. Llamar al motor unificado
    try:
        doc = _ERP.fetch_document(tido, nudo)
    except Exception as e:
        # Re-raise para que el caller maneje (ConnectionError típico)
        raise

    if not doc:
        return None, []

    raw_lineas = doc.get("lineas_raw") or []

    # 2. Construir header legacy desde el doc del motor (compat con templates)
    header = {
        "tido":             doc.get("tido"),
        "nudo":             doc.get("nudo"),
        "nudo_display":     doc.get("nudo_display"),
        "fecha":            doc.get("fecha"),
        "valor_neto":       doc.get("valor_neto"),
        "valor_iva":        doc.get("valor_iva"),
        "valor_bruto":      doc.get("valor_bruto"),
        "cliente_nombre":   doc.get("cliente_nombre"),
        "cliente_rut":      doc.get("cliente_rut"),
        # Campos de contacto enriquecidos desde /entidades (motor)
        "email":            doc.get("email"),
        "telefono":         doc.get("telefono"),
        "direccion":        doc.get("direccion"),
        "comuna":           doc.get("comuna"),
        "observaciones":    doc.get("observaciones"),
        # Tipo de operación derivado de SKUs ZZ del documento
        "tipo_operacion":   doc.get("tipo_operacion", ""),
        "tipo_codigo":      doc.get("tipo_codigo", ""),
        # Diagnóstico para el frontend (botón "Diagnosticar ERP")
        "all_fields":       doc.get("all_fields", []),
        "raw_sample":       doc.get("raw_sample", {}),
        "raw_linea_sample": doc.get("raw_linea_sample", {}),
        "n_lineas":         doc.get("n_lineas", 0),
        "datos_completos":  doc.get("datos_completos", False),
        "diagnostics":      doc.get("diagnostics", {}),  # nudo_tried, rut_tried, latency_ms
    }

    # ── Cruzar líneas con BD local (bultos/peso) — BATCH SQL ───────────
    # OPTIMIZACIÓN: una sola query con WHERE sku IN (...) para todos los
    # SKUs del documento. Antes hacíamos N queries seriadas (lento en docs
    # con muchos productos).
    skus_set = set()
    for l in raw_lineas:
        s = (l.get("KOPRCT") or "").strip().upper()
        if s:
            skus_set.add(s)

    sku_data_map = {}
    if skus_set:
        skus_list = list(skus_set)
        ph = ",".join(["%s"] * len(skus_list))
        rows_sku = mysql_fetchall(f"""
            SELECT
                UPPER(TRIM(p.sku))                                     AS sku_norm,
                p.id                                                   AS app_id,
                p.nombre                                               AS nombre_app,
                COUNT(DISTINCT b.id)                                   AS total_bultos,
                COALESCE(SUM(b.peso), 0)                               AS peso_total,
                COALESCE(SUM(b.largo * b.ancho * b.alto), 0)           AS volumen_cm3,
                ROUND(COALESCE(SUM(b.largo * b.ancho * b.alto) / 4000.0, 0), 4)
                                                                       AS peso_vol
            FROM `{PRODUCTS_TABLE}` p
            LEFT JOIN `{BULTOS_TABLE}` b ON b.product_id = p.id
            WHERE UPPER(TRIM(p.sku)) IN ({ph})
            GROUP BY p.id, p.nombre
        """, tuple(skus_list)) or []
        for r in rows_sku:
            sku_data_map[r["sku_norm"]] = dict(r)

    ZZ_CODES = {"ZZENVIO","ZZINGREPUESTO","ZZSERVTEC","ZZRETIRO","ZZINSTALACION","ZZINGARREQUIP"}
    lineas = []
    for l in raw_lineas:
        sku         = (l.get("KOPRCT") or "").strip().upper()
        descripcion = (l.get("NOKOPR") or "").strip()
        qty          = float(l.get("CAPRCO1") or 0)
        qty_desp     = float(l.get("CAPRAD1") or 0)
        saldo_linea  = max(qty - qty_desp, 0)
        es_zz        = sku in ZZ_CODES

        if not sku:
            continue

        app_data = sku_data_map.get(sku)
        tiene_ficha  = app_data is not None
        total_bultos = int(app_data["total_bultos"]) if tiene_ficha else 0
        tiene_bultos = tiene_ficha and float(app_data.get("volumen_cm3") or 0) > 0

        peso_kg_u  = float(app_data["peso_total"])  if tiene_ficha else 0
        peso_vol_u = float(app_data["peso_vol"])    if tiene_ficha else 0
        vol_u      = float(app_data["volumen_cm3"]) if tiene_ficha else 0
        pred_u     = max(peso_kg_u, peso_vol_u)

        nombre_app = (app_data["nombre_app"] if tiene_ficha else "") or ""
        diferencia = tiene_ficha and nombre_app.strip().upper() != descripcion.upper()

        lineas.append({
            "sku":              sku,
            "descripcion_erp":  descripcion,
            "nombre_app":       nombre_app,
            "cantidad":         qty,
            "total_bultos":     total_bultos,
            "app_id":           app_data["app_id"] if tiene_ficha else None,
            "tiene_ficha":      tiene_ficha,
            "tiene_bultos":     tiene_bultos,
            "peso_kg_u":        round(peso_kg_u,  4),
            "peso_vol_u":       round(peso_vol_u, 4),
            "vol_u":            round(vol_u,      2),
            "pred_u":           round(pred_u,     4),
            "peso_kg_tot":      round(peso_kg_u  * qty, 4),
            "peso_vol_tot":     round(peso_vol_u * qty, 4),
            "vol_tot":          round(vol_u      * qty, 2),
            "pred_tot":         round(pred_u     * qty, 4),
            "diferencia":       diferencia,
            "cantidad_despachada": qty_desp,
            "saldo":               saldo_linea,
            "es_zz":               es_zz,
            "vaneli":              float(l.get("VANELI") or 0),
        })

    # NOTA: La caché del documento la maneja erp_engine.ERPClient internamente.
    # _ERP_DOC_CACHE (legacy) se mantiene como dict vacío para compat con
    # cualquier código que pudiera consultarlo, pero no se usa aquí.

    return header, lineas


def _parse_docs_from_form(form):
    """Parse tido_0/nudo_0, tido_1/nudo_1, … from a form. Returns [(tido, nudo), …]."""
    import re as _re
    docs = []
    i = 0
    while i < 20:   # safety cap
        tido_i = (form.get(f"tido_{i}") or "").strip().upper()
        nudo_i = (form.get(f"nudo_{i}") or "").strip()
        if tido_i == "" and nudo_i == "":
            break
        if nudo_i:
            for _cod, _ in TIPOS_DOC_CUBICADOR:
                _m = _re.match(r"^" + _cod + r"\s*(\S+)$", nudo_i.upper())
                if _m:
                    tido_i = _cod
                    nudo_i = _m.group(1)
                    break
            docs.append((tido_i or "FCV", nudo_i))
        i += 1
    return docs


def _fetch_multi_docs(docs):
    """Fetch each doc and merge lines by SKU. Returns (headers, merged_lineas, errors)."""
    headers = []
    merged  = {}      # sku → line dict
    errors  = []

    for tido_i, nudo_i in docs:
        try:
            hdr, lineas = _cubicador_fetch(tido_i, nudo_i)
            if hdr is None:
                errors.append(f"No se encontró {tido_i} N° {nudo_i} en el ERP.")
            else:
                headers.append(hdr)
                for l in lineas:
                    sku = l["sku"]
                    if sku in merged:
                        m = merged[sku]
                        m["cantidad"]     += l["cantidad"]
                        m["peso_kg_tot"]  += l["peso_kg_tot"]
                        m["peso_vol_tot"] += l["peso_vol_tot"]
                        m["vol_tot"]      += l["vol_tot"]
                        m["pred_tot"]     += l["pred_tot"]
                    else:
                        merged[sku] = dict(l)
        except ConnectionError as ce:
            errors.append(str(ce))
        except Exception as ex:
            errors.append(f"Error {tido_i} N° {nudo_i}: {ex}")

    return headers, list(merged.values()), errors


def _cubicador_num(value, default=0.0):
    """Return a float for Excel/JSON, even when values arrive as Decimal or CLP text."""
    if value is None:
        return default
    try:
        return float(value)
    except Exception:
        try:
            txt = str(value).strip().replace("$", "").replace(" ", "")
            if "," in txt and "." in txt:
                txt = txt.replace(".", "").replace(",", ".")
            elif "," in txt:
                txt = txt.replace(",", ".")
            return float(txt)
        except Exception:
            return default


def _cubicador_export_payload(headers, lineas, docs):
    """JSON-safe snapshot of the current cubicador table for fast multidoc exports.

    Incluye TODOS los datos del cliente (dirección, comuna, email, teléfono,
    observaciones, ZZ envío) para generar informes PDF/Excel completos.
    """
    # Calcular ZZ envío por documento
    zz_por_doc = {}
    for l in (lineas or []):
        if (l.get("sku") or "").upper() == "ZZENVIO":
            key = f"{l.get('_doc_tido','')}-{l.get('_doc_nudo','')}"
            zz_por_doc[key] = _cubicador_num(l.get("vaneli"))

    return {
        "docs": [[str(t), str(n)] for t, n in (docs or [])],
        "headers": [
            {
                "tido": str(h.get("tido", "")),
                "nudo": str(h.get("nudo", "")),
                "nudo_display": str(h.get("nudo_display", h.get("nudo", ""))),
                "fecha": str(h.get("fecha", "")),
                "cliente_nombre": str(h.get("cliente_nombre", "")),
                "cliente_rut": str(h.get("cliente_rut", "")),
                # ★★★ Datos de contacto + dirección + comuna + obs ★★★
                "email":          str(h.get("email", "")),
                "telefono":       str(h.get("telefono", "")),
                "direccion":      str(h.get("direccion", "")),
                "comuna":         str(h.get("comuna", "")),
                "observaciones":  str(h.get("observaciones", "")),
                "zzenvio":        zz_por_doc.get(f"{h.get('tido','')}-{h.get('nudo','')}", 0),
                "valor_neto": _cubicador_num(h.get("valor_neto")),
                "valor_iva": _cubicador_num(h.get("valor_iva")),
                "valor_bruto": _cubicador_num(h.get("valor_bruto")),
            }
            for h in (headers or [])
        ],
        "lineas": [
            {
                "sku": str(l.get("sku", "")),
                "descripcion_erp": str(l.get("descripcion_erp", "")),
                "cantidad": _cubicador_num(l.get("cantidad")),
                "total_bultos": int(_cubicador_num(l.get("total_bultos"))),
                "tiene_ficha": bool(l.get("tiene_ficha")),
                "tiene_bultos": bool(l.get("tiene_bultos")),
                "peso_kg_u": _cubicador_num(l.get("peso_kg_u")),
                "peso_vol_u": _cubicador_num(l.get("peso_vol_u")),
                "vol_u": _cubicador_num(l.get("vol_u")),
                "pred_u": _cubicador_num(l.get("pred_u")),
                "pred_tot": _cubicador_num(l.get("pred_tot")),
                "peso_kg_tot": _cubicador_num(l.get("peso_kg_tot")),
                "peso_vol_tot": _cubicador_num(l.get("peso_vol_tot")),
                "vol_tot": _cubicador_num(l.get("vol_tot")),
            }
            for l in (lineas or [])
        ],
    }


@app.route("/cubicador", methods=["GET", "POST"])
@login_required
def cubicador():
    if not g.permissions.get("cubicador"):
        flash("No tienes acceso al módulo Cubicador.", "danger")
        return redirect(url_for("index"))

    docs      = []
    resultado = None
    error_msg = None

    if request.method == "POST":
        docs = _parse_docs_from_form(request.form)
    else:
        # GET: backward compat (?tido=FCV&nudo=9344)
        _tido = (request.args.get("tido") or "FCV").strip().upper()
        _nudo = (request.args.get("nudo") or "").strip()
        if _nudo:
            docs = [(_tido, _nudo)]

    if docs:
        headers, lineas, errors = _fetch_multi_docs(docs)
        if errors:
            error_msg = " · ".join(errors)
        if headers:
            resultado = {
                "headers": headers,
                "lineas":  lineas,
                "docs":    docs,
                "multi":   len(docs) > 1,
                "export_payload": _cubicador_export_payload(headers, lineas, docs),
            }

    resp = make_response(render_template(
        "cubicador/index.html",
        tipos_doc=TIPOS_DOC_CUBICADOR,
        docs=docs,
        resultado=resultado,
        error_msg=error_msg,
    ))
    # Anti-caché agresivo: nunca servir HTML viejo cacheado por el navegador.
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/cubicador/export/excel", methods=["POST"])
@login_required
def cubicador_export_excel():
    """Descarga el resultado del cubicador como Excel (.xlsx) — soporta múltiples documentos."""
    if not g.permissions.get("cubicador"):
        flash("No tienes acceso al módulo Cubicador.", "danger")
        return redirect(url_for("index"))

    docs = _parse_docs_from_form(request.form)
    if not docs:
        flash("Debes ingresar al menos un documento.", "warning")
        return redirect(url_for("cubicador"))

    payload_raw = (request.form.get("payload_json") or "").strip()
    headers = lineas = errors = None
    if payload_raw:
        try:
            payload = json.loads(payload_raw)
            payload_docs = payload.get("docs") or []
            docs = [
                (str(d[0]).strip().upper(), str(d[1]).strip())
                for d in payload_docs
                if isinstance(d, (list, tuple)) and len(d) >= 2 and str(d[1]).strip()
            ] or docs
            parsed_headers = payload.get("headers") or []
            if parsed_headers:
                headers = parsed_headers
                lineas = payload.get("lineas") or []
                errors = []
            else:
                headers = lineas = errors = None
        except Exception:
            headers = lineas = errors = None

    if headers is None:
        try:
            headers, lineas, errors = _fetch_multi_docs(docs)
        except Exception as ex:
            flash(f"Error al consultar el ERP: {ex}", "danger")
            return redirect(url_for("cubicador"))

    if not headers:
        flash(" · ".join(errors) if errors else "No se encontraron documentos.", "warning")
        return redirect(url_for("cubicador"))

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from io import BytesIO

    wb = openpyxl.Workbook()
    ws = wb.active

    BLACK, RED, LGRAY = "1A1A1A", "CC0000", "F5F5F5"

    def _hdr_cell(cell, val):
        cell.value = val
        cell.font = Font(bold=True, color="FFFFFF", size=9)
        cell.fill = PatternFill("solid", fgColor=BLACK)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def _money_excel(value):
        return f"${_cubicador_num(value):,.0f}"

    # ── Fila 1: título ───────────────────────────────────────────────
    if len(docs) == 1:
        tido0, nudo0 = docs[0]
        ws.title    = f"{tido0}{nudo0}"[:31]
        title_text  = f"CUBICADOR ILUS  ·  {tido0} N° {nudo0}"
    else:
        ws.title   = "Cubicador Múltiple"
        title_text = f"CUBICADOR ILUS  ·  {len(docs)} documentos combinados"

    ws.merge_cells("A1:J1")
    c = ws["A1"]
    c.value = title_text
    c.font  = Font(bold=True, size=13, color="FFFFFF")
    c.fill  = PatternFill("solid", fgColor=BLACK)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # ── Filas 2…N: una por documento ────────────────────────────────
    row_offset = 2
    for hdr in headers:
        ws.merge_cells(f"A{row_offset}:D{row_offset}")
        ws[f"A{row_offset}"].value = (
            f"{hdr['tido']} N° {hdr.get('nudo_display', hdr.get('nudo',''))}  ·  "
            f"{hdr.get('cliente_nombre','—')}   RUT {hdr.get('cliente_rut','—')}   {hdr.get('fecha','')}"
        )
        ws[f"A{row_offset}"].font = Font(size=9)
        ws[f"A{row_offset}"].alignment = Alignment(horizontal="left", vertical="center")

        ws.merge_cells(f"E{row_offset}:J{row_offset}")
        ws[f"E{row_offset}"].value = (
            f"Neto: {_money_excel(hdr.get('valor_neto'))}   "
            f"IVA: {_money_excel(hdr.get('valor_iva'))}   "
            f"Bruto: {_money_excel(hdr.get('valor_bruto'))}"
        )
        ws[f"E{row_offset}"].font = Font(bold=True, size=9)
        ws[f"E{row_offset}"].alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row_offset].height = 18
        row_offset += 1

    # ── Fila de encabezados de columna ───────────────────────────────
    hdr_row = row_offset
    cols = ["SKU", "Descripción ERP", "Cant", "Bultos",
            "Kg/u", "PV/u", "Vol cm³/u", "Predom/u", "Total Predom", "Tipo"]
    for ci, h in enumerate(cols, 1):
        _hdr_cell(ws.cell(row=hdr_row, column=ci), h)
    ws.row_dimensions[hdr_row].height = 20
    row_offset += 1

    # ── Filas de datos ───────────────────────────────────────────────
    for ri, l in enumerate(lineas, row_offset):
        bg = LGRAY if ri % 2 == 0 else "FFFFFF"
        vals = [
            l["sku"],
            l["descripcion_erp"],
            int(l["cantidad"]),
            l["total_bultos"] if l["tiene_ficha"] else "s/f",
            l["peso_kg_u"]  if l["tiene_bultos"] else None,
            l["peso_vol_u"] if l["tiene_bultos"] else None,
            l["vol_u"]      if l["tiene_bultos"] else None,
            l["pred_u"]     if l["tiene_bultos"] else None,
            l["pred_tot"]   if l["tiene_bultos"] else None,
            ("kg" if l["peso_kg_u"] >= l["peso_vol_u"] else "pv") if l["tiene_bultos"] else None,
        ]
        for ci, val in enumerate(vals, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.fill = PatternFill("solid", fgColor=bg)
            cell.font  = Font(size=9, bold=(ci == 9), color=(RED if ci == 9 else "000000"))
            cell.alignment = Alignment(
                horizontal="center" if ci in (3, 4, 10) else ("right" if ci >= 5 else "left"),
                vertical="center",
            )
            if ci in (5, 6, 7, 8, 9) and val is not None:
                cell.number_format = "#,##0.0"

    # ── Fila de totales ──────────────────────────────────────────────
    tr = row_offset + len(lineas)
    ws.merge_cells(f"A{tr}:B{tr}")
    ws.cell(row=tr, column=1, value="TOTALES").font = Font(bold=True, color="FFFFFF", size=9)
    ws.cell(row=tr, column=1).fill = PatternFill("solid", fgColor=BLACK)
    ws.cell(row=tr, column=1).alignment = Alignment(horizontal="right")

    totales = {
        3: int(sum(l["cantidad"]     for l in lineas)),
        5: sum(l["peso_kg_tot"]      for l in lineas),
        6: sum(l["peso_vol_tot"]     for l in lineas),
        7: sum(l["vol_tot"]          for l in lineas),
        9: sum(l["pred_tot"]         for l in lineas),
    }
    for ci in range(1, 11):
        cell = ws.cell(row=tr, column=ci)
        cell.fill = PatternFill("solid", fgColor=BLACK)
        if ci in totales:
            cell.value = totales[ci]
            cell.font  = Font(bold=True, color=("CC0000" if ci == 9 else "FFFFFF"), size=9)
            cell.alignment = Alignment(horizontal="center" if ci == 3 else "right")
            if ci != 3:
                cell.number_format = "#,##0.0"

    # ── Anchos de columna ────────────────────────────────────────────
    for ci, w in enumerate([14, 42, 7, 8, 10, 10, 12, 12, 14, 7], 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    ws.freeze_panes = f"A{hdr_row + 1}"

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    fname = "cubicador_" + "_".join(f"{t}{n}" for t, n in docs) + ".xlsx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname[:80],
    )


def _cubicador_pdf_response_ilus(headers, lineas, docs):
    """PDF comercial ILUS para el cubicador."""
    import html as _html
    import os as _os
    from io import BytesIO as _BytesIO
    from datetime import datetime as _dt

    def _esc(value):
        return _html.escape("" if value is None else str(value))

    def _money(value):
        return "$" + fvol_filter(value or 0)

    logo_path = _os.path.join(_os.path.dirname(__file__), "static", "logo_pdf.txt")
    try:
        with open(logo_path) as f:
            logo_b64 = f.read().strip()
        logo_tag = f'<img src="data:image/png;base64,{logo_b64}" style="height:50px;display:block">'
    except Exception:
        logo_tag = '<div style="font-size:18pt;font-weight:900;color:#fff;line-height:1">ILUS</div>'

    fecha_gen = _dt.now().strftime("%d/%m/%Y %H:%M")
    generated_by = current_username() or "Daniel Aguilar"
    first = headers[0] if headers else {}
    fecha_doc = first.get("fecha", "")
    docs_count = len(headers)
    sku_count = len(lineas)
    if len(docs) == 1 and headers:
        main_title = f"{first.get('tido','')} Nro {first.get('nudo_display', first.get('nudo',''))}"
        report_type = "REPORTE DOCUMENTO"
    else:
        main_title = f"{len(docs)} documentos comerciales"
        report_type = "REPORTE MULTI-DOCUMENTO"

    total_neto = sum(float(h.get("valor_neto") or 0) for h in headers)
    total_iva = sum(float(h.get("valor_iva") or 0) for h in headers)
    total_bruto = sum(float(h.get("valor_bruto") or 0) for h in headers)
    tot_qty  = sum(l["cantidad"]     for l in lineas)
    tot_kg   = sum(l["peso_kg_tot"]  for l in lineas)
    tot_pv   = sum(l["peso_vol_tot"] for l in lineas)
    tot_vol  = sum(l["vol_tot"]      for l in lineas)
    tot_pred = sum(l["pred_tot"]     for l in lineas)
    tot_bult = sum(l["total_bultos"] for l in lineas)

    docs_rows = ""
    for hdr in headers:
        docs_rows += f"""
        <tr>
          <td><span class="pdf-icon">F</span> {_esc(hdr.get('tido'))}</td>
          <td class="mono">{_esc(hdr.get('nudo_display', hdr.get('nudo','')))}</td>
          <td>{_esc(hdr.get('fecha'))}</td>
          <td class="r">{_money(hdr.get('valor_neto'))}</td>
          <td class="r">{_money(hdr.get('valor_iva'))}</td>
          <td class="r strong">{_money(hdr.get('valor_bruto'))}</td>
          <td class="c">{sku_count if docs_count == 1 else "-"}</td>
          <td class="state">Vigente</td>
        </tr>"""

    rows_html = ""
    for l in lineas:
        sf = not l["tiene_bultos"]
        pred_type = "kg" if (l.get("peso_kg_u") or 0) >= (l.get("peso_vol_u") or 0) else "pv"
        pred_class = "red" if pred_type == "kg" else "green"
        doc_ref = l.get("doc_ref") or (str(first.get("tido", "")) + " " + str(first.get("nudo_display", first.get("nudo", ""))))
        rows_html += f"""
        <tr>
          <td class="mono">{_esc(l['sku'])}</td>
          <td>{_esc(l['descripcion_erp'])}</td>
          <td class="mono small">{_esc(doc_ref)}</td>
          <td class="c">{int(l['cantidad'])}</td>
          <td class="c">{l['total_bultos'] if l['tiene_ficha'] else 's/f'}</td>
          <td class="r">{'-' if sf else fkg_filter(l['peso_kg_u'])}</td>
          <td class="r">{'-' if sf else fkg_filter(l['peso_vol_u'])}</td>
          <td class="r">{'-' if sf else fvol_filter(l['vol_u'])}</td>
          <td class="r {pred_class}">{'-' if sf else fkg_filter(l['pred_u'])} <span>{'' if sf else pred_type}</span></td>
          <td class="r strong red">{'-' if sf else fkg_filter(l['pred_tot'])}</td>
        </tr>"""

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
@page{{size:A4;margin:10mm 9mm 11mm}}
*{{box-sizing:border-box}}
body{{margin:0;background:#f3f4f6;color:#151515;font-family:Arial,Helvetica,sans-serif;font-size:7.3pt;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
.report{{background:#fff;min-height:100vh;border-bottom:8px solid #e60000}}
.hero{{height:76px;background:#050505;color:#fff;display:flex;align-items:center;justify-content:space-between;padding:0 22px;position:relative;overflow:hidden}}
.hero:after{{content:"";position:absolute;left:255px;top:-8px;width:82px;height:92px;background:#e60000;transform:skewX(-34deg)}}
.hero-left{{position:relative;z-index:1;min-width:210px}}.hero-right{{text-align:right;position:relative;z-index:1;text-transform:uppercase;font-weight:900;font-size:11pt;letter-spacing:.35px;line-height:1.05}}
.hero-right .report-type{{display:block;color:#e60000;font-size:6.4pt;margin-top:5px}}
.hero-right .hero-meta{{display:block;color:#d1d5db;font-size:6pt;margin-top:6px;text-transform:none;font-weight:700;letter-spacing:0}}
.wrap{{padding:14px 16px 10px}}
.top-grid{{display:grid;grid-template-columns:1.08fr 1fr 1fr;gap:8px;margin-bottom:10px}}
.card{{border:1px solid #e5e7eb;border-radius:7px;background:#fff;padding:10px;min-height:88px}}
.card-title{{font-weight:900;text-transform:uppercase;font-size:6.4pt;color:#374151;margin-bottom:8px;display:flex;align-items:center;gap:5px}}
.mini-icon{{width:14px;height:14px;border-radius:3px;background:#ffe8e8;color:#e60000;display:inline-grid;place-items:center;font-size:7pt}}
.client-name{{font-size:9pt;font-weight:900;margin-bottom:6px}}
.line{{font-size:6.6pt;color:#374151;margin:3px 0}}
.ok{{display:inline-block;border-radius:9px;background:#e8f7ee;color:#0a8a38;padding:3px 6px;font-size:5.8pt;font-weight:800;margin-top:4px}}
.summary-row{{display:flex;justify-content:space-between;margin:4px 0;font-size:6.8pt}}
.summary-row b{{font-size:7.5pt}}.summary-row .danger{{color:#e60000}}
.section-title{{background:#111;color:#fff;border-left:5px solid #e60000;padding:7px 9px;border-radius:5px 5px 0 0;font-size:7.1pt;text-transform:uppercase;font-weight:900;margin-top:10px}}
table{{width:100%;border-collapse:collapse;background:#fff}}
th{{background:#f7f8fa;color:#111;font-size:5.8pt;text-transform:uppercase;padding:7px 5px;border:1px solid #edf0f3;text-align:left}}
td{{padding:6px 5px;border:1px solid #edf0f3;vertical-align:middle}}
tbody tr:nth-child(even){{background:#fafafa}}
.detail th{{background:#111;color:#fff;border-color:#111}}
.c{{text-align:center}}.r{{text-align:right}}.mono{{font-family:Consolas,monospace;font-weight:800}}.small{{font-size:5.9pt}}.strong{{font-weight:900}}.red{{color:#e60000}}.green{{color:#098a28}}.muted{{color:#6b7280}}
.pdf-icon{{display:inline-grid;place-items:center;width:13px;height:13px;background:#ffe8e8;color:#e60000;border-radius:3px;font-size:6pt;font-weight:900;margin-right:3px}}
.state{{color:#0a8a38;font-weight:800;text-align:center}}
.totals-grid{{display:grid;grid-template-columns:1.6fr 1fr;gap:8px;margin-top:10px}}
.metric-grid{{display:grid;grid-template-columns:repeat(5,1fr);gap:0;border:1px solid #e5e7eb;border-radius:7px;overflow:hidden;background:#fff}}
.metric{{padding:9px 7px;text-align:center;border-right:1px solid #e5e7eb}}.metric:last-child{{border-right:0}}
.metric-label{{font-size:5.6pt;color:#6b7280;text-transform:uppercase;font-weight:800;margin-bottom:4px}}
.metric-val{{font-size:11pt;font-weight:900}}
.pred-box{{margin-top:8px;background:#050505;color:#fff;border-radius:6px;padding:10px 12px;display:flex;align-items:center;justify-content:space-between}}
.pred-box .value{{font-size:14pt;color:#e60000;font-weight:900}}
.obs{{border:1px solid #e5e7eb;border-radius:7px;padding:10px;background:#fff;min-height:88px}}
.footer{{margin-top:10px;border-top:1px solid #e5e7eb;padding:10px 16px;display:grid;grid-template-columns:125px 1.2fr 1fr 55px;gap:12px;align-items:center;font-size:5.8pt;color:#6b7280;background:#fff}}
.footer-logo img{{height:34px}}.qr{{width:48px;height:48px;border:2px solid #e60000;display:grid;grid-template-columns:repeat(5,1fr);gap:2px;padding:4px}}
.qr i{{background:#111}}.qr i:nth-child(2n){{background:#fff}}.qr i:nth-child(3n){{background:#e60000}}
</style></head><body>
<div class="report">
  <div class="hero">
    <div class="hero-left">{logo_tag}</div>
    <div class="hero-right">Sistema ILUS ERP<span class="report-type">Modulo Transporte Cubicador</span><span class="hero-meta">Documento generado por {_esc(generated_by)}</span></div>
  </div>
  <div class="wrap">
    <div class="top-grid">
      <div class="card">
        <div class="card-title"><span class="mini-icon">C</span>Cliente</div>
        <div class="client-name">{_esc(first.get('cliente_nombre') or 'Cliente no informado')}</div>
        <div class="line">RUT: {_esc(first.get('cliente_rut') or '-')}</div>
        <div class="line">Documento principal: {_esc(main_title)}</div>
        <div class="ok">Cliente sincronizado</div>
      </div>
      <div class="card">
        <div class="card-title"><span class="mini-icon">R</span>Resumen general</div>
        <div class="summary-row"><span>Total documentos:</span><b>{docs_count}</b></div>
        <div class="summary-row"><span>Total neto:</span><b>{_money(total_neto)}</b></div>
        <div class="summary-row"><span>IVA:</span><b>{_money(total_iva)}</b></div>
        <div class="summary-row"><span>Total bruto:</span><b class="danger">{_money(total_bruto)}</b></div>
        <div class="summary-row"><span>Total SKU:</span><b>{sku_count}</b></div>
      </div>
      <div class="card">
        <div class="card-title"><span class="mini-icon">P</span>Periodo del reporte</div>
        <div class="summary-row"><span>Fecha de emision:</span><b>{_esc(fecha_doc or '-')}</b></div>
        <div class="summary-row"><span>Generado por:</span><b>{_esc(generated_by)}</b></div>
        <div class="summary-row"><span>Fecha generacion:</span><b>{_esc(fecha_gen)}</b></div>
      </div>
    </div>
    <div class="section-title">Documentos incluidos en el reporte</div>
    <table>
      <thead><tr>
        <th>Tipo documento</th><th>Nro documento</th><th>Fecha emision</th><th class="r">Neto</th><th class="r">IVA</th><th class="r">Bruto</th><th class="c">SKUs</th><th class="c">Estado</th>
      </tr></thead>
      <tbody>{docs_rows}</tbody>
      <tfoot><tr>
        <td colspan="3" class="r strong">Totales generales</td><td class="r strong">{_money(total_neto)}</td><td class="r strong">{_money(total_iva)}</td><td class="r strong red">{_money(total_bruto)}</td><td class="c strong">{sku_count}</td><td></td>
      </tr></tfoot>
    </table>
    <div class="section-title">Detalle de cubicaje general</div>
    <table class="detail">
      <thead><tr>
        <th>SKU</th><th>Descripcion ERP</th><th>Doc.</th><th class="c">Cant.</th><th class="c">Bultos</th><th class="r">Kg/u</th><th class="r">PV/u</th><th class="r">Vol cm3/u</th><th class="r">Predom/u</th><th class="r">Total predom</th>
      </tr></thead>
      <tbody>{rows_html}</tbody>
      <tfoot><tr>
        <td colspan="3" class="r strong">Totales</td><td class="c strong">{int(tot_qty)}</td><td class="c strong">{tot_bult}</td><td class="r strong">{fkg_filter(tot_kg)}</td><td class="r strong">{fkg_filter(tot_pv)}</td><td class="r strong">{fvol_filter(tot_vol)}</td><td class="r">-</td><td class="r strong red">{fkg_filter(tot_pred)}</td>
      </tr></tfoot>
    </table>
    <div class="totals-grid">
      <div>
        <div class="section-title">Resumen de cubicaje</div>
        <div class="metric-grid">
          <div class="metric"><div class="metric-label">Unidades</div><div class="metric-val">{int(tot_qty)}</div></div>
          <div class="metric"><div class="metric-label">Bultos</div><div class="metric-val">{tot_bult}</div></div>
          <div class="metric"><div class="metric-label">Peso real</div><div class="metric-val">{fkg_filter(tot_kg)}</div></div>
          <div class="metric"><div class="metric-label">Peso vol.</div><div class="metric-val">{fkg_filter(tot_pv)}</div></div>
          <div class="metric"><div class="metric-label">Volumen</div><div class="metric-val">{fvol_filter(tot_vol)}</div></div>
        </div>
        <div class="pred-box"><span>Peso predominante total</span><span class="value">{fkg_filter(tot_pred)} kg</span></div>
      </div>
      <div>
        <div class="section-title">Observaciones</div>
        <div class="obs">Reporte generado automaticamente desde el Cubicador ILUS. Los pesos se calculan con la ficha logistica disponible para cada SKU.</div>
      </div>
    </div>
  </div>
  <div class="footer">
    <div class="footer-logo">{logo_tag}</div>
    <div>Sistema de ILUS ERP<br>Modulo de Transporte Cubicador</div>
    <div>Documento generado por:<br>{_esc(generated_by)}</div>
    <div class="qr"><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i><i></i></div>
  </div>
</div>
</body></html>"""

    pdf_bytes = _pw_pdf(
        html,
        page_format="A4",
        margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
    )
    fname = "cubicador_" + "_".join(f"{t}{n}" for t, n in docs) + ".pdf"
    return send_file(
        _BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=fname[:80],
    )


@app.route("/cubicador/export/pdf", methods=["POST"])
@login_required
def cubicador_export_pdf():
    """Descarga el resultado del cubicador como PDF elegante — soporta múltiples documentos."""
    if not g.permissions.get("cubicador"):
        flash("No tienes acceso al módulo Cubicador.", "danger")
        return redirect(url_for("index"))

    docs = _parse_docs_from_form(request.form)
    if not docs:
        flash("Debes ingresar al menos un documento.", "warning")
        return redirect(url_for("cubicador"))

    try:
        headers, lineas, errors = _fetch_multi_docs(docs)
    except Exception as ex:
        flash(f"Error al consultar el ERP: {ex}", "danger")
        return redirect(url_for("cubicador"))

    if not headers:
        flash(" · ".join(errors) if errors else "No se encontraron documentos.", "warning")
        return redirect(url_for("cubicador"))

    # ── Logo ILUS ────────────────────────────────────────────────────
    return _cubicador_pdf_response_ilus(headers, lineas, docs)

    import os as _os
    _logo_path = _os.path.join(_os.path.dirname(__file__), "static", "logo_pdf.txt")
    try:
        with open(_logo_path) as _f:
            logo_b64 = _f.read().strip()
        logo_tag = f'<img src="data:image/png;base64,{logo_b64}" style="height:46px;display:block">'
    except Exception:
        logo_tag = '<div style="font-size:18pt;font-weight:900;color:#CC0000;line-height:1">ILUS</div>'

    from datetime import datetime as _dt
    fecha_gen = _dt.now().strftime("%d/%m/%Y %H:%M")

    # ── Título principal ─────────────────────────────────────────────
    if len(docs) == 1 and headers:
        h0 = headers[0]
        main_title = f"{h0['tido']} N° {h0.get('nudo_display', h0.get('nudo',''))} — Cubicador ILUS"
    else:
        main_title = f"Cubicador ILUS · {len(docs)} documentos"

    # ── Tarjetas de documentos ───────────────────────────────────────
    docs_html = ""
    for hdr in headers:
        docs_html += f"""
        <div class="doc-card">
          <div class="doc-title">{hdr['tido']} N° {hdr.get('nudo_display', hdr.get('nudo',''))}</div>
          <div class="doc-fecha">{hdr.get('fecha','')}</div>
          <div class="doc-cliente">{hdr.get('cliente_nombre','—')}</div>
          <div class="doc-rut">RUT {hdr.get('cliente_rut','—')}</div>
          <div class="doc-bruto">${fvol_filter(hdr.get('valor_bruto',0))}</div>
        </div>"""

    # ── Totales ──────────────────────────────────────────────────────
    tot_qty  = sum(l["cantidad"]     for l in lineas)
    tot_kg   = sum(l["peso_kg_tot"]  for l in lineas)
    tot_pv   = sum(l["peso_vol_tot"] for l in lineas)
    tot_vol  = sum(l["vol_tot"]      for l in lineas)
    tot_pred = sum(l["pred_tot"]     for l in lineas)
    tot_bult = sum(l["total_bultos"] for l in lineas)

    # ── Filas de la tabla ────────────────────────────────────────────
    rows_html = ""
    for i, l in enumerate(lineas):
        bg = "#f7f7f7" if i % 2 == 0 else "#ffffff"
        sf = not l["tiene_bultos"]
        rows_html += f"""
        <tr style="background:{bg}">
          <td class="mono">{l['sku']}</td>
          <td>{l['descripcion_erp']}</td>
          <td class="c">{int(l['cantidad'])}</td>
          <td class="c">{l['total_bultos'] if l['tiene_ficha'] else 's/f'}</td>
          <td class="r">{'—' if sf else fkg_filter(l['peso_kg_u'])}</td>
          <td class="r">{'—' if sf else fkg_filter(l['peso_vol_u'])}</td>
          <td class="r">{'—' if sf else fvol_filter(l['vol_u'])}</td>
          <td class="r fw red">{'—' if sf else fkg_filter(l['pred_tot'])}</td>
        </tr>"""

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:Arial,sans-serif;font-size:8.5pt;color:#1a1a1a;padding:16px 18px}}

/* ── Cabecera de página ── */
.page-hdr{{background:#1a1a1a;padding:12px 16px;border-radius:6px 6px 0 0;
           display:flex;justify-content:space-between;align-items:center}}
.page-hdr .main-title{{font-size:12.5pt;color:#fff;font-weight:900;margin:0}}
.page-hdr .sub{{font-size:7pt;color:#aaa;margin-top:3px}}
.page-hdr .red-bar{{width:4px;background:#CC0000;height:36px;border-radius:2px;margin-right:10px;flex-shrink:0}}
.page-hdr .left{{display:flex;align-items:center}}

/* ── Franja de documentos ── */
.docs-strip{{background:#f2f2f2;border:1px solid #ddd;border-top:none;
             border-radius:0 0 6px 6px;padding:8px 14px;margin-bottom:14px;
             display:flex;flex-wrap:wrap;gap:8px}}
.doc-card{{background:#fff;border:1px solid #ddd;border-left:3px solid #CC0000;
           border-radius:0 4px 4px 0;padding:7px 10px;flex:1;min-width:150px;max-width:250px}}
.doc-title{{font-weight:900;font-size:8.5pt;color:#CC0000}}
.doc-fecha{{font-size:6.8pt;color:#888;margin-bottom:4px}}
.doc-cliente{{font-weight:700;font-size:8pt;color:#1a1a1a}}
.doc-rut{{font-size:6.8pt;color:#888}}
.doc-bruto{{font-size:9.5pt;font-weight:900;color:#1a1a1a;margin-top:4px}}

/* ── Tabla ── */
table{{width:100%;border-collapse:collapse;font-size:7.8pt}}
th{{background:#1a1a1a;color:#fff;padding:6px 5px;text-align:left;
    font-size:7pt;text-transform:uppercase;letter-spacing:.3px}}
td{{padding:5px;border-bottom:1px solid #ebebeb;vertical-align:middle}}
.c{{text-align:center}}.r{{text-align:right}}
.mono{{font-family:monospace;font-weight:bold}}.fw{{font-weight:bold}}.red{{color:#CC0000}}
tfoot tr{{background:#1a1a1a!important;color:#fff;font-weight:bold}}
tfoot td{{border:none;padding:7px 5px}}

/* ── Barra de totales ── */
.totals-bar{{margin-top:12px;background:#1a1a1a;border-radius:6px;
             padding:10px 16px;display:flex;gap:0;align-items:center}}
.t-item{{text-align:center;flex:1;padding:0 8px;border-right:1px solid #333}}
.t-item:last-child{{border-right:none}}
.t-label{{font-size:6pt;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-bottom:2px}}
.t-val{{font-size:11pt;font-weight:900;color:#fff;line-height:1.1}}
.t-pred .t-val{{color:#CC0000;font-size:13pt}}

/* ── Footer ── */
.foot{{margin-top:10px;font-size:6.5pt;color:#aaa;text-align:right}}
</style></head><body>

<div class="page-hdr">
  <div class="left">
    <div class="red-bar"></div>
    <div>
      <div class="main-title">{main_title}</div>
      <div class="sub">Generado el {fecha_gen}</div>
    </div>
  </div>
  <div>{logo_tag}</div>
</div>

<div class="docs-strip">
  {docs_html}
</div>

<table>
  <thead><tr>
    <th>SKU</th><th>Descripción ERP</th>
    <th class="c">Cant</th><th class="c">Bultos</th>
    <th class="r">Kg/u</th><th class="r">PV/u</th>
    <th class="r">Vol cm³</th>
    <th class="r" style="color:#CC0000">Total Predom</th>
  </tr></thead>
  <tbody>{rows_html}</tbody>
  <tfoot><tr>
    <td colspan="2" style="text-align:right;color:#bbb;font-size:7pt;letter-spacing:.5px">TOTALES</td>
    <td class="c">{int(tot_qty)}</td>
    <td class="c">{tot_bult}</td>
    <td class="r">{fkg_filter(tot_kg)}</td>
    <td class="r">{fkg_filter(tot_pv)}</td>
    <td class="r">{fvol_filter(tot_vol)}</td>
    <td class="r red" style="font-size:9pt">{fkg_filter(tot_pred)}</td>
  </tr></tfoot>
</table>

<div class="totals-bar">
  <div class="t-item">
    <div class="t-label">Unidades</div>
    <div class="t-val">{int(tot_qty)}</div>
  </div>
  <div class="t-item">
    <div class="t-label">Bultos</div>
    <div class="t-val">{tot_bult}</div>
  </div>
  <div class="t-item">
    <div class="t-label">Total Kg</div>
    <div class="t-val">{fkg_filter(tot_kg)}</div>
  </div>
  <div class="t-item">
    <div class="t-label">Total PV</div>
    <div class="t-val">{fkg_filter(tot_pv)}</div>
  </div>
  <div class="t-item">
    <div class="t-label">Vol cm³</div>
    <div class="t-val">{fvol_filter(tot_vol)}</div>
  </div>
  <div class="t-item t-pred">
    <div class="t-label">Predominante</div>
    <div class="t-val">{fkg_filter(tot_pred)}</div>
  </div>
</div>

<div class="foot">ILUS Sport &amp; Health · Sistema de Gestión de Productos</div>
</body></html>"""

    from io import BytesIO as _BytesIO

    pdf_bytes = _pw_pdf(
        html,
        page_format = "A4",
        margin = {"top": "12mm", "bottom": "12mm", "left": "14mm", "right": "14mm"},
    )

    fname = "cubicador_" + "_".join(f"{t}{n}" for t, n in docs) + ".pdf"
    return send_file(
        _BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=fname[:80],
    )


@app.route("/cubicador/sync-nombre", methods=["POST"])
@login_required
def cubicador_sync_nombre():
    """Actualiza el nombre de un producto en nuestra BD con el nombre del ERP."""
    if not g.permissions.get("cubicador"):
        return jsonify({"error": "sin permiso"}), 403
    sku        = request.form.get("sku", "").strip().upper()
    nombre_erp = request.form.get("nombre_erp", "").strip()
    if not sku or not nombre_erp:
        return jsonify({"error": "datos incompletos"}), 400
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            f"UPDATE `{PRODUCTS_TABLE}` SET nombre=%s WHERE UPPER(TRIM(sku))=%s",
            (nombre_erp, sku)
        )
    conn.commit()
    _invalidate_listing_cache()
    return jsonify({"ok": True, "sku": sku, "nombre": nombre_erp})


# ═══════════════════════════════════════════════════════════════
#  MÓDULO: ASIGNAR Y COTIZAR
# ═══════════════════════════════════════════════════════════════

# ── Credenciales FedEx Rate API ─────────────────────────────────
FEDEX_RATE_CLIENT_ID     = "l74e144461994249a0abc124abef203e10"
FEDEX_RATE_CLIENT_SECRET = "2ce7d089fbf642e89c0748389cfda22d"
FEDEX_ACCOUNT            = "204155375"
FEDEX_ORIGIN_POSTAL      = "9276181"
FEDEX_ORIGIN_CITY        = "Maipu"
FEDEX_OAUTH_URL          = "https://apis.fedex.com/oauth/token"
FEDEX_RATE_URL           = "https://apis.fedex.com/rate/v1/rates/quotes"

# Cache del token OAuth (dura ~3600 s)
_fedex_token_cache = {"token": None, "expires_at": 0}
_fedex_token_lock  = threading.Lock()


def _fedex_get_token() -> str:
    """Obtiene un Bearer token de FedEx (con cache para evitar re-autenticar en cada request)."""
    import requests as _req
    with _fedex_token_lock:
        now = time.time()
        if _fedex_token_cache["token"] and now < _fedex_token_cache["expires_at"] - 30:
            return _fedex_token_cache["token"]
        resp = _req.post(
            FEDEX_OAUTH_URL,
            data={
                "grant_type":    "client_credentials",
                "client_id":     FEDEX_RATE_CLIENT_ID,
                "client_secret": FEDEX_RATE_CLIENT_SECRET,
            },
            timeout=15,
        )
        resp.raise_for_status()
        j = resp.json()
        _fedex_token_cache["token"]      = j["access_token"]
        _fedex_token_cache["expires_at"] = now + int(j.get("expires_in", 3600))
        return _fedex_token_cache["token"]


def _fedex_calc_rate(peso_pred_kg: float, postal_destino: str,
                     es_residencial: bool = False) -> dict:
    """
    Llama a la FedEx Rate API y devuelve dict con tarifa, tiempo de tránsito, etc.
    Peso facturable = max(peso_pred_kg, 0.5) redondeado a 1 decimal.
    """
    import requests as _req
    token     = _fedex_get_token()
    peso_fact = max(round(peso_pred_kg, 1), 0.5)

    body = {
        "accountNumber": {"value": FEDEX_ACCOUNT},
        "requestedShipment": {
            "shipper": {
                "address": {
                    "city":                FEDEX_ORIGIN_CITY,
                    "postalCode":          FEDEX_ORIGIN_POSTAL,
                    "countryCode":         "CL",
                    "streetLines":         ["Aurora de Chile 2486"],
                    "residential":         False,
                    "stateOrProvinceCode": "",
                }
            },
            "recipient": {
                "address": {
                    "city":                "",
                    "postalCode":          str(postal_destino),
                    "countryCode":         "CL",
                    "streetLines":         [""],
                    "residential":         es_residencial,
                    "stateOrProvinceCode": "",
                }
            },
            "pickupType":      "DROPOFF_AT_FEDEX_LOCATION",
            "packagingType":   "YOUR_PACKAGING",
            "rateRequestType": ["ACCOUNT", "LIST"],
            "requestedPackageLineItems": [{
                "groupPackageCount": 1,
                "physicalPackaging": "YOUR_PACKAGING",
                "weight": {"units": "KG", "value": peso_fact},
            }],
        },
        "carrierCodes":            ["FDXE"],
        "returnLocalizedDateTime": True,
        "webSiteCountryCode":      "CL",
    }

    resp = _req.post(
        FEDEX_RATE_URL,
        json=body,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type":  "application/json",
            "X-locale":      "es_CL",
        },
        timeout=20,
    )
    resp.raise_for_status()
    data  = resp.json()
    rates = data.get("output", {}).get("rateReplyDetails", [])
    if not rates:
        return {"error": "Sin tarifas disponibles en respuesta FedEx"}

    # Elegir primer rate con detalles
    best = next((r for r in rates if r.get("ratedShipmentDetails")), rates[0])
    rated_details = best.get("ratedShipmentDetails", [{}])
    detail = (
        next((d for d in rated_details if d.get("rateType") == "ACCOUNT"), None)
        or next((d for d in rated_details if d.get("rateType") == "LIST"), None)
        or (rated_details[0] if rated_details else {})
    )

    total  = float(detail.get("totalNetCharge", detail.get("totalNetFedExCharge", 0)))
    moneda = detail.get("currency", "CLP")

    recargos_raw = detail.get("shipmentRateDetail", {}).get("surCharges", [])
    recargos = [
        {"nombre": rc.get("description", rc.get("type", "")), "monto": float(rc.get("amount", 0))}
        for rc in recargos_raw
    ]

    transit_raw = best.get("operationalDetail", {}).get("transitTime", "")
    transit_map = {
        "ONE_DAY":    "1 día hábil",
        "TWO_DAYS":   "2 días hábiles",
        "THREE_DAYS": "3 días hábiles",
        "FOUR_DAYS":  "4 días hábiles",
        "FIVE_DAYS":  "5 días hábiles",
    }

    return {
        "servicio":        best.get("serviceType", "FEDEX_GROUND"),
        "tarifa":          total,
        "moneda":          moneda,
        "tiempo_transito": transit_map.get(transit_raw, transit_raw or "—"),
        "peso_facturable": peso_fact,
        "recargos":        recargos,
    }


# Mapa comunas CL → código postal (expandible)
_COMUNA_POSTAL = {
    "SANTIAGO": "8320000", "PROVIDENCIA": "7500000", "LAS CONDES": "7550000",
    "VITACURA": "7630000", "NUNOA": "7750000", "LA FLORIDA": "8240000",
    "MAIPU": "9276181", "PUDAHUEL": "9020000", "QUILICURA": "8711000",
    "RENCA": "8580000", "CERRO NAVIA": "9000000", "ESTACION CENTRAL": "9110000",
    "PENALOLEN": "7941000", "MACUL": "7901000", "SAN JOAQUIN": "8810000",
    "LA GRANJA": "8890000", "LA PINTANA": "8310000", "EL BOSQUE": "8080000",
    "SAN BERNARDO": "8060000", "BUIN": "9620000", "PIRQUE": "9650000",
    "TALAGANTE": "9760000", "PAINE": "9680000", "MELIPILLA": "9810000",
    "COLINA": "8500000", "LAMPA": "9520000", "PENARFLOR": "9760000",
    "CURACAVI": "9860000", "VALPARAISO": "2340000", "VINA DEL MAR": "2520000",
    "CONCON": "2521000", "QUILPUE": "2430000", "VILLA ALEMANA": "2490000",
    "SAN ANTONIO": "2830000", "RANCAGUA": "2820000", "TALCA": "3460000",
    "CURICO": "3340000", "LINARES": "3580000", "CHILLAN": "3780000",
    "CONCEPCION": "4030000", "TALCAHUANO": "4040000", "LOS ANGELES": "4440000",
    "TEMUCO": "4780000", "VALDIVIA": "5090000", "OSORNO": "5290000",
    "PUERTO MONTT": "5480000", "PUERTO VARAS": "5550000",
    "COYHAIQUE": "5950000", "PUNTA ARENAS": "6200000",
    "IQUIQUE": "1100000", "ARICA": "1000000", "ANTOFAGASTA": "1240000",
    "CALAMA": "1390000", "COPIAPO": "1530000", "LA SERENA": "1700000",
    "COQUIMBO": "1780000",
}


def _comuna_to_postal(comuna: str):
    """Convierte nombre de comuna a código postal CL. Devuelve None si no se conoce."""
    if not comuna:
        return None
    key = (comuna.strip().upper()
           .replace("Ñ","N").replace("Á","A").replace("É","E")
           .replace("Í","I").replace("Ó","O").replace("Ú","U"))
    return _COMUNA_POSTAL.get(key)


@app.route("/asignar", methods=["GET"])
@login_required
def asignar_cotizar():
    """Página principal del módulo Asignar y Cotizar.

    Forzamos headers anti-caché para que el navegador NUNCA sirva una
    versión vieja del HTML. Crítico porque este template es el que más
    JavaScript inline tiene y donde más sufrimos por caché del usuario.
    """
    if not g.permissions.get("cubicador"):
        flash("No tienes acceso al módulo Asignar y Cotizar.", "danger")
        return redirect(url_for("index"))
    resp = make_response(render_template("cubicador/asignar.html"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/api/asignar/documento", methods=["POST"])
@login_required
def api_asignar_documento():
    """
    Busca un documento en el ERP y devuelve header + líneas con datos de cubicaje.
    POST JSON: { tido, nudo }
    Campos de respuesta alineados con lo que usa asignar.html.
    """
    if not g.permissions.get("cubicador"):
        return jsonify({"error": "Sin permiso"}), 403

    data = request.get_json(silent=True) or {}
    tido = (data.get("tido") or "FCV").strip().upper()
    nudo = str(data.get("nudo") or "").strip()

    if not nudo:
        return jsonify({"error": "Número de documento requerido"}), 400

    try:
        hdr, lineas = _cubicador_fetch(tido, nudo)
    except ConnectionError as ce:
        return jsonify({"error": str(ce)}), 503
    except Exception as ex:
        return jsonify({"error": f"Error al consultar ERP: {ex}"}), 500

    if hdr is None:
        # Devolvemos diagnóstico del motor para que el frontend pueda mostrar
        # exactamente qué variantes de NUDO se probaron contra el ERP.
        try:
            erp_diag = _ERP._doc_cached(f"{tido}|{nudo}")
        except Exception:
            erp_diag = None
        return jsonify({
            "error": f"No se encontró {tido} N° {nudo} en el ERP Random",
            "diagnostics": {
                "nudo_tried": erp_engine.nudo_variants(
                    erp_engine.TIDO_NUDO_MAP[tido][1](nudo) if tido in erp_engine.TIDO_NUDO_MAP else nudo
                ),
                "tido": tido,
                "nudo": nudo,
                "hint": "Verifica que el documento existe en el ERP Random (la app probó todas las variantes de padding que conoce)."
            }
        }), 404

    # Extraer valor de ZZ Envío (tarifa de despacho cargada en el documento)
    zzenvio_valor = 0.0
    for l in lineas:
        if l["sku"].upper() == "ZZENVIO":
            zzenvio_valor = float(l.get("vaneli", 0))
            break

    lineas_out = []
    tot_qty = tot_kg = tot_pv = tot_vol = tot_pred = tot_bultos = 0.0

    for l in lineas:
        if l.get("es_zz"):
            continue
        qty          = l["cantidad"]
        peso_kg_u    = l["peso_kg_u"]
        peso_vol_u   = l["peso_vol_u"]
        pred_u       = l["pred_u"]
        pred_tot     = l["pred_tot"]
        bultos_u     = int(l["total_bultos"])
        bultos_total = bultos_u * int(qty) if l["tiene_bultos"] else 0

        tot_qty    += qty
        tot_kg     += l["peso_kg_tot"]
        tot_pv     += l["peso_vol_tot"]
        tot_vol    += l["vol_tot"]
        tot_pred   += pred_tot
        tot_bultos += bultos_total

        lineas_out.append({
            # Nombres exactos que lee asignar.html
            "sku":             l["sku"],
            "descripcion_erp": l["descripcion_erp"] or l["nombre_app"],
            "cantidad":        qty,
            "total_bultos":    bultos_u,          # bultos por unidad
            "bultos_tot":      bultos_total,       # bultos totales
            "tiene_ficha":     l["tiene_ficha"],
            "tiene_bultos":    l["tiene_bultos"],
            "peso_kg_u":       round(peso_kg_u,  3),
            "peso_vol_u":      round(peso_vol_u, 3),
            "vol_u":           round(l["vol_u"], 1),   # cm³/u
            "pred_u":          round(pred_u,     3),
            "pred_tot":        round(pred_tot,   3),
        })

    postal_destino = _comuna_to_postal(hdr.get("comuna", ""))

    return jsonify({
        "ok":     True,
        "header": {**hdr, "postal_destino": postal_destino},
        "lineas": lineas_out,
        "totales": {
            "total_qty":    int(tot_qty),
            "total_bultos": int(tot_bultos),
            "peso_kg":      round(tot_kg,   3),
            "peso_pv":      round(tot_pv,   3),
            "vol_cm3":      round(tot_vol,  1),
            "peso_pred":    round(tot_pred, 3),
        },
        "tipos_doc":     TIPOS_DOC_CUBICADOR,
        "zzenvio_valor": round(zzenvio_valor, 0),
    })


@app.route("/api/asignar/tarifa-fedex", methods=["POST"])
@login_required
def api_asignar_tarifa_fedex():
    """
    Consulta la tarifa FedEx para el envío.
    POST JSON: { peso_pred, zona_id, es_residencial, es_remoto, valor_neto }
      zona_id: 1=RM, 2=V/VI/VII, 3=VIII/IX, 4=XIV/X, 5=XI/XII, 6=I/II/III/IV
    Responde: { ok, costo_total, detalle:{servicio, tiempo_transito, peso_facturable, recargos} }
    """
    if not g.permissions.get("cubicador"):
        return jsonify({"error": "Sin permiso"}), 403

    data       = request.get_json(silent=True) or {}
    peso_pred  = float(data.get("peso_pred", 0) or 0)
    zona_id    = str(data.get("zona_id", "1")).strip()
    es_resid   = bool(data.get("es_residencial", False))
    es_remoto  = bool(data.get("es_remoto", False))
    valor_neto = float(data.get("valor_neto", 0) or 0)

    if peso_pred <= 0:
        return jsonify({"error": "Peso predominante debe ser mayor a 0"}), 400

    # Mapear zona_id → postal de referencia para la API FedEx
    _ZONA_POSTAL = {
        "1": "8320000",   # RM Santiago
        "2": "2340000",   # Valparaíso
        "3": "4030000",   # Concepción
        "4": "5480000",   # Puerto Montt
        "5": "6200000",   # Punta Arenas
        "6": "1240000",   # Antofagasta
    }
    postal_dest = _ZONA_POSTAL.get(zona_id, "8320000")

    try:
        result = _fedex_calc_rate(peso_pred, postal_dest, es_residencial=es_resid)
        if "error" in result:
            return jsonify({"error": result["error"]}), 422

        costo_base = result["tarifa"]
        # Recargos adicionales según opciones marcadas
        recargo_resid  = 4200 if es_resid  else 0
        recargo_remoto = 6800 if es_remoto else 0
        costo_total    = costo_base + recargo_resid + recargo_remoto

        return jsonify({
            "ok":          True,
            "costo_total": round(costo_total, 0),
            "detalle": {
                "servicio":        result["servicio"],
                "tiempo_transito": result["tiempo_transito"],
                "peso_facturable": result["peso_facturable"],
                "costo_base":      round(costo_base, 0),
                "recargo_resid":   recargo_resid,
                "recargo_remoto":  recargo_remoto,
                "recargos_fedex":  result.get("recargos", []),
                "moneda":          result["moneda"],
                "zona_id":         zona_id,
            },
        })
    except Exception as ex:
        import traceback
        print(f"[FEDEX ERROR] {ex}\n{traceback.format_exc()}")
        return jsonify({"error": f"Error FedEx API: {str(ex)}"}), 502


# ═══════════════════════════════════════════════════════════════
#  MÓDULO: TRANSPORTE Y DISTRIBUCIÓN
# ═══════════════════════════════════════════════════════════════

ZZ_SKUS = {'ZZenvio', 'ZZINGREPUESTO', 'ZZSERVTEC', 'ZZRetiro', 'ZZINSTALACION', 'ZZINGARREQUIP'}

ESTADOS_COMPROMISO = [
    'Pendiente', 'En proceso', 'Despachado', 'Problema',
    'Pedido de vuelta', 'Preventa', 'Indemnización', 'Garantía',
    'Logística inversa', 'Prioridad', 'Indemnización revisada',
    'Indemnización rechazada', 'Regalo', 'Reentrega',
]
COURIERS = [
    'FedEx', 'Envíame', 'Transportes Milling', 'Starken',
    'Daniel Pulgar', 'Servicio Técnico', 'Transportes Felca', 'Dropit', 'Clickex',
]
ESTADOS_ENTREGA = [
    'En preparación', 'Entregado a transporte',
    'En ruta', 'Entregado', 'Entrega fallida', 'Devolución',
]

ESTADO_COLORS = {
    'Pendiente':              'warning',
    'En proceso':             'primary',
    'Despachado':             'success',
    'Problema':               'danger',
    'Pedido de vuelta':       'danger',
    'Preventa':               'secondary',
    'Indemnización':          'danger',
    'Garantía':               'info',
    'Logística inversa':      'secondary',
    'Prioridad':              'danger',
    'Indemnización revisada': 'warning',
    'Indemnización rechazada':'danger',
    'Regalo':                 'info',
    'Reentrega':              'warning',
}


def _tr_log(entity_type, entity_id, accion, detalle=""):
    """Registra un evento de trazabilidad en transport_logs."""
    try:
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO transport_logs (entity_type,entity_id,accion,detalle,usuario) "
                "VALUES (%s,%s,%s,%s,%s)",
                (entity_type, entity_id, accion, detalle, current_username())
            )
        conn.commit()
    except Exception:
        pass


def _parse_obdo(obdo: str) -> dict:
    """Parsea texto libre OBDO: 'Dirección - teléfono - email'"""
    import re as _re
    result = {"direccion": "", "telefono": "", "email": ""}
    if not obdo:
        return result
    email_m = _re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', obdo)
    if email_m:
        result["email"] = email_m.group(0)
        obdo = obdo.replace(email_m.group(0), "")
    phone_m = _re.search(r'\b[\+56]*\s*[29]\d{7,8}\b', obdo)
    if phone_m:
        result["telefono"] = phone_m.group(0).strip()
        obdo = obdo.replace(phone_m.group(0), "")
    result["direccion"] = _re.sub(r'[-–]+\s*$', '', obdo).strip(' -–')
    return result


def _clasif_from_skus(skus):
    """Determina la clasificación de un compromiso según sus SKUs ZZ."""
    skus = [s.strip().upper() for s in skus if s]
    if not skus: return "despacho"
    if all(s == "ZZRETIRO" for s in skus): return "retiro"
    if any(s == "ZZINSTALACION" for s in skus): return "instalacion"
    if any(s in ("ZZSERVTEC", "ZZINGREPUESTO", "ZZINGARREQUIP") for s in skus): return "mantencion"
    return "despacho"


def _tr_fetch_from_erp(tido, nudo):
    """
    Obtiene un documento del ERP vía API y lo guarda/actualiza en transport_commitments.
    Usa la misma lógica de _cubicador_fetch pero orientada a transporte.
    Retorna el id del commitment o None.
    """
    from datetime import datetime as _dt
    TOKEN  = ERP_CONFIG.get("api_token", "")
    nudos  = _nudo_variants(nudo)

    # Mapear VD/WEB → NVV igual que cubicador
    if tido in _ERP_TIDO_NUDO_MAP:
        erp_tido, nudo_fn = _ERP_TIDO_NUDO_MAP[tido]
        erp_nudo = nudo_fn(nudo)
        nudos = _nudo_variants(erp_nudo)
    else:
        erp_tido = tido

    raw_header, raw_lineas = None, []
    for nv in nudos:
        try:
            body = _erp_get("/documentos/render",
                            {"tido": erp_tido, "nudo": nv, "empresa": "01"},
                            TOKEN, timeout=12)
            data = body.get("data") or []
            if data:
                raw_header = data[0].get("maeedo") or {}
                raw_lineas = data[0].get("maeddo") or []
                break
        except Exception as e:
            raise ConnectionError(f"ERP no responde: {e}")

    if not raw_header:
        return None, "No encontrado en ERP"

    # Filtrar sólo líneas ZZ
    zz_lines = [l for l in raw_lineas
                if (l.get("KOPRCT") or "").strip().upper() in {s.upper() for s in ZZ_SKUS}]
    if not zz_lines:
        return None, "Documento sin líneas ZZ"

    # Calcular saldo
    saldo_total = sum(
        float(l.get("CAPRCO1") or 0) - float(l.get("CAPRAD1") or 0)
        for l in zz_lines
    )
    tiene_saldo = 1 if saldo_total > 0 else 0

    # Parsear OBDO
    obdo_str = (raw_header.get("OBDO") or raw_header.get("TEXTO1") or "").strip()
    parsed   = _parse_obdo(obdo_str)
    direccion = parsed["direccion"] or (raw_header.get("DIENDESP") or "").strip()

    # Nombre cliente
    endo = (raw_header.get("ENDO") or "").strip()
    cliente_nombre = (raw_header.get("NOKOEN") or "").strip().title()
    if not cliente_nombre and endo:
        try:
            ent = _erp_get("/entidades", {"rten": endo}, TOKEN, timeout=6)
            ed  = (ent.get("data") or [{}])[0]
            cliente_nombre = (ed.get("NOKOEN") or "").strip().title()
        except Exception:
            pass

    # Fecha
    from datetime import datetime as _dt
    def _parse_date(s):
        if not s: return None
        try: return _dt.fromisoformat(s.replace("Z", "+00:00")).date()
        except: return None

    fecha_em  = _parse_date(raw_header.get("FEEMDO"))
    fecha_ent = _parse_date(raw_header.get("FEER"))

    # Clasificación — basado en los SKUs ZZ predominantes
    skus_upper = [(l.get("KOPRCT") or "").strip().upper() for l in zz_lines]
    clasificacion = _clasif_from_skus(skus_upper)

    # Costo ZZ (suma de PPPRNE de líneas ZZ)
    costo_zz = sum(float(l.get("PPPRNE") or 0) for l in zz_lines)

    # Guía (si CAPRAD1 >= CAPRCO1 en todas → tiene guía)
    guia_numero = (raw_header.get("NUDO_GIA") or raw_header.get("NUDGIA") or "").strip() or None

    # NOTA: get_db() devuelve conexión del pool via g. NO llamar conn.close()
    # al final — teardown_appcontext la cierra. Cerrar aquí deja g._db
    # apuntando a una conexión cerrada.
    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO transport_commitments
                  (tido,nudo,endo,fecha_emision,fecha_entrega,cliente_nombre,cliente_rut,
                   comuna,direccion,telefono,email,valor_neto,valor_bruto,costo_zz,
                   tiene_saldo,guia_numero,clasificacion,erp_synced_at,created_by,updated_by)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,NOW(),%s,%s)
                ON DUPLICATE KEY UPDATE
                  fecha_emision=VALUES(fecha_emision), fecha_entrega=VALUES(fecha_entrega),
                  cliente_nombre=VALUES(cliente_nombre), cliente_rut=VALUES(cliente_rut),
                  comuna=VALUES(comuna), direccion=VALUES(direccion),
                  telefono=VALUES(telefono), email=VALUES(email),
                  valor_neto=VALUES(valor_neto), valor_bruto=VALUES(valor_bruto),
                  costo_zz=CASE WHEN costo_zz=0 THEN VALUES(costo_zz) ELSE costo_zz END,
                  tiene_saldo=VALUES(tiene_saldo), guia_numero=VALUES(guia_numero),
                  clasificacion=VALUES(clasificacion), erp_synced_at=NOW(),
                  updated_by=VALUES(updated_by)
            """, (
                tido, str(nudo), endo, fecha_em, fecha_ent,
                cliente_nombre, endo,
                (raw_header.get("CMEN") or raw_header.get("NOKOZO") or
                 raw_header.get("NOKOCOMU") or raw_header.get("NOKOCOMUNADE") or
                 raw_header.get("NOKOMUENDE") or raw_header.get("NOKOMUNEN") or
                 raw_header.get("NOKCOMENDESP") or "").strip(),
                direccion, parsed["telefono"], parsed["email"],
                float(raw_header.get("VANEDO") or 0),
                float(raw_header.get("VABRDO") or 0),
                costo_zz, tiene_saldo, guia_numero, clasificacion,
                current_username(), current_username()
            ))
            comm_id = cur.lastrowid or mysql_fetchone(
                "SELECT id FROM transport_commitments WHERE tido=%s AND nudo=%s",
                (tido, str(nudo))
            )["id"]

            # Líneas ZZ
            cur.execute("DELETE FROM transport_commitment_lines WHERE commitment_id=%s", (comm_id,))
            for l in zz_lines:
                cant  = float(l.get("CAPRCO1") or 0)
                cantd = float(l.get("CAPRAD1")  or 0)
                cur.execute("""
                    INSERT INTO transport_commitment_lines
                      (commitment_id,koprct,nokopr,cantidad,cant_despachada,saldo,bodega)
                    VALUES (%s,%s,%s,%s,%s,%s,%s)
                """, (comm_id,
                      (l.get("KOPRCT") or "").strip().upper(),
                      (l.get("NOKOPR") or "").strip(),
                      cant, cantd, cant - cantd,
                      (l.get("BOSULIDO") or "").strip()))
        conn.commit()
    except Exception:
        try: conn.rollback()
        except Exception: pass
        raise
    # Importante: NO cerrar conn aquí (es del pool). teardown_appcontext lo hace.

    return comm_id, None


# ── SYNC MASIVO ERP → TRANSPORTE ────────────────────────────────

def _tr_bulk_sync_erp_mysql(fecha_desde, fecha_hasta):
    """
    Sincronización masiva desde el ERP vía MySQL directo.
    Trae todos los documentos con líneas ZZ y saldo pendiente en el rango.
    Retorna (count_ok, list_errors).
    """
    conn_erp = get_erp_conn()
    if not conn_erp:
        return 0, ["No se pudo conectar al ERP (MySQL)"]

    zz_list      = list(ZZ_SKUS)
    zz_in        = ",".join(["%s"] * len(zz_list))
    TIDOS_VALIDOS = ("FCV", "BLV", "GDV", "NVV", "NVI", "COV")
    tido_in      = ",".join(["%s"] * len(TIDOS_VALIDOS))

    try:
        with conn_erp.cursor() as cur:
            cur.execute(f"""
                SELECT
                    h.TIDO, h.NUDO, h.ENDO, h.FEEMDO, h.FEER,
                    h.NOKOEN,
                    COALESCE(h.OBDO, h.TEXTO1, '') AS OBDO,
                    COALESCE(h.DIENDESP, '')        AS DIENDESP,
                    COALESCE(h.VANEDO, 0)           AS VANEDO,
                    COALESCE(h.VABRDO, 0)           AS VABRDO,
                    COALESCE(h.NOKOZO, h.CMEN, h.NOKOCOMU, h.NOKOCOMUNADE,
                             h.NOKOMUENDE, h.NOKOMUNEN, h.NOKCOMENDESP, '') AS COMUNA,
                    COALESCE(h.NUDGIA, '')           AS NUDGIA,
                    SUM(GREATEST(d.CAPRCO1 - COALESCE(d.CAPRAD1, 0), 0)) AS saldo_zz,
                    SUM(COALESCE(d.PPPRNE, 0))       AS costo_zz_sum,
                    GROUP_CONCAT(DISTINCT UPPER(d.KOPRCT) ORDER BY d.KOPRCT) AS zz_skus
                FROM MAEEDO h
                JOIN MAEDDO d ON d.TIDO = h.TIDO AND d.NUDO = h.NUDO
                WHERE d.KOPRCT IN ({zz_in})
                  AND h.TIDO IN ({tido_in})
                  AND h.FEEMDO BETWEEN %s AND %s
                  AND (d.CAPRCO1 - COALESCE(d.CAPRAD1, 0)) > 0
                GROUP BY h.TIDO, h.NUDO
                HAVING saldo_zz > 0
                ORDER BY h.FEEMDO DESC
                LIMIT 500
            """, zz_list + list(TIDOS_VALIDOS) + [fecha_desde, fecha_hasta])
            rows = cur.fetchall()
    except Exception as exc:
        return 0, [f"Error al consultar ERP: {exc}"]
    finally:
        conn_erp.close()

    count, errs = 0, []
    local_conn = get_mysql()
    try:
        for row in rows:
            try:
                tido  = (row.get("TIDO") or "").strip()
                nudo  = (row.get("NUDO") or "").strip()
                endo  = (row.get("ENDO") or "").strip()
                nombre = (row.get("NOKOEN") or "").strip().title()
                obdo  = (row.get("OBDO") or "").strip()
                parsed = _parse_obdo(obdo)
                dir_  = parsed["direccion"] or (row.get("DIENDESP") or "").strip()
                tel   = parsed["telefono"]
                mail  = parsed["email"]
                comuna = (row.get("COMUNA") or "").strip()
                vneto  = float(row.get("VANEDO") or 0)
                vbruto = float(row.get("VABRDO") or 0)
                costo_zz = float(row.get("costo_zz_sum") or 0)
                guia  = (row.get("NUDGIA") or "").strip() or None
                zz_present = (row.get("zz_skus") or "").upper()
                _zz_list_bulk = [s.strip() for s in zz_present.split(",") if s.strip()]
                clasif = _clasif_from_skus(_zz_list_bulk) if _zz_list_bulk else "despacho"

                def _pd(val):
                    if not val: return None
                    if hasattr(val, "date"): return val.date()
                    try:
                        from datetime import datetime as _dt
                        return _dt.fromisoformat(str(val).replace("Z", "")).date()
                    except Exception:
                        return None

                fecha_em  = _pd(row.get("FEEMDO"))
                fecha_ent = _pd(row.get("FEER"))

                with local_conn.cursor() as cur:
                    cur.execute("""
                        INSERT INTO transport_commitments
                          (tido,nudo,endo,fecha_emision,fecha_entrega,
                           cliente_nombre,cliente_rut,
                           comuna,direccion,telefono,email,
                           valor_neto,valor_bruto,costo_zz,
                           tiene_saldo,guia_numero,clasificacion,
                           erp_synced_at,created_by,updated_by)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,1,%s,%s,NOW(),'sync','sync')
                        ON DUPLICATE KEY UPDATE
                          fecha_emision =VALUES(fecha_emision),
                          fecha_entrega =VALUES(fecha_entrega),
                          cliente_nombre=VALUES(cliente_nombre),
                          cliente_rut   =VALUES(cliente_rut),
                          comuna        =VALUES(comuna),
                          direccion     =VALUES(direccion),
                          telefono      =VALUES(telefono),
                          email         =VALUES(email),
                          valor_neto    =VALUES(valor_neto),
                          valor_bruto   =VALUES(valor_bruto),
                          costo_zz      =CASE WHEN costo_zz=0 THEN VALUES(costo_zz) ELSE costo_zz END,
                          tiene_saldo   =1,
                          guia_numero   =VALUES(guia_numero),
                          clasificacion =VALUES(clasificacion),
                          erp_synced_at =NOW()
                    """, (tido, nudo, endo, fecha_em, fecha_ent,
                          nombre, endo, comuna, dir_, tel, mail,
                          vneto, vbruto, costo_zz, guia, clasif))
                local_conn.commit()
                count += 1
            except Exception as e2:
                errs.append(f"{row.get('TIDO')} {row.get('NUDO')}: {e2}")
    finally:
        local_conn.close()

    return count, errs or None


def _tr_import_from_excel(file_bytes, filename):
    """
    Importa documentos desde un Excel/CSV exportado del ERP.
    Detecta columnas automáticamente.
    Retorna (count_ok, list_errors, preview_rows).
    """
    import io

    # ── intentar openpyxl (xlsx) o csv ──
    rows_raw = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws = wb.active
        headers = [str(c.value or "").strip().upper() for c in next(ws.iter_rows(min_row=1, max_row=1))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            rows_raw.append(dict(zip(headers, [str(v or "").strip() for v in row])))
    except Exception:
        # Fallback: CSV
        import csv
        decoded = file_bytes.decode("utf-8-sig", errors="replace")
        reader  = csv.DictReader(io.StringIO(decoded))
        for row in reader:
            rows_raw.append({k.strip().upper(): str(v or "").strip() for k, v in row.items()})

    if not rows_raw:
        return 0, ["Archivo vacío o formato no reconocido"], []

    # ── mapeo flexible de columnas ──
    COL_MAP = {
        "tido":   ["TIDO", "TIPO", "TIPO DOCUMENTO", "TIPODOC", "TIPO DOC"],
        "nudo":   ["NUDO", "NUMERO", "N°", "NDOC", "NUM DOC", "NUMERO DOC", "FOLIO"],
        "nombre": ["NOKOEN", "NOMBRE", "CLIENTE", "RAZON SOCIAL", "NOMBRE CLIENTE"],
        "fecha":  ["FEEMDO", "FECHA", "FECHA EMISION", "FECHA EMISIÓN", "FECHA DOC"],
        "sku":    ["KOPRCT", "SKU", "CODIGO", "CÓDIGO", "PRODUCTO"],
        "cant":   ["CAPRCO1", "CANTIDAD", "QTY", "CANT", "CANT PEDIDA"],
        "cantd":  ["CAPRAD1", "DESPACHADO", "DESPACHO", "CANT DESPACHADA", "ENTREGADO"],
        "costo":  ["PPPRNE", "PRECIO", "VALOR", "COSTO", "MONTO"],
        "comuna": ["CMEN", "NOKOZO", "COMUNA", "CIUDAD"],
        "dir":    ["DIENDESP", "DIRECCION", "DIRECCIÓN", "OBDO"],
        "rut":    ["ENDO", "RUT", "RUTEN", "RUT CLIENTE"],
    }

    def find_col(row, candidates):
        for c in candidates:
            if c in row:
                return row[c]
        return ""

    # Agrupar por TIDO+NUDO
    from collections import defaultdict
    docs = defaultdict(lambda: {"lineas": [], "header": {}})
    skipped = 0
    for row in rows_raw:
        tido  = find_col(row, COL_MAP["tido"]).upper().strip()
        nudo  = find_col(row, COL_MAP["nudo"]).strip()
        sku   = find_col(row, COL_MAP["sku"]).upper().strip()
        if not tido or not nudo:
            skipped += 1
            continue
        key = (tido, nudo)
        docs[key]["header"] = row
        if sku in {s.upper() for s in ZZ_SKUS}:
            docs[key]["lineas"].append(row)

    # Filtrar docs que tienen al menos una línea ZZ con saldo
    count, errs, preview = 0, [], []
    local_conn = get_mysql()
    try:
        for (tido, nudo), doc in docs.items():
            if not doc["lineas"]:
                continue  # sin líneas ZZ → ignorar
            hrow = doc["header"]
            nombre  = find_col(hrow, COL_MAP["nombre"]).title()
            fecha_s = find_col(hrow, COL_MAP["fecha"])
            comuna  = find_col(hrow, COL_MAP["comuna"])
            dir_    = find_col(hrow, COL_MAP["dir"])
            rut     = find_col(hrow, COL_MAP["rut"])
            costo   = sum(
                _safe_float(find_col(l, COL_MAP["costo"]))
                for l in doc["lineas"]
            )
            saldo   = sum(
                max(_safe_float(find_col(l, COL_MAP["cant"])) -
                    _safe_float(find_col(l, COL_MAP["cantd"])), 0)
                for l in doc["lineas"]
            )
            if saldo <= 0:
                continue
            zz_skus = {find_col(l, COL_MAP["sku"]).upper() for l in doc["lineas"]}
            clasif  = "retiro" if zz_skus == {"ZZRETIRO"} else "despacho"

            # Parsear fecha
            fecha_em = None
            if fecha_s:
                for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y"):
                    try:
                        from datetime import datetime as _dt
                        fecha_em = _dt.strptime(fecha_s[:10], fmt).date()
                        break
                    except Exception:
                        continue

            try:
                with local_conn.cursor() as cur:
                    cur.execute("""
                        INSERT INTO transport_commitments
                          (tido,nudo,endo,fecha_emision,cliente_nombre,cliente_rut,
                           comuna,direccion,costo_zz,tiene_saldo,clasificacion,
                           erp_synced_at,created_by,updated_by)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,1,%s,NOW(),'excel','excel')
                        ON DUPLICATE KEY UPDATE
                          fecha_emision =VALUES(fecha_emision),
                          cliente_nombre=VALUES(cliente_nombre),
                          cliente_rut   =VALUES(cliente_rut),
                          comuna        =VALUES(comuna),
                          direccion     =VALUES(direccion),
                          costo_zz      =CASE WHEN costo_zz=0 THEN VALUES(costo_zz) ELSE costo_zz END,
                          tiene_saldo   =1,
                          clasificacion =VALUES(clasificacion),
                          erp_synced_at =NOW()
                    """, (tido, nudo, rut, fecha_em, nombre, rut,
                          comuna, dir_, costo, clasif))
                local_conn.commit()
                count += 1
                if len(preview) < 5:
                    preview.append({"tido": tido, "nudo": nudo, "cliente": nombre,
                                    "clasif": clasif, "saldo_lineas": saldo})
            except Exception as e2:
                errs.append(f"{tido} {nudo}: {e2}")
    finally:
        local_conn.close()

    return count, errs or None, preview


def _safe_float(v):
    try:
        return float(str(v).replace(",", ".").strip() or 0)
    except Exception:
        return 0.0


# ── RUTAS TRANSPORTE ─────────────────────────────────────────────

def _tr_required(fn):
    """Decorator: requiere permiso transporte."""
    from functools import wraps
    @wraps(fn)
    @login_required
    def wrapper(*a, **kw):
        if not g.permissions.get("transporte"):
            flash("Sin acceso al módulo Transporte.", "danger")
            return redirect(url_for("index"))
        return fn(*a, **kw)
    return wrapper


@app.route("/transporte/api/sync", methods=["POST"])
@_tr_required
def tr_sync():
    """Sincronización masiva desde ERP MySQL. Recibe fecha_desde y fecha_hasta."""
    data = request.get_json(silent=True) or {}
    fecha_desde = data.get("fecha_desde") or ""
    fecha_hasta = data.get("fecha_hasta") or ""

    # Validar fechas
    from datetime import datetime as _dt, date as _date, timedelta as _td
    try:
        fd = _dt.strptime(fecha_desde, "%Y-%m-%d").date()
        fh = _dt.strptime(fecha_hasta, "%Y-%m-%d").date()
    except Exception:
        # Default: últimos 60 días
        fh = _date.today()
        fd = fh - _td(days=60)

    count, errs = _tr_bulk_sync_erp_mysql(str(fd), str(fh))
    _tr_log("commitment", 0, "sync_masivo",
            f"ERP MySQL {fd}→{fh}: {count} importados, {len(errs or [])} errores")
    return jsonify({
        "ok": True,
        "importados": count,
        "errores": (errs or [])[:10],
        "rango": f"{fd.strftime('%d/%m/%Y')} → {fh.strftime('%d/%m/%Y')}",
    })


@app.route("/transporte/api/compromisos")
@_tr_required
def tr_compromisos_json():
    """Versión JSON del monitor para carga AJAX."""
    from datetime import date as _date, timedelta as _td
    periodo = request.args.get("periodo", "")
    hoy = _date.today()
    _def_desde = (hoy - _td(days=30)).strftime("%Y-%m-%d")
    _def_hasta = hoy.strftime("%Y-%m-%d")
    if periodo == "mes":   fecha_desde, fecha_hasta = _def_desde, _def_hasta
    elif periodo == "3m":  fecha_desde, fecha_hasta = (hoy-_td(days=90)).strftime("%Y-%m-%d"), _def_hasta
    elif periodo == "año": fecha_desde, fecha_hasta = hoy.strftime("%Y-01-01"), _def_hasta
    elif periodo == "todo": fecha_desde = fecha_hasta = ""
    else:
        fecha_desde = request.args.get("fecha_desde", _def_desde).strip()
        fecha_hasta = request.args.get("fecha_hasta", _def_hasta).strip()

    estado = request.args.get("estado","")
    clasif = request.args.get("clasificacion","")
    q      = request.args.get("q","").strip()

    where, params = ["tiene_saldo=1"], []
    if estado: where.append("estado=%s"); params.append(estado)
    if clasif: where.append("clasificacion=%s"); params.append(clasif)
    if q:
        where.append("(cliente_nombre LIKE %s OR nudo LIKE %s OR tido LIKE %s OR comuna LIKE %s)")
        qp = f"%{q}%"; params += [qp,qp,qp,qp]
    if fecha_desde: where.append("fecha_emision >= %s"); params.append(fecha_desde)
    if fecha_hasta: where.append("fecha_emision <= %s"); params.append(fecha_hasta)

    rows = mysql_fetchall(
        "SELECT * FROM transport_commitments WHERE " + " AND ".join(where) +
        " ORDER BY fecha_emision DESC LIMIT 500", tuple(params)
    )

    result = []
    for r in rows:
        result.append({
            "id":           r["id"],
            "tido":         r["tido"],
            "nudo":         r["nudo"],
            "fecha":        r["fecha_emision"].strftime("%d/%m/%Y") if r["fecha_emision"] else "",
            "cliente":      r["cliente_nombre"] or "—",
            "rut":          r["cliente_rut"] or "",
            "comuna":       r["comuna"] or "—",
            "estado":       r["estado"] or "Pendiente",
            "costo_zz":     float(r["costo_zz"] or 0),
            "clasificacion":r["clasificacion"] or "despacho",
        })
    return jsonify({"ok": True, "compromisos": result, "total": len(result)})


@app.route("/transporte/api/inline-bulto", methods=["POST"])
@_tr_required
def tr_inline_bulto():
    """Guarda medidas de un producto inline desde el modal de detalle.
    Acepta 'bultos': [{largo,ancho,alto,peso}, ...] o campos sueltos (compat.).
    """
    data   = request.get_json(silent=True) or {}
    sku    = (data.get("sku") or "").strip().upper()
    nombre = (data.get("nombre") or "").strip()

    if not sku:
        return jsonify({"error": "Falta el SKU"}), 400

    # Aceptar lista de bultos O campos sueltos (retrocompat.)
    bultos_raw = data.get("bultos") or []
    if not bultos_raw:
        # compatibilidad con versión anterior (campos individuales)
        if data.get("largo"):
            bultos_raw = [{
                "largo": data.get("largo"), "ancho": data.get("ancho"),
                "alto":  data.get("alto"),  "peso":  data.get("peso"),
            }]

    bultos = []
    for b in bultos_raw:
        l = float(b.get("largo") or 0)
        a = float(b.get("ancho") or 0)
        h = float(b.get("alto")  or 0)
        p = float(b.get("peso")  or 0)
        if l and a and h and p:
            bultos.append((l, a, h, p))

    if not bultos:
        return jsonify({"error": "Debes ingresar al menos un bulto con todos sus campos"}), 400

    conn = get_db()
    try:
        with conn.cursor() as cur:
            # Buscar o crear producto
            cur.execute(f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE UPPER(TRIM(sku))=%s LIMIT 1", (sku,))
            row = cur.fetchone()
            if row:
                pid = row["id"]
            else:
                cur.execute(
                    f"INSERT INTO `{PRODUCTS_TABLE}` (sku, nombre, estado, created_by, updated_by) "
                    f"VALUES (%s,%s,'activo',%s,%s)",
                    (sku, nombre or sku, current_username(), current_username())
                )
                pid = cur.lastrowid

            # Reemplazar bultos existentes
            cur.execute(f"DELETE FROM `{BULTOS_TABLE}` WHERE product_id=%s", (pid,))
            for idx, (l, a, h, p) in enumerate(bultos, start=1):
                cur.execute(
                    f"INSERT INTO `{BULTOS_TABLE}` (product_id, bulto_num, largo, ancho, alto, peso) "
                    f"VALUES (%s,%s,%s,%s,%s,%s)",
                    (pid, idx, l, a, h, p)
                )
        conn.commit()
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # Calcular y retornar totales agregados de todos los bultos
    total_kg  = round(sum(b[3] for b in bultos), 4)
    total_pv  = round(sum(b[0] * b[1] * b[2] for b in bultos) / 4000.0, 4)
    pred      = round(max(total_kg, total_pv), 4)
    return jsonify({
        "ok":      True,
        "peso_kg": total_kg,
        "peso_vol": total_pv,
        "pred":    pred,
    })


@app.route("/transporte/api/upload-foto", methods=["POST"])
@_tr_required
def tr_upload_foto():
    """Sube hasta 2 fotos a un producto (por SKU) desde el modal de transporte."""
    sku = (request.form.get("sku") or "").strip().upper()
    if not sku:
        return jsonify({"error": "Falta el SKU"}), 400
    file = request.files.get("foto")
    if not file or not file.filename:
        return jsonify({"error": "Sin archivo"}), 400
    if not allowed_file(file.filename):
        return jsonify({"error": "Formato no permitido (JPG, PNG, WEBP)"}), 400

    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute(f"SELECT id FROM `{PRODUCTS_TABLE}` WHERE UPPER(TRIM(sku))=%s LIMIT 1", (sku,))
            row = cur.fetchone()
            if row:
                pid = row["id"]
            else:
                cur.execute(
                    f"INSERT INTO `{PRODUCTS_TABLE}` (sku,nombre,estado,created_by,updated_by) VALUES (%s,%s,'activo',%s,%s)",
                    (sku, sku, current_username(), current_username()),
                )
                pid = cur.lastrowid

            cur.execute(f"SELECT COUNT(*) AS n FROM `{PHOTOS_TABLE}` WHERE product_id=%s", (pid,))
            n = (cur.fetchone() or {}).get("n", 0)
            if n >= MAX_PHOTOS:
                return jsonify({"error": f"Máximo {MAX_PHOTOS} fotos por producto"}), 400

            ext = secure_filename(file.filename).rsplit(".", 1)[-1].lower()
            ts  = int(datetime.now().timestamp())
            if _CLD_READY:
                try:
                    filename = _cloud_upload(file, public_id=f"p{pid}_{ts}", folder="ilus/products")
                except Exception as exc:
                    return jsonify({"error": f"Error Cloudinary: {exc}"}), 500
            else:
                filename = f"p{pid}_{ts}.{ext}"
                file.save(os.path.join(UPLOAD_FOLDER, filename))

            cur.execute(
                f"INSERT INTO `{PHOTOS_TABLE}` (product_id,filename,orden) VALUES (%s,%s,%s)",
                (pid, filename, n + 1),
            )
        conn.commit()
        return jsonify({"ok": True, "url": _photo_src(filename), "total": n + 1})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/transporte/api/importar-excel", methods=["POST"])
@_tr_required
def tr_importar_excel():
    """Importa documentos desde un Excel/CSV exportado del ERP."""
    f = request.files.get("archivo")
    if not f or not f.filename:
        return jsonify({"error": "Sin archivo"}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("xlsx", "xls", "csv"):
        return jsonify({"error": "Solo se aceptan .xlsx, .xls o .csv"}), 400
    data = f.read()
    count, errs, preview = _tr_import_from_excel(data, f.filename)
    _tr_log("commitment", 0, "importar_excel",
            f"Archivo {f.filename}: {count} importados")
    return jsonify({
        "ok": True,
        "importados": count,
        "errores": (errs or [])[:10],
        "preview": preview,
    })


@app.route("/transporte/")
@_tr_required
def transporte_index():
    from datetime import date as _date, timedelta as _td

    periodo = request.args.get("periodo", "")
    hoy     = _date.today()

    # Fechas por defecto: último mes
    _default_desde = (hoy - _td(days=30)).strftime("%Y-%m-%d")
    _default_hasta = hoy.strftime("%Y-%m-%d")

    if periodo == "mes":
        fecha_desde, fecha_hasta = _default_desde, _default_hasta
    elif periodo == "3m":
        fecha_desde = (hoy - _td(days=90)).strftime("%Y-%m-%d")
        fecha_hasta = _default_hasta
    elif periodo == "año":
        fecha_desde = hoy.strftime("%Y-01-01")
        fecha_hasta = _default_hasta
    elif periodo == "todo":
        fecha_desde, fecha_hasta = "", ""
    else:
        fecha_desde = request.args.get("fecha_desde", _default_desde).strip()
        fecha_hasta = request.args.get("fecha_hasta", _default_hasta).strip()

    filtros = {
        "estado":        request.args.get("estado", ""),
        "clasificacion": request.args.get("clasificacion", ""),
        "q":             request.args.get("q", "").strip(),
        "fecha_desde":   fecha_desde,
        "fecha_hasta":   fecha_hasta,
    }

    where, params = ["tiene_saldo=1"], []
    if filtros["estado"]:
        where.append("estado=%s"); params.append(filtros["estado"])
    if filtros["clasificacion"]:
        where.append("clasificacion=%s"); params.append(filtros["clasificacion"])
    if filtros["q"]:
        where.append("(cliente_nombre LIKE %s OR nudo LIKE %s OR tido LIKE %s OR comuna LIKE %s)")
        q = f"%{filtros['q']}%"; params += [q, q, q, q]
    if filtros["fecha_desde"]:
        where.append("fecha_emision >= %s"); params.append(filtros["fecha_desde"])
    if filtros["fecha_hasta"]:
        where.append("fecha_emision <= %s"); params.append(filtros["fecha_hasta"])

    sql = ("SELECT * FROM transport_commitments WHERE " +
           " AND ".join(where) + " ORDER BY fecha_emision DESC LIMIT 500")
    compromisos = mysql_fetchall(sql, tuple(params))

    # Manifiestos activos para asignación rápida
    manifiestos = mysql_fetchall(
        "SELECT id,correlativo,courier,estado FROM transport_manifests "
        "WHERE estado IN ('En preparación','En curso') ORDER BY id DESC"
    )
    return render_template(
        "transporte/index.html",
        compromisos=compromisos,
        filtros=filtros,
        estados=ESTADOS_COMPROMISO,
        estado_colors=ESTADO_COLORS,
        couriers=COURIERS,
        manifiestos=manifiestos,
    )


@app.route("/transporte/api/agregar", methods=["POST"])
@_tr_required
def tr_agregar():
    """Agrega un documento al monitor desde el ERP."""
    tido = (request.form.get("tido") or "FCV").strip().upper()
    nudo = (request.form.get("nudo") or "").strip()
    if not nudo:
        return jsonify({"error": "Ingresa el número de documento"}), 400
    comm_id, err = _tr_fetch_from_erp(tido, nudo)
    if err:
        return jsonify({"error": err}), 404
    _tr_log("commitment", comm_id, "agregado",
            f"Documento {tido} {nudo} importado desde ERP")
    return jsonify({"ok": True, "id": comm_id})


@app.route("/transporte/api/compromisos/<int:cid>", methods=["PUT"])
@_tr_required
def tr_update_compromiso(cid):
    """Actualiza campos operativos: estado, costo_zz, notas, fecha_agenda."""
    data   = request.get_json(silent=True) or {}
    campos = {}
    if "estado" in data and data["estado"] in ESTADOS_COMPROMISO:
        campos["estado"] = data["estado"]
    if "costo_zz" in data:
        try: campos["costo_zz"] = float(data["costo_zz"])
        except: pass
    if "notas" in data:
        campos["notas"] = data["notas"]
    if "fecha_agenda" in data:
        campos["fecha_agenda"] = data["fecha_agenda"] or None
    if not campos:
        return jsonify({"error": "sin campos válidos"}), 400

    campos["updated_by"] = current_username()
    sets   = ", ".join(f"{k}=%s" for k in campos)
    vals   = list(campos.values()) + [cid]
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(f"UPDATE transport_commitments SET {sets} WHERE id=%s", vals)
    conn.commit()

    detalle = "; ".join(f"{k}={v}" for k, v in data.items() if k != "notas")
    _tr_log("commitment", cid, "actualizado", detalle)
    return jsonify({"ok": True})


@app.route("/transporte/api/compromisos/<int:cid>/detalle")
@_tr_required
def tr_detalle(cid):
    """
    Detalle completo de un compromiso para el modal de vista.
    Llama a _cubicador_fetch para obtener líneas con pesos y cruza con fotos locales.
    """
    c = mysql_fetchone("SELECT * FROM transport_commitments WHERE id=%s", (cid,))
    if not c:
        return jsonify({"error": "No encontrado"}), 404

    tido = (c["tido"] or "").strip()
    nudo = (c["nudo"] or "").strip()

    try:
        header, lineas = _cubicador_fetch(tido, nudo)
    except Exception as e:
        return jsonify({"error": f"ERP no disponible: {e}"}), 503

    if not header:
        return jsonify({"error": "Documento no encontrado en ERP"}), 404

    # Separar líneas ZZ de líneas de producto
    ZZ_UP = {s.upper() for s in {"ZZENVIO","ZZINGREPUESTO","ZZSERVTEC","ZZRETIRO","ZZINSTALACION","ZZINGARREQUIP"}}
    lineas_prod = []
    lineas_zz   = []
    for l in lineas:
        if l.get("es_zz") or l["sku"].upper() in ZZ_UP:
            lineas_zz.append(l)
        elif l.get("saldo", l["cantidad"]) > 0:
            # Solo incluir líneas con saldo pendiente
            lineas_prod.append(l)

    # Adjuntar fotos solo a líneas de producto.
    # Antes: 1 query por línea (N+1 hasta 20+ queries). Ahora: 1 sola query con IN(...).
    prod_ids = [l["app_id"] for l in lineas_prod if l.get("app_id")]
    fotos_por_prod = {}
    if prod_ids:
        placeholders = ",".join(["%s"] * len(prod_ids))
        ph_rows = mysql_fetchall(
            f"SELECT product_id, filename FROM `{PHOTOS_TABLE}` "
            f"WHERE product_id IN ({placeholders}) ORDER BY product_id, orden",
            tuple(prod_ids)
        )
        for p in ph_rows:
            if p["filename"] and len(fotos_por_prod.setdefault(p["product_id"], [])) < 3:
                fotos_por_prod[p["product_id"]].append(_photo_src(p["filename"]))
    for l in lineas_prod:
        l["fotos"] = fotos_por_prod.get(l.get("app_id"), [])

    # Totales solo de líneas con saldo
    tot_kg   = round(sum(l["peso_kg_tot"]  for l in lineas_prod), 3)
    tot_pv   = round(sum(l["peso_vol_tot"] for l in lineas_prod), 3)
    tot_pred = round(sum(l["pred_tot"]     for l in lineas_prod), 3)
    pred_tipo = "kg" if tot_kg >= tot_pv else "pv"

    # ZZenvio cost
    costo_zz_envio = sum(
        float(l.get("pred_u") or 0) for l in lineas_zz
        if l["sku"].upper() == "ZZENVIO"
    )

    # Commune: prefer DB stored, then ERP header
    comuna = c["comuna"] or header.get("comuna") or ""
    if not comuna:
        # Try to extract from direccion
        dir_val = c["direccion"] or header.get("direccion") or ""
        if " - " in dir_val:
            parts = [p.strip() for p in dir_val.split(" - ")]
            if parts:
                comuna = parts[0].split(",")[-1].strip()

    return jsonify({
        "ok": True,
        "debug_fields": header.get("all_fields", []),
        "compromiso": {
            "id":            c["id"],
            "tido":          tido,
            "nudo":          nudo,
            "cliente":       c["cliente_nombre"] or header.get("cliente_nombre") or "—",
            "rut":           c["cliente_rut"] or header.get("cliente_rut") or "",
            "comuna":        comuna,
            "direccion":     c["direccion"] or header.get("direccion") or "",
            "telefono":      c["telefono"] or header.get("telefono") or "",
            "email":         c["email"] or header.get("email") or "",
            "costo_zz":      float(c["costo_zz"] or 0),
            "costo_zz_envio":costo_zz_envio,
            "clasificacion": c["clasificacion"] or "despacho",
            "fecha_emision": c["fecha_emision"].strftime("%d/%m/%Y") if c["fecha_emision"] else "",
            "estado":        c["estado"] or "",
        },
        "lineas": lineas_prod,
        "lineas_zz": lineas_zz,
        "totales": {
            "kg":   tot_kg,
            "pv":   tot_pv,
            "pred": tot_pred,
            "pred_tipo": pred_tipo,
        },
    })


@app.route("/transporte/api/compromisos/<int:cid>/lineas")
@_tr_required
def tr_lineas(cid):
    lineas = mysql_fetchall(
        "SELECT * FROM transport_commitment_lines WHERE commitment_id=%s", (cid,)
    )
    return jsonify([dict(l) for l in lineas])


@app.route("/transporte/api/compromisos/<int:cid>/logs")
@_tr_required
def tr_commitment_logs(cid):
    logs = mysql_fetchall(
        "SELECT * FROM transport_logs WHERE entity_type='commitment' AND entity_id=%s "
        "ORDER BY created_at DESC LIMIT 50", (cid,)
    )
    return jsonify([dict(l) for l in logs])


# ── MANIFIESTOS ──────────────────────────────────────────────────

@app.route("/transporte/manifiestos")
@_tr_required
def tr_manifiestos():
    filtros = {
        "courier": request.args.get("courier", ""),
        "estado":  request.args.get("estado", ""),
    }
    where, params = ["1=1"], []
    if filtros["courier"]:
        where.append("courier=%s"); params.append(filtros["courier"])
    if filtros["estado"]:
        where.append("estado=%s"); params.append(filtros["estado"])
    manifiestos = mysql_fetchall(
        "SELECT * FROM transport_manifests WHERE " + " AND ".join(where) +
        " ORDER BY fecha DESC, id DESC", tuple(params)
    )
    return render_template(
        "transporte/manifiestos.html",
        manifiestos=manifiestos,
        filtros=filtros,
        couriers=COURIERS,
        estados_manifest=["En preparación", "En curso", "Cerrado", "Entregado completo"],
    )


@app.route("/transporte/manifiestos/nuevo", methods=["POST"])
@_tr_required
def tr_crear_manifiesto():
    courier = request.form.get("courier", "").strip()
    fecha   = request.form.get("fecha", "").strip()
    notas   = request.form.get("notas", "").strip()
    if not courier or not fecha:
        return jsonify({"error": "courier y fecha son obligatorios"}), 400

    conn = get_db()
    try:
        with conn.cursor() as cur:
            # Correlativo auto: MAN-YYYY-NNNN
            cur.execute(
                "SELECT COUNT(*)+1 AS n FROM transport_manifests "
                "WHERE YEAR(created_at)=YEAR(NOW())"
            )
            n = (cur.fetchone() or {}).get("n", 1)
            from datetime import datetime as _dt
            corr = f"MAN-{_dt.now().year}-{int(n):04d}"
            cur.execute(
                "INSERT INTO transport_manifests (correlativo,fecha,courier,notas,created_by) "
                "VALUES (%s,%s,%s,%s,%s)",
                (corr, fecha, courier, notas, current_username())
            )
            mid = cur.lastrowid
        conn.commit()
    finally:
        conn.close()

    _tr_log("manifest", mid, "creado", f"courier={courier} fecha={fecha}")
    return jsonify({"ok": True, "id": mid, "correlativo": corr})


@app.route("/transporte/manifiestos/<int:mid>")
@_tr_required
def tr_manifiesto_detalle(mid):
    manifiesto = mysql_fetchone(
        "SELECT * FROM transport_manifests WHERE id=%s", (mid,)
    )
    if not manifiesto:
        flash("Manifiesto no encontrado", "danger")
        return redirect(url_for("tr_manifiestos"))

    items = mysql_fetchall("""
        SELECT mi.*, c.tido, c.nudo, c.cliente_nombre, c.comuna,
               c.direccion, c.valor_bruto, c.costo_zz, c.clasificacion
        FROM transport_manifest_items mi
        JOIN transport_commitments c ON c.id = mi.commitment_id
        WHERE mi.manifest_id=%s
        ORDER BY mi.orden, mi.id
    """, (mid,))

    logs = mysql_fetchall(
        "SELECT * FROM transport_logs WHERE entity_type='manifest' AND entity_id=%s "
        "ORDER BY created_at DESC LIMIT 30", (mid,)
    )
    return render_template(
        "transporte/manifiesto_detalle.html",
        manifiesto=manifiesto,
        items=items,
        logs=logs,
        estados_entrega=ESTADOS_ENTREGA,
        estados_manifest=["En preparación", "En curso", "Cerrado", "Entregado completo"],
        couriers=COURIERS,
    )


@app.route("/transporte/manifiestos/<int:mid>/items", methods=["POST"])
@_tr_required
def tr_agregar_item(mid):
    cid = request.get_json(silent=True, force=True).get("commitment_id")
    if not cid:
        return jsonify({"error": "commitment_id requerido"}), 400
    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT IGNORE INTO transport_manifest_items (manifest_id,commitment_id) "
                "VALUES (%s,%s)", (mid, cid)
            )
            # Recalcular totales
            cur.execute("""
                UPDATE transport_manifests m SET
                  total_items=(SELECT COUNT(*) FROM transport_manifest_items WHERE manifest_id=m.id),
                  costo_total=(SELECT COALESCE(SUM(c.costo_zz),0)
                               FROM transport_manifest_items mi
                               JOIN transport_commitments c ON c.id=mi.commitment_id
                               WHERE mi.manifest_id=m.id)
                WHERE m.id=%s
            """, (mid,))
        conn.commit()
    finally:
        conn.close()
    _tr_log("manifest", mid, "item agregado", f"commitment_id={cid}")
    return jsonify({"ok": True})


@app.route("/transporte/manifiestos/<int:mid>/items/<int:item_id>", methods=["DELETE"])
@_tr_required
def tr_quitar_item(mid, item_id):
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            "DELETE FROM transport_manifest_items WHERE id=%s AND manifest_id=%s",
            (item_id, mid)
        )
        cur.execute(
            "UPDATE transport_manifests SET "
            "total_items=(SELECT COUNT(*) FROM transport_manifest_items WHERE manifest_id=%s) "
            "WHERE id=%s", (mid, mid)
        )
    conn.commit()
    _tr_log("manifest", mid, "item eliminado", f"item_id={item_id}")
    return jsonify({"ok": True})


@app.route("/transporte/manifiestos/<int:mid>/items/<int:item_id>/estado", methods=["PUT"])
@_tr_required
def tr_estado_entrega(mid, item_id):
    estado = (request.get_json(silent=True) or {}).get("estado_entrega", "")
    if estado not in ESTADOS_ENTREGA:
        return jsonify({"error": "estado inválido"}), 400
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            "UPDATE transport_manifest_items SET estado_entrega=%s WHERE id=%s AND manifest_id=%s",
            (estado, item_id, mid)
        )
    conn.commit()
    _tr_log("manifest_item", item_id, "estado_entrega", estado)
    return jsonify({"ok": True})


@app.route("/transporte/manifiestos/<int:mid>/estado", methods=["PUT"])
@_tr_required
def tr_estado_manifiesto(mid):
    estado = (request.get_json(silent=True) or {}).get("estado", "")
    validos = ["En preparación", "En curso", "Cerrado", "Entregado completo"]
    if estado not in validos:
        return jsonify({"error": "estado inválido"}), 400
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute("UPDATE transport_manifests SET estado=%s WHERE id=%s", (estado, mid))
    conn.commit()
    _tr_log("manifest", mid, "estado cambiado", estado)
    return jsonify({"ok": True})


# ── MANIFIESTOS: asignar compromiso desde drag-drop ──────────────────────────

@app.route("/transporte/api/manifiestos/asignar", methods=["POST"])
@_tr_required
def tr_asignar_a_manifiesto():
    """Asigna uno o más compromisos a un manifiesto (crea si mid=None)."""
    data         = request.get_json(silent=True) or {}
    commitment_ids = data.get("commitment_ids", [])
    mid          = data.get("manifest_id")     # None → crear nuevo

    if not commitment_ids:
        return jsonify({"error": "sin compromisos"}), 400

    conn = get_db()
    with conn.cursor() as cur:
        # crear manifiesto nuevo si no se indicó uno
        if not mid:
            courier = data.get("courier", "Por asignar")
            fecha   = data.get("fecha") or __import__("datetime").date.today().isoformat()
            cur.execute(
                "SELECT COALESCE(MAX(CAST(SUBSTRING(correlativo,4) AS UNSIGNED)),0)+1 FROM transport_manifests"
            )
            num = cur.fetchone()[0] or 1
            correlativo = f"MAN{num:04d}"
            cur.execute(
                """INSERT INTO transport_manifests (correlativo, fecha, courier, created_by)
                   VALUES (%s,%s,%s,%s)""",
                (correlativo, fecha, courier, current_user.correo),
            )
            mid = cur.lastrowid
        else:
            correlativo = None

        added, dupes = 0, 0
        for cid in commitment_ids:
            try:
                cur.execute(
                    """INSERT IGNORE INTO transport_manifest_items
                       (manifest_id, commitment_id) VALUES (%s,%s)""",
                    (mid, cid),
                )
                if cur.rowcount:
                    added += 1
                else:
                    dupes += 1
            except Exception:
                dupes += 1

        # recalcular total_items
        cur.execute(
            "UPDATE transport_manifests SET total_items=(SELECT COUNT(*) FROM transport_manifest_items WHERE manifest_id=%s) WHERE id=%s",
            (mid, mid),
        )

    conn.commit()
    return jsonify({"ok": True, "manifest_id": mid, "correlativo": correlativo, "added": added, "duplicados": dupes})


# ── COURIERS — helper functions ──────────────────────────────────────────────

def _parse_peso_upper(col_header) -> float:
    """Converts a weight column header to its upper bound in kg."""
    s = str(col_header).strip().replace(',', '.').replace(' ', '').lower()
    # Special cases
    if '+' in s or 'mas' in s or 'more' in s:
        return 999999.0
    # Range like "100al499", "500-1999", "1001-5000"
    for sep in ['-', 'al']:
        if sep in s:
            parts = s.split(sep)
            try:
                return float(parts[-1])
            except Exception:
                pass
    # Single number
    try:
        return float(s)
    except Exception:
        return 999999.0


def _courier_tarifa_lookup(courier_id: int, comuna: str, peso_kg: float):
    """
    Returns the price (float) for a given courier+commune+weight, or None.
    """
    row = mysql_fetchone(
        "SELECT precios_json FROM transport_courier_comunas "
        "WHERE courier_id=%s AND LOWER(TRIM(comuna))=LOWER(TRIM(%s))",
        (courier_id, comuna)
    )
    if not row or not row.get('precios_json'):
        return None
    try:
        precios = json.loads(row['precios_json'])
    except Exception:
        return None
    # Build sorted bracket list
    brackets = []
    for key, price in precios.items():
        if price is None:
            continue
        upper = _parse_peso_upper(key)
        brackets.append((upper, float(price)))
    brackets.sort(key=lambda x: x[0])
    if not brackets:
        return None
    for upper, price in brackets:
        if peso_kg <= upper:
            return price
    return brackets[-1][1]  # over max weight → return last price


# ── COURIERS ─────────────────────────────────────────────────────────────────

@app.route("/transporte/couriers")
@_tr_required
def tr_couriers():
    couriers = mysql_fetchall(
        """SELECT c.*,
           COUNT(DISTINCT t.id) AS total_tarifas,
           COUNT(DISTINCT cc.id) AS total_comunas
           FROM transport_couriers c
           LEFT JOIN transport_courier_tarifas t ON t.courier_id=c.id AND t.activo=1
           LEFT JOIN transport_courier_comunas cc ON cc.courier_id=c.id
           GROUP BY c.id ORDER BY c.activo DESC, c.nombre""",
        ()
    )
    return render_template("transporte/couriers.html", couriers=couriers)


@app.route("/transporte/couriers/nuevo", methods=["POST"])
@_tr_required
def tr_courier_nuevo():
    d = request.form
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO transport_couriers
               (nombre,rut,contacto,telefono,email,tipo,notas,
                peso_max_bulto,peso_max_guia,vol_max_bulto,factor_vol)
               VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
            (
                d.get("nombre","").strip(),
                d.get("rut","").strip(),
                d.get("contacto","").strip(),
                d.get("telefono","").strip(),
                d.get("email","").strip(),
                d.get("tipo","nacional"),
                d.get("notas","").strip(),
                float(d.get("peso_max_bulto") or 0),
                float(d.get("peso_max_guia") or 0),
                float(d.get("vol_max_bulto") or 0),
                float(d.get("factor_vol") or 5000),
            ),
        )
    conn.commit()
    _invalidate_couriers_cache()
    flash("Courier creado correctamente.", "success")
    return redirect(url_for("tr_couriers"))


@app.route("/transporte/couriers/import-tariffs", methods=["POST"])
@_tr_required
def tr_couriers_import_tariffs():
    """Importa tarifas de couriers desde Excel multi-hoja.

    Sube un archivo Excel (.xlsx) con una hoja por courier. Cada hoja debe
    tener columnas: Comuna/Destino, Sucursal, Zona, Días Tránsito, y luego
    columnas numéricas (1, 2, ..., 99) y/o rangos ("100-499", "500-1999"...).

    Soporta el formato de Libro1.xlsx (FedEx/Clickex/Felca/Milling).
    Reemplaza todas las tarifas existentes del courier (DELETE + INSERT).
    """
    if not g.permissions.get("transporte_admin") and not g.permissions.get("superadmin"):
        flash("Necesitás permiso de transporte admin.", "danger")
        return redirect(url_for("tr_couriers"))

    if "file" not in request.files:
        flash("Debes seleccionar un archivo Excel.", "warning")
        return redirect(url_for("tr_couriers"))

    f = request.files["file"]
    if not f.filename:
        flash("Archivo vacío.", "warning")
        return redirect(url_for("tr_couriers"))
    if not f.filename.lower().endswith((".xlsx", ".xlsm")):
        flash("Solo .xlsx / .xlsm.", "warning")
        return redirect(url_for("tr_couriers"))

    # Guardar temporal
    import tempfile
    import courier_tariff_import as cti
    try:
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            f.save(tmp.name)
            tmp_path = tmp.name
        conn = get_db()
        result = cti.import_excel_to_db(tmp_path, conn)
        os.unlink(tmp_path)
        _invalidate_couriers_cache()
        total = sum(result.values())
        flash(
            f"Tarifas importadas: " +
            ", ".join(f"{k}={v}" for k, v in result.items()) +
            f" ({total} comunas total).",
            "success"
        )
    except Exception as e:
        flash(f"Error importando tarifas: {e}", "danger")
    return redirect(url_for("tr_couriers"))


@app.route("/transporte/couriers/<int:cid>", methods=["GET"])
@_tr_required
def tr_courier_ficha(cid):
    courier = mysql_fetchone("""
        SELECT c.*,
               COUNT(DISTINCT cc.id) AS total_comunas,
               COUNT(DISTINCT ct.id) AS total_contratos,
               MAX(cc.updated_at)    AS ultima_importacion
        FROM transport_couriers c
        LEFT JOIN transport_courier_comunas cc ON cc.courier_id=c.id
        LEFT JOIN transport_courier_contratos ct ON ct.courier_id=c.id AND ct.vigente=1
        WHERE c.id=%s GROUP BY c.id
    """, (cid,))
    if not courier:
        flash("Courier no encontrado.", "danger")
        return redirect(url_for("tr_couriers"))

    contratos = mysql_fetchall(
        "SELECT * FROM transport_courier_contratos WHERE courier_id=%s ORDER BY vigente DESC, created_at DESC",
        (cid,)
    ) or []

    comunas_sample = mysql_fetchall(
        "SELECT comuna, zona, region, dias_transito, precios_json FROM transport_courier_comunas "
        "WHERE courier_id=%s ORDER BY region, comuna LIMIT 10", (cid,)
    ) or []

    regions = mysql_fetchall(
        "SELECT DISTINCT region FROM transport_courier_comunas WHERE courier_id=%s AND region!='' ORDER BY region", (cid,)
    ) or []
    zonas = mysql_fetchall(
        "SELECT DISTINCT zona FROM transport_courier_comunas WHERE courier_id=%s AND zona!='' ORDER BY zona", (cid,)
    ) or []

    return render_template("transporte/courier_ficha.html",
        courier=courier,
        contratos=contratos,
        comunas_sample=comunas_sample,
        regions=[r['region'] for r in regions],
        zonas=[z['zona'] for z in zonas],
    )


@app.route("/transporte/couriers/<int:cid>", methods=["PUT"])
@_tr_required
def tr_courier_editar(cid):
    d = request.get_json(silent=True) or {}
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """UPDATE transport_couriers SET
               nombre=%s, nombre_fantasia=%s, rut=%s, giro=%s,
               contacto=%s, contacto_cargo=%s, telefono=%s, email=%s,
               tipo=%s, notas=%s, activo=%s,
               peso_max_bulto=%s, peso_max_guia=%s, vol_max_bulto=%s, factor_vol=%s,
               logo_url=%s, website=%s, direccion=%s
               WHERE id=%s""",
            (
                d.get("nombre","").strip(),
                (d.get("nombre_fantasia") or "").strip(),
                d.get("rut","").strip(),
                (d.get("giro") or "").strip(),
                d.get("contacto","").strip(),
                (d.get("contacto_cargo") or "").strip(),
                d.get("telefono","").strip(),
                d.get("email","").strip(),
                d.get("tipo","nacional"),
                d.get("notas","").strip(),
                1 if d.get("activo", True) else 0,
                float(d.get("peso_max_bulto") or 0),
                float(d.get("peso_max_guia") or 0),
                float(d.get("vol_max_bulto") or 0),
                float(d.get("factor_vol") or 5000),
                d.get("logo_url","").strip(),
                (d.get("website") or "").strip(),
                (d.get("direccion") or "").strip(),
                cid,
            ),
        )
    conn.commit()
    _invalidate_couriers_cache()
    return jsonify({"ok": True})


@app.route("/transporte/couriers/<int:cid>", methods=["DELETE"])
@_tr_required
def tr_courier_eliminar(cid):
    conn = get_db()
    with conn.cursor() as cur:
        # desactivar en lugar de borrar (preserva historial)
        cur.execute("UPDATE transport_couriers SET activo=0 WHERE id=%s", (cid,))
    conn.commit()
    _invalidate_couriers_cache()
    return jsonify({"ok": True})


@app.route("/transporte/couriers/<int:cid>/api", methods=["GET"])
@_tr_required
def tr_courier_api_data(cid):
    """Returns full courier data as JSON for the ficha page."""
    courier = mysql_fetchone("SELECT * FROM transport_couriers WHERE id=%s", (cid,))
    if not courier:
        return jsonify({"error": "Not found"}), 404
    return jsonify(dict(courier))


@app.route("/transporte/couriers/<int:cid>/contratos", methods=["POST"])
@_tr_required
def tr_courier_contrato_nuevo(cid):
    """Save/update contract info (no file, just metadata)."""
    d = request.get_json(silent=True) or {}
    conn = get_db()
    try:
        with conn.cursor() as cur:
            if d.get("vigente"):
                cur.execute("UPDATE transport_courier_contratos SET vigente=0 WHERE courier_id=%s", (cid,))
            cur.execute("""
                INSERT INTO transport_courier_contratos
                    (courier_id, nombre, descripcion, archivo_url, tipo, vigente, fecha_inicio, fecha_fin, subido_por)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
            """, (
                cid,
                (d.get("nombre") or "Contrato").strip(),
                (d.get("descripcion") or "").strip(),
                (d.get("archivo_url") or "").strip(),
                d.get("tipo", "contrato"),
                1 if d.get("vigente") else 0,
                d.get("fecha_inicio") or None,
                d.get("fecha_fin") or None,
                current_username(),
            ))
            new_id = cur.lastrowid
        conn.commit()
        return jsonify({"ok": True, "id": new_id})
    finally:
        conn.close()


@app.route("/transporte/couriers/<int:cid>/contratos/<int:kid>", methods=["PUT"])
@_tr_required
def tr_courier_contrato_update(cid, kid):
    d = request.get_json(silent=True) or {}
    conn = get_db()
    try:
        with conn.cursor() as cur:
            if d.get("vigente"):
                cur.execute("UPDATE transport_courier_contratos SET vigente=0 WHERE courier_id=%s", (cid,))
            cur.execute("""
                UPDATE transport_courier_contratos SET
                    nombre=%s, descripcion=%s, archivo_url=%s, tipo=%s,
                    vigente=%s, fecha_inicio=%s, fecha_fin=%s
                WHERE id=%s AND courier_id=%s
            """, (
                (d.get("nombre") or "Contrato").strip(),
                (d.get("descripcion") or "").strip(),
                (d.get("archivo_url") or "").strip(),
                d.get("tipo", "contrato"),
                1 if d.get("vigente") else 0,
                d.get("fecha_inicio") or None,
                d.get("fecha_fin") or None,
                kid, cid,
            ))
        conn.commit()
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/transporte/couriers/<int:cid>/contratos/<int:kid>", methods=["DELETE"])
@_tr_required
def tr_courier_contrato_delete(cid, kid):
    conn = get_db()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM transport_courier_contratos WHERE id=%s AND courier_id=%s", (kid, cid))
        conn.commit()
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/transporte/couriers/<int:cid>/logo", methods=["POST"])
@_tr_required
def tr_courier_logo(cid):
    """Upload logo via URL. Accepts JSON {logo_url, logo_type}."""
    d = request.get_json(silent=True) or {}
    logo_url  = (d.get("logo_url") or "").strip()
    logo_type = d.get("logo_type", "principal")

    col_map = {"principal": "logo_url", "square": "logo_square_url", "label": "logo_label_url"}
    col = col_map.get(logo_type, "logo_url")

    conn = get_db()
    try:
        with conn.cursor() as cur:
            for c_def in ["logo_square_url VARCHAR(400)", "logo_label_url VARCHAR(400)",
                          "nombre_fantasia VARCHAR(120)", "giro VARCHAR(150)",
                          "contacto_cargo VARCHAR(120)", "renovacion_automatica TINYINT(1) DEFAULT 0"]:
                try:
                    cur.execute(f"ALTER TABLE transport_couriers ADD COLUMN {c_def}")
                except Exception:
                    pass
            cur.execute(f"UPDATE transport_couriers SET {col}=%s WHERE id=%s", (logo_url, cid))
        conn.commit()
        return jsonify({"ok": True, "logo_url": logo_url})
    finally:
        conn.close()


@app.route("/transporte/couriers/<int:cid>/comunas/paginated", methods=["GET"])
@_tr_required
def tr_courier_comunas_paginated(cid):
    """Returns paginated, filterable commune pricing."""
    page   = max(1, request.args.get("page", 1, type=int))
    per    = 25
    search = (request.args.get("q") or "").strip()
    region = (request.args.get("region") or "").strip()
    zona   = (request.args.get("zona") or "").strip()

    conditions = ["courier_id=%s"]
    params = [cid]
    if search:
        conditions.append("LOWER(comuna) LIKE LOWER(%s)")
        params.append(f"%{search}%")
    if region:
        conditions.append("region=%s")
        params.append(region)
    if zona:
        conditions.append("zona=%s")
        params.append(zona)

    where  = " AND ".join(conditions)
    total  = (mysql_fetchone(f"SELECT COUNT(*) AS n FROM transport_courier_comunas WHERE {where}", params) or {}).get("n", 0)
    offset = (page - 1) * per
    rows   = mysql_fetchall(
        f"SELECT codigo, sucursal, comuna, zona, region, dias_transito, precios_json "
        f"FROM transport_courier_comunas WHERE {where} ORDER BY region, comuna LIMIT %s OFFSET %s",
        params + [per, offset]
    ) or []

    result = []
    for r in rows:
        item = dict(r)
        if r.get("precios_json"):
            try:
                item["precios"] = json.loads(r["precios_json"])
            except Exception:
                item["precios"] = {}
        else:
            item["precios"] = {}
        del item["precios_json"]
        result.append(item)

    return jsonify({
        "rows":  result,
        "total": total,
        "page":  page,
        "pages": max(1, (total + per - 1) // per),
        "per":   per,
    })


# ── COURIERS: IMPORT / EXPORT / LOOKUP / COMUNAS ─────────────────────────────

@app.route("/transporte/couriers/import", methods=["POST"])
@_tr_required
def transporte_couriers_import():
    """Import commune-based pricing from Excel."""
    import openpyxl, math

    file = request.files.get('file')
    if not file:
        return jsonify({"ok": False, "error": "No se recibió archivo"}), 400

    try:
        wb = openpyxl.load_workbook(file, data_only=True)
    except Exception as exc:
        return jsonify({"ok": False, "error": f"No se pudo leer el Excel: {exc}"}), 400

    # Map sheet names to courier names
    SHEET_MAP = {
        'felca':   'Transporte Felca',
        'milling': 'Transportes Melling',
        'clickex': 'Clickex',
    }

    conn = get_db()
    total_inserted = 0
    results = []

    try:
        with conn.cursor() as cur:
            for sheet_name in wb.sheetnames:
                courier_name = SHEET_MAP.get(sheet_name.lower().strip())
                if not courier_name:
                    results.append({"sheet": sheet_name, "skipped": True, "reason": "Hoja no reconocida"})
                    continue

                # Get or create courier
                cur.execute("SELECT id FROM transport_couriers WHERE LOWER(nombre)=LOWER(%s)", (courier_name,))
                row = cur.fetchone()
                if row:
                    courier_id = row['id']
                else:
                    cur.execute("INSERT INTO transport_couriers (nombre, tipo) VALUES (%s, 'nacional')", (courier_name,))
                    courier_id = cur.lastrowid

                ws = wb[sheet_name]
                header = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]

                is_clickex = sheet_name.lower().strip() == 'clickex'

                if is_clickex:
                    # Clickex: Región(0), Destino(1), Días(2), Lu-Sa(3-8), then prices from col 9
                    meta_cols = 9
                    comunas_col = 1    # "Destino"
                    region_col  = 0
                    dias_col    = 2
                    dias_entrega_cols = list(range(3, 9))  # Lu Ma Mi Ju Vi Sa
                    dias_names = ['Lu','Ma','Mi','Ju','Vi','Sa']
                else:
                    # Felca/Milling: Codigo(0), Sucursal(1), Comuna(2), ZONA(3), Dias(4), prices from col 5
                    meta_cols = 5
                    comunas_col = 2
                    sucursal_col = 1
                    codigo_col = 0
                    zona_col = 3
                    dias_col = 4

                price_headers = header[meta_cols:]

                # Delete existing data for this courier
                cur.execute("DELETE FROM transport_courier_comunas WHERE courier_id=%s", (courier_id,))

                inserted = 0
                for row_data in ws.iter_rows(min_row=2, values_only=True):
                    if all(v is None for v in row_data):
                        continue

                    if is_clickex:
                        region = str(row_data[region_col] or '').strip()
                        comuna = str(row_data[comunas_col] or '').strip()
                        dias_str = str(row_data[dias_col] or '').strip() if row_data[dias_col] is not None else ''
                        # Dias entrega
                        dias_list = []
                        for i, dn in zip(dias_entrega_cols, dias_names):
                            v = row_data[i]
                            if v and str(v).strip() not in ('', 'None'):
                                dias_list.append(dn)
                        dias_entrega = ','.join(dias_list)
                        sucursal = ''
                        codigo = ''
                        zona = ''
                    else:
                        codigo = str(row_data[codigo_col] or '').strip()
                        sucursal = str(row_data[sucursal_col] or '').strip()
                        comuna = str(row_data[comunas_col] or '').strip()
                        zona = str(row_data[zona_col] or '').strip()
                        dias_val = row_data[dias_col]
                        dias_str = '' if dias_val in ('#N/A', None) else str(dias_val).strip()
                        region = ''
                        dias_entrega = ''

                    if not comuna:
                        continue

                    # Build prices dict
                    price_data = row_data[meta_cols:]
                    precios = {}
                    for ph, pv in zip(price_headers, price_data):
                        if ph is None:
                            continue
                        key = str(ph).strip()
                        if pv is None or str(pv).strip() in ('#N/A', 'N/A', ''):
                            continue
                        try:
                            val_f = float(pv)
                            if not math.isnan(val_f):
                                precios[key] = round(val_f, 2)
                        except (TypeError, ValueError):
                            pass

                    if not precios:
                        continue

                    cur.execute("""
                        INSERT INTO transport_courier_comunas
                            (courier_id, codigo, sucursal, comuna, zona, region, dias_transito, dias_entrega, precios_json)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                        ON DUPLICATE KEY UPDATE
                            codigo=VALUES(codigo), sucursal=VALUES(sucursal),
                            zona=VALUES(zona), region=VALUES(region),
                            dias_transito=VALUES(dias_transito), dias_entrega=VALUES(dias_entrega),
                            precios_json=VALUES(precios_json)
                    """, (courier_id, codigo, sucursal, comuna, zona, region, dias_str, dias_entrega, json.dumps(precios, ensure_ascii=False)))
                    inserted += 1

                total_inserted += inserted
                results.append({"sheet": sheet_name, "courier": courier_name, "courier_id": courier_id, "inserted": inserted})

        conn.commit()
        return jsonify({"ok": True, "total": total_inserted, "results": results})
    except Exception as exc:
        import traceback; traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500
    finally:
        conn.close()


@app.route("/transporte/couriers/export", methods=["GET"])
@_tr_required
def transporte_couriers_export():
    """Export all courier pricing as Excel, one sheet per courier."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from io import BytesIO

    couriers = mysql_fetchall(
        "SELECT id, nombre FROM transport_couriers WHERE activo=1 ORDER BY nombre"
    ) or []

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    RED_FILL = PatternFill("solid", fgColor="CC0000")
    WHITE_FONT = Font(color="FFFFFF", bold=True)

    for c in couriers:
        rows = mysql_fetchall(
            "SELECT codigo, sucursal, comuna, zona, region, dias_transito, dias_entrega, precios_json "
            "FROM transport_courier_comunas WHERE courier_id=%s ORDER BY region, comuna",
            (c['id'],)
        ) or []
        if not rows:
            continue

        ws = wb.create_sheet(title=c['nombre'][:31])

        # Collect all weight keys
        all_keys = []
        for r in rows:
            if r.get('precios_json'):
                try:
                    keys = list(json.loads(r['precios_json']).keys())
                    for k in keys:
                        if k not in all_keys:
                            all_keys.append(k)
                except Exception:
                    pass

        # Headers
        headers = ['Codigo','Sucursal','Comuna','Zona','Region','Dias Transito','Dias Entrega'] + all_keys
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(1, ci, h)
            cell.font = WHITE_FONT
            cell.fill = RED_FILL
            cell.alignment = Alignment(horizontal='center')

        for ri, r in enumerate(rows, 2):
            precios = {}
            if r.get('precios_json'):
                try: precios = json.loads(r['precios_json'])
                except Exception: pass

            row_vals = [
                r.get('codigo',''), r.get('sucursal',''), r.get('comuna',''),
                r.get('zona',''), r.get('region',''), r.get('dias_transito',''), r.get('dias_entrega','')
            ] + [precios.get(k,'') for k in all_keys]

            for ci, v in enumerate(row_vals, 1):
                ws.cell(ri, ci, v)

    if not wb.sheetnames:
        ws = wb.create_sheet("Sin datos")
        ws.cell(1,1,"No hay tarifas importadas aún.")

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    from flask import send_file
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='ILUS_Tarifas_Couriers.xlsx'
    )


@app.route("/transporte/couriers/lookup", methods=["GET"])
@_tr_required
def transporte_couriers_lookup():
    """API: GET ?courier_id=X&comuna=Y&peso=Z → returns price."""
    courier_id = request.args.get('courier_id', type=int)
    comuna     = (request.args.get('comuna') or '').strip()
    peso       = request.args.get('peso', type=float)

    if not all([courier_id, comuna, peso]):
        return jsonify({"ok": False, "error": "Faltan parámetros: courier_id, comuna, peso"}), 400

    price = _courier_tarifa_lookup(courier_id, comuna, peso)
    if price is None:
        # Try partial match
        row = mysql_fetchone(
            "SELECT comuna, precios_json FROM transport_courier_comunas "
            "WHERE courier_id=%s AND LOWER(comuna) LIKE LOWER(%s) LIMIT 1",
            (courier_id, f"%{comuna}%")
        )
        if row:
            try:
                precios = json.loads(row['precios_json'])
                brackets = sorted(
                    [(float(_parse_peso_upper(k)), float(v)) for k, v in precios.items() if v is not None],
                    key=lambda x: x[0]
                )
                for upper, price in brackets:
                    if peso <= upper:
                        return jsonify({"ok": True, "precio": price, "comuna_matched": row['comuna'], "partial_match": True})
                return jsonify({"ok": True, "precio": brackets[-1][1], "comuna_matched": row['comuna'], "partial_match": True})
            except Exception:
                pass
        return jsonify({"ok": False, "error": f"No se encontró tarifa para '{comuna}' en este courier"}), 404

    return jsonify({"ok": True, "precio": price, "comuna": comuna, "peso": peso})


@app.route("/transporte/couriers/comunas", methods=["GET"])
@_tr_required
def transporte_couriers_comunas():
    """API: GET ?courier_id=X → list of communes for that courier."""
    courier_id = request.args.get('courier_id', type=int)
    if not courier_id:
        return jsonify([])
    rows = mysql_fetchall(
        "SELECT comuna, zona, region, dias_transito FROM transport_courier_comunas "
        "WHERE courier_id=%s ORDER BY region, comuna",
        (courier_id,)
    ) or []
    return jsonify(rows)


# ── TARIFAS DE COURIER ────────────────────────────────────────────────────────

@app.route("/transporte/couriers/<int:cid>/tarifas", methods=["POST"])
@_tr_required
def tr_tarifa_nueva(cid):
    d = request.get_json(silent=True) or {}
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO transport_courier_tarifas
               (courier_id,zona,peso_desde,peso_hasta,precio_base,precio_kg_extra,moneda)
               VALUES (%s,%s,%s,%s,%s,%s,%s)""",
            (
                cid,
                d.get("zona","General").strip(),
                float(d.get("peso_desde") or 0),
                float(d.get("peso_hasta") or 0),
                float(d.get("precio_base") or 0),
                float(d.get("precio_kg_extra") or 0),
                d.get("moneda","CLP"),
            ),
        )
        tid = cur.lastrowid
    conn.commit()
    _invalidate_couriers_cache()
    return jsonify({"ok": True, "id": tid})


@app.route("/transporte/couriers/<int:cid>/tarifas/<int:tid>", methods=["PUT"])
@_tr_required
def tr_tarifa_editar(cid, tid):
    d = request.get_json(silent=True) or {}
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """UPDATE transport_courier_tarifas SET
               zona=%s,peso_desde=%s,peso_hasta=%s,
               precio_base=%s,precio_kg_extra=%s,moneda=%s,activo=%s
               WHERE id=%s AND courier_id=%s""",
            (
                d.get("zona","General").strip(),
                float(d.get("peso_desde") or 0),
                float(d.get("peso_hasta") or 0),
                float(d.get("precio_base") or 0),
                float(d.get("precio_kg_extra") or 0),
                d.get("moneda","CLP"),
                1 if d.get("activo", True) else 0,
                tid, cid,
            ),
        )
    conn.commit()
    _invalidate_couriers_cache()
    return jsonify({"ok": True})


@app.route("/transporte/couriers/<int:cid>/tarifas/<int:tid>", methods=["DELETE"])
@_tr_required
def tr_tarifa_eliminar(cid, tid):
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            "DELETE FROM transport_courier_tarifas WHERE id=%s AND courier_id=%s",
            (tid, cid),
        )
    conn.commit()
    return jsonify({"ok": True})


# ── COTIZADOR: calcular costo de envío ────────────────────────────────────────

@app.route("/transporte/api/cotizar", methods=["POST"])
@_tr_required
def tr_cotizar():
    """
    Calcula el costo estimado de envío con cada courier activo.
    Body JSON: {commitment_id, courier_id (opcional, si no devuelve todos)}
    """
    data = request.get_json(silent=True) or {}
    cid  = data.get("commitment_id")
    solo_courier = data.get("courier_id")

    if not cid:
        return jsonify({"error": "commitment_id requerido"}), 400

    # OPTIMIZACIÓN: filtramos SKUs antes de calcular peso/vol (evita JOIN sin filtro)
    skus_rows = mysql_fetchall(
        """SELECT DISTINCT UPPER(TRIM(koprct)) AS sku, saldo
            FROM transport_commitment_lines
            WHERE commitment_id=%s AND saldo>0 AND koprct IS NOT NULL AND koprct<>''""",
        (cid,),
    ) or []
    if not skus_rows:
        peso_real = vol_total = 0
    else:
        skus_unicos = list({r["sku"] for r in skus_rows})
        ph_sku = ",".join(["%s"] * len(skus_unicos))
        bk_rows = mysql_fetchall(
            f"""SELECT UPPER(TRIM(p.sku)) AS sku,
                       COALESCE(SUM(b.peso),0) AS peso_total,
                       COALESCE(SUM(b.largo * b.ancho * b.alto),0) AS volumen_cm3
                FROM `{PRODUCTS_TABLE}` p
                LEFT JOIN `{BULTOS_TABLE}` b ON b.product_id=p.id
                WHERE UPPER(TRIM(p.sku)) IN ({ph_sku})
                GROUP BY p.sku""",
            tuple(skus_unicos),
        ) or []
        bk_map = { r["sku"]: r for r in bk_rows }
        peso_real = 0.0
        vol_total = 0.0
        for r in skus_rows:
            bk = bk_map.get(r["sku"])
            if bk:
                peso_real += float(r["saldo"] or 0) * float(bk["peso_total"] or 0)
                vol_total += float(r["saldo"] or 0) * float(bk["volumen_cm3"] or 0)

    # Couriers: cacheado si pedimos todos, query directa si filtramos uno
    if solo_courier:
        couriers = mysql_fetchall(
            """SELECT c.*, GROUP_CONCAT(
                   CONCAT_WS('|',t.zona,t.peso_desde,t.peso_hasta,t.precio_base,t.precio_kg_extra)
                   ORDER BY t.zona,t.peso_desde SEPARATOR ';;'
               ) AS tarifas_raw
               FROM transport_couriers c
               LEFT JOIN transport_courier_tarifas t ON t.courier_id=c.id AND t.activo=1
               WHERE c.activo=1 AND c.id=%s
               GROUP BY c.id""",
            (solo_courier,),
        )
    else:
        couriers = _get_couriers_cached()

    resultados = []
    for co in couriers:
        factor_vol = float(co["factor_vol"] or 5000)
        peso_vol   = vol_total / factor_vol if vol_total else 0
        peso_pred  = max(peso_real, peso_vol)

        # validar restricciones
        advertencias = []
        if co["peso_max_bulto"] and peso_real > float(co["peso_max_bulto"]):
            advertencias.append(f"Excede peso máx por bulto ({co['peso_max_bulto']} kg)")
        if co["peso_max_guia"] and peso_pred > float(co["peso_max_guia"]):
            advertencias.append(f"Excede peso máx por guía ({co['peso_max_guia']} kg)")

        # calcular tarifa
        costo = None
        zona_aplicada = None
        zona_req = data.get("zona", "General")
        if co["tarifas_raw"]:
            for bloque in co["tarifas_raw"].split(";;"):
                partes = bloque.split("|")
                if len(partes) < 5:
                    continue
                zona_t, pd, ph, pb, pke = partes
                pd, ph, pb, pke = float(pd), float(ph), float(pb), float(pke)
                if zona_t not in (zona_req, "General"):
                    continue
                if pd <= peso_pred and (ph == 0 or peso_pred <= ph):
                    extra = max(0, peso_pred - pd) * pke if pke else 0
                    costo = pb + extra
                    zona_aplicada = zona_t
                    break

        resultados.append({
            "courier_id"   : co["id"],
            "nombre"       : co["nombre"],
            "peso_real"    : round(peso_real, 3),
            "peso_vol"     : round(peso_vol, 3),
            "peso_pred"    : round(peso_pred, 3),
            "costo"        : round(costo, 0) if costo is not None else None,
            "zona"         : zona_aplicada,
            "advertencias" : advertencias,
        })

    resultados.sort(key=lambda x: (x["costo"] is None, x["costo"] or 999999))
    return jsonify({"ok": True, "resultados": resultados})


# ── COTIZAR COLA (batch: varios commitment_ids) ───────────────────────────────

@app.route("/transporte/api/cola/cotizar", methods=["POST"])
@_tr_required
def tr_cola_cotizar():
    """
    Calcula peso/volumen acumulado para un array de commitment_ids
    y devuelve cotización con todos los couriers activos.
    Usado por el panel manifiesto para mostrar prevale en tiempo real.
    """
    data = request.get_json(silent=True) or {}
    cids = [int(c) for c in (data.get("commitment_ids") or []) if str(c).isdigit()]

    if not cids:
        return jsonify({"ok": True, "peso_real": 0, "vol_total": 0,
                        "peso_vol": 0, "peso_pred": 0, "resultados": []})

    ph = ",".join(["%s"] * len(cids))

    # OPTIMIZACIÓN: en vez de un subquery que materializa TODO products×bultos,
    # primero traemos los SKUs únicos del compromiso y filtramos directo.
    # Pasa de ~tabla completa de products (cientos/miles) a solo los SKUs
    # presentes en estos commitments (típicamente <30).
    skus_rows = mysql_fetchall(
        f"""SELECT DISTINCT UPPER(TRIM(koprct)) AS sku, saldo
            FROM transport_commitment_lines
            WHERE commitment_id IN ({ph}) AND saldo > 0 AND koprct IS NOT NULL AND koprct <> ''""",
        tuple(cids),
    ) or []

    if not skus_rows:
        peso_real = vol_total = 0
    else:
        skus_unicos = list({r["sku"] for r in skus_rows})
        ph_sku = ",".join(["%s"] * len(skus_unicos))

        # Una sola query batch: trae peso/vol por SKU solo de los relevantes
        bk_rows = mysql_fetchall(
            f"""SELECT UPPER(TRIM(p.sku)) AS sku,
                       COALESCE(SUM(b.peso), 0)                     AS peso_total,
                       COALESCE(SUM(b.largo * b.ancho * b.alto), 0) AS vol_cm3
                FROM `{PRODUCTS_TABLE}` p
                LEFT JOIN `{BULTOS_TABLE}` b ON b.product_id = p.id
                WHERE UPPER(TRIM(p.sku)) IN ({ph_sku})
                GROUP BY p.sku""",
            tuple(skus_unicos),
        ) or []
        bk_map = { r["sku"]: r for r in bk_rows }

        # Sumar en memoria — más rápido que JOIN sin filtro
        peso_real = 0.0
        vol_total = 0.0
        for r in skus_rows:
            bk = bk_map.get(r["sku"])
            if bk:
                peso_real += float(r["saldo"] or 0) * float(bk["peso_total"] or 0)
                vol_total += float(r["saldo"] or 0) * float(bk["vol_cm3"] or 0)

    # Couriers activos con sus tarifas — cacheado 60s (cambian rara vez)
    couriers = _get_couriers_cached()

    zona_req = data.get("zona", "General")
    resultados = []
    factor_base = 5000  # factor volumen por defecto

    for co in couriers:
        factor_vol = float(co["factor_vol"] or 5000)
        factor_base = factor_vol  # usar último valor para resumen
        peso_vol  = vol_total / factor_vol if vol_total else 0
        peso_pred = max(peso_real, peso_vol)

        advertencias = []
        if co["peso_max_guia"] and peso_pred > float(co["peso_max_guia"]):
            advertencias.append(f"Excede máx/guía ({co['peso_max_guia']} kg)")

        costo = None
        if co["tarifas_raw"]:
            for bloque in co["tarifas_raw"].split(";;"):
                partes = bloque.split("|")
                if len(partes) < 5:
                    continue
                zona_t, pd, ph2, pb, pke = partes
                pd, ph2, pb, pke = float(pd), float(ph2), float(pb), float(pke)
                if zona_t not in (zona_req, "General"):
                    continue
                if pd <= peso_pred and (ph2 == 0 or peso_pred <= ph2):
                    extra = max(0, peso_pred - pd) * pke if pke else 0
                    costo = pb + extra
                    break

        resultados.append({
            "courier_id"   : co["id"],
            "nombre"       : co["nombre"],
            "peso_max_guia": float(co["peso_max_guia"] or 0),
            "costo"        : round(costo, 0) if costo is not None else None,
            "advertencias" : advertencias,
        })

    resultados.sort(key=lambda x: (x["costo"] is None, x["costo"] or 999_999))

    peso_vol_def  = vol_total / factor_base if vol_total else 0
    peso_pred_def = max(peso_real, peso_vol_def)

    # Conteo de ítems por compromiso (para mostrar en el panel)
    items_rows = mysql_fetchall(
        f"""SELECT commitment_id,
               COUNT(*)                    AS n_lineas,
               COALESCE(SUM(saldo), 0)     AS total_saldo
           FROM transport_commitment_lines
           WHERE commitment_id IN ({",".join(["%s"]*len(cids))}) AND saldo > 0
           GROUP BY commitment_id""",
        tuple(cids),
    )
    items_by_cid = {r["commitment_id"]: {"n": int(r["n_lineas"]), "saldo": int(r["total_saldo"])}
                    for r in items_rows}

    return jsonify({
        "ok"        : True,
        "peso_real" : round(peso_real, 2),
        "vol_total" : round(vol_total, 0),
        "peso_vol"  : round(peso_vol_def, 2),
        "peso_pred" : round(peso_pred_def, 2),
        "resultados": resultados,
        "items"     : items_by_cid,
    })


# ── MANIFIESTOS: listar activos para panel drag-drop ─────────────────────────

@app.route("/transporte/api/manifiestos/activos")
@_tr_required
def tr_manifiestos_activos():
    rows = mysql_fetchall(
        """SELECT m.id, m.correlativo, m.fecha, m.courier,
                  m.estado, m.total_items
           FROM transport_manifests m
           WHERE m.estado IN ('En preparación','En curso')
           ORDER BY m.fecha DESC, m.id DESC LIMIT 50""",
        (),
    )
    return jsonify([dict(r) for r in rows])


# ══════════════════════════════════════════════════════════════
#  MÓDULO: COMUNICACIONES (solo superadmin)
#  Email (SMTP dinámico) + WhatsApp (Twilio)
# ══════════════════════════════════════════════════════════════

def init_comunicaciones_tables():
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_smtp_config (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    smtp_host   VARCHAR(200) DEFAULT 'smtp.gmail.com',
                    smtp_port   INT DEFAULT 587,
                    smtp_user   VARCHAR(200),
                    smtp_pass   VARCHAR(500),
                    from_name   VARCHAR(200) DEFAULT 'ILUS Sport & Health',
                    from_addr   VARCHAR(200),
                    secure      TINYINT(1) DEFAULT 0,
                    updated_by  VARCHAR(190),
                    updated_at  DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_client_config (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    company_name    VARCHAR(200) DEFAULT 'ILUS Sport & Health',
                    reply_to        VARCHAR(200),
                    support_email   VARCHAR(200),
                    support_phone   VARCHAR(50),
                    tracking_url    VARCHAR(400),
                    logo_url        MEDIUMTEXT,
                    corp_color      VARCHAR(20) DEFAULT '#CC0000',
                    email_cc        VARCHAR(500),
                    email_bcc       VARCHAR(500),
                    updated_by      VARCHAR(190),
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_whatsapp_config (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    account_sid     VARCHAR(120),
                    auth_token      VARCHAR(250),
                    from_number     VARCHAR(50),
                    biz_number      VARCHAR(50),
                    updated_by      VARCHAR(190),
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_resend_config (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    api_key         VARCHAR(250),
                    from_addr       VARCHAR(200),
                    updated_by      VARCHAR(190),
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_log (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    canal           ENUM('email','whatsapp') NOT NULL,
                    destinatario    VARCHAR(300),
                    asunto          VARCHAR(500),
                    estado          ENUM('ok','error') NOT NULL,
                    detalle         TEXT,
                    enviado_por     VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_canal (canal),
                    INDEX idx_fecha (created_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── PLANTILLAS POR ESTADO ─────────────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS comm_templates (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    modulo      VARCHAR(40) NOT NULL DEFAULT 'transporte'
                                COMMENT 'transporte | retiros | mantenciones',
                    estado      VARCHAR(60) NOT NULL,
                    canal       ENUM('email','whatsapp') NOT NULL,
                    asunto      VARCHAR(300),
                    cuerpo      MEDIUMTEXT,
                    activo      TINYINT(1) DEFAULT 1,
                    updated_by  VARCHAR(190),
                    updated_at  DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    UNIQUE KEY uq_mod_estado_canal (modulo, estado, canal),
                    INDEX idx_modulo (modulo)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migración para tablas existentes (que no tenían columna 'modulo')
            for _mig in [
                "ALTER TABLE comm_templates ADD COLUMN modulo VARCHAR(40) NOT NULL DEFAULT 'transporte' AFTER id",
                "ALTER TABLE comm_templates ADD INDEX idx_modulo (modulo)",
                # Eliminar la UNIQUE vieja (estado, canal) para permitir duplicados entre módulos
                "ALTER TABLE comm_templates DROP INDEX uq_estado_canal",
                "ALTER TABLE comm_templates ADD UNIQUE KEY uq_mod_estado_canal (modulo, estado, canal)",
            ]:
                try: cur.execute(_mig)
                except Exception: pass
            # Insertar plantillas por defecto si la tabla está vacía
            cur.execute("SELECT COUNT(*) AS n FROM comm_templates")
            tpl_row = cur.fetchone()
            if (tpl_row or {}).get('n', 1) == 0:
                _ESTADOS = [
                    ('programado',       'Pedido programado'),
                    ('en_ruta',          'Pedido en ruta'),
                    ('en_camino',        'En camino — próxima entrega'),
                    ('entregado',        'Pedido entregado'),
                    ('fallido',          'Intento de entrega fallido'),
                    ('inicio_sesion',    'Inicio de sesión en tu cuenta'),
                    ('cambio_pass',      'Cambio de contraseña'),
                ]
                _EMAIL_DEFAULTS = {
                    'programado':    ('Tu pedido {{id_pedido}} fue programado',
                                     'Hola {{nombre_cliente}}, tu pedido <strong>{{id_pedido}}</strong> ha sido programado y está siendo preparado. Te avisaremos cuando salga a despacho.'),
                    'en_ruta':       ('Tu pedido {{id_pedido}} está en ruta',
                                     'Hola {{nombre_cliente}}, tu pedido <strong>{{id_pedido}}</strong> ya salió de bodega y está en camino con {{courier}}. N° de seguimiento: <strong>{{numero_seguimiento}}</strong>.'),
                    'en_camino':     ('¡Tu pedido llega hoy! — {{id_pedido}}',
                                     'Hola {{nombre_cliente}}, tu pedido <strong>{{id_pedido}}</strong> está en reparto y llegará hoy a <em>{{direccion_entrega}}</em>.'),
                    'entregado':     ('Pedido {{id_pedido}} entregado ✓',
                                     'Hola {{nombre_cliente}}, tu pedido <strong>{{id_pedido}}</strong> fue entregado exitosamente. ¡Gracias por tu confianza en ILUS!'),
                    'fallido':       ('No pudimos entregar tu pedido {{id_pedido}}',
                                     'Hola {{nombre_cliente}}, intentamos entregar tu pedido <strong>{{id_pedido}}</strong> pero no fue posible. Nos pondremos en contacto para coordinar una nueva entrega.'),
                    'inicio_sesion': ('Nuevo inicio de sesión en tu cuenta',
                                     'Hola {{nombre_usuario}}, detectamos un inicio de sesión en tu cuenta el {{fecha_hora}}. Si no fuiste tú, contáctanos de inmediato.'),
                    'cambio_pass':   ('Tu contraseña fue cambiada',
                                     'Hola {{nombre_usuario}}, tu contraseña fue actualizada el {{fecha_hora}}. Si no realizaste este cambio, contáctanos inmediatamente.'),
                }
                _WA_DEFAULTS = {
                    'programado':    '🟡 *ILUS* — Hola {{nombre_cliente}}, tu pedido *{{id_pedido}}* fue programado y está en preparación.',
                    'en_ruta':       '🚚 *ILUS* — Tu pedido *{{id_pedido}}* salió de bodega con {{courier}}. Seguimiento: {{numero_seguimiento}}',
                    'en_camino':     '📦 *ILUS* — ¡Tu pedido *{{id_pedido}}* llega hoy! Dirección: {{direccion_entrega}}',
                    'entregado':     '✅ *ILUS* — Tu pedido *{{id_pedido}}* fue entregado. ¡Gracias {{nombre_cliente}}!',
                    'fallido':       '❌ *ILUS* — No pudimos entregar tu pedido *{{id_pedido}}*. Te contactaremos para reagendar.',
                    'inicio_sesion': '🔐 *ILUS* — Inicio de sesión detectado en tu cuenta el {{fecha_hora}}.',
                    'cambio_pass':   '🔑 *ILUS* — Tu contraseña fue cambiada el {{fecha_hora}}.',
                }
                for estado_key, _ in _ESTADOS:
                    asunto, cuerpo = _EMAIL_DEFAULTS.get(estado_key, ('', ''))
                    cur.execute(
                        "INSERT INTO comm_templates (modulo, estado, canal, asunto, cuerpo) VALUES ('transporte',%s,'email',%s,%s)",
                        (estado_key, asunto, cuerpo)
                    )
                    wa_body = _WA_DEFAULTS.get(estado_key, '')
                    cur.execute(
                        "INSERT INTO comm_templates (modulo, estado, canal, asunto, cuerpo) VALUES ('transporte',%s,'whatsapp',%s,%s)",
                        (estado_key, '', wa_body)
                    )

            # ── PLANTILLAS DE RETIROS (siembra idempotente: solo si no existen) ──
            # Variables disponibles: {{code}}, {{cliente}}, {{persona_retira}},
            # {{fecha_solicitada}}, {{fecha_propuesta}}, {{fecha_confirmada}},
            # {{horario}}, {{documento}}, {{n_bultos}}, {{kg}}, {{m3}},
            # {{link_seguimiento}}, {{warehouse_name}}, {{warehouse_addr}}
            _RETIRO_TPL = [
                # (estado, canal, asunto, cuerpo)
                ('solicitud_recibida', 'email',
                 '✓ Recibimos tu solicitud de retiro {{code}}',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Recibimos tu solicitud de retiro <strong>{{code}}</strong> para el {{fecha_solicitada}} entre {{horario}}.</p>'
                 '<p>Estamos validando la disponibilidad y te confirmaremos en menos de 2 horas hábiles.</p>'
                 '<p>Puedes revisar el estado en cualquier momento aquí: <a href="{{link_seguimiento}}">Ver mi retiro</a></p>'
                 '<p>Saludos,<br>Equipo ILUS</p>'),
                ('solicitud_recibida', 'whatsapp',
                 '',
                 '🟢 *ILUS* — Recibimos tu solicitud de retiro *{{code}}* para el {{fecha_solicitada}} ({{horario}}). '
                 'Te confirmaremos pronto. Seguimiento: {{link_seguimiento}}'),

                ('propuesta_enviada', 'email',
                 '📅 ILUS te propone una fecha para el retiro {{code}}',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Para tu retiro <strong>{{code}}</strong> ({{documento}} de {{cliente}}), '
                 'ILUS propone la siguiente fecha:</p>'
                 '<p style="background:#fff7ed;padding:14px;border-left:4px solid #fb923c;border-radius:6px">'
                 '<strong>📆 {{fecha_propuesta}}</strong> · ⏰ {{horario}}</p>'
                 '<p>Por favor confirma o propone otra fecha en este enlace:<br>'
                 '<a href="{{link_seguimiento}}" style="background:#e60000;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;font-weight:bold">Responder propuesta</a></p>'
                 '<p>Saludos,<br>Equipo ILUS</p>'),
                ('propuesta_enviada', 'whatsapp',
                 '',
                 '📅 *ILUS* — Para tu retiro *{{code}}* proponemos: {{fecha_propuesta}} ({{horario}}). '
                 'Confirma aquí: {{link_seguimiento}}'),

                ('agenda_confirmada', 'email',
                 '✅ Retiro {{code}} confirmado para el {{fecha_confirmada}}',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Tu retiro <strong>{{code}}</strong> está confirmado:</p>'
                 '<p style="background:#dcfce7;padding:14px;border-left:4px solid #22c55e;border-radius:6px">'
                 '<strong>📆 {{fecha_confirmada}}</strong> · ⏰ {{horario}}<br>'
                 '<i>📍 {{warehouse_name}}</i><br>'
                 '<small>{{warehouse_addr}}</small></p>'
                 '<p><strong>Recuerda llevar:</strong></p>'
                 '<ul><li>Tu cédula de identidad (debe coincidir con la persona autorizada)</li>'
                 '<li>Vehículo apto para {{n_bultos}} bulto(s) ({{kg}} kg / {{m3}} m³)</li></ul>'
                 '<p>Seguimiento: <a href="{{link_seguimiento}}">{{link_seguimiento}}</a></p>'
                 '<p>Te esperamos.<br>Equipo ILUS</p>'),
                ('agenda_confirmada', 'whatsapp',
                 '',
                 '✅ *ILUS* — Retiro *{{code}}* confirmado: {{fecha_confirmada}} ({{horario}}). '
                 'Lugar: {{warehouse_name}}. Lleva tu carnet. Seguimiento: {{link_seguimiento}}'),

                ('en_preparacion', 'email',
                 '📦 Tu retiro {{code}} está siendo preparado',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Bodega ya está preparando tu retiro <strong>{{code}}</strong> para el {{fecha_confirmada}} ({{horario}}).</p>'
                 '<p>Cuando llegues, presenta tu cédula al ingreso. Te esperamos a tiempo.</p>'
                 '<p>Saludos,<br>Equipo ILUS</p>'),
                ('en_preparacion', 'whatsapp',
                 '',
                 '📦 *ILUS* — Bodega ya está preparando tu retiro *{{code}}* para mañana ({{horario}}). Trae tu carnet.'),

                ('retirada', 'email',
                 '✅ Retiro {{code}} completado — gracias',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Confirmamos la entrega de tu retiro <strong>{{code}}</strong>. '
                 'Esperamos que disfrutes tu equipo ILUS.</p>'
                 '<p>Si necesitas asesoría o un service post-venta, escríbenos. Estamos para ayudarte.</p>'
                 '<p>¡Gracias por confiar en ILUS!<br>Equipo ILUS</p>'),
                ('retirada', 'whatsapp',
                 '',
                 '✅ *ILUS* — Retiro *{{code}}* completado. ¡Gracias por confiar en nosotros! 💪'),

                ('rechazada', 'email',
                 'Solicitud de retiro {{code}} cancelada',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Tu solicitud de retiro <strong>{{code}}</strong> fue cancelada. '
                 'Si fue un error o quieres reagendar, escríbenos respondiendo este correo.</p>'
                 '<p>Saludos,<br>Equipo ILUS</p>'),
                ('rechazada', 'whatsapp',
                 '',
                 '⚠ *ILUS* — Tu solicitud de retiro *{{code}}* fue cancelada. Si fue error, escríbenos.'),

                ('reagendada', 'email',
                 '🔄 Retiro {{code}} reagendado',
                 '<p>Hola <strong>{{persona_retira}}</strong>,</p>'
                 '<p>Tu retiro <strong>{{code}}</strong> ha sido reagendado.</p>'
                 '<p>Nueva fecha propuesta: <strong>{{fecha_propuesta}}</strong> · {{horario}}</p>'
                 '<p>Confirma aquí: <a href="{{link_seguimiento}}">{{link_seguimiento}}</a></p>'
                 '<p>Saludos,<br>Equipo ILUS</p>'),
                ('reagendada', 'whatsapp',
                 '',
                 '🔄 *ILUS* — Retiro *{{code}}* reagendado para {{fecha_propuesta}} ({{horario}}). Confirma: {{link_seguimiento}}'),
            ]
            for _est, _can, _asu, _cue in _RETIRO_TPL:
                try:
                    cur.execute(
                        "INSERT IGNORE INTO comm_templates (modulo, estado, canal, asunto, cuerpo) "
                        "VALUES ('retiros', %s, %s, %s, %s)",
                        (_est, _can, _asu, _cue)
                    )
                except Exception: pass
            # ── Normalización: dejar máximo UNA fila por tabla (id=1) ────
            # Si por bugs anteriores quedaron filas duplicadas, conservamos
            # la más reciente y la forzamos a id=1 para que la lectura
            # siempre devuelva el dato correcto.
            for tbl in ("comm_smtp_config", "comm_client_config",
                        "comm_whatsapp_config", "comm_resend_config"):
                try:
                    cur.execute(f"SELECT COUNT(*) AS n FROM {tbl}")
                    r = cur.fetchone() or {}
                    n = r.get("n", 0) if isinstance(r, dict) else (r[0] if r else 0)
                    if n > 1:
                        cur.execute(f"SELECT * FROM {tbl} ORDER BY id DESC LIMIT 1")
                        keep = cur.fetchone()
                        if keep:
                            cols = [c for c in keep.keys() if c not in ("id","updated_at")]
                            cur.execute(f"DELETE FROM {tbl}")
                            cols_sql  = ", ".join(["id"] + cols)
                            place_sql = ", ".join(["%s"] * (len(cols) + 1))
                            vals = [1] + [keep[c] for c in cols]
                            cur.execute(
                                f"INSERT INTO {tbl} ({cols_sql}) VALUES ({place_sql})",
                                vals
                            )
                            print(f"[ILUS] {tbl}: normalizada a fila unica (id=1)")
                    elif n == 1:
                        cur.execute(f"UPDATE {tbl} SET id=1 WHERE id<>1")
                except Exception as _norm_err:
                    print(f"[ILUS][WARN] Normalizacion {tbl}: {_norm_err}")
        conn.commit()
    finally:
        conn.close()


# ── Helpers de config ─────────────────────────────────────────

def _comm_template_defaults():
    def despacho_body(intro, rows, note="", color="#CC0000"):
        row_html = "\n".join(
            "<tr><td style='padding:6px 0;font-size:13px;color:#555'>"
            f"<strong style='color:#222'>{label}:</strong>&nbsp; {value}</td></tr>"
            for label, value in rows
        )
        note_html = (
            f"<p style='margin:18px 0 0;font-size:13px;color:#777;line-height:1.55'>{note}</p>"
            if note else ""
        )
        return (
            f"<p style='margin:0 0 16px;font-size:15px;color:{color};font-weight:700'>Hola, {{{{nombre_cliente}}}}</p>\n"
            f"<p style='margin:0 0 14px;font-size:14px;color:#444;line-height:1.65'>{intro}</p>\n"
            f"<table cellpadding='0' cellspacing='0' width='100%' style='background:#f5f5f7;border-left:4px solid {color};"
            "border-radius:4px;padding:14px 18px;margin:18px 0'>\n"
            f"{row_html}\n</table>\n{note_html}"
        )

    def sistema_body(intro, rows, note="", color="#CC0000"):
        row_html = "\n".join(
            "<tr><td style='padding:6px 0;font-size:13px;color:#555'>"
            f"<strong style='color:#222'>{label}:</strong>&nbsp; {value}</td></tr>"
            for label, value in rows
        )
        note_html = (
            f"<div style='background:#fff3cd;border:1px solid #ffc107;border-radius:6px;padding:12px 16px;"
            f"font-size:13px;color:#664d03;margin-top:16px'>{note}</div>"
            if note else ""
        )
        return (
            f"<p style='margin:0 0 16px;font-size:15px;color:{color};font-weight:700'>Hola, {{{{nombre_usuario}}}}</p>\n"
            f"<p style='margin:0 0 14px;font-size:14px;color:#444;line-height:1.65'>{intro}</p>\n"
            f"<table cellpadding='0' cellspacing='0' width='100%' style='background:#f5f5f7;border-left:4px solid {color};"
            "border-radius:4px;padding:14px 18px;margin:18px 0'>\n"
            f"{row_html}\n</table>\n{note_html}"
        )

    return {
        "programado": {
            "email": (
                "Tu pedido {{id_pedido}} ha sido programado - ILUS",
                despacho_body(
                    "Tu pedido ha sido <strong>programado exitosamente</strong> y esta siendo preparado en nuestra bodega.",
                    [
                        ("Nro. pedido", "{{id_pedido}}"),
                        ("Courier", "{{courier}}"),
                        ("Entrega estimada", "{{fecha_entrega}}"),
                        ("Direccion", "{{direccion_entrega}}"),
                        ("Costo despacho", "{{costo_envio}}"),
                    ],
                    "Te avisaremos cuando salga de bodega con su numero de seguimiento.",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS Sport & Health*\n\nHola *{{nombre_cliente}}*, tu pedido *{{id_pedido}}* fue programado correctamente.\n\nCourier: {{courier}}\nEntrega estimada: {{fecha_entrega}}\nDireccion: {{direccion_entrega}}\nCosto despacho: {{costo_envio}}\n\nTe avisaremos cuando salga de bodega.",
            ),
        },
        "en_ruta": {
            "email": (
                "Tu pedido {{id_pedido}} esta en ruta - ILUS",
                despacho_body(
                    "Tu pedido <strong>{{id_pedido}}</strong> ya salio de bodega y esta en camino contigo.",
                    [
                        ("Nro. pedido", "{{id_pedido}}"),
                        ("Courier", "{{courier}}"),
                        ("Nro. seguimiento", "{{numero_seguimiento}}"),
                        ("Entrega estimada", "{{fecha_entrega}}"),
                        ("Direccion", "{{direccion_entrega}}"),
                    ],
                    "Puedes rastrear tu envio usando el numero de seguimiento del courier.",
                    "#007bff",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS Sport & Health*\n\nHola *{{nombre_cliente}}*, tu pedido *{{id_pedido}}* ya esta en ruta.\n\nCourier: {{courier}}\nSeguimiento: {{numero_seguimiento}}\nEntrega estimada: {{fecha_entrega}}\nDestino: {{direccion_entrega}}",
            ),
        },
        "en_camino": {
            "email": (
                "Tu pedido {{id_pedido}} llega hoy - ILUS",
                despacho_body(
                    "Tu pedido <strong>{{id_pedido}}</strong> esta en reparto y llegara hoy a tu domicilio.",
                    [
                        ("Nro. pedido", "{{id_pedido}}"),
                        ("Courier", "{{courier}}"),
                        ("Direccion de entrega", "{{direccion_entrega}}"),
                        ("Nro. seguimiento", "{{numero_seguimiento}}"),
                    ],
                    "Asegurate de que alguien pueda recibir el pedido.",
                    "#17a2b8",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS Sport & Health*\n\nHola *{{nombre_cliente}}*, tu pedido *{{id_pedido}}* esta en reparto y llegara hoy.\n\nDireccion: {{direccion_entrega}}\nCourier: {{courier}}\nSeguimiento: {{numero_seguimiento}}\n\nAsegurate de que alguien pueda recibirlo.",
            ),
        },
        "entregado": {
            "email": (
                "Pedido {{id_pedido}} entregado con exito - ILUS",
                despacho_body(
                    "Tu pedido <strong>{{id_pedido}}</strong> fue <strong>entregado exitosamente</strong>.",
                    [
                        ("Nro. pedido", "{{id_pedido}}"),
                        ("Entregado en", "{{direccion_entrega}}"),
                        ("Courier", "{{courier}}"),
                        ("Nro. seguimiento", "{{numero_seguimiento}}"),
                    ],
                    "Gracias por confiar en ILUS Sport & Health. Estamos disponibles si necesitas apoyo con tu equipo.",
                    "#20c997",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS Sport & Health*\n\nHola *{{nombre_cliente}}*, tu pedido *{{id_pedido}}* fue entregado exitosamente.\n\nEntregado en: {{direccion_entrega}}\nCourier: {{courier}}\nSeguimiento: {{numero_seguimiento}}\n\nGracias por confiar en ILUS.",
            ),
        },
        "fallido": {
            "email": (
                "No pudimos entregar tu pedido {{id_pedido}} - ILUS",
                despacho_body(
                    "Lamentamos informarte que no fue posible entregar tu pedido <strong>{{id_pedido}}</strong> en el intento realizado.",
                    [
                        ("Nro. pedido", "{{id_pedido}}"),
                        ("Motivo", "{{motivo_falla}}"),
                        ("Courier", "{{courier}}"),
                        ("Direccion intentada", "{{direccion_entrega}}"),
                    ],
                    "Nuestro equipo se pondra en contacto para coordinar una nueva fecha de entrega.",
                    "#dc3545",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS Sport & Health*\n\nHola *{{nombre_cliente}}*, no pudimos entregar tu pedido *{{id_pedido}}*.\n\nMotivo: {{motivo_falla}}\nDireccion: {{direccion_entrega}}\nCourier: {{courier}}\n\nNos comunicaremos contigo para reagendar la entrega.",
            ),
        },
        "inicio_sesion": {
            "email": (
                "Nuevo inicio de sesion detectado - ILUS",
                sistema_body(
                    "Detectamos un nuevo inicio de sesion en tu cuenta del sistema ILUS.",
                    [
                        ("Cuenta", "{{email_usuario}}"),
                        ("Fecha y hora", "{{fecha_hora}}"),
                        ("IP de acceso", "{{ip_acceso}}"),
                    ],
                    "Si reconoces esta actividad, puedes ignorar este mensaje. Si no fuiste tu, cambia tu contrasena y contacta al administrador.",
                    "#6f42c1",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS - Alerta de seguridad*\n\nHola *{{nombre_usuario}}*, se detecto un inicio de sesion en tu cuenta.\n\nFecha: {{fecha_hora}}\nIP: {{ip_acceso}}\nCuenta: {{email_usuario}}\n\nSi no fuiste tu, cambia tu contrasena de inmediato.",
            ),
        },
        "cambio_pass": {
            "email": (
                "Tu contrasena fue actualizada - ILUS",
                sistema_body(
                    "La contrasena de tu cuenta en el sistema ILUS fue actualizada exitosamente.",
                    [
                        ("Cuenta", "{{email_usuario}}"),
                        ("Fecha y hora", "{{fecha_hora}}"),
                    ],
                    "Si no reconoces este cambio, contacta al administrador de inmediato.",
                    "#fd7e14",
                ),
            ),
            "whatsapp": (
                "",
                "*ILUS - Cambio de contrasena*\n\nHola *{{nombre_usuario}}*, tu contrasena fue actualizada el *{{fecha_hora}}*.\n\nCuenta: {{email_usuario}}\n\nSi no realizaste este cambio, contacta al administrador.",
            ),
        },
    }


def _comm_seed_default_templates(overwrite=False):
    try:
        user = current_username() or "system"
    except Exception:
        user = "system"
    conn = get_db()
    try:
        with conn.cursor() as cur:
            for estado, channels in _comm_template_defaults().items():
                for canal, (asunto, cuerpo) in channels.items():
                    if overwrite:
                        cur.execute(
                            """INSERT INTO comm_templates (estado, canal, asunto, cuerpo, updated_by)
                               VALUES (%s,%s,%s,%s,%s)
                               ON DUPLICATE KEY UPDATE
                                 asunto=VALUES(asunto),
                                 cuerpo=VALUES(cuerpo),
                                 updated_by=VALUES(updated_by),
                                 activo=1""",
                            (estado, canal, asunto, cuerpo, user),
                        )
                    else:
                        cur.execute(
                            """INSERT IGNORE INTO comm_templates
                               (estado, canal, asunto, cuerpo, updated_by)
                               VALUES (%s,%s,%s,%s,%s)""",
                            (estado, canal, asunto, cuerpo, user),
                        )
        conn.commit()
    finally:
        conn.close()


def _get_smtp_cfg():
    """
    Config SMTP — prioridad:
      1. Variables de entorno Railway (SMTP_HOST / SMTP_USER / SMTP_PASS …)
      2. Fila guardada en BD (comm_smtp_config)
      3. EMAIL_CONFIG de config.py
    Las env vars garantizan que la config sobreviva cualquier redeploy o reset de BD.
    """
    # ── 1. Fila en BD (máxima prioridad: lo que el usuario configura) ───
    try:
        row = mysql_fetchone(
            "SELECT * FROM comm_smtp_config WHERE id=1 LIMIT 1"
        )
        if row and row.get("smtp_user"):
            return {
                "smtp_host": row["smtp_host"] or "smtp.gmail.com",
                "smtp_port": int(row["smtp_port"] or 587),
                "smtp_user": row["smtp_user"],
                "smtp_pass": row["smtp_pass"] or "",
                "from_name": row["from_name"] or "ILUS Sport & Health",
                "from_addr": row["from_addr"] or row["smtp_user"],
                "secure":    bool(row.get("secure")),
                "_source":   "db",
            }
    except Exception:
        pass
    # ── 2. Variables de entorno (respaldo para despliegue) ──────────────
    env_user = os.environ.get("SMTP_USER", "").strip()
    env_pass = os.environ.get("SMTP_PASS", "").strip()
    if env_user and env_pass:
        return {
            "smtp_host": os.environ.get("SMTP_HOST", "smtp.gmail.com").strip(),
            "smtp_port": int(os.environ.get("SMTP_PORT", "587")),
            "smtp_user": env_user,
            "smtp_pass": env_pass,
            "from_name": os.environ.get("SMTP_FROM_NAME", "ILUS Sport & Health").strip(),
            "from_addr": os.environ.get("SMTP_FROM_ADDR", env_user).strip(),
            "secure":    os.environ.get("SMTP_SECURE", "").lower() in ("1","true","yes"),
            "_source":   "env",
        }
    # ── 3. config.py ────────────────────────────────────────────────────
    cfg = dict(EMAIL_CONFIG)
    cfg["_source"] = "config"
    return cfg


def _safe_smtp_cfg(cfg=None):
    cfg = dict(cfg or _get_smtp_cfg())
    if cfg.get("smtp_pass"):
        cfg["smtp_pass"] = "••••••••"
    return cfg


def _get_client_cfg():
    try:
        row = mysql_fetchone(
            "SELECT * FROM comm_client_config WHERE id=1 LIMIT 1"
        )
        if row:
            return dict(row)
    except Exception:
        pass
    return {"company_name": "ILUS Sport & Health", "corp_color": "#CC0000"}


def _get_wa_cfg():
    try:
        row = mysql_fetchone(
            "SELECT * FROM comm_whatsapp_config WHERE id=1 LIMIT 1"
        )
        if row:
            return dict(row)
    except Exception:
        pass
    return {}


def _comm_log_entry(canal, dest, asunto, estado, detalle=""):
    try:
        conn = get_db()
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO comm_log (canal,destinatario,asunto,estado,detalle,enviado_por) "
                "VALUES (%s,%s,%s,%s,%s,%s)",
                (canal, dest[:300], (asunto or "")[:500], estado,
                 (detalle or "")[:2000], current_username()),
            )
        conn.commit()
    except Exception:
        pass


def _send_email_dinamico(to, subject, html_body, cfg=None):
    """Envía email usando config SMTP dinámica (DB o config.py)."""
    import ssl as _ssl
    cfg = cfg or _get_smtp_cfg()
    cc = _get_client_cfg()
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"]    = f"{cfg['from_name']} <{cfg.get('from_addr', cfg['smtp_user'])}>"
    msg["To"]      = to
    recipients = [to]
    if cc.get("reply_to"):
        msg["Reply-To"] = cc["reply_to"]
    cc_addrs = [a.strip() for a in (cc.get("email_cc") or "").replace(";", ",").split(",") if a.strip()]
    bcc_addrs = [a.strip() for a in (cc.get("email_bcc") or "").replace(";", ",").split(",") if a.strip()]
    if cc_addrs:
        msg["Cc"] = ", ".join(cc_addrs)
    recipients.extend(cc_addrs + bcc_addrs)
    msg.attach(MIMEText(html_body, "html", "utf-8"))
    host   = cfg["smtp_host"]
    port   = int(cfg.get("smtp_port", 587))
    secure = cfg.get("secure", False)
    if secure:
        ctx = _ssl.create_default_context()
        with _open_smtp_client(host, port, True, timeout=15, context=ctx) as srv:
            srv.login(cfg["smtp_user"], cfg["smtp_pass"])
            srv.sendmail(cfg.get("from_addr", cfg["smtp_user"]), recipients, msg.as_string())
    else:
        with _open_smtp_client(host, port, False, timeout=15) as srv:
            srv.ehlo()
            srv.starttls()
            srv.login(cfg["smtp_user"], cfg["smtp_pass"])
            srv.sendmail(cfg.get("from_addr", cfg["smtp_user"]), recipients, msg.as_string())


def _smtp_cfg_from_request(data=None):
    data = data or {}
    prev = _get_smtp_cfg()
    smtp_pass = (data.get("pass") or "").strip()
    is_masked = smtp_pass.startswith("â€¢") or smtp_pass == "••••••••" or set(smtp_pass or "") == {"•"}
    bullet = chr(8226)
    is_masked = (
        is_masked
        or smtp_pass == (bullet * 8)
        or set(smtp_pass or "") == {bullet}
        or (bool(smtp_pass) and not any(ch.isalnum() for ch in smtp_pass))
    )
    if is_masked:
        smtp_pass = ""
    if not smtp_pass or smtp_pass == "••••••••":
        smtp_pass = prev.get("smtp_pass") or ""
    host = (data.get("host") or prev.get("smtp_host") or "smtp.gmail.com").strip()
    port = int(data.get("port") or prev.get("smtp_port") or 587)
    secure = bool(data.get("secure")) if "secure" in data else bool(prev.get("secure"))
    if "gmail.com" in host.lower():
        if port == 587 and secure:
            secure = False
        elif port == 465 and not secure:
            secure = True
    return {
        "smtp_host": host,
        "smtp_port": port,
        "smtp_user": (data.get("user") or prev.get("smtp_user") or "").strip(),
        "smtp_pass": smtp_pass,
        "from_name": (data.get("fromName") or prev.get("from_name") or "ILUS Sport & Health").strip(),
        "from_addr": (data.get("fromAddr") or prev.get("from_addr") or data.get("user") or prev.get("smtp_user") or "").strip(),
        "secure": secure,
    }


def _smtp_connection_diagnose(cfg):
    """Valida conexion SMTP y login sin enviar ningun correo."""
    import socket
    import ssl as _ssl

    host = (cfg.get("smtp_host") or "").strip()
    port = int(cfg.get("smtp_port") or 0)
    user = (cfg.get("smtp_user") or "").strip()
    password = cfg.get("smtp_pass") or ""
    secure = bool(cfg.get("secure"))

    checks = []
    suggestions = []
    if not host:
        return {"ok": False, "stage": "config", "message": "Falta el servidor SMTP.", "suggestions": ["Ingresa smtp.gmail.com para Gmail."]}
    if not port:
        return {"ok": False, "stage": "config", "message": "Falta el puerto SMTP.", "suggestions": ["Usa 587 con STARTTLS o 465 con SSL."]}
    if not user:
        return {"ok": False, "stage": "config", "message": "Falta el email usuario SMTP.", "suggestions": ["Ingresa la cuenta que autentica el correo."]}
    if not password:
        return {"ok": False, "stage": "config", "message": "Falta la contraseña / App Password.", "suggestions": ["Guarda una App Password de Gmail de 16 caracteres."]}

    gmail = "gmail.com" in host.lower()
    if gmail and secure and port != 465:
        suggestions.append("Para Gmail con SSL/TLS marcado usa puerto 465, o desmarca SSL/TLS y usa puerto 587.")
    if gmail and not secure and port != 587:
        suggestions.append("Para Gmail con STARTTLS normalmente se usa puerto 587 y SSL/TLS desmarcado.")
    if gmail and " " in password:
        suggestions.append("Gmail permite mostrar la App Password con espacios, pero si falla prueba guardarla sin espacios.")

    try:
        ipv4 = _smtp_ipv4(host)
        if ipv4 and ipv4 != host:
            checks.append(f"DNS IPv4 resuelto: {ipv4}")
        if secure:
            checks.append(f"Conectando con SSL a {host}:{port}")
            ctx = _ssl.create_default_context()
            with _open_smtp_client(host, port, True, timeout=10, context=ctx) as srv:
                checks.append("Servidor SSL respondio correctamente")
                srv.ehlo()
                checks.append("Autenticando usuario SMTP")
                srv.login(user, password)
        else:
            checks.append(f"Conectando a {host}:{port}")
            with _open_smtp_client(host, port, False, timeout=10) as srv:
                srv.ehlo()
                if srv.has_extn("starttls"):
                    checks.append("STARTTLS disponible; negociando cifrado")
                    ctx = _ssl.create_default_context()
                    srv.starttls(context=ctx)
                    srv.ehlo()
                else:
                    checks.append("El servidor no anuncia STARTTLS")
                    suggestions.append("Si el servidor exige cifrado, activa SSL/TLS con puerto 465 o usa puerto 587 con STARTTLS.")
                checks.append("Autenticando usuario SMTP")
                srv.login(user, password)
        checks.append("Credenciales aceptadas. La aplicacion puede enviar correos.")
        return {
            "ok": True,
            "stage": "auth",
            "message": "Conexion SMTP verificada. Host, puerto, cifrado y credenciales estan correctos.",
            "checks": checks,
            "suggestions": suggestions,
        }
    except smtplib.SMTPAuthenticationError as exc:
        code = getattr(exc, "smtp_code", "")
        detail = (getattr(exc, "smtp_error", b"") or b"").decode("utf-8", "ignore")
        hints = [
            "Revisa que el email usuario sea correcto.",
            "Si usas Gmail, no uses tu clave normal: crea una App Password de 16 caracteres.",
            "Verifica que la verificacion en 2 pasos este activa en la cuenta Gmail.",
        ]
        return {"ok": False, "stage": "auth", "message": f"Autenticacion rechazada por el servidor SMTP ({code}).", "detail": detail, "checks": checks, "suggestions": hints + suggestions}
    except (smtplib.SMTPConnectError, smtplib.SMTPServerDisconnected, socket.timeout, TimeoutError) as exc:
        hints = [
            "Revisa host y puerto.",
            "Para Gmail usa 587 sin SSL/TLS marcado, o 465 con SSL/TLS marcado.",
            "Si esta en Railway o hosting, confirma que el proveedor permita conexiones SMTP salientes.",
        ]
        return {"ok": False, "stage": "connect", "message": "No se pudo conectar correctamente al servidor SMTP.", "detail": str(exc), "checks": checks, "suggestions": hints + suggestions}
    except _ssl.SSLError as exc:
        hints = [
            "Hay una mezcla incorrecta entre puerto y tipo de cifrado.",
            "Para 587 desmarca SSL/TLS; para 465 marca SSL/TLS.",
        ]
        return {"ok": False, "stage": "tls", "message": "Error de SSL/TLS al negociar la conexion.", "detail": str(exc), "checks": checks, "suggestions": hints + suggestions}
    except OSError as exc:
        hints = [
            "El host o puerto no responde desde esta aplicacion.",
            "Revisa si el firewall o el hosting bloquea SMTP saliente.",
            "Prueba Gmail con smtp.gmail.com puerto 587 y SSL/TLS desmarcado.",
        ]
        return {"ok": False, "stage": "network", "message": "Error de red conectando al SMTP.", "detail": str(exc), "checks": checks, "suggestions": hints + suggestions}
    except smtplib.SMTPException as exc:
        return {"ok": False, "stage": "smtp", "message": "El servidor SMTP respondio con un error.", "detail": str(exc), "checks": checks, "suggestions": suggestions}


def _comm_render_email_document(title, body_html, subtitle=""):
    """Renderiza un fragmento HTML dentro del diseÃ±o corporativo actual."""
    cc = _get_client_cfg()
    color = cc.get("corp_color") or "#CC0000"
    company = cc.get("company_name") or "ILUS Sport & Health"
    logo = cc.get("logo_url") or ""
    header = _email_header_ilus(title, subtitle, color, logo, company)
    body = _email_body_section(body_html)
    footer = _email_footer_ilus(company)
    return _email_wrapper(_email_card(header, body, footer), company)


def _send_whatsapp(account_sid, auth_token, from_num, to_num, body):
    """Envía WhatsApp vía Twilio. Lanza RuntimeError si twilio no está instalado."""
    import re as _re
    try:
        from twilio.rest import Client as _TwilioClient
    except ImportError:
        raise RuntimeError(
            "El paquete 'twilio' no está instalado. "
            "Ejecuta: pip install twilio"
        )
    m = _re.search(r'(AC[a-f0-9]{32})', account_sid or "", _re.I)
    if m:
        account_sid = m.group(1)
    if not account_sid.upper().startswith("AC"):
        raise ValueError("El Account SID debe comenzar con 'AC'")
    if auth_token in ("[AuthToken]", "", None):
        raise ValueError("Ingresa un Auth Token válido")
    from_num = from_num if from_num.startswith("whatsapp:") else f"whatsapp:{from_num}"
    to_num   = to_num   if to_num.startswith("whatsapp:")   else f"whatsapp:{to_num}"
    client   = _TwilioClient(account_sid, auth_token)
    msg      = client.messages.create(from_=from_num, to=to_num, body=body)
    return msg.sid


# ── Templates de email ────────────────────────────────────────

def _email_wrapper(inner, company="ILUS Sport & Health"):
    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Comunicacion - {company}</title></head>
<body style="margin:0;padding:0;background:#f1f2f4;font-family:Arial,Helvetica,sans-serif;color:#111827">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f1f2f4;padding:28px 12px">
<tr><td align="center">{inner}</td></tr>
</table></body></html>"""


def _email_card(header_html, body_html, footer_html="", width=580):
    return f"""
<table width="{width}" cellpadding="0" cellspacing="0"
       style="background:#ffffff;border-radius:10px;overflow:hidden;
              max-width:{width}px;width:100%;box-shadow:0 6px 20px rgba(15,23,42,.10)">
  <tr><td>{header_html}</td></tr>
  <tr><td style="background:#fff">{body_html}</td></tr>
  {"<tr><td>" + footer_html + "</td></tr>" if footer_html else ""}
</table>"""


def _email_header_ilus(title, subtitle="", corp_color="#CC0000", logo_url=None, company="ILUS"):
    logo_url = logo_url or "https://ilusfitness.com/cdn/shop/files/Logo_ILUS_Fitness_Blanco_equipamiento_para_gimnasios.png"
    return f"""
<table width="100%" cellpadding="0" cellspacing="0">
  <tr>
    <td style="background:#000;padding:24px 28px;text-align:center">
      <img src="{logo_url}" alt="{company}"
           style="height:48px;max-width:230px;width:auto;display:block;margin:0 auto;object-fit:contain">
    </td>
  </tr>
  <tr>
    <td style="background:#111;padding:28px 32px;text-align:center;border-top:1px solid #202020">
      <div style="color:#ffffff;font-size:22px;line-height:1.25;font-weight:800">{title}</div>
      {"<div style='color:#f3f4f6;font-size:13px;line-height:1.45;margin-top:8px'>" + subtitle + "</div>" if subtitle else ""}
    </td>
  </tr>
</table>"""


def _email_body_section(content):
    return f'<div style="padding:32px 30px;font-size:14px;line-height:1.6;color:#111827">{content}</div>'


def _email_footer_ilus(company="ILUS Sport & Health"):
    return f"""
<table width="100%" cellpadding="0" cellspacing="0"
       style="background:#000">
  <tr>
    <td style="padding:24px 32px;text-align:center;color:#6b7280;font-size:11px;line-height:1.6">
      <div style="color:#DC143C;font-size:13px;font-weight:700;text-transform:uppercase">{company}</div>
      <div style="color:#9ca3af;margin-top:5px">Equipamiento profesional para alto rendimiento</div>
      <div style="color:#6b7280;margin-top:14px">
        Este correo fue generado automaticamente. Para soporte, utiliza nuestros canales oficiales.
      </div>
    </td>
  </tr>
</table>"""


def _email_info_box(rows, corp_color="#CC0000"):
    """rows = [['Label', 'Valor'], ...]"""
    items = "".join(
        f"<tr><td style='padding:6px 0;color:#888;font-size:12px;width:38%'>{r[0]}</td>"
        f"<td style='padding:6px 0;font-size:13px;font-weight:600;color:#1a1a1a'>{r[1]}</td></tr>"
        for r in rows
    )
    return f"""
<table width="100%" cellpadding="0" cellspacing="0"
       style="background:#fff8f5;border-left:4px solid {corp_color};
              border-radius:4px;padding:12px 16px;margin:16px 0">
  <tr><td><table width="100%" cellpadding="0" cellspacing="0">{items}</table></td></tr>
</table>"""


def _email_btn(text, url, corp_color="#CC0000"):
    return f"""
<table cellpadding="0" cellspacing="0" style="margin:20px 0">
  <tr>
    <td style="border-radius:8px;background:{corp_color};padding:12px 28px;text-align:center">
      <a href="{url}" style="color:#fff;font-size:14px;font-weight:700;
                             text-decoration:none;display:block">{text}</a>
    </td>
  </tr>
</table>"""


def tpl_email_prueba(sender_name, company="ILUS Sport & Health", corp_color="#CC0000"):
    header = _email_header_ilus("✅ Conexión SMTP verificada", "Prueba de correo", corp_color)
    body   = _email_body_section(f"""
      <p style="font-size:15px;color:#1a1a1a;margin:0 0 16px">
        ¡Hola! Este correo confirma que la configuración SMTP está funcionando correctamente.
      </p>
      {_email_info_box([
          ["Enviado por", sender_name],
          ["Servidor", "SMTP dinámico — ILUS Comunicaciones"],
      ], corp_color)}
      <p style="font-size:12px;color:#999;margin:20px 0 0">
        Si recibiste este mensaje, la integración de email está activa y lista para usar.
      </p>""")
    footer = _email_footer_ilus(company)
    return _email_wrapper(_email_card(header, body, footer), company)


def tpl_email_estado_pedido(data):
    """Template: actualización de estado a cliente."""
    cc    = _get_client_cfg()
    color = cc.get("corp_color", "#CC0000")
    co    = cc.get("company_name", "ILUS Sport & Health")
    logo  = cc.get("logo_url", "")
    estado_badge = {
        "En preparación": "🔵", "En ruta": "🚚", "Entregado": "✅",
        "Entrega fallida": "❌", "Pendiente": "⏳",
    }.get(data.get("status", ""), "📦")
    header = _email_header_ilus(
        f"{estado_badge} {data.get('status','Actualización')}",
        f"Pedido #{data.get('trackingCode','')}", color, logo, co
    )
    rows = [
        ["Tracking", data.get("trackingCode", "—")],
        ["Estado",   data.get("status", "—")],
    ]
    if data.get("eta"):
        rows.append(["Entrega estimada", data["eta"]])
    if data.get("conductorName"):
        rows.append(["Conductor", data["conductorName"]])
    if data.get("conductorPhone"):
        rows.append(["Teléfono conductor", data["conductorPhone"]])
    btn = _email_btn("Rastrear pedido", data.get("trackingUrl", "#"), color) if data.get("trackingUrl") else ""
    body = _email_body_section(
        f"<p style='font-size:15px;color:#1a1a1a;margin:0 0 16px'>"
        f"Hola <strong>{data.get('customerName','')}</strong>, te informamos sobre el estado de tu pedido.</p>"
        + _email_info_box(rows, color) + btn
    )
    return _email_wrapper(_email_card(header, body, _email_footer_ilus(co)), co)


# ── RUTAS: COMUNICACIONES ─────────────────────────────────────

@app.route("/comunicaciones/")
@_require_superadmin
def comm_index():
    smtp_cfg  = _get_smtp_cfg()
    client_cfg = _get_client_cfg()
    wa_cfg    = _get_wa_cfg()
    log_rows  = []
    try:
        log_rows = mysql_fetchall(
            "SELECT * FROM comm_log ORDER BY created_at DESC LIMIT 80"
        )
    except Exception:
        pass
    # Convertir timestamps UTC → America/Santiago para visualización
    try:
        from zoneinfo import ZoneInfo
        _tz_scl = ZoneInfo("America/Santiago")
        _tz_utc = ZoneInfo("UTC")
        log_rows = [dict(r) for r in log_rows]
        for r in log_rows:
            ts = r.get("created_at")
            if ts and isinstance(ts, datetime):
                # Si naive, asumir UTC (default MySQL en Railway/Clever Cloud)
                if ts.tzinfo is None:
                    ts = ts.replace(tzinfo=_tz_utc)
                r["created_at"] = ts.astimezone(_tz_scl)
    except Exception:
        pass
    smtp_cfg_safe = _safe_smtp_cfg(smtp_cfg)
    resp = make_response(render_template(
        "comunicaciones/index.html",
        smtp_cfg=smtp_cfg_safe,
        client_cfg=client_cfg,
        wa_cfg=wa_cfg,
        log_rows=log_rows,
    ))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/comunicaciones/smtp/config", methods=["POST"])
@_require_superadmin
def comm_smtp_save():
    d = request.get_json(silent=True) or {}
    cfg = _smtp_cfg_from_request(d)
    # Si la contraseña viene en blanco o enmascarada, recuperamos la guardada en BD
    prev = _get_smtp_cfg()
    smtp_pass = cfg["smtp_pass"]
    if not smtp_pass or smtp_pass == "••••••••" or set(smtp_pass) == {"•"}:
        smtp_pass = prev.get("smtp_pass") or ""
    host = cfg["smtp_host"]
    port = cfg["smtp_port"]
    user = cfg["smtp_user"]
    from_name = cfg["from_name"]
    from_addr = (cfg["from_addr"] or user).strip()
    secure = 1 if cfg["secure"] else 0
    if not user:
        return jsonify({"error": "Ingresa el email usuario SMTP"}), 400
    if not smtp_pass:
        return jsonify({"error": "Ingresa y guarda la App Password"}), 400
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO comm_smtp_config
               (id,smtp_host,smtp_port,smtp_user,smtp_pass,from_name,from_addr,secure,updated_by)
               VALUES (1,%s,%s,%s,%s,%s,%s,%s,%s)
               ON DUPLICATE KEY UPDATE
                 smtp_host=VALUES(smtp_host),
                 smtp_port=VALUES(smtp_port),
                 smtp_user=VALUES(smtp_user),
                 smtp_pass=VALUES(smtp_pass),
                 from_name=VALUES(from_name),
                 from_addr=VALUES(from_addr),
                 secure=VALUES(secure),
                 updated_by=VALUES(updated_by)""",
            (host, port, user, smtp_pass, from_name, from_addr, secure, current_username()),
        )
        cur.execute("DELETE FROM comm_smtp_config WHERE id <> 1")
    conn.commit()
    warning = ""
    if "gmail.com" in host.lower() and port == 587 and d.get("secure"):
        warning = "Para Gmail con puerto 587 se usa STARTTLS; se desmarco SSL/TLS automaticamente."
    return jsonify({"ok": True, "warning": warning, "smtp": _safe_smtp_cfg({
        "smtp_host": host,
        "smtp_port": port,
        "smtp_user": user,
        "smtp_pass": smtp_pass,
        "from_name": from_name,
        "from_addr": from_addr,
        "secure": bool(secure),
    })})


@app.route("/comunicaciones/smtp/test", methods=["POST"])
@_require_superadmin
def comm_smtp_test():
    d = request.get_json(silent=True) or {}
    cfg = _smtp_cfg_from_request(d)
    try:
        result = _smtp_connection_diagnose(cfg)
        _comm_log_entry(
            "email",
            cfg.get("smtp_user") or "",
            "Diagnostico SMTP",
            "ok" if result.get("ok") else "error",
            result.get("message", ""),
        )
        return jsonify(result), (200 if result.get("ok") else 422)
    except Exception as exc:
        _comm_log_entry("email", cfg.get("smtp_user") or "", "Diagnostico SMTP", "error", str(exc))
        return jsonify({
            "ok": False,
            "stage": "internal",
            "message": "No se pudo ejecutar el diagnostico SMTP.",
            "detail": str(exc),
        }), 500


# ══════════════════════════════════════════════════════════════════
# RESEND — GESTIÓN DE DOMINIOS (verificar dominio propio)
# ══════════════════════════════════════════════════════════════════

def _resend_api_call(method: str, path: str, body: dict = None) -> tuple:
    """
    Llamada genérica a la API de Resend.
    Devuelve (ok, status_code, response_dict).
    """
    import urllib.request as _ur
    import urllib.error  as _ue

    cfg = _get_resend_cfg()
    api_key = cfg.get("api_key", "")
    if not api_key:
        return False, 0, {"error": "No hay API Key configurada"}

    data = json.dumps(body).encode("utf-8") if body else None
    headers = {"Authorization": f"Bearer {api_key}"}
    if body:
        headers["Content-Type"] = "application/json"

    req = _ur.Request(f"https://api.resend.com{path}", data=data, headers=headers, method=method)
    try:
        with _ur.urlopen(req, timeout=20) as resp:
            raw = resp.read().decode("utf-8", "ignore")
            return True, resp.status, (json.loads(raw) if raw else {})
    except _ue.HTTPError as exc:
        body_str = ""
        try: body_str = exc.read().decode("utf-8", "ignore")
        except: pass
        try:
            parsed = json.loads(body_str) if body_str else {}
        except Exception:
            parsed = {"raw": body_str}
        return False, exc.code, parsed
    except Exception as exc:
        return False, 0, {"error": str(exc)}


def _comm_resend_test_inner():
    d  = request.get_json(silent=True) or {}
    to = (d.get("to") or "").strip()
    if not to:
        return jsonify({"ok": False, "message": "Ingresa un destinatario"}), 400

    cfg = _get_resend_cfg()
    if not cfg.get("api_key"):
        return jsonify({
            "ok": False,
            "message": "No hay API Key de Resend configurada.",
            "suggestions": [
                "Pega tu API Key en el campo de arriba y presiona Guardar.",
                "O agrégala como RESEND_API_KEY en Railway Variables.",
                "Ve a resend.com/api-keys para obtener una.",
            ],
        }), 422

    # Limpiar error anterior
    try:
        g._last_resend_error = ""
    except Exception:
        pass

    html = _ilus_email_html(
        titulo      = "Resend API verificado",
        subtitulo   = "Prueba de integración — ILUS Comunicaciones",
        saludo      = "Email entregado correctamente",
        parrafos    = [
            "La integración con Resend API está funcionando.",
            "Los correos se enviarán vía HTTPS y no dependen de puertos SMTP.",
        ],
        info_lineas = [("", "Enviado por", current_username() or "ILUS"),
                       ("", "Fecha", datetime.now().strftime("%d/%m/%Y %H:%M"))],
    )
    ok = _send_via_resend(to, "✅ Prueba Resend API — ILUS", html)
    err = getattr(g, "_last_resend_error", None) or {}
    if not isinstance(err, dict):
        err = {"raw_body": str(err), "message": str(err), "name": "", "http_code": 0, "status_code": 0}

    if ok:
        _comm_log_entry("email", to, "Prueba Resend API", "ok", f"Enviado via Resend (source={cfg.get('_source')})")
        return jsonify({"ok": True, "message": f"✅ Email enviado a {to} vía Resend API."})

    # ── No envió: armar diagnóstico claro ────────────────────────────────
    raw_body  = err.get("raw_body", "")
    resend_msg = err.get("message", "")
    err_name   = err.get("name", "")
    http_code  = err.get("http_code", 0)
    _comm_log_entry("email", to, "Prueba Resend API", "error", (raw_body or resend_msg)[:400])

    # Detectar tipo de error
    suggestions    = []
    diagnosis      = ""
    action_url     = ""
    action_label   = ""
    is_domain_issue = False
    is_self_only_issue = False

    body_lower = (raw_body + " " + resend_msg + " " + err_name).lower()

    if "validation_error" in body_lower and ("only send testing" in body_lower or "your own email" in body_lower):
        # ESTE es el famoso 1010: "you can only send to your own email"
        is_self_only_issue = True
        diagnosis = ("Resend exige verificar un dominio para enviar a CUALQUIER destinatario. "
                     "Sin dominio verificado, solo permite enviar al email registrado en tu cuenta de Resend.")
        suggestions = [
            "OPCIÓN A (definitiva): Verifica un dominio en Resend → resend.com/domains. Una vez verificado, podrás usar 'noreply@tudominio.cl' como remitente y enviar a cualquier persona.",
            f"OPCIÓN B (temporal): Envía solo al email con que registraste tu cuenta Resend. Prueba primero ver con qué email entras a resend.com.",
            "OPCIÓN C: Cambia el destinatario a una dirección @resend.dev de prueba.",
        ]
        action_url = "https://resend.com/domains"
        action_label = "Ir a verificar dominio en Resend"
    elif http_code == 401 or "invalid_api_key" in body_lower or "unauthorized" in body_lower:
        diagnosis = "La API Key es inválida o fue revocada en Resend."
        suggestions = [
            "Ve a resend.com/api-keys → genera una key NUEVA con 'Full access'.",
            "Copia la key completa (empieza con 're_') sin espacios.",
            "Actualízala en Railway Variables (RESEND_API_KEY) y redeploy.",
        ]
        action_url = "https://resend.com/api-keys"
        action_label = "Generar nueva API Key"
    elif "domain" in body_lower and ("not verified" in body_lower or "not_verified" in body_lower):
        is_domain_issue = True
        diagnosis = "El dominio del remitente no está verificado en Resend."
        suggestions = [
            f"Verifica el dominio del remitente actual ({cfg.get('from_addr','onboarding@resend.dev')}) en Resend.",
            "O cambia el remitente a 'onboarding@resend.dev' (plan gratuito).",
        ]
        action_url = "https://resend.com/domains"
        action_label = "Verificar dominio"
    elif http_code == 429 or "rate" in body_lower:
        diagnosis = "Has superado el límite de envíos por minuto/día."
        suggestions = ["Espera unos minutos y reintenta.", "Plan gratuito: 100/día, 10/segundo."]
    else:
        diagnosis = resend_msg or "Error desconocido de Resend."
        suggestions = [
            "Revisa que la cuenta Resend tenga el email verificado.",
            "Verifica que la API Key esté activa en resend.com/api-keys.",
        ]
        action_url = "https://resend.com/domains"
        action_label = "Ver mis dominios"

    return jsonify({
        "ok": False,
        "message": diagnosis or "Resend no pudo enviar el correo.",
        "resend_message": resend_msg,
        "resend_name":    err_name,
        "http_code":      http_code,
        "raw_body":       raw_body[:500],
        "suggestions":    suggestions,
        "action_url":     action_url,
        "action_label":   action_label,
        "is_self_only_issue": is_self_only_issue,
        "is_domain_issue":    is_domain_issue,
    }), 422


@app.route("/comunicaciones/email/status", methods=["GET"])
@_require_superadmin
def comm_email_status():
    """Estado del sistema de email: SMTP (único método)."""
    smtp_cfg = _get_smtp_cfg()
    configured = bool(smtp_cfg.get("smtp_user") and smtp_cfg.get("smtp_pass"))
    return jsonify({
        "configured":    configured,
        "active_method": "smtp",
        "smtp": {
            "host":   smtp_cfg.get("smtp_host", ""),
            "port":   smtp_cfg.get("smtp_port", 587),
            "user":   smtp_cfg.get("smtp_user", ""),
            "from":   smtp_cfg.get("from_addr", "") or smtp_cfg.get("smtp_user", ""),
            "secure": bool(smtp_cfg.get("secure")),
            "source": smtp_cfg.get("_source", "config"),
        },
    })


def _comm_smtp_test_send_legacy():
    d  = request.get_json(silent=True) or {}
    to = (d.get("to") or "").strip()
    if not to:
        return jsonify({"error": "Ingresa un destinatario"}), 400
    html = _ilus_email_html(
        titulo           = "✅ Prueba SMTP",
        subtitulo        = "Verificación de conexión — ILUS Comunicaciones",
        saludo           = "¡Conexión verificada!",
        parrafos         = [
            "Este correo confirma que la configuración SMTP está funcionando correctamente.",
            "Si lo recibiste, la integración de email está activa y lista para usar.",
        ],
        info_lineas      = [
            ("", "Enviado por", current_username()),
            ("", "Servidor",    "SMTP dinámico — ILUS Comunicaciones"),
        ],
    )
    try:
        _send_ilus_email(to, "🧪 Prueba SMTP — ILUS Comunicaciones", html)
        _comm_log_entry("email", to, "Prueba SMTP", "ok")
        return jsonify({"ok": True})
    except Exception as exc:
        _comm_log_entry("email", to, "Prueba SMTP", "error", str(exc))
        return jsonify({"error": str(exc)}), 500


@app.route("/comunicaciones/email/enviar", methods=["POST"])
@_require_superadmin
def comm_email_enviar():
    d       = request.get_json(silent=True) or {}
    to      = (d.get("to") or "").strip()
    subject = (d.get("subject") or "").strip()
    body    = (d.get("html") or "").strip()
    if not all([to, subject, body]):
        return jsonify({"error": "Faltan campos: to, subject, html"}), 400
    # SIEMPRE envolver con la plantilla corporativa ILUS (formato del preview)
    html = _comm_render_email_document(subject, body, "Comunicaciones")
    try:
        _send_email_dinamico(to, subject, html)
        _comm_log_entry("email", to, subject, "ok")
        return jsonify({"ok": True})
    except Exception as exc:
        _comm_log_entry("email", to, subject, "error", str(exc))
        return jsonify({"error": str(exc)}), 500


@app.route("/comunicaciones/email/preview", methods=["POST"])
@_require_superadmin
def comm_email_preview():
    d = request.get_json(silent=True) or {}
    title = (d.get("title") or "Vista previa de comunicacion").strip()
    body = (d.get("html") or "").strip() or (
        "<p style='margin:0 0 14px;font-size:14px;color:#111827;line-height:1.65'>"
        "Hola <strong>{{nombre_cliente}}</strong>, este es un ejemplo del mensaje que recibira el destinatario usando la plantilla oficial ILUS.</p>"
        "<table cellpadding='0' cellspacing='0' width='100%' style='background:#f5f5f7;border-left:4px solid #CC0000;border-radius:4px;padding:14px 18px;margin:18px 0'>"
        "<tr><td style='padding:5px 0;font-size:13px;color:#555'><strong style='color:#222'>Pedido:</strong>&nbsp; {{id_pedido}}</td></tr>"
        "<tr><td style='padding:5px 0;font-size:13px;color:#555'><strong style='color:#222'>Estado:</strong>&nbsp; En ruta</td></tr>"
        "<tr><td style='padding:5px 0;font-size:13px;color:#555'><strong style='color:#222'>Courier:</strong>&nbsp; {{courier}}</td></tr>"
        "</table>"
    )
    resp = jsonify({"ok": True, "html": _comm_render_email_document(
        title,
        body,
        "Gestion de solicitudes y seguimiento"
    )})
    resp.headers["Cache-Control"] = "no-store"
    return resp


@app.route("/comunicaciones/cliente/config", methods=["POST"])
@_require_superadmin
def comm_client_save():
    d = request.get_json(silent=True) or {}
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO comm_client_config
               (id,company_name,reply_to,support_email,support_phone,
                tracking_url,logo_url,corp_color,email_cc,email_bcc,updated_by)
               VALUES (1,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
               ON DUPLICATE KEY UPDATE
                 company_name=VALUES(company_name),
                 reply_to=VALUES(reply_to),
                 support_email=VALUES(support_email),
                 support_phone=VALUES(support_phone),
                 tracking_url=VALUES(tracking_url),
                 logo_url=VALUES(logo_url),
                 corp_color=VALUES(corp_color),
                 email_cc=VALUES(email_cc),
                 email_bcc=VALUES(email_bcc),
                 updated_by=VALUES(updated_by)""",
            (
                (d.get("company_name") or "ILUS Sport & Health").strip(),
                (d.get("reply_to") or "").strip(),
                (d.get("support_email") or "").strip(),
                (d.get("support_phone") or "").strip(),
                (d.get("tracking_url") or "").strip(),
                (d.get("logo_url") or "").strip(),
                (d.get("corp_color") or "#CC0000").strip(),
                (d.get("email_cc") or "").strip(),
                (d.get("email_bcc") or "").strip(),
                current_username(),
            ),
        )
        cur.execute("DELETE FROM comm_client_config WHERE id <> 1")
    conn.commit()
    return jsonify({"ok": True, "client": _get_client_cfg()})


@app.route("/comunicaciones/whatsapp/config", methods=["POST"])
@_require_superadmin
def comm_wa_save():
    d = request.get_json(silent=True) or {}
    conn = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO comm_whatsapp_config
               (id,account_sid,auth_token,from_number,biz_number,updated_by)
               VALUES (1,%s,%s,%s,%s,%s)
               ON DUPLICATE KEY UPDATE
                 account_sid=VALUES(account_sid),
                 auth_token=VALUES(auth_token),
                 from_number=VALUES(from_number),
                 biz_number=VALUES(biz_number),
                 updated_by=VALUES(updated_by)""",
            (
                (d.get("account_sid") or "").strip(),
                (d.get("auth_token") or "").strip(),
                (d.get("from_number") or "").strip(),
                (d.get("biz_number") or "").strip(),
                current_username(),
            ),
        )
        cur.execute("DELETE FROM comm_whatsapp_config WHERE id <> 1")
    conn.commit()
    return jsonify({"ok": True})


@app.route("/comunicaciones/whatsapp/test", methods=["POST"])
@_require_superadmin
def comm_wa_test():
    d     = request.get_json(silent=True) or {}
    saved = _get_wa_cfg()
    sid   = (d.get("account_sid") or saved.get("account_sid") or "").strip()
    tok   = (d.get("auth_token")  or saved.get("auth_token")  or "").strip()
    frm   = (d.get("from_number") or saved.get("from_number") or "").strip()
    to    = (d.get("to") or "").strip()
    if not all([sid, tok, frm, to]):
        return jsonify({"error": "Completa: Account SID, Auth Token, From y destinatario"}), 400
    body = (f"✅ Prueba de WhatsApp — ILUS Comunicaciones\n"
            f"Enviado por: {current_username()}")
    body = (d.get("body") or "").strip() or body
    try:
        msg_sid = _send_whatsapp(sid, tok, frm, to, body)
        _comm_log_entry("whatsapp", to, "Prueba WA", "ok", msg_sid)
        return jsonify({"ok": True, "msg_sid": msg_sid})
    except Exception as exc:
        detail = str(exc)
        code = getattr(exc, "code", None) or getattr(exc, "status", None)
        more_info = getattr(exc, "more_info", "") or ""
        suggestions = []
        if "not installed" in detail.lower() or "twilio" in detail.lower() and "instal" in detail.lower():
            suggestions.append("Instala Twilio en el mismo Python que ejecuta la app y reinicia el servidor.")
        if str(code) in {"21211", "21614"} or "not a valid phone number" in detail.lower():
            suggestions.append("Revisa que el destinatario tenga formato internacional, por ejemplo +56912345678.")
        if str(code) in {"63015", "63016"} or "sandbox" in detail.lower():
            suggestions.append("En sandbox, el destinatario debe enviar primero el codigo join al numero +1 415 523 8886.")
        if str(code) in {"20003", "20404"} or "authenticate" in detail.lower() or "credentials" in detail.lower():
            suggestions.append("Revisa Account SID y Auth Token desde el Dashboard de Twilio.")
        if frm.startswith("whatsapp:+14155238886"):
            suggestions.append("Estas usando sandbox Twilio; confirma que el telefono destino ya se unio al sandbox.")
        _comm_log_entry("whatsapp", to, "Prueba WA", "error", detail)
        return jsonify({
            "error": detail,
            "code": code,
            "more_info": more_info,
            "suggestions": suggestions,
        }), 500


@app.route("/comunicaciones/templates", methods=["GET"])
@_require_superadmin
def comm_templates_get():
    """Devuelve plantillas agrupadas por estado.
    Filtra por ?modulo=transporte (default), retiros, mantenciones."""
    modulo = (request.args.get("modulo") or "transporte").strip()
    if modulo not in ("transporte", "retiros", "mantenciones"):
        modulo = "transporte"
    rows = mysql_fetchall(
        "SELECT * FROM comm_templates WHERE modulo=%s ORDER BY estado, canal",
        (modulo,),
    ) or []
    legacy_tokens = ("Dropit", "direccion_origen", "direccion_destino", "link_tracking", "nombre_conductor")
    if modulo == "transporte" and any(
        any(tok in ((r.get("asunto") or "") + " " + (r.get("cuerpo") or "")) for tok in legacy_tokens)
        for r in rows
    ):
        _comm_seed_default_templates(overwrite=True)
        rows = mysql_fetchall(
            "SELECT * FROM comm_templates WHERE modulo=%s ORDER BY estado, canal",
            (modulo,),
        ) or []
    data = {}
    for r in rows:
        est = r["estado"]
        if est not in data:
            data[est] = {}
        data[est][r["canal"]] = {
            "id":     r["id"],
            "asunto": r.get("asunto") or "",
            "cuerpo": r.get("cuerpo") or "",
            "activo": r.get("activo", 1),
        }
    return jsonify(data)


@app.route("/comunicaciones/templates/<estado>/<canal>", methods=["PUT"])
@_require_superadmin
def comm_template_save(estado, canal):
    """Guarda/actualiza una plantilla para estado + canal en un módulo dado.
    Body opcional: { modulo: 'transporte' | 'retiros' | 'mantenciones' }"""
    if canal not in ("email", "whatsapp"):
        return jsonify({"error": "Canal inválido"}), 400
    d      = request.get_json(silent=True) or {}
    modulo = (d.get("modulo") or "transporte").strip()
    if modulo not in ("transporte", "retiros", "mantenciones"):
        modulo = "transporte"
    asunto = d.get("asunto", "")
    cuerpo = d.get("cuerpo", "")
    user   = current_username()
    conn   = get_db()
    with conn.cursor() as cur:
        cur.execute(
            """INSERT INTO comm_templates (modulo, estado, canal, asunto, cuerpo, updated_by)
               VALUES (%s,%s,%s,%s,%s,%s)
               ON DUPLICATE KEY UPDATE
                 asunto=VALUES(asunto), cuerpo=VALUES(cuerpo), updated_by=VALUES(updated_by)""",
            (modulo, estado, canal, asunto, cuerpo, user)
        )
    conn.commit()
    return jsonify({"ok": True})


@app.route("/comunicaciones/templates/restaurar-todo", methods=["POST"])
@_require_superadmin
def comm_templates_restore_all():
    """Restaura todas las plantillas a la base oficial ILUS."""
    _comm_seed_default_templates(overwrite=True)
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════════════
#  COMUNICACIONES — TRAZABILIDAD: log de emails + página de estado
# ══════════════════════════════════════════════════════════════════════

@app.route("/comunicaciones/log")
@require_permission("admin")
def comm_email_log():
    """Redirige al tab Historial dentro de Comunicaciones (centralizado).

    Antes era una página separada. Ahora todo se concentra en la pestaña
    Historial de /comunicaciones para evitar duplicidad. Si vino con filtros
    en query, se preservan en el hash para que el JS los procese.
    """
    qs = []
    for k in ("evento", "estado", "limit"):
        v = request.args.get(k)
        if v: qs.append(f"{k}={v}")
    hash_part = "#tabLog"
    if qs:
        hash_part += "?" + "&".join(qs)
    return redirect(url_for("comm_index") + hash_part)


@app.route("/comunicaciones/log/legacy")
@require_permission("admin")
def comm_email_log_legacy():
    """Vista legacy del log (mantenida por compatibilidad con scripts/links viejos)."""
    limit = min(int(request.args.get("limit") or 100), 500)
    evento = request.args.get("evento","")
    estado = request.args.get("estado","")
    where, params = ["1=1"], []
    if evento: where.append("evento=%s"); params.append(evento)
    if estado: where.append("estado=%s"); params.append(estado)
    rows = mysql_fetchall(
        f"SELECT * FROM email_log WHERE {' AND '.join(where)} "
        f"ORDER BY created_at DESC LIMIT {limit}",
        tuple(params)
    )
    # Stats últimos 30 días
    stats = mysql_fetchone("""
        SELECT
          COUNT(*) AS total,
          SUM(CASE WHEN estado='enviado' THEN 1 ELSE 0 END) AS exitosos,
          SUM(CASE WHEN estado='fallido' THEN 1 ELSE 0 END) AS fallidos,
          COUNT(DISTINCT destinatario) AS destinatarios_unicos
        FROM email_log
        WHERE created_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)
    """) or {}
    eventos_top = mysql_fetchall("""
        SELECT evento, COUNT(*) AS n FROM email_log
        WHERE evento IS NOT NULL AND created_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)
        GROUP BY evento ORDER BY n DESC LIMIT 10
    """)
    return render_template("comunicaciones/log.html",
        rows=[dict(r) for r in rows],
        stats=dict(stats),
        eventos_top=[dict(e) for e in eventos_top],
        filtros={"evento":evento,"estado":estado},
    )


@app.route("/comunicaciones/log/<int:lid>/reintentar", methods=["POST"])
@require_permission("admin")
def comm_log_reintentar(lid):
    """Reintenta un email que falló."""
    row = mysql_fetchone("SELECT * FROM email_log WHERE id=%s",(lid,))
    if not row: return jsonify({"error":"Log no encontrado"}), 404
    # Reintento envuelto con la plantilla corporativa
    body = (
        f"<p>Reintento del envío original: <strong>{row['asunto']}</strong></p>"
        f"<p>Si crees que es un error, contacta al administrador.</p>"
    )
    html = _comm_render_email_document(row["asunto"], body, "Reintento")
    sent = _send_ilus_email(row["destinatario"], row["asunto"], html, evento="retry_"+(row.get("evento") or "manual"))
    return jsonify({"ok": bool(sent)})


# ══════════════════════════════════════════════════════════════
#  MÓDULO: MANTENCIONES
#  Acceso: superadmin + ejecutivo
# ══════════════════════════════════════════════════════════════

def _mant_required(view):
    """Decorador: requiere permiso 'mantenciones'."""
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not g.get("permissions", {}).get("mantenciones"):
            return redirect(url_for("index"))
        return view(*args, **kwargs)
    return login_required(wrapped)


def init_mantenciones_tables():
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # ── Clientes de mantención ──────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_clientes (
                    id               INT AUTO_INCREMENT PRIMARY KEY,
                    razon_social     VARCHAR(200) NOT NULL,
                    rut              VARCHAR(20),
                    contacto_nombre  VARCHAR(200),
                    contacto_tel     VARCHAR(50),
                    contacto_email   VARCHAR(200),
                    direccion        TEXT,
                    comuna           VARCHAR(100),
                    ciudad           VARCHAR(100),
                    notas            TEXT,
                    estado           ENUM('activo','inactivo','prospecto') DEFAULT 'activo',
                    created_by       VARCHAR(190),
                    created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at       DATETIME DEFAULT CURRENT_TIMESTAMP
                                     ON UPDATE CURRENT_TIMESTAMP,
                    INDEX idx_rut    (rut),
                    INDEX idx_estado (estado)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migraciones: añadir columnas nuevas si todavía no existen
            for _mig_sql in [
                "ALTER TABLE mant_clientes ADD COLUMN updated_by VARCHAR(190)",
                "ALTER TABLE mant_clientes ADD COLUMN region VARCHAR(100)",
                "ALTER TABLE mant_clientes ADD COLUMN giro VARCHAR(200) COMMENT 'Rubro/actividad económica'",
                "ALTER TABLE mant_clientes ADD COLUMN email_empresa VARCHAR(200) COMMENT 'Email institucional (desde ERP)'",
                "ALTER TABLE mant_clientes ADD COLUMN tel_empresa VARCHAR(50) COMMENT 'Teléfono institucional (desde ERP)'",
                "ALTER TABLE mant_clientes ADD COLUMN contacto_cargo VARCHAR(120) COMMENT 'Cargo del contacto principal'",
                "ALTER TABLE mant_clientes ADD COLUMN contacto2_nombre VARCHAR(200) COMMENT 'Contacto secundario'",
                "ALTER TABLE mant_clientes ADD COLUMN contacto2_cargo VARCHAR(120)",
                "ALTER TABLE mant_clientes ADD COLUMN contacto2_tel VARCHAR(50)",
                "ALTER TABLE mant_clientes ADD COLUMN contacto2_email VARCHAR(200)",
                "ALTER TABLE mant_clientes ADD COLUMN notas_confidenciales TEXT COMMENT 'Notas internas confidenciales para análisis IA'",
                "ALTER TABLE mant_clientes MODIFY estado ENUM('activo','inactivo','prospecto','suspendido') DEFAULT 'activo'",
            ]:
                try:
                    cur.execute(_mig_sql)
                except Exception:
                    pass   # ya existe → ignorar
            # ── Máquinas por cliente (de docs ERP) ─────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_maquinas (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id   INT NOT NULL,
                    sku          VARCHAR(100),
                    nombre       VARCHAR(400),
                    serie        VARCHAR(100),
                    doc_origen   VARCHAR(80),
                    doc_fecha    DATE,
                    cantidad     INT DEFAULT 1,
                    notas        TEXT,
                    estado       ENUM('activo','baja','garantia') DEFAULT 'activo',
                    created_by   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE CASCADE,
                    INDEX idx_cliente (cliente_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── Contratos ───────────────────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_contratos (
                    id                  INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id          INT NOT NULL,
                    nombre              VARCHAR(200),
                    archivo_nombre      VARCHAR(300),
                    archivo_path        VARCHAR(500),
                    archivo_tipo        ENUM('pdf','word','otro') DEFAULT 'pdf',
                    fecha_inicio        DATE,
                    fecha_vencimiento   DATE,
                    es_indefinido       TINYINT(1) DEFAULT 0,
                    monto_mensual       DECIMAL(12,2),
                    monto_anual         DECIMAL(12,2),
                    frecuencia_meses    INT COMMENT 'Frecuencia mantencion en meses',
                    notas               TEXT,
                    ai_analizado        TINYINT(1) DEFAULT 0,
                    ai_fecha            DATETIME,
                    ai_resumen          TEXT,
                    ai_puntos_criticos  TEXT COMMENT 'JSON array',
                    ai_alertas          TEXT COMMENT 'JSON array',
                    ai_frecuencia_sug   INT,
                    ai_score            INT COMMENT '0-100 calidad contrato',
                    estado              ENUM('vigente','vencido','por_vencer','indefinido')
                                        DEFAULT 'vigente',
                    created_by          VARCHAR(190),
                    created_at          DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at          DATETIME DEFAULT CURRENT_TIMESTAMP
                                        ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE CASCADE,
                    INDEX idx_cliente     (cliente_id),
                    INDEX idx_vencimiento (fecha_vencimiento)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── Visitas / agenda ────────────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_visitas (
                    id                INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id        INT NOT NULL,
                    contrato_id       INT,
                    titulo            VARCHAR(200),
                    fecha_programada  DATE NOT NULL,
                    fecha_realizada   DATE,
                    hora_inicio       TIME,
                    hora_fin          TIME,
                    tecnico           VARCHAR(200),
                    tipo              ENUM('preventiva','correctiva','garantia','inspeccion')
                                      DEFAULT 'preventiva',
                    estado            ENUM('programada','completada','cancelada','reagendada')
                                      DEFAULT 'programada',
                    descripcion       TEXT,
                    observaciones     TEXT,
                    costo             DECIMAL(10,2),
                    created_by        VARCHAR(190),
                    created_at        DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at        DATETIME DEFAULT CURRENT_TIMESTAMP
                                      ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (cliente_id)  REFERENCES mant_clientes(id)  ON DELETE CASCADE,
                    FOREIGN KEY (contrato_id) REFERENCES mant_contratos(id) ON DELETE SET NULL,
                    INDEX idx_cliente (cliente_id),
                    INDEX idx_fecha   (fecha_programada),
                    INDEX idx_estado  (estado)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # ── Log de actividad ────────────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_logs (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    entidad     ENUM('cliente','maquina','contrato','visita') NOT NULL,
                    entidad_id  INT NOT NULL,
                    accion      VARCHAR(100) NOT NULL,
                    detalle     TEXT,
                    usuario     VARCHAR(190),
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_entidad (entidad, entidad_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Adjuntos de contrato (hasta 4 extra) ───────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_adjuntos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    contrato_id     INT NOT NULL,
                    nombre_original VARCHAR(300),
                    archivo_path    VARCHAR(500),
                    tipo            VARCHAR(50),
                    created_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (contrato_id) REFERENCES mant_contratos(id) ON DELETE CASCADE,
                    INDEX idx_contrato (contrato_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Columnas extra mant_maquinas (migracion) ───────────
            for col_sql in [
                "ALTER TABLE mant_maquinas ADD COLUMN ubicacion_cliente VARCHAR(200)",
                "ALTER TABLE mant_maquinas ADD COLUMN fecha_instalacion DATE",
                "ALTER TABLE mant_maquinas ADD COLUMN estado_op ENUM('operativo','critico','en_mantencion') DEFAULT 'operativo'",
                # Código interno ILUS — DEPRECADO en favor de campo `serie` auto-generado
                # Se mantiene la columna por compatibilidad con datos previos
                "ALTER TABLE mant_maquinas ADD COLUMN codigo_interno VARCHAR(80) NULL",
                "ALTER TABLE mant_maquinas ADD COLUMN justif_fecha_inst TEXT NULL COMMENT 'Justificación si fecha_instalacion difiere de doc_fecha'",
                "ALTER TABLE mant_maquinas ADD COLUMN justif_doc_mismatch TEXT NULL COMMENT 'Justificación si el doc_origen pertenece a otro RUT'",
                # 2 etiquetas libres editables — el usuario las usa para filtrar/categorizar
                # Ejemplos: "área cardio", "color rojo", "modelo 2024", "garantía extendida", etc.
                "ALTER TABLE mant_maquinas ADD COLUMN tag_1 VARCHAR(120) NULL COMMENT 'Etiqueta libre 1'",
                "ALTER TABLE mant_maquinas ADD COLUMN tag_2 VARCHAR(120) NULL COMMENT 'Etiqueta libre 2'",
                # Asegurar largo del serie para soportar series largas del fabricante
                "ALTER TABLE mant_maquinas MODIFY COLUMN serie VARCHAR(120)",
                # Tabla de sucursales (información adicional opcional)
                """CREATE TABLE IF NOT EXISTS mant_sucursales (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id      INT NOT NULL,
                    nombre          VARCHAR(200) NOT NULL,
                    direccion       VARCHAR(300),
                    comuna          VARCHAR(100),
                    ciudad          VARCHAR(100),
                    region          VARCHAR(100),
                    encargado_nombre VARCHAR(200),
                    encargado_cargo  VARCHAR(120),
                    encargado_tel    VARCHAR(50),
                    encargado_email  VARCHAR(200),
                    contacto2_nombre VARCHAR(200),
                    contacto2_cargo  VARCHAR(120),
                    contacto2_tel    VARCHAR(50),
                    contacto2_email  VARCHAR(200),
                    notas            TEXT,
                    activo           TINYINT(1) DEFAULT 1,
                    es_principal     TINYINT(1) DEFAULT 0 COMMENT 'Si TRUE, esta sucursal predomina sobre la dirección base del cliente',
                    created_by       VARCHAR(190),
                    created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at       DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    INDEX idx_cliente (cliente_id),
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4""",
                # Migración para BDs viejas que ya tenían la tabla sin es_principal
                "ALTER TABLE mant_sucursales ADD COLUMN es_principal TINYINT(1) DEFAULT 0",
                # Tabla de auditoría de cambios sensibles en equipos (N° serie, etc.)
                """CREATE TABLE IF NOT EXISTS mant_maquina_audit (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    maquina_id  INT NOT NULL,
                    cliente_id  INT,
                    campo       VARCHAR(60) NOT NULL,
                    valor_antes TEXT,
                    valor_nuevo TEXT,
                    motivo      TEXT,
                    usuario     VARCHAR(190),
                    fecha       DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_maquina (maquina_id),
                    INDEX idx_fecha   (fecha)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4""",
            ]:
                try:
                    cur.execute(col_sql)
                except Exception:
                    pass  # columna ya existe

            # ── Columnas extra mant_contratos (migracion) ──────────
            for col_sql in [
                "ALTER TABLE mant_contratos ADD COLUMN sla_horas INT",
                "ALTER TABLE mant_contratos ADD COLUMN incluye_repuestos TINYINT(1) DEFAULT 0",
                "ALTER TABLE mant_contratos ADD COLUMN incluye_mant_gratis TINYINT(1) DEFAULT 0",
                "ALTER TABLE mant_contratos ADD COLUMN costo_por_mant DECIMAL(12,2)",
                "ALTER TABLE mant_contratos ADD COLUMN costo_total DECIMAL(12,2)",
                "ALTER TABLE mant_contratos ADD COLUMN nivel_riesgo ENUM('alto','medio','bajo')",
                "ALTER TABLE mant_contratos ADD COLUMN ai_tipo_contrato VARCHAR(200)",
                "ALTER TABLE mant_contratos ADD COLUMN ai_clausulas TEXT",
                "ALTER TABLE mant_contratos ADD COLUMN ai_mejoras TEXT",
                "ALTER TABLE mant_contratos ADD COLUMN ai_cobertura TEXT",
                "ALTER TABLE mant_contratos ADD COLUMN ai_editable TEXT COMMENT 'JSON campos editados por usuario'",
                "ALTER TABLE mant_contratos ADD COLUMN ai_vigencia_inicio DATE",
                "ALTER TABLE mant_contratos ADD COLUMN ai_vigencia_fin DATE",
                # v2 — trazabilidad y gestión avanzada
                "ALTER TABLE mant_contratos ADD COLUMN ai_usuario VARCHAR(190)",
                "ALTER TABLE mant_contratos ADD COLUMN clausulas_custom TEXT COMMENT 'JSON clausulas personalizadas del contrato'",
                "ALTER TABLE mant_contratos ADD COLUMN variables_extra TEXT COMMENT 'JSON variables adicionales editadas por usuario'",
            ]:
                try:
                    cur.execute(col_sql)
                except Exception:
                    pass  # columna ya existe

            # ── Índices de performance (idempotentes) ──────────────────
            # Sustancialmente acelera /mantenciones/clientes, /mantenciones/, ficha, etc.
            for idx_sql in [
                # mant_contratos: lookup más reciente por cliente + filtrar por estado
                "CREATE INDEX idx_ct_cliente_created ON mant_contratos (cliente_id, created_at)",
                "CREATE INDEX idx_ct_estado ON mant_contratos (estado)",
                "CREATE INDEX idx_ct_estado_inicio ON mant_contratos (estado, fecha_inicio)",
                # mant_visitas: filtrar próximas visitas por cliente/estado/fecha
                "CREATE INDEX idx_v_cliente_estado_fecha ON mant_visitas (cliente_id, estado, fecha_programada)",
                "CREATE INDEX idx_v_estado_fecha ON mant_visitas (estado, fecha_programada)",
                # mant_clientes: orden por razón social filtrando por estado, búsqueda por created_by
                "CREATE INDEX idx_mc_estado_razon ON mant_clientes (estado, razon_social)",
                "CREATE INDEX idx_mc_created_by ON mant_clientes (created_by)",
                # app_products: filtrar por estado / created_by sin full scan
                f"CREATE INDEX idx_prod_estado ON `{PRODUCTS_TABLE}` (estado)",
                f"CREATE INDEX idx_prod_created_by ON `{PRODUCTS_TABLE}` (created_by)",
                # ── REFORZAMIENTOS DE INTEGRIDAD (datos reales) ────────────────
                # Previene duplicados de cliente con mismo RUT
                # NOTA: en MySQL, índice UNIQUE permite múltiples NULL (lo que queremos
                # para clientes sin RUT). Solo bloquea duplicados con RUT presente.
                "CREATE UNIQUE INDEX uq_mc_rut ON mant_clientes (rut)",
                # Performance en módulo OT (queries frecuentes)
                "CREATE INDEX idx_vt_visita_completada ON mant_visita_tareas (visita_id, completada)",
                "CREATE INDEX idx_vf_visita_tipo ON mant_visita_fotos (visita_id, tipo_foto)",
                "CREATE INDEX idx_vfo_visita_tomada ON mant_visita_fotos (visita_id, tomada_at)",
                # mant_maquinas: buscar por serie (búsqueda diaria en sala)
                "CREATE INDEX idx_mm_serie ON mant_maquinas (serie)",
                "CREATE INDEX idx_mm_sku ON mant_maquinas (sku)",
                "CREATE INDEX idx_mm_cliente_estado ON mant_maquinas (cliente_id, estado)",
                # mant_logs: trazabilidad por usuario
                "CREATE INDEX idx_logs_usuario_fecha ON mant_logs (usuario, created_at)",
                # mant_visitas: lookup por numero_ot
                "CREATE INDEX idx_v_numero_ot ON mant_visitas (numero_ot)",
                # mant_sucursales: lookup principal por cliente
                "CREATE INDEX idx_ms_cliente_principal ON mant_sucursales (cliente_id, es_principal)",
            ]:
                try:
                    cur.execute(idx_sql)
                except Exception:
                    pass  # índice ya existe

            # ── Tabla: Reportes de servicio (Informe Post Servicio) ─────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_reportes (
                    id                INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id        INT NOT NULL,
                    tipo              ENUM('mantencion','instalacion','inspeccion','garantia','otro')
                                      DEFAULT 'mantencion',
                    estado            ENUM('borrador','emitido','entregado') DEFAULT 'borrador',
                    ticket_num        VARCHAR(30),
                    asunto            VARCHAR(300),
                    tecnico_junior    VARCHAR(190),
                    tecnico_senior    VARCHAR(190),
                    fecha_solicitado  DATE,
                    fecha_inicio      DATE,
                    fecha_cierre      DATE,
                    antecedentes      TEXT,
                    objetivos         TEXT COMMENT 'JSON lista de strings',
                    trabajos          TEXT COMMENT 'JSON lista de strings',
                    observaciones     TEXT COMMENT 'JSON lista de strings',
                    maquinas_json     TEXT COMMENT 'JSON lista {sku,descripcion,cantidad,modelo,serie,repuesto,garantia,observacion}',
                    fotos_json        TEXT COMMENT 'JSON lista rutas de imágenes',
                    ai_diagnostico    TEXT COMMENT 'Diagnóstico IA del reporte',
                    ai_acciones       TEXT COMMENT 'JSON acciones recomendadas por IA',
                    ai_fecha          DATETIME,
                    ai_usuario        VARCHAR(190),
                    created_by        VARCHAR(190),
                    created_at        DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at        DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE CASCADE,
                    INDEX idx_cliente (cliente_id),
                    INDEX idx_fecha   (fecha_inicio)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migración: html_path para snapshots HTML de reportes
            for _mig in [
                "ALTER TABLE mant_reportes ADD COLUMN html_path VARCHAR(500)",
                "ALTER TABLE mant_reportes ADD COLUMN html_generated_at DATETIME",
            ]:
                try: cur.execute(_mig)
                except Exception: pass

            # ── Tabla: Notificaciones de mantenciones ──────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_notificaciones (
                    id            INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id    INT,
                    entidad       ENUM('cliente','contrato','maquina','visita','reporte') DEFAULT 'cliente',
                    entidad_id    INT,
                    tipo          ENUM('vencimiento','sla','visita_proxima','garantia',
                                       'sin_mantencion','contrato_riesgo','ai_alerta','otro')
                                  DEFAULT 'otro',
                    titulo        VARCHAR(300),
                    mensaje       TEXT,
                    estado        ENUM('pendiente','enviada','leida','ignorada') DEFAULT 'pendiente',
                    canal         ENUM('email','sistema','whatsapp') DEFAULT 'sistema',
                    destinatario  VARCHAR(190),
                    fecha_envio   DATETIME,
                    fecha_lectura DATETIME,
                    created_at    DATETIME DEFAULT CURRENT_TIMESTAMP,
                    created_by    VARCHAR(190),
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE SET NULL,
                    INDEX idx_estado  (estado),
                    INDEX idx_cliente (cliente_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Tabla: Adjuntos de contratos (multi-archivo) ────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_contrato_adjuntos (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    contrato_id  INT NOT NULL,
                    cliente_id   INT NOT NULL,
                    tipo         ENUM('contrato','imagen','solicitud','cotizacion','reporte','otro')
                                 DEFAULT 'otro',
                    nombre       VARCHAR(300),
                    archivo_nombre VARCHAR(300),
                    archivo_path VARCHAR(500),
                    mime_type    VARCHAR(100),
                    tamaño_bytes INT,
                    descripcion  TEXT,
                    created_by   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (contrato_id) REFERENCES mant_contratos(id) ON DELETE CASCADE,
                    INDEX idx_contrato (contrato_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Repuestos por cliente / visita ─────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_repuestos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    cliente_id      INT NOT NULL,
                    visita_id       INT NULL,
                    reporte_id      INT NULL,
                    maquina_id      INT NULL,
                    sku             VARCHAR(120),
                    nombre          VARCHAR(400) NOT NULL,
                    descripcion     TEXT,
                    cantidad        DECIMAL(10,2) DEFAULT 1,
                    costo_unitario  DECIMAL(12,2) DEFAULT 0,
                    precio_venta    DECIMAL(12,2) DEFAULT 0,
                    moneda          VARCHAR(8) DEFAULT 'CLP',
                    tipo            ENUM('venta','garantia','reposicion','consumo')
                                    DEFAULT 'venta',
                    estado          ENUM('cotizado','aprobado','instalado','facturado','cancelado')
                                    DEFAULT 'cotizado',
                    proveedor       VARCHAR(200),
                    documento       VARCHAR(120) COMMENT 'OC / FCV / boleta proveedor',
                    fecha           DATE,
                    observacion     TEXT,
                    created_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP
                                    ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (cliente_id) REFERENCES mant_clientes(id) ON DELETE CASCADE,
                    INDEX idx_cliente (cliente_id),
                    INDEX idx_visita  (visita_id),
                    INDEX idx_reporte (reporte_id),
                    INDEX idx_tipo    (tipo),
                    INDEX idx_estado  (estado)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Técnicos de mantención (catálogo + ficha) ──────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_tecnicos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    nombre          VARCHAR(200) NOT NULL,
                    rut             VARCHAR(20),
                    especialidad    VARCHAR(160) COMMENT 'Cardio, fuerza, eléctrico, etc.',
                    nivel           ENUM('junior','senior','externo') DEFAULT 'junior',
                    telefono        VARCHAR(50),
                    email           VARCHAR(200),
                    direccion       VARCHAR(300),
                    comuna          VARCHAR(100),
                    region          VARCHAR(100),
                    foto_url        VARCHAR(500),
                    notas           TEXT,
                    tarifa_visita   DECIMAL(12,2) COMMENT 'Costo base por visita (CLP)',
                    activo          TINYINT(1) DEFAULT 1,
                    es_externo      TINYINT(1) DEFAULT 0 COMMENT 'Subcontratado vs interno',
                    empresa_externa VARCHAR(200) COMMENT 'Si es_externo=1, nombre del proveedor',
                    fecha_ingreso   DATE,
                    created_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP
                                    ON UPDATE CURRENT_TIMESTAMP,
                    INDEX idx_activo (activo),
                    INDEX idx_nombre (nombre)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            # Migración: agregar tecnico_id en mant_visitas (FK opcional)
            for _mig in [
                "ALTER TABLE mant_visitas ADD COLUMN tecnico_id INT NULL AFTER tecnico",
                "ALTER TABLE mant_visitas ADD INDEX idx_tecnico_id (tecnico_id)",
            ]:
                try: cur.execute(_mig)
                except Exception: pass

            # Migración: número de OT (Orden de Trabajo) único por visita
            for _mig in [
                "ALTER TABLE mant_visitas ADD COLUMN numero_ot VARCHAR(30) NULL AFTER id",
                "ALTER TABLE mant_visitas ADD UNIQUE KEY uq_numero_ot (numero_ot)",
            ]:
                try: cur.execute(_mig)
                except Exception: pass

            # Tabla de repuestos asociados a una visita (manual o desde ERP)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_visita_repuestos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    visita_id       INT NOT NULL,
                    sku             VARCHAR(80) NULL
                                    COMMENT 'SKU del producto si vino del ERP',
                    producto_id     INT NULL
                                    COMMENT 'id del catalogo local si fue desde productos',
                    descripcion     VARCHAR(300) NOT NULL,
                    cantidad        DECIMAL(10,2) DEFAULT 1.00,
                    costo_unitario  DECIMAL(12,2) DEFAULT 0.00,
                    costo_total     DECIMAL(12,2) DEFAULT 0.00,
                    origen          ENUM('manual','erp','catalogo') DEFAULT 'manual',
                    notas           VARCHAR(500) DEFAULT NULL,
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_visita  (visita_id),
                    INDEX idx_sku     (sku)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ════════════════════════════════════════════════════════════
            # CHECKLIST DE TAREAS DENTRO DE UNA OT/VISITA
            # Cada visita puede tener N tareas concretas (cambio de
            # trotadora, inspección, limpieza, foto, etc.) que el técnico
            # marca como completadas a medida que ejecuta.
            # ════════════════════════════════════════════════════════════
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_visita_tareas (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    visita_id       INT NOT NULL,
                    orden           INT DEFAULT 0,
                    titulo          VARCHAR(300) NOT NULL,
                    descripcion     TEXT,
                    tipo            ENUM('inspeccion','cambio','reparacion','limpieza',
                                          'levantamiento','instalacion','garantia','otro')
                                    DEFAULT 'otro',
                    maquina_id      INT NULL
                                    COMMENT 'FK opcional a la máquina afectada',
                    cantidad        INT DEFAULT 1
                                    COMMENT 'ej. cambio de 4 trotadoras → cantidad=4',
                    completada      TINYINT(1) DEFAULT 0,
                    completada_at   DATETIME NULL,
                    completada_por  VARCHAR(190) NULL,
                    observaciones   TEXT,
                    created_by      VARCHAR(190),
                    created_at      DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (visita_id)  REFERENCES mant_visitas(id)  ON DELETE CASCADE,
                    FOREIGN KEY (maquina_id) REFERENCES mant_maquinas(id) ON DELETE SET NULL,
                    INDEX idx_visita   (visita_id),
                    INDEX idx_completa (completada)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ════════════════════════════════════════════════════════════
            # FOTOS DE LA VISITA / OT
            # Galería con clasificación: antes / durante / después / serie /
            # falla / reparación / general. Cada foto puede asociarse a una
            # tarea específica (mant_visita_tareas) o a una máquina.
            # ════════════════════════════════════════════════════════════
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_visita_fotos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    visita_id       INT NOT NULL,
                    tarea_id        INT NULL,
                    maquina_id      INT NULL,
                    archivo_path    VARCHAR(500) NOT NULL,
                    archivo_nombre  VARCHAR(300),
                    tipo_foto       ENUM('antes','durante','despues','serie',
                                         'falla','reparacion','general','levantamiento')
                                    DEFAULT 'general',
                    descripcion     VARCHAR(500),
                    tomada_por      VARCHAR(190),
                    tomada_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                    file_size_kb    INT,
                    FOREIGN KEY (visita_id)  REFERENCES mant_visitas(id)       ON DELETE CASCADE,
                    FOREIGN KEY (tarea_id)   REFERENCES mant_visita_tareas(id) ON DELETE SET NULL,
                    FOREIGN KEY (maquina_id) REFERENCES mant_maquinas(id)      ON DELETE SET NULL,
                    INDEX idx_visita  (visita_id),
                    INDEX idx_maquina (maquina_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ════════════════════════════════════════════════════════════
            # GALERÍA POR MÁQUINA (inventario fotográfico permanente)
            # Independiente de las visitas: foto de la máquina en sí, su
            # número de serie, estado al ingresar, etc. Útil para tener
            # un "perfil visual" de cada equipo en la ficha del cliente.
            # ════════════════════════════════════════════════════════════
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_maquina_fotos (
                    id              INT AUTO_INCREMENT PRIMARY KEY,
                    maquina_id      INT NOT NULL,
                    archivo_path    VARCHAR(500) NOT NULL,
                    archivo_nombre  VARCHAR(300),
                    tipo_foto       ENUM('principal','serie','marca','detalle',
                                         'instalada','daño','antes_reparacion','despues_reparacion')
                                    DEFAULT 'principal',
                    descripcion     VARCHAR(500),
                    es_principal    TINYINT(1) DEFAULT 0
                                    COMMENT 'Foto que aparece en cards y listados',
                    visita_origen   INT NULL
                                    COMMENT 'Si la foto se tomó durante una visita específica',
                    tomada_por      VARCHAR(190),
                    tomada_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (maquina_id)    REFERENCES mant_maquinas(id) ON DELETE CASCADE,
                    FOREIGN KEY (visita_origen) REFERENCES mant_visitas(id)  ON DELETE SET NULL,
                    INDEX idx_maquina  (maquina_id),
                    INDEX idx_principal(es_principal)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # Tabla N:N para múltiples técnicos asignados a una visita
            cur.execute("""
                CREATE TABLE IF NOT EXISTS mant_visita_tecnicos (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    visita_id   INT NOT NULL,
                    tecnico_id  INT NOT NULL,
                    rol         VARCHAR(40) DEFAULT 'tecnico'
                                COMMENT 'tecnico|lider|supervisor',
                    horas       DECIMAL(5,2) DEFAULT 0
                                COMMENT 'Horas reales trabajadas (post-visita)',
                    costo       DECIMAL(12,2) DEFAULT 0
                                COMMENT 'Costo individual del técnico para esta visita',
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uq_visita_tecnico (visita_id, tecnico_id),
                    INDEX idx_visita  (visita_id),
                    INDEX idx_tecnico (tecnico_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Log de emails enviados (trazabilidad global) ───────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS email_log (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    destinatario VARCHAR(300) NOT NULL,
                    asunto       VARCHAR(500),
                    evento       VARCHAR(100) COMMENT 'crear_usuario, cambio_clave, retiro, reporte, manual, test, etc',
                    canal        VARCHAR(50) DEFAULT 'email',
                    estado       ENUM('enviado','fallido') DEFAULT 'enviado',
                    error_msg    TEXT,
                    actor        VARCHAR(190),
                    metadata     TEXT COMMENT 'JSON con info adicional',
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_evento (evento),
                    INDEX idx_dest   (destinatario),
                    INDEX idx_created(created_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Imágenes carrusel del login ────────────────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS login_images (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    archivo_path VARCHAR(500) NOT NULL,
                    titulo       VARCHAR(200),
                    subtitulo    VARCHAR(300),
                    orden        INT DEFAULT 0,
                    activa       TINYINT(1) DEFAULT 1,
                    created_by   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_orden (orden),
                    INDEX idx_activa (activa)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Imágenes carrusel público de RETIROS ───────────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS retiros_carousel (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    archivo_path VARCHAR(500) NOT NULL,
                    titulo       VARCHAR(200),
                    subtitulo    VARCHAR(300),
                    orden        INT DEFAULT 0,
                    activa       TINYINT(1) DEFAULT 1,
                    created_by   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_orden (orden),
                    INDEX idx_activa (activa)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Banner / Avisos en la página pública de retiros ────────
            # Muestra banner amarillo/rojo/azul/verde encima del hero
            # cuando hay info importante (cierre anticipado, inventario, etc.)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS retiros_announcements (
                    id           INT AUTO_INCREMENT PRIMARY KEY,
                    titulo       VARCHAR(200) NOT NULL,
                    mensaje      TEXT,
                    tipo         ENUM('info','warning','danger','success') DEFAULT 'info',
                    icon         VARCHAR(40) DEFAULT 'info-circle',
                    fecha_desde  DATETIME NULL COMMENT 'NULL=siempre vigente',
                    fecha_hasta  DATETIME NULL COMMENT 'NULL=sin caducidad',
                    activa       TINYINT(1) DEFAULT 1,
                    orden        INT DEFAULT 0,
                    created_by   VARCHAR(190),
                    created_at   DATETIME DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_activa (activa),
                    INDEX idx_fechas (fecha_desde, fecha_hasta)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # ── Roles dinámicos (matriz módulo × acción) ───────────────
            cur.execute("""
                CREATE TABLE IF NOT EXISTS roles_dinamicos (
                    id          INT AUTO_INCREMENT PRIMARY KEY,
                    slug        VARCHAR(60) UNIQUE NOT NULL,
                    nombre      VARCHAR(120) NOT NULL,
                    descripcion VARCHAR(300),
                    color       VARCHAR(20) DEFAULT '#6b7280',
                    is_system   TINYINT(1) DEFAULT 0 COMMENT 'rol nativo no eliminable',
                    activo      TINYINT(1) DEFAULT 1,
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS rol_permisos (
                    id        INT AUTO_INCREMENT PRIMARY KEY,
                    rol_slug  VARCHAR(60) NOT NULL,
                    modulo    VARCHAR(60) NOT NULL,
                    accion    VARCHAR(60) NOT NULL,
                    permitido TINYINT(1) DEFAULT 0,
                    UNIQUE KEY unique_perm (rol_slug, modulo, accion),
                    INDEX idx_rol (rol_slug)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)

            # Seed: roles del sistema (solo si no existen)
            for slug, nombre, color, is_sys in [
                ('superadmin', 'Super Administrador', '#dc2626', 1),
                ('admin',      'Administrador',      '#2563eb', 1),
                ('editor',     'Editor',             '#ea580c', 1),
                ('mantenciones','Mantenciones',      '#16a34a', 1),
                ('transporte', 'Transporte',         '#d97706', 1),
                ('vendedor',   'Vendedor',           '#7c3aed', 1),
                ('lector',     'Solo lectura',       '#6b7280', 1),
            ]:
                try:
                    cur.execute(
                        "INSERT IGNORE INTO roles_dinamicos (slug,nombre,color,is_system) VALUES (%s,%s,%s,%s)",
                        (slug, nombre, color, is_sys)
                    )
                except Exception:
                    pass

            # Regla: no se pueden repetir nombres de rol (ya hay UNIQUE en slug)
            try:
                cur.execute("ALTER TABLE roles_dinamicos ADD UNIQUE KEY uq_rol_nombre (nombre)")
            except Exception:
                pass  # ya existe el índice

        conn.commit()
    finally:
        conn.close()


def _mant_log(entidad, entidad_id, accion, detalle=""):
    """Registra acción en mant_logs. Usa mysql_execute (pool) — evita
    el TCP handshake de get_mysql() en cada llamada (~250-400ms en cloud)."""
    try:
        mysql_execute(
            "INSERT INTO mant_logs (entidad,entidad_id,accion,detalle,usuario) "
            "VALUES (%s,%s,%s,%s,%s)",
            (entidad, entidad_id, accion, detalle, current_username())
        )
    except Exception:
        pass


def _next_ot_number():
    """
    Genera el proximo numero de Orden de Trabajo correlativo.
    Formato: OT-YYYY-NNNNN (ej. OT-2026-00042)
    """
    year = datetime.now().year
    prefix = f"OT-{year}-"
    last = mysql_fetchone(
        "SELECT numero_ot FROM mant_visitas "
        "WHERE numero_ot LIKE %s "
        "ORDER BY id DESC LIMIT 1",
        (f"{prefix}%",)
    )
    next_n = 1
    if last and last.get("numero_ot"):
        try:
            seq = last["numero_ot"].split("-")[-1]
            next_n = int(seq) + 1
        except Exception:
            next_n = 1
    return f"{prefix}{next_n:05d}"


def _normalize_hora(s):
    """
    '8' -> '08:00', '12' -> '12:00', '8:30' -> '08:30', '8:5' -> '08:05'.
    Devuelve None si vacio o invalido.
    """
    if not s: return None
    s = str(s).strip()
    if not s: return None
    if s.isdigit():
        h = int(s)
        if 0 <= h <= 23:
            return f"{h:02d}:00"
        return None
    if ":" in s:
        parts = s.split(":")
        try:
            h = int(parts[0]) if parts[0] else 0
            m = int(parts[1]) if len(parts) > 1 and parts[1] else 0
            if 0 <= h <= 23 and 0 <= m <= 59:
                return f"{h:02d}:{m:02d}"
        except (ValueError, IndexError):
            return None
    return None


# Timestamp del último UPDATE de estado de contratos.
# El UPDATE es deterministic (depende sólo de la fecha) — basta 1 vez/día.
_MANT_ESTADOS_LAST_RUN = 0.0


def _mant_actualizar_estado_contratos(force=False):
    """Actualiza estado de contratos según fechas.

    Optimización: ejecuta como mucho 1 vez cada hora por worker. El usuario
    final no paga 100-300ms en cada hit a /mantenciones. El UPDATE es
    deterministic (depende solo de hoy), así que basta refrescarlo
    periódicamente.

    Para forzar (ej. tras crear/editar un contrato): pasar force=True.
    """
    global _MANT_ESTADOS_LAST_RUN
    now = time.time()
    if not force and (now - _MANT_ESTADOS_LAST_RUN) < 3600:
        return  # ya se corrió hace menos de 1h
    try:
        conn = get_mysql()
        hoy = datetime.now().date()
        pronto = hoy + timedelta(days=60)
        with conn.cursor() as cur:
            cur.execute("""
                UPDATE mant_contratos SET estado =
                  CASE
                    WHEN es_indefinido = 1 THEN 'indefinido'
                    WHEN fecha_vencimiento IS NULL THEN 'vigente'
                    WHEN fecha_vencimiento < %s THEN 'vencido'
                    WHEN fecha_vencimiento <= %s THEN 'por_vencer'
                    ELSE 'vigente'
                  END
                WHERE estado NOT IN ('vencido')
                   OR fecha_vencimiento >= %s
            """, (hoy, pronto, hoy))
        conn.commit()
        conn.close()
        _MANT_ESTADOS_LAST_RUN = now
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════
# RESET DATOS DE MANTENCIONES (solo superadmin)
#
# Wipe completo y atómico para transición de datos demo → datos reales.
# Borra: clientes (CASCADE: contratos, visitas, maquinas, sucursales,
#         visita_tareas, visita_fotos, visita_tecnicos, visita_repuestos,
#         adjuntos, maquina_fotos)
#       + logs + maquina_audit
# Conserva: técnicos (trabajadores reales), tablas de configuración.
# Resetea AUTO_INCREMENT para empezar IDs desde 1.
# Borra opcionalmente los archivos físicos de static/uploads/mantenciones/
#
# Doble confirmación requerida: body JSON {"confirm":"RESET"}.
# ══════════════════════════════════════════════════════════════════════

@app.route("/admin/mantenciones/reset", methods=["POST"])
@login_required
def admin_mantenciones_reset():
    """Borra todos los datos transaccionales de mantenciones. SOLO superadmin."""
    # Doble check de permisos (no usamos require_permission porque queremos
    # ser estrictos: SOLO superadmin)
    if not g.permissions.get("superadmin"):
        return jsonify({"ok": False, "error": "Solo superadmin puede ejecutar este reset."}), 403

    body = request.get_json(silent=True) or {}
    if (body.get("confirm") or "").strip() != "RESET":
        return jsonify({"ok": False,
                        "error": "Confirmación incorrecta. Enviar {\"confirm\":\"RESET\"} en el body."}), 400

    user = current_username() or "sistema"
    incluir_archivos = bool(body.get("borrar_archivos", True))

    # Snapshot de contadores ANTES de borrar (para reporte al usuario)
    counts_before = {}
    try:
        for t in ("mant_clientes", "mant_contratos", "mant_visitas", "mant_maquinas",
                  "mant_visita_tareas", "mant_visita_fotos", "mant_visita_tecnicos",
                  "mant_visita_repuestos", "mant_sucursales", "mant_adjuntos",
                  "mant_maquina_fotos", "mant_maquina_audit", "mant_logs"):
            try:
                r = mysql_fetchone(f"SELECT COUNT(*) AS n FROM {t}")
                counts_before[t] = int(r.get("n", 0)) if r else 0
            except Exception:
                counts_before[t] = None  # tabla no existe

    except Exception:
        pass

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Apagar FK checks para limpiar en cualquier orden de manera segura
            cur.execute("SET FOREIGN_KEY_CHECKS = 0")
            try:
                # Orden: hijo → padre. Por si alguna FK no es CASCADE perfecta.
                for t in (
                    "mant_logs",
                    "mant_maquina_audit",
                    "mant_visita_repuestos",
                    "mant_visita_fotos",
                    "mant_visita_tareas",
                    "mant_visita_tecnicos",
                    "mant_maquina_fotos",
                    "mant_adjuntos",
                    "mant_visitas",
                    "mant_maquinas",
                    "mant_contratos",
                    "mant_sucursales",
                    "mant_clientes",
                ):
                    try:
                        cur.execute(f"DELETE FROM {t}")
                    except Exception:
                        pass  # tabla no existe → skip
                    try:
                        cur.execute(f"ALTER TABLE {t} AUTO_INCREMENT = 1")
                    except Exception:
                        pass

                # Re-activar FK checks
                cur.execute("SET FOREIGN_KEY_CHECKS = 1")

                # Log de auditoría (en mant_logs ya vaciada, pero queda como primera entrada)
                cur.execute(
                    "INSERT INTO mant_logs (entidad,entidad_id,accion,detalle,usuario) "
                    "VALUES ('cliente', 0, 'reset_datos', %s, %s)",
                    (f"Reset completo ejecutado. Snapshot anterior: {json.dumps(counts_before)}", user)
                )
            finally:
                # Garantizar que FK checks queden activos pase lo que pase
                try: cur.execute("SET FOREIGN_KEY_CHECKS = 1")
                except Exception: pass
        conn.commit()
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({"ok": False, "error": f"Error al borrar: {e}"}), 500
    finally:
        conn.close()

    # Borrar archivos físicos de fotos (opcional pero recomendado)
    archivos_borrados = 0
    if incluir_archivos:
        try:
            import shutil
            fotos_dir = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                "static", "uploads", "mantenciones"
            )
            if os.path.isdir(fotos_dir):
                # Contar archivos antes de borrar
                for root, _, files in os.walk(fotos_dir):
                    archivos_borrados += len(files)
                shutil.rmtree(fotos_dir, ignore_errors=True)
                os.makedirs(fotos_dir, exist_ok=True)
        except Exception:
            pass

    # Recargar tablas vacías para garantizar schema
    try:
        init_mantenciones_tables()
    except Exception:
        pass

    return jsonify({
        "ok": True,
        "mensaje": "Datos de mantenciones limpiados. Listo para entrar datos reales.",
        "borrado": counts_before,
        "archivos_borrados": archivos_borrados,
        "usuario": user,
    })


# ── DECORATOR de acceso ────────────────────────────────────────────────

# ── RUTAS PRINCIPALES ─────────────────────────────────────────────────

@app.route("/mantenciones")
@_mant_required
def mant_index():
    _mant_actualizar_estado_contratos()
    hoy   = datetime.now().date()
    pronto = hoy + timedelta(days=60)

    # KPIs
    clientes    = mysql_fetchone("SELECT COUNT(*) AS n FROM mant_clientes WHERE estado='activo'", ()) or {}
    contratos   = mysql_fetchone("SELECT COUNT(*) AS n FROM mant_contratos WHERE estado IN ('vigente','indefinido')", ()) or {}
    vencen      = mysql_fetchone("SELECT COUNT(*) AS n FROM mant_contratos WHERE estado='por_vencer'", ()) or {}
    prox_visitas = mysql_fetchall(
        "SELECT v.*, c.razon_social FROM mant_visitas v "
        "JOIN mant_clientes c ON c.id=v.cliente_id "
        "WHERE v.estado='programada' AND v.fecha_programada BETWEEN %s AND %s "
        "ORDER BY v.fecha_programada LIMIT 10",
        (hoy, hoy + timedelta(days=30))
    )
    alertas_contratos = mysql_fetchall(
        "SELECT ct.*, cl.razon_social FROM mant_contratos ct "
        "JOIN mant_clientes cl ON cl.id=ct.cliente_id "
        "WHERE ct.estado IN ('por_vencer','vencido') ORDER BY ct.fecha_vencimiento LIMIT 8",
        ()
    )
    sin_visita = mysql_fetchall(
        """SELECT c.id, c.razon_social,
                  MAX(v.fecha_realizada) AS ultima_visita
           FROM mant_clientes c
           LEFT JOIN mant_visitas v ON v.cliente_id=c.id AND v.estado='completada'
           WHERE c.estado='activo'
           GROUP BY c.id
           HAVING ultima_visita IS NULL OR ultima_visita < %s
           ORDER BY ultima_visita LIMIT 6""",
        (hoy - timedelta(days=180),)
    )
    # Ingresos del mes
    ingresos_mes = mysql_fetchone(
        "SELECT COALESCE(SUM(monto_mensual),0) AS total FROM mant_contratos "
        "WHERE estado IN ('vigente','indefinido')", ()
    ) or {}

    return render_template("mantenciones/index.html",
        kpi_clientes   = clientes.get("n", 0),
        kpi_contratos  = contratos.get("n", 0),
        kpi_vencen     = vencen.get("n", 0),
        kpi_ingresos   = float(ingresos_mes.get("total", 0)),
        prox_visitas   = [dict(r) for r in prox_visitas],
        alertas        = [dict(r) for r in alertas_contratos],
        sin_visita     = [dict(r) for r in sin_visita],
        hoy            = hoy,
    )


@app.route("/mantenciones/clientes")
@_mant_required
def mant_clientes():
    q            = request.args.get("q", "").strip()
    estado       = request.args.get("estado", "activo")
    contrato_fil = request.args.get("contrato", "")
    vista        = request.args.get("vista", "grid")

    where, params = ["1=1"], []
    if estado:
        where.append("c.estado=%s"); params.append(estado)
    if q:
        where.append("(c.razon_social LIKE %s OR c.rut LIKE %s OR c.contacto_email LIKE %s)")
        qp = f"%{q}%"; params += [qp, qp, qp]
    wstr = " AND ".join(where)

    # Reescritura: 5 subqueries correlacionadas (N×5 ejecuciones) → 4 derived tables agregadas (1 ejecución cada una).
    # Resultado: pasa de ~1500 queries lógicas con 300 clientes a ~5 queries totales.
    rows = mysql_fetchall(f"""
        SELECT c.*,
               COALESCE(m_agg.cnt, 0)       AS maquinas_count,
               COALESCE(c_agg.cnt, 0)       AS contratos_count,
               c_latest.estado              AS contrato_estado,
               c_latest.fecha_vencimiento   AS contrato_vencimiento,
               v_next.fecha_programada      AS prox_visita
        FROM mant_clientes c
        LEFT JOIN (
            SELECT cliente_id, COUNT(*) AS cnt
            FROM mant_maquinas WHERE estado='activo' GROUP BY cliente_id
        ) m_agg ON m_agg.cliente_id = c.id
        LEFT JOIN (
            SELECT cliente_id, COUNT(*) AS cnt
            FROM mant_contratos GROUP BY cliente_id
        ) c_agg ON c_agg.cliente_id = c.id
        LEFT JOIN (
            SELECT ct.cliente_id, ct.estado, ct.fecha_vencimiento
            FROM mant_contratos ct
            INNER JOIN (
                SELECT cliente_id, MAX(id) AS max_id
                FROM mant_contratos GROUP BY cliente_id
            ) latest ON latest.cliente_id = ct.cliente_id AND latest.max_id = ct.id
        ) c_latest ON c_latest.cliente_id = c.id
        LEFT JOIN (
            SELECT cliente_id, MIN(fecha_programada) AS fecha_programada
            FROM mant_visitas
            WHERE estado='programada' AND fecha_programada >= CURDATE()
            GROUP BY cliente_id
        ) v_next ON v_next.cliente_id = c.id
        WHERE {wstr}
        ORDER BY c.razon_social LIMIT 300
    """, tuple(params))
    clientes = [dict(r) for r in rows]

    # Post-filtro contrato
    if contrato_fil == 'vigente':
        clientes = [c for c in clientes if c.get('contrato_estado') == 'vigente']
    elif contrato_fil == 'vencido':
        clientes = [c for c in clientes if c.get('contrato_estado') in ('vencido','por_vencer')]
    elif contrato_fil == 'sin':
        clientes = [c for c in clientes if not c.get('contratos_count')]

    # Enriquecer cada cliente: completaje, badge, días a próx visita
    today_d = datetime.today().date()
    for c in clientes:
        # Normalizar fechas
        for fld in ('contrato_vencimiento', 'prox_visita'):
            if c.get(fld) and isinstance(c[fld], datetime):
                c[fld] = c[fld].date()
        c['prox_visita_dias'] = (c['prox_visita'] - today_d).days if c.get('prox_visita') else None
        # Completaje 0–100 (perfil + operacional)
        sc = 0
        if c.get('rut'):             sc += 12
        if c.get('contacto_nombre'): sc += 12
        if c.get('contacto_email'):  sc += 12
        if c.get('contacto_tel'):    sc += 10
        if c.get('direccion'):       sc += 10
        if c.get('comuna'):          sc += 9
        if c.get('contrato_estado'): sc += 20
        if (c.get('maquinas_count') or 0) > 0: sc += 15
        c['completaje'] = sc
        # Badge alerta combinado
        ce = c.get('contrato_estado')
        if c.get('estado') == 'inactivo':
            c['badge'] = 'inactivo'
        elif ce in ('vencido', 'por_vencer') and c.get('estado') == 'activo':
            c['badge'] = 'advertencia'
        else:
            c['badge'] = c.get('estado') or 'activo'
        # Normalizar a 'sin_contrato' para display uniforme
        if not c.get('contrato_estado'):
            c['contrato_estado'] = 'sin_contrato'

    # Stats globales (sin filtros)
    gs = mysql_fetchone("""
        SELECT
          (SELECT COUNT(*) FROM mant_clientes)                                              AS total,
          (SELECT COUNT(*) FROM mant_clientes WHERE estado='activo')                        AS activos,
          (SELECT COUNT(DISTINCT ct.cliente_id) FROM mant_contratos ct
           WHERE ct.estado='vigente')                                                        AS con_contrato,
          (SELECT COUNT(DISTINCT m.cliente_id)  FROM mant_maquinas m)                       AS con_equipos,
          (SELECT COUNT(*) FROM mant_visitas v
           WHERE v.estado='programada'
             AND v.fecha_programada BETWEEN CURDATE()
             AND DATE_ADD(CURDATE(), INTERVAL 30 DAY))                                      AS visitas_30d,
          (SELECT COUNT(*) FROM mant_contratos ct2
           WHERE ct2.estado IN ('vencido','por_vencer'))                                     AS contratos_alerta
    """)
    global_stats = dict(gs) if gs else {}

    # Completaje promedio + notificaciones pendientes
    if clientes:
        global_stats['completaje_avg'] = round(sum(c['completaje'] for c in clientes) / len(clientes))
    else:
        global_stats['completaje_avg'] = 0
    notif_r = mysql_fetchone(
        "SELECT COUNT(*) AS n FROM mant_notificaciones WHERE estado='pendiente'", ()
    )
    global_stats['notificaciones'] = (notif_r['n'] if notif_r else 0)

    return render_template("mantenciones/clientes.html",
        clientes     = clientes,
        filtros      = {"q": q, "estado": estado, "contrato": contrato_fil, "vista": vista},
        global_stats = global_stats,
    )


@app.route("/mantenciones/clientes/wizard")
@_mant_required
def mant_cliente_wizard():
    """Wizard inteligente de 4 pasos para crear cliente de mantención."""
    return render_template("mantenciones/cliente_wizard.html")


def _erp_buscar_clientes(q, limit=20):
    """
    Busca clientes en el ERP por RUT o razón social.
    Normaliza el RUT (con/sin puntos) para máxima compatibilidad.
    """
    ERP_SALES = ERP_CONFIG.get("table_sales", "HEBDOC")
    # Normalizar RUT: buscar con y sin puntos
    q_like       = f"%{q}%"
    q_sin_puntos = q.replace(".", "").replace(" ", "")
    q_sin_like   = f"%{q_sin_puntos}%"
    erp_conn = get_erp_conn()
    if not erp_conn:
        return []
    try:
        with erp_conn.cursor() as cur:
            cur.execute(
                f"""SELECT DISTINCT
                       TRIM(d.NRAZON) AS razon_social,
                       TRIM(d.NRUC)   AS rut
                    FROM `{ERP_SALES}` d
                    WHERE (
                        TRIM(d.NRAZON) LIKE %s
                        OR TRIM(d.NRUC)   LIKE %s
                        OR REPLACE(REPLACE(TRIM(d.NRUC),'.',''),' ','') LIKE %s
                    )
                      AND d.TIDO IN ('FCV','BLV','NVV','VD','WEB','FCO','NVI','GDV')
                    ORDER BY d.NRAZON
                    LIMIT {int(limit)}""",
                (q_like, q_like, q_sin_like)
            )
            rows = cur.fetchall()
        return [{"razon_social": r["razon_social"], "rut": r["rut"]} for r in rows if r.get("razon_social")]
    except Exception:
        return []
    finally:
        try: erp_conn.close()
        except: pass


@app.route("/mantenciones/api/clientes/autocomplete")
@_mant_required
def mant_clientes_autocomplete():
    """
    Autocomplete unificado: busca en clientes locales (mant_clientes) + ERP.
    Devuelve lista ordenada con origen (local/erp) para mostrar en dropdown.
    """
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify([])

    # 1) Buscar en clientes locales (siempre disponible)
    q_like = f"%{q}%"
    locales = mysql_fetchall(
        """SELECT id, razon_social, rut, contacto_email, estado,
                  ciudad AS region, comuna, direccion, contacto_tel AS telefono
           FROM mant_clientes
           WHERE razon_social LIKE %s OR rut LIKE %s
           ORDER BY razon_social LIMIT 15""",
        (q_like, q_like)
    )
    resultados = []
    ids_rut_vistos = set()
    for r in locales:
        rut = (r.get("rut") or "").strip()
        resultados.append({
            "id":           r["id"],
            "razon_social": r["razon_social"],
            "rut":          rut,
            "email":        r.get("contacto_email",""),
            "region":       r.get("region",""),
            "comuna":       r.get("comuna",""),
            "direccion":    r.get("direccion",""),
            "telefono":     r.get("telefono",""),
            "estado":       r.get("estado",""),
            "origen":       "local",
        })
        if rut: ids_rut_vistos.add(rut)

    # 2) Buscar en ERP vía REST API /entidades (funciona desde Railway)
    TOKEN = ERP_CONFIG.get("api_token", "")
    try:
        body     = _erp_get("/entidades", {"search": q, "empresa": "01", "limit": "30"}, TOKEN, timeout=8)
        ent_list = body.get("data") or (body if isinstance(body, list) else [])
        for e in ent_list[:25]:
            rut    = (e.get("RTEN")   or "").strip()
            nombre = (e.get("NOKOEN") or "").strip()
            if not nombre:
                continue
            if rut and rut in ids_rut_vistos:
                continue
            # Capturar región, comuna, dirección y teléfono si el ERP los entrega
            region  = (e.get("NOKOREG")  or e.get("NOKOREGIO") or e.get("REGION") or "").strip()
            comuna  = (e.get("NOKOCOMU") or e.get("NOKOCOMUNADE") or e.get("COMUNA") or e.get("NOKOMUNNE") or "").strip()
            dir_    = (e.get("DIEN")     or e.get("DIRESP") or e.get("DIENDESP") or e.get("DIENDE") or "").strip()
            tel     = (e.get("FOEN")     or e.get("FONOEN") or e.get("TELEN")  or "").strip()
            resultados.append({
                "id":           None,
                "razon_social": nombre,
                "rut":          rut,
                "email":        (e.get("EMAIL") or e.get("COREN") or "").strip(),
                "region":       region,
                "comuna":       comuna,
                "direccion":    dir_,
                "telefono":     tel,
                "giro":         (e.get("GIEN") or e.get("GIRO") or "").strip(),
                "estado":       "",
                "origen":       "erp",
            })
            if rut: ids_rut_vistos.add(rut)
    except Exception:
        # Fallback: MySQL directo si REST falla (p.ej. entorno local)
        erp_rows = _erp_buscar_clientes(q, limit=15)
        for r in erp_rows:
            rut = (r.get("rut") or "").strip()
            if rut in ids_rut_vistos:
                continue
            resultados.append({
                "id":           None,
                "razon_social": r["razon_social"],
                "rut":          rut,
                "email":        "",
                "region":       "",
                "comuna":       "",
                "direccion":    "",
                "telefono":     "",
                "estado":       "",
                "origen":       "erp",
            })
            if rut: ids_rut_vistos.add(rut)

    # Ordenar locales primero, luego ERP
    resultados.sort(key=lambda x: (0 if x["origen"]=="local" else 1, x["razon_social"].lower()))
    return jsonify(resultados[:20])


@app.route("/mantenciones/api/erp-rut", methods=["POST"])
@_mant_required
def mant_erp_rut_lookup():
    """Busca un cliente en el ERP/local por RUT y devuelve sus datos básicos."""
    d   = request.get_json(silent=True) or {}
    rut = d.get("rut", "").strip()
    if not rut:
        return jsonify({"error": "RUT requerido"}), 400

    # 1) Buscar en clientes locales primero
    local = mysql_fetchone(
        "SELECT * FROM mant_clientes WHERE rut=%s OR rut LIKE %s LIMIT 1",
        (rut, f"%{rut.split('-')[0]}%")
    )
    if local:
        return jsonify({
            "encontrado":    True,
            "origen":        "local",
            "id":            local["id"],
            "razon_social":  local["razon_social"],
            "rut":           local["rut"] or rut,
            "direccion":     local.get("direccion",""),
            "comuna":        local.get("comuna",""),
            "ciudad":        local.get("ciudad",""),
            "email":         local.get("contacto_email",""),
            "contacto":      local.get("contacto_nombre",""),
            "tel":           local.get("contacto_tel",""),
        })

    # 2) Buscar en ERP
    rows = _erp_buscar_clientes(rut, limit=3)
    if rows:
        r = rows[0]
        return jsonify({
            "encontrado":   True,
            "origen":       "erp",
            "id":           None,
            "razon_social": r["razon_social"],
            "rut":          r["rut"],
            "direccion":    "",
            "comuna":       "",
            "ciudad":       "",
            "email":        "",
        })

    return jsonify({"encontrado": False})


_MANT_IMG_EXTS = ("jpg", "jpeg", "png", "webp")

@app.route("/mantenciones/api/agente-contrato", methods=["POST"])
@_mant_required
def mant_agente_contrato():
    """
    AGENTE IA DE CONTRATOS — lee PDF, Word o FOTO (cámara móvil) y extrae:
    - Datos del CLIENTE: RUT, razón social, dirección, contacto
    - Análisis del CONTRATO: tipo, vigencia, SLA, cláusulas, costos, riesgos
    - Equipos mencionados en el contrato
    Soporta: PDF, DOCX, JPG, PNG (foto del contrato desde celular).
    """
    f = request.files.get("archivo")
    if not f or not f.filename:
        return jsonify({"error": "Sin archivo"}), 400

    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ("pdf", "doc", "docx") + _MANT_IMG_EXTS:
        return jsonify({"error": "Formato no válido. Usa PDF, Word o imagen (JPG/PNG)."}), 400

    # Verificar API key antes de procesar el archivo
    ai_key = _get_ai_key()
    if not ai_key:
        return jsonify({"error": (
            "⚠️ API de IA no configurada. "
            "Agrega tu ANTHROPIC_API_KEY en Railway → Variables de entorno."
        )}), 503

    # Guardar temporalmente
    tmp_path = os.path.join(MANT_UPLOADS, f"tmp_{int(time.time())}_{secure_filename(f.filename)}")
    f.save(tmp_path)

    texto = ""
    img_b64 = None
    img_media_type = None
    is_image = ext in _MANT_IMG_EXTS

    try:
        if is_image:
            with open(tmp_path, "rb") as img_f:
                img_b64 = base64.b64encode(img_f.read()).decode()
            img_media_type = f"image/{'jpeg' if ext == 'jpg' else ext}"
        elif ext == "pdf":
            import pdfplumber
            with pdfplumber.open(tmp_path) as pdf:
                texto = "\n".join(p.extract_text() or "" for p in pdf.pages[:20])
        elif ext in ("doc", "docx"):
            import docx as _docx
            doc = _docx.Document(tmp_path)
            texto = "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return jsonify({"error": f"No se pudo leer el archivo: {e}"}), 500
    finally:
        try: os.remove(tmp_path)
        except: pass

    if not is_image and not texto.strip():
        return jsonify({"error": "No se pudo extraer texto del documento. Prueba subir el PDF o una foto del contrato."}), 422

    prompt_agente = """Eres un agente experto en análisis de contratos de mantención de equipos fitness en Chile (ILUS Fitness).
Tu tarea es extraer TODA la información del contrato y responder ÚNICAMENTE con JSON estructurado, sin texto adicional.

Estructura requerida:
{
  "cliente": {
    "razon_social": "nombre legal completo o null",
    "rut": "RUT con formato XX.XXX.XXX-X o null",
    "direccion": "dirección completa o null",
    "comuna": "comuna o null",
    "ciudad": "ciudad o null",
    "region": "región o null",
    "contacto_nombre": "nombre del contacto principal o null",
    "contacto_cargo": "cargo del contacto o null",
    "contacto_email": "email o null",
    "contacto_tel": "teléfono o null"
  },
  "prestador": {
    "razon_social": "empresa prestadora (usualmente ILUS) o null",
    "rut": "RUT prestador o null",
    "contacto_nombre": "técnico/representante o null"
  },
  "contrato": {
    "nombre": "nombre o título del contrato",
    "numero": "número de contrato o null",
    "tipo_contrato": "Preventivo|Correctivo|Full|Garantía|Mixto|Otro",
    "vigencia_inicio": "YYYY-MM-DD o null",
    "vigencia_fin": "YYYY-MM-DD o null",
    "es_indefinido": true_o_false,
    "renovacion_automatica": true_o_false,
    "dias_aviso_termino": número_entero_o_null,
    "frecuencia_meses": número_entero_o_null,
    "visitas_anuales": número_entero_o_null,
    "horario_atencion": "descripción horario o null",
    "sla_horas": número_o_null,
    "tiempo_respuesta_urgente_horas": número_o_null,
    "monto_mensual": número_o_null,
    "costo_por_mant": número_o_null,
    "costo_total": número_o_null,
    "moneda": "CLP|UF|USD",
    "forma_pago": "descripción o null",
    "incluye_mant_gratis": true_o_false,
    "incluye_repuestos": true_o_false,
    "limite_repuestos": "descripción límite de repuestos o null",
    "cobertura_descripcion": "qué cubre el contrato exactamente",
    "exclusiones": ["exclusión1", "exclusión2"],
    "penalidades": "descripción de penalidades por incumplimiento o null",
    "nivel_riesgo": "alto|medio|bajo",
    "score": número_0_a_100,
    "resumen": "2-3 oraciones resumiendo el contrato para el prestador",
    "clausulas_criticas": ["clausula crítica 1", "clausula crítica 2"],
    "alertas": ["alerta operativa 1", "alerta operativa 2"],
    "mejoras_prioritarias": ["mejora sugerida 1", "mejora sugerida 2"]
  },
  "equipos": [
    {
      "nombre": "nombre del equipo fitness",
      "marca": "marca o null",
      "modelo": "modelo o null",
      "sku": "código o null",
      "cantidad": número_entero,
      "ubicacion": "sala/piso/zona o null",
      "notas": "observaciones o null"
    }
  ],
  "instalaciones": [
    {
      "nombre": "nombre de la instalación/sede o null",
      "direccion": "dirección o null",
      "comuna": "comuna o null"
    }
  ]
}

Criterios de evaluación:
- score 80-100: contrato muy favorable para el prestador (buenas tarifas, SLA razonable, cobertura clara)
- score 50-79: contrato aceptable con algunas condiciones a revisar
- score 20-49: contrato desfavorable (tarifas bajas, SLA exigente, sin límite de repuestos)
- score 0-19: contrato de alto riesgo financiero u operativo

Si el documento es una fotografía del contrato, extrae la información visible con la misma rigurosidad.
Si un dato no aparece en el documento, usa null. No inventes información."""

    prompt_usuario = f"""Analiza este contrato de mantención de equipos fitness y extrae TODA la información:

{texto[:8000] if not is_image else '[Contrato en imagen adjunta — analiza todo el texto visible]'}

Incluye: datos del cliente, condiciones contractuales, costos, equipos mencionados, cláusulas críticas, riesgos y alertas operativas."""

    try:
        import anthropic as _anthropic
        cliente_ia = _anthropic.Anthropic(api_key=ai_key)

        if is_image:
            # Análisis por visión (foto del contrato desde celular)
            msg = cliente_ia.messages.create(
                model="claude-opus-4-5",
                max_tokens=2500,
                system=prompt_agente,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {
                        "type": "base64",
                        "media_type": img_media_type,
                        "data": img_b64
                    }},
                    {"type": "text", "text": prompt_usuario}
                ]}]
            )
        else:
            msg = cliente_ia.messages.create(
                model="claude-opus-4-5",
                max_tokens=2500,
                system=prompt_agente,
                messages=[{"role": "user", "content": prompt_usuario}]
            )

        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
        resultado = json.loads(raw)
    except Exception as e:
        return jsonify({"error": f"Error en análisis IA: {e}"}), 500

    # Cruzar RUT detectado con ERP/local para enriquecer
    rut_detectado = (resultado.get("cliente") or {}).get("rut","")
    erp_data = {}
    if rut_detectado:
        local = mysql_fetchone(
            "SELECT * FROM mant_clientes WHERE rut LIKE %s LIMIT 1",
            (f"%{rut_detectado.split('-')[0]}%",)
        )
        if local:
            erp_data = {"origen":"local","id":local["id"],
                        "razon_social":local["razon_social"],
                        "rut":local["rut"],"estado":local["estado"]}
        else:
            rows = _erp_buscar_clientes(rut_detectado.split("-")[0], limit=1)
            if rows:
                erp_data = {"origen":"erp","id":None,
                            "razon_social":rows[0]["razon_social"],
                            "rut":rows[0]["rut"]}

    resultado["_erp_match"] = erp_data
    return jsonify({"ok": True, "resultado": resultado})


@app.route("/mantenciones/api/clientes/<int:cid>/generar-calendario", methods=["POST"])
@_mant_required
def mant_generar_calendario(cid):
    """
    Genera visitas de mantención automáticas basadas en el contrato activo del cliente.
    Devuelve preview (dry_run=true) o guarda en DB.
    """
    d        = request.get_json(silent=True) or {}
    dry_run  = d.get("dry_run", True)
    desde_str= d.get("desde")     # YYYY-MM-DD
    meses    = int(d.get("meses", 12))
    tipo     = d.get("tipo", "preventiva")
    tecnico  = d.get("tecnico", "")

    # Obtener frecuencia del contrato activo
    ct = mysql_fetchone(
        """SELECT * FROM mant_contratos
           WHERE cliente_id=%s AND estado IN ('vigente','indefinido')
           ORDER BY created_at DESC LIMIT 1""",
        (cid,)
    )
    cliente = mysql_fetchone("SELECT razon_social FROM mant_clientes WHERE id=%s", (cid,))
    if not ct:
        return jsonify({"error": "Sin contrato activo para este cliente"}), 404

    frecuencia = ct.get("ai_frecuencia_sug") or ct.get("frecuencia_meses") or 3
    desde = datetime.strptime(desde_str, "%Y-%m-%d").date() if desde_str else datetime.now().date()

    visitas_preview = []
    fecha_actual = desde
    while fecha_actual <= desde + timedelta(days=meses * 30):
        visitas_preview.append({
            "fecha":    str(fecha_actual),
            "titulo":   f"Mantención {tipo.capitalize()} — {cliente['razon_social'] if cliente else ''}",
            "tipo":     tipo,
            "tecnico":  tecnico,
            "contrato_id": ct["id"],
        })
        # Siguiente visita
        mes = fecha_actual.month + frecuencia
        anio = fecha_actual.year
        while mes > 12:
            mes -= 12
            anio += 1
        try:
            fecha_actual = fecha_actual.replace(year=anio, month=mes)
        except ValueError:
            fecha_actual = fecha_actual.replace(year=anio, month=mes, day=28)

    if dry_run:
        return jsonify({"ok": True, "preview": visitas_preview, "frecuencia": frecuencia})

    # Guardar visitas
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            for v in visitas_preview:
                cur.execute(
                    """INSERT INTO mant_visitas
                       (cliente_id,contrato_id,titulo,fecha_programada,tipo,estado,created_by)
                       VALUES (%s,%s,%s,%s,%s,'programada',%s)""",
                    (cid, v["contrato_id"], v["titulo"], v["fecha"], v["tipo"], current_username())
                )
        conn.commit()
        _mant_log("cliente", cid, "calendario_generado", f"{len(visitas_preview)} visitas")
        return jsonify({"ok": True, "creadas": len(visitas_preview)})
    finally:
        conn.close()


@app.route("/mantenciones/api/contratos/<int:ctid>/ai-editar", methods=["PUT"])
@_mant_required
def mant_contrato_ai_editar(ctid):
    """Guarda los campos del análisis IA editados manualmente por el usuario."""
    d = request.get_json(silent=True) or {}
    editable_fields = [
        "nombre", "fecha_inicio", "fecha_vencimiento", "es_indefinido",
        "monto_mensual", "monto_anual", "frecuencia_meses", "notas", "estado",
        "sla_horas", "incluye_repuestos", "incluye_mant_gratis",
        "costo_por_mant", "costo_total", "nivel_riesgo",
        "ai_tipo_contrato", "ai_cobertura", "ai_vigencia_inicio", "ai_vigencia_fin",
    ]
    sets = [f"{f}=%s" for f in editable_fields if f in d]
    vals = [d[f] for f in editable_fields if f in d]
    # Guardar snapshot editable como JSON también
    sets.append("ai_editable=%s")
    vals.append(json.dumps(d, ensure_ascii=False))
    if not sets:
        return jsonify({"error": "Sin campos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_contratos SET {','.join(sets)} WHERE id=%s", vals + [ctid])
        conn.commit()
        _mant_log("contrato", ctid, "ai_editado_manual")
        return jsonify({"ok": True})
    finally:
        conn.close()


# (mant_adjunto_subir y mant_adjunto_ver migrados a mant_contrato_adjuntos — ver sección ADJUNTOS)


@app.route("/mantenciones/clientes/nuevo", methods=["GET", "POST"])
@_mant_required
def mant_cliente_nuevo():
    if request.method == "POST":
        d = request.form
        is_wizard = request.headers.get("X-Wizard") == "1"

        # ── VALIDACIONES de datos reales ─────────────────────────────────
        razon = (d.get("razon_social") or "").strip()
        if not razon:
            err = "La razón social es obligatoria."
            if is_wizard: return jsonify({"ok": False, "error": err}), 400
            flash(err, "danger")
            return render_template("mantenciones/cliente_form.html", cliente=None)
        if len(razon) > 200:
            err = "La razón social excede 200 caracteres."
            if is_wizard: return jsonify({"ok": False, "error": err}), 400
            flash(err, "danger")
            return render_template("mantenciones/cliente_form.html", cliente=None)

        # RUT (opcional, pero si viene debe ser válido + único)
        rut_input = (d.get("rut") or "").strip()
        rut_norm = None
        if rut_input:
            ok, val_or_err = validar_rut(rut_input)
            if not ok:
                err = f"RUT inválido: {val_or_err}"
                if is_wizard: return jsonify({"ok": False, "error": err}), 400
                flash(err, "danger")
                return render_template("mantenciones/cliente_form.html", cliente=None)
            rut_norm = val_or_err
            # Check duplicado (UNIQUE index nos protege, pero damos error amigable antes)
            existing = mysql_fetchone(
                "SELECT id, razon_social FROM mant_clientes WHERE rut=%s LIMIT 1",
                (rut_norm,)
            )
            if existing:
                err = (f"Ya existe un cliente con ese RUT: "
                       f"«{existing.get('razon_social','')}» (ID {existing.get('id')}).")
                if is_wizard: return jsonify({"ok": False, "error": err,
                                              "duplicate_id": existing.get("id")}), 409
                flash(err, "warning")
                return redirect(url_for("mant_ficha", cid=existing.get("id")))

        # Emails (institucional + 2 contactos)
        emails_to_check = [
            ("email_empresa",   d.get("email_empresa")),
            ("contacto_email",  d.get("contacto_email")),
            ("contacto2_email", d.get("contacto2_email")),
        ]
        emails_norm = {}
        for campo, val in emails_to_check:
            ok, val_or_err = validar_email(val)
            if not ok:
                err = f"Email inválido ({campo}): {val_or_err}"
                if is_wizard: return jsonify({"ok": False, "error": err}), 400
                flash(err, "danger")
                return render_template("mantenciones/cliente_form.html", cliente=None)
            emails_norm[campo] = val_or_err

        # Teléfonos (normalizados pero no rechazamos por formato — formato chileno es flexible)
        tel_empresa     = normalizar_telefono(d.get("tel_empresa"))
        tel_contacto    = normalizar_telefono(d.get("contacto_tel"))
        tel_contacto2   = normalizar_telefono(d.get("contacto2_tel"))

        conn = get_mysql()
        try:
            with conn.cursor() as cur:
                try:
                    cur.execute(
                        """INSERT INTO mant_clientes
                           (razon_social, rut,
                            email_empresa, tel_empresa,
                            contacto_nombre, contacto_cargo, contacto_tel, contacto_email,
                            contacto2_nombre, contacto2_cargo, contacto2_tel, contacto2_email,
                            direccion, comuna, ciudad, region, giro,
                            notas, notas_confidenciales,
                            estado, created_by, updated_by)
                           VALUES (%s,%s, %s,%s,
                                   %s,%s,%s,%s,
                                   %s,%s,%s,%s,
                                   %s,%s,%s,%s,%s,
                                   %s,%s,
                                   %s,%s,%s)""",
                        (razon, rut_norm,
                         emails_norm["email_empresa"], tel_empresa,
                         (d.get("contacto_nombre") or "").strip()[:200],
                         (d.get("contacto_cargo")  or "").strip()[:120],
                         tel_contacto, emails_norm["contacto_email"],
                         (d.get("contacto2_nombre") or "").strip()[:200],
                         (d.get("contacto2_cargo")  or "").strip()[:120],
                         tel_contacto2, emails_norm["contacto2_email"],
                         (d.get("direccion") or "").strip(),
                         (d.get("comuna")    or "").strip()[:100],
                         (d.get("ciudad")    or "").strip()[:100],
                         (d.get("region")    or "").strip()[:100],
                         (d.get("giro")      or "").strip()[:200],
                         (d.get("notas")     or "").strip(),
                         (d.get("notas_confidenciales") or "").strip(),
                         (d.get("estado") or "activo"),
                         current_username(), current_username())
                    )
                    cid = cur.lastrowid
                except Exception as e:
                    # Manejo amigable del error de UNIQUE rut (carrera con check anterior)
                    if "Duplicate entry" in str(e) and "uq_mc_rut" in str(e):
                        err = "Ya existe un cliente con ese RUT (revisar lista de clientes)."
                        if is_wizard: return jsonify({"ok": False, "error": err}), 409
                        flash(err, "warning")
                        return redirect(url_for("mant_clientes"))
                    raise
            conn.commit()
            _mant_log("cliente", cid, "creado", razon)
            if is_wizard:
                return jsonify({"ok": True, "id": cid})
            return redirect(url_for("mant_ficha", cid=cid))
        finally:
            conn.close()
    return render_template("mantenciones/cliente_form.html", cliente=None)


@app.route("/mantenciones/api/clientes/enriquecer")
@_mant_required
def mant_enriquecer_cliente():
    """
    Consulta /entidades del ERP por RUT exacto y devuelve datos normalizados.
    GET /mantenciones/api/clientes/enriquecer?rut=77.123.456-7
    Usa _normalize_phone_cl y _cmen_to_comuna para máxima calidad.
    """
    rut = (request.args.get("rut") or "").strip()
    if not rut:
        return jsonify({"error": "RUT requerido"}), 400
    TOKEN = ERP_CONFIG.get("api_token", "")
    try:
        body = _erp_get("/entidades", {"rten": rut}, TOKEN, timeout=8)
        data = body.get("data") or []
        if not data:
            return jsonify({"encontrado": False, "rut": rut})
        e    = data[0]
        cien = (e.get("CIEN") or "").strip()
        cmen = (e.get("CMEN") or "").strip()
        raw_tel = (e.get("FOEN") or e.get("FAEN") or "").strip()
        region  = _REGION_NOMBRES.get(str(cien).zfill(3), "")
        return jsonify({
            "encontrado":    True,
            "razon_social":  (e.get("NOKOEN") or "").strip().title(),
            "rut":           (e.get("RTEN")   or rut).strip(),
            "email":         (e.get("EMAIL")  or e.get("EMAILCOMER") or "").strip(),
            "telefono":      _normalize_phone_cl(raw_tel),
            "direccion":     (e.get("DIEN")   or "").strip().title(),
            "comuna":        _cmen_to_comuna(cien, cmen),
            "region":        region,
            "giro":          (e.get("GIEN")   or "").strip(),
            "observaciones": (e.get("OBEN")   or "").strip(),
            "cien":          cien,
            "cmen":          cmen,
        })
    except Exception as ex:
        return jsonify({"error": str(ex), "encontrado": False}), 503


@app.route("/mantenciones/api/ultimo-cliente")
@_mant_required
def mant_ultimo_cliente():
    """Devuelve el último cliente creado por el usuario actual (para el wizard)."""
    row = mysql_fetchone(
        "SELECT id, razon_social FROM mant_clientes WHERE created_by=%s ORDER BY created_at DESC LIMIT 1",
        (current_username(),)
    )
    if not row:
        return jsonify({"error": "No encontrado"}), 404
    return jsonify({"id": row["id"], "razon_social": row["razon_social"]})


@app.route("/mantenciones/clientes/<int:cid>")
@_mant_required
def mant_ficha(cid):
    cliente   = mysql_fetchone("SELECT * FROM mant_clientes WHERE id=%s", (cid,))
    if not cliente:
        return redirect(url_for("mant_clientes"))

    # ── Helper: normaliza datetime/date de MySQL a datetime.date ──────────
    def _d(val):
        """Convierte datetime o date a date para comparaciones seguras."""
        if val is None:
            return None
        return val.date() if hasattr(val, 'date') else val

    # ── Normaliza filas de DB para evitar mezcla datetime/date ───────────
    def _norm_maquina(row):
        r = dict(row)
        for k in ('doc_fecha', 'fecha_instalacion', 'created_at', 'updated_at'):
            if k in r:
                r[k] = _d(r[k])
        return r

    def _norm_contrato(row):
        r = dict(row)
        for k in ('fecha_inicio', 'fecha_vencimiento', 'ai_fecha',
                  'ai_vigencia_inicio', 'ai_vigencia_fin',
                  'created_at', 'updated_at'):
            if k in r:
                r[k] = _d(r[k])
        return r

    def _norm_visita(row):
        r = dict(row)
        r['fecha_programada'] = _d(r.get('fecha_programada'))
        if 'created_at' in r:
            r['created_at'] = _d(r['created_at'])
        return r

    maquinas_raw  = mysql_fetchall("SELECT * FROM mant_maquinas WHERE cliente_id=%s ORDER BY created_at DESC", (cid,))
    contratos_raw = mysql_fetchall("SELECT * FROM mant_contratos WHERE cliente_id=%s ORDER BY created_at DESC", (cid,))
    visitas_raw   = mysql_fetchall(
        "SELECT * FROM mant_visitas WHERE cliente_id=%s ORDER BY fecha_programada DESC", (cid,)
    )
    # Historial: traer hasta 250 registros del cliente, sus contratos, equipos, visitas y reportes
    # asociados — el frontend luego permite filtrar por acción, usuario y texto.
    logs = mysql_fetchall(
        """SELECT * FROM mant_logs
           WHERE (entidad='cliente' AND entidad_id=%s)
              OR (entidad='contrato' AND entidad_id IN (SELECT id FROM mant_contratos WHERE cliente_id=%s))
              OR (entidad='maquina'  AND entidad_id IN (SELECT id FROM mant_maquinas  WHERE cliente_id=%s))
              OR (entidad='visita'   AND entidad_id IN (SELECT id FROM mant_visitas   WHERE cliente_id=%s))
              OR (entidad='reporte'  AND entidad_id IN (SELECT id FROM mant_reportes  WHERE cliente_id=%s))
           ORDER BY created_at DESC LIMIT 250""",
        (cid, cid, cid, cid, cid)
    )

    # Normalizar todas las fechas a datetime.date
    maquinas     = [_norm_maquina(r)  for r in maquinas_raw]
    contratos    = [_norm_contrato(r) for r in contratos_raw]
    visitas_full = [_norm_visita(r)   for r in visitas_raw]

    # ── ESTADÍSTICAS PARA SIDEBAR / GRÁFICOS ──────────────────────────────
    hoy = datetime.now().date()
    fecha_corte_12m = hoy - timedelta(days=365)

    visitas_12m = [v for v in visitas_full if v.get("fecha_programada") and v["fecha_programada"] >= fecha_corte_12m]
    visitas_realizadas = [v for v in visitas_12m if v.get("estado") == "completada"]
    visitas_programadas_30d = [
        v for v in visitas_full
        if v.get("estado") == "programada"
        and v.get("fecha_programada")
        and hoy <= v["fecha_programada"] <= (hoy + timedelta(days=30))
    ]
    visitas_correctivas = [v for v in visitas_12m if v.get("tipo") == "correctiva"]

    total_programadas_12m  = len([v for v in visitas_12m if v.get("estado") in ("programada","completada","reagendada")])
    total_realizadas_12m   = len(visitas_realizadas)
    cumplimiento_pct       = round(100 * total_realizadas_12m / total_programadas_12m) if total_programadas_12m else 0

    # Mantenciones por mes (últimos 12 meses) — para barra
    import calendar
    months_es = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
    mant_por_mes = []
    for i in range(11, -1, -1):
        ref_year  = hoy.year
        ref_month = hoy.month - i
        while ref_month <= 0:
            ref_month += 12
            ref_year  -= 1
        cnt = sum(1 for v in visitas_12m
                  if v.get("fecha_programada")
                  and v["fecha_programada"].month == ref_month
                  and v["fecha_programada"].year  == ref_year
                  and v.get("estado") == "completada")
        mant_por_mes.append({
            "mes":   months_es[ref_month-1],
            "year":  ref_year,
            "count": cnt,
        })

    # Estado de equipos (donut)
    eq_activos      = len([m for m in maquinas if (m.get("estado") or "activo") == "activo"])
    eq_advertencia  = len([m for m in maquinas if m.get("estado") == "advertencia"])
    eq_inactivos    = len([m for m in maquinas if m.get("estado") == "inactivo"])
    eq_total        = max(1, len(maquinas))

    # Próxima visita
    prox_visita = next((v for v in visitas_full
                        if v.get("estado") == "programada"
                        and v.get("fecha_programada")
                        and v["fecha_programada"] >= hoy), None)
    # Última visita realizada
    ult_visita = next((v for v in visitas_full
                       if v.get("estado") == "completada"
                       and v.get("fecha_programada")
                       and v["fecha_programada"] <= hoy), None)
    dias_desde_ultima = (hoy - ult_visita["fecha_programada"]).days if ult_visita else None
    dias_hasta_prox   = (prox_visita["fecha_programada"] - hoy).days if prox_visita else None

    # Alertas inteligentes (calculadas server-side)
    alertas_smart = []
    # 1. Equipos con garantía próxima a vencer
    for m in maquinas:
        if m.get("doc_fecha"):
            try:
                dias_doc = (hoy - m["doc_fecha"]).days
                if 150 <= dias_doc <= 180:
                    alertas_smart.append({
                        "tipo": "warning",
                        "icon": "shield-exclamation",
                        "titulo": f"{m['nombre']} — Garantía próxima a vencer",
                        "detalle": f"Quedan {180 - dias_doc} días de garantía",
                    })
            except Exception:
                pass
    # 2. Contratos por vencer
    for ct in contratos:
        if ct.get("fecha_vencimiento") and not ct.get("es_indefinido"):
            try:
                dias_v = (ct["fecha_vencimiento"] - hoy).days
                if 0 < dias_v <= 60:
                    alertas_smart.append({
                        "tipo": "warning",
                        "icon": "file-earmark-text",
                        "titulo": f"Contrato vence en {dias_v} días",
                        "detalle": ct.get("nombre") or "Contrato sin nombre",
                    })
                elif dias_v <= 0:
                    alertas_smart.append({
                        "tipo": "danger",
                        "icon": "file-earmark-x",
                        "titulo": "Contrato VENCIDO",
                        "detalle": ct.get("nombre") or "Contrato sin nombre",
                    })
            except Exception:
                pass
    # 3. Mucho tiempo sin visita
    if dias_desde_ultima is not None and dias_desde_ultima > 90:
        alertas_smart.append({
            "tipo": "warning",
            "icon": "calendar-x",
            "titulo": f"Sin visita hace {dias_desde_ultima} días",
            "detalle": "Considera programar una mantención preventiva",
        })
    # 4. Próxima mantención cercana
    if dias_hasta_prox is not None and 0 <= dias_hasta_prox <= 7:
        alertas_smart.append({
            "tipo": "info",
            "icon": "calendar-check",
            "titulo": f"Próxima visita en {dias_hasta_prox} día{'s' if dias_hasta_prox != 1 else ''}",
            "detalle": prox_visita.get("titulo") or prox_visita.get("tipo","").title(),
        })

    stats = {
        "realizadas_12m":          total_realizadas_12m,
        "programadas_30d":         len(visitas_programadas_30d),
        "incidencias_abiertas":    len([v for v in visitas_full if v.get("tipo") == "correctiva" and v.get("estado") == "programada"]),
        "cumplimiento_pct":        cumplimiento_pct,
        "mant_por_mes":            mant_por_mes,
        "max_mes":                 max((m["count"] for m in mant_por_mes), default=1) or 1,
        "eq_activos":              eq_activos,
        "eq_advertencia":          eq_advertencia,
        "eq_inactivos":            eq_inactivos,
        "eq_total":                len(maquinas),
        "eq_pct_activos":          round(100 * eq_activos / eq_total),
        "eq_pct_advertencia":      round(100 * eq_advertencia / eq_total),
        "eq_pct_inactivos":        round(100 * eq_inactivos / eq_total),
        "ult_visita_fecha":        ult_visita["fecha_programada"] if ult_visita else None,
        "ult_visita_dias":         dias_desde_ultima,
        "prox_visita_fecha":       prox_visita["fecha_programada"] if prox_visita else None,
        "prox_visita_dias":        dias_hasta_prox,
        "alertas_smart":           alertas_smart,
        "total_correctivas_12m":   len(visitas_correctivas),
    }

    # Normaliza logs: conserva datetime convertido a hora Chile (America/Santiago).
    # MySQL en cloud guarda en UTC; sin conversión la hora mostrada está mal.
    try:
        from zoneinfo import ZoneInfo
        _tz_scl = ZoneInfo("America/Santiago")
        _tz_utc = ZoneInfo("UTC")
    except Exception:
        _tz_scl = _tz_utc = None

    def _to_chile(dt):
        if dt is None or _tz_scl is None or not isinstance(dt, datetime):
            return dt
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=_tz_utc)
        return dt.astimezone(_tz_scl)

    def _norm_log(row):
        r = dict(row)
        r['created_at'] = _to_chile(r.get('created_at'))
        return r

    # Sucursales (información adicional opcional)
    try:
        sucursales = mysql_fetchall(
            "SELECT * FROM mant_sucursales WHERE cliente_id=%s AND activo=1 ORDER BY nombre",
            (cid,)
        )
        sucursales = [dict(s) for s in (sucursales or [])]
    except Exception:
        sucursales = []

    return render_template("mantenciones/ficha.html",
        cliente   = dict(cliente),
        maquinas  = maquinas,
        contratos = contratos,
        visitas   = visitas_full[:50],
        logs      = [_norm_log(r) for r in logs],
        hoy       = hoy,
        stats     = stats,
        sucursales = sucursales,
    )


@app.route("/mantenciones/api/clientes/<int:cid>", methods=["PUT"])
@_mant_required
def mant_cliente_update(cid):
    d = request.get_json(silent=True) or {}
    fields = ["razon_social","rut","email_empresa","tel_empresa","giro",
              "contacto_nombre","contacto_cargo","contacto_tel","contacto_email",
              "contacto2_nombre","contacto2_cargo","contacto2_tel","contacto2_email",
              "direccion","comuna","ciudad","region",
              "notas","notas_confidenciales","estado"]
    sets   = [f"{f}=%s" for f in fields if f in d]
    vals   = [d[f] for f in fields if f in d]
    if not sets:
        return jsonify({"error": "Sin campos"}), 400
    # Auditoría: quién actualizó
    sets.append("updated_by=%s")
    vals.append(current_username())
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_clientes SET {','.join(sets)} WHERE id=%s",
                        vals + [cid])
        conn.commit()
        _mant_log("cliente", cid, "editado")
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/clientes/<int:cid>", methods=["DELETE"])
@_mant_required
def mant_cliente_delete(cid):
    """Elimina un cliente y TODOS sus datos asociados.

    Doble protección:
      1. Solo admin/superadmin puede ejecutar (verificado por permisos).
      2. Requiere confirm_text en el body que coincida con razón social o RUT.
    """
    # Verificación de rol
    if not (g.permissions.get("admin") or g.permissions.get("superadmin")):
        return jsonify({"error":"Solo administradores pueden eliminar clientes"}), 403

    cliente = mysql_fetchone(
        "SELECT id, razon_social, rut FROM mant_clientes WHERE id=%s", (cid,)
    )
    if not cliente:
        return jsonify({"error":"Cliente no encontrado"}), 404

    d = request.get_json(silent=True) or {}
    confirm = (d.get("confirm_text") or "").strip().lower()
    rs = (cliente.get("razon_social") or "").strip().lower()
    rut = re.sub(r"[^0-9kK]","",(cliente.get("rut") or "")).lower()
    confirm_norm = re.sub(r"[^0-9kK]","",confirm) if confirm.replace(".","").replace("-","").replace("k","").isdigit() else confirm
    if confirm not in (rs, (cliente.get("rut") or "").lower()) and confirm_norm != rut:
        return jsonify({
            "error": f"Para confirmar, escribe exactamente la razón social o el RUT del cliente.",
            "expected_rs": cliente.get("razon_social"),
            "expected_rut": cliente.get("rut"),
        }), 400

    # Inventario de lo que se va a eliminar (para auditoría)
    counts = {}
    try:
        for tbl, label in [
            ("mant_maquinas","equipos"),
            ("mant_contratos","contratos"),
            ("mant_visitas","visitas"),
            ("mant_reportes","reportes"),
            ("mant_repuestos","repuestos"),
            ("mant_notificaciones","notificaciones"),
        ]:
            row = mysql_fetchone(f"SELECT COUNT(*) AS n FROM {tbl} WHERE cliente_id=%s",(cid,))
            counts[label] = (row or {}).get("n", 0)
    except Exception: pass

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Borrar dependencias que NO tienen ON DELETE CASCADE explícito
            cur.execute("DELETE FROM mant_repuestos WHERE cliente_id=%s", (cid,))
            cur.execute("DELETE FROM mant_reportes WHERE cliente_id=%s", (cid,))
            cur.execute("DELETE FROM mant_notificaciones WHERE cliente_id=%s", (cid,))
            cur.execute("DELETE FROM mant_logs WHERE entidad='cliente' AND entidad_id=%s", (cid,))
            # mant_maquinas, mant_contratos, mant_visitas tienen ON DELETE CASCADE
            cur.execute("DELETE FROM mant_clientes WHERE id=%s", (cid,))
            # AUDIT: registrar la eliminación con entidad='global' para que persista
            cur.execute(
                "INSERT INTO mant_logs (entidad,entidad_id,accion,detalle,usuario) "
                "VALUES ('global', 0, 'cliente_eliminado', %s, %s)",
                (f"#{cid} {cliente.get('razon_social') or ''} (RUT {cliente.get('rut') or '—'}) "
                 f"— {counts}",
                 current_username())
            )
        conn.commit()
        print(f"[MANT][DEL] cliente#{cid} '{cliente.get('razon_social')}' eliminado por {current_username()} — {counts}")
        return jsonify({
            "ok": True,
            "razon_social": cliente.get("razon_social"),
            "rut": cliente.get("rut"),
            "eliminado": counts,
        })
    except Exception as exc:
        return jsonify({"error": f"No se pudo eliminar: {exc}"}), 500
    finally:
        conn.close()


# ── MÁQUINAS ──────────────────────────────────────────────────────────

def _generar_serie_ilus(cid: int, sku: str = "") -> str:
    """
    Genera un N° Serie único para un equipo físico cuando el fabricante no
    proporciona uno.

    Formato amigable: {RUT}-{SKU4}-{n}
      RUT  = RUT del cliente sin DV ni puntos (ej: 65206047)
      SKU4 = últimos 4 caracteres alfanuméricos del SKU del modelo
      n    = secuencial dentro del cliente para ese SKU

    Ejemplos:
      65206047-0905-1   (Trotadora ILUS, SKU 1027100905, primera unidad)
      65206047-0905-2   (segunda trotadora del mismo modelo)
      65206047-4640-1   (otro equipo, SKU termina en 4640)

    Único globalmente porque incluye RUT del cliente (identidad legal).
    Editable: el usuario puede reemplazar con el serial real del fabricante.
    """
    # Obtener RUT del cliente (sin DV, sin puntos)
    rut_clean = "00000000"
    try:
        row = mysql_fetchone("SELECT rut FROM mant_clientes WHERE id=%s", (cid,))
        if row and row.get("rut"):
            raw = str(row["rut"]).replace(".", "").replace(" ", "").replace("-", "").upper()
            rut_clean = raw[:-1] if len(raw) >= 8 else raw
    except Exception:
        pass

    # Últimos 4 chars alfanum del SKU
    sku_clean = "".join(c for c in (sku or "AUTO").upper() if c.isalnum())
    sku4 = sku_clean[-4:] if len(sku_clean) >= 4 else (sku_clean.rjust(4, "0") if sku_clean else "AUTO")
    base = f"{rut_clean}-{sku4}"
    rows = mysql_fetchall(
        "SELECT serie FROM mant_maquinas WHERE cliente_id=%s AND serie LIKE %s",
        (cid, f"{base}-%")
    )
    usados = set()
    for r in rows or []:
        suf = (r.get("serie") or "").rsplit("-", 1)[-1]
        try: usados.add(int(suf))
        except Exception: pass
    seq = 1
    while seq in usados:
        seq += 1
    return f"{base}-{seq}"


@app.route("/mantenciones/api/clientes/<int:cid>/maquinas", methods=["POST"])
@_mant_required
def mant_maquina_add(cid):
    d = request.get_json(silent=True) or {}
    # N° Serie: si viene del front lo usamos, si no auto-generamos uno único.
    # Esto cubre los 2 casos: (a) representación trae serial, (b) sistema lo crea
    serie = (d.get("serie") or "").strip()
    if not serie or serie.startswith("(auto"):
        serie = _generar_serie_ilus(cid, d.get("sku", ""))
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_maquinas
                   (cliente_id,sku,nombre,serie,doc_origen,doc_fecha,cantidad,notas,
                    ubicacion_cliente,estado_op,fecha_instalacion,
                    tag_1,tag_2,
                    justif_fecha_inst,justif_doc_mismatch,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (cid, d.get("sku",""), d.get("nombre",""), serie,
                 d.get("doc_origen",""), d.get("doc_fecha") or None,
                 int(d.get("cantidad",1)), d.get("notas",""),
                 d.get("ubicacion_cliente",""),
                 d.get("estado_op","operativo"),
                 d.get("fecha_instalacion") or None,
                 (d.get("tag_1") or "").strip()[:120] or None,
                 (d.get("tag_2") or "").strip()[:120] or None,
                 d.get("justif_fecha_inst") or None,
                 d.get("justif_doc_mismatch") or None,
                 current_username())
            )
            mid = cur.lastrowid
        conn.commit()
        _mant_log("maquina", mid, "agregada", d.get("nombre",""))
        return jsonify({"ok": True, "id": mid, "serie": serie})
    finally:
        conn.close()


@app.route("/mantenciones/api/maquinas/<int:mid>", methods=["DELETE"])
@_mant_required
def mant_maquina_del(mid):
    m_info = mysql_fetchone(
        "SELECT cliente_id, nombre, sku, serie FROM mant_maquinas WHERE id=%s", (mid,)
    )
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM mant_maquinas WHERE id=%s", (mid,))
        conn.commit()
        if m_info:
            detalle = f"{m_info.get('nombre') or ''} (SKU {m_info.get('sku') or '—'} · Serie {m_info.get('serie') or '—'})"
            _mant_log("maquina", mid, "eliminada", detalle)
            _mant_log("cliente", m_info.get("cliente_id"), "equipo_eliminado", detalle)
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/maquinas/<int:mid>/serie", methods=["PUT"])
@_mant_required
def mant_maquina_actualizar_serie(mid):
    """
    Actualiza el N° serie de un equipo con auditoría completa.
    Cada cambio queda registrado en mant_maquina_audit:
      - valor anterior, valor nuevo, motivo, usuario, fecha

    Body JSON:
      serie:  str (nuevo valor — vacío vuelve a auto-generar)
      motivo: str (mín 5 chars, justificación del cambio)
    """
    d = request.get_json(silent=True) or {}
    serie_nueva_raw = (d.get("serie") or "").strip()[:120]
    motivo = (d.get("motivo") or "").strip()[:500]
    if len(motivo) < 5:
        return jsonify({"error": "El motivo debe tener al menos 5 caracteres"}), 400

    # Cargar equipo actual
    maq = mysql_fetchone(
        "SELECT id, cliente_id, sku, serie FROM mant_maquinas WHERE id=%s", (mid,)
    )
    if not maq:
        return jsonify({"error": "Equipo no encontrado"}), 404

    serie_anterior = maq.get("serie") or ""

    # Si quedó vacío → regenerar auto
    if not serie_nueva_raw:
        serie_nueva = _generar_serie_ilus(maq["cliente_id"], maq.get("sku", ""))
    else:
        serie_nueva = serie_nueva_raw

    if serie_nueva == serie_anterior:
        return jsonify({"ok": True, "sin_cambios": True, "serie": serie_nueva})

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Actualizar el equipo
            cur.execute("UPDATE mant_maquinas SET serie=%s WHERE id=%s",
                        (serie_nueva, mid))
            # Registrar en auditoría dedicada
            cur.execute(
                """INSERT INTO mant_maquina_audit
                   (maquina_id, cliente_id, campo, valor_antes, valor_nuevo, motivo, usuario)
                   VALUES (%s,%s,'serie',%s,%s,%s,%s)""",
                (mid, maq["cliente_id"], serie_anterior, serie_nueva, motivo, current_username())
            )
        conn.commit()
        # También al log general para que aparezca en historial de la ficha
        _mant_log("maquina", mid, "serie_cambiada",
                  f"'{serie_anterior}' → '{serie_nueva}'. Motivo: {motivo}")
        return jsonify({
            "ok": True,
            "serie": serie_nueva,
            "serie_anterior": serie_anterior,
            "usuario": current_username(),
        })
    finally:
        conn.close()


@app.route("/mantenciones/api/maquinas/<int:mid>/audit", methods=["GET"])
@_mant_required
def mant_maquina_audit_list(mid):
    """Devuelve el historial de cambios sensibles de un equipo (serie, etc)."""
    rows = mysql_fetchall(
        "SELECT * FROM mant_maquina_audit WHERE maquina_id=%s ORDER BY fecha DESC LIMIT 50",
        (mid,)
    )
    # Convertir a hora Chile
    try:
        from zoneinfo import ZoneInfo
        tz_scl = ZoneInfo("America/Santiago")
        tz_utc = ZoneInfo("UTC")
        out = []
        for r in (rows or []):
            r = dict(r)
            f = r.get("fecha")
            if isinstance(f, datetime):
                if f.tzinfo is None: f = f.replace(tzinfo=tz_utc)
                r["fecha"] = f.astimezone(tz_scl).strftime("%d/%m/%Y %H:%M")
            out.append(r)
        return jsonify({"ok": True, "audit": out})
    except Exception:
        return jsonify({"ok": True, "audit": [dict(r) for r in (rows or [])]})


@app.route("/mantenciones/api/clientes/<int:cid>/visita-multi", methods=["POST"])
@_mant_required
def mant_visita_multi(cid):
    """
    Crea UNA sola visita asociada a N equipos del cliente.
    Caso de uso real: "el lunes voy a Vitacura a cambiar 4 trotadoras + revisar 2 bicis"
    → 1 visita técnica con 6 equipos involucrados, en lugar de 6 visitas separadas.

    Body JSON:
      maquina_ids: list[int]      — equipos a incluir
      tipo_visita: str             — garantia|correctiva|preventiva|inspeccion
      fecha_programada: 'YYYY-MM-DD'
      motivo: str (mín 8 chars)
      estado_nuevo: 'critico'|'en_mantencion'|'operativo'  (se aplica a todos)
      tecnico: str (opcional, nombre libre — DEPRECADO; usar tecnico_id)
      tecnico_id: int (opcional, FK a mant_tecnicos)
      titulo: str (opcional, generado si no viene)
      hora_inicio: 'HH:MM' (opcional)
      hora_fin: 'HH:MM' (opcional)
      costo: float (opcional, CLP)
      observaciones: str (opcional, hasta 1000 chars)
    """
    d = request.get_json(silent=True) or {}
    mids = d.get("maquina_ids") or []
    if not isinstance(mids, list) or not mids:
        return jsonify({"error": "Debes seleccionar al menos un equipo"}), 400
    try:
        mids = [int(m) for m in mids]
    except (TypeError, ValueError):
        return jsonify({"error": "IDs de equipo inválidos"}), 400

    motivo = (d.get("motivo") or "").strip()
    if len(motivo) < 8:
        return jsonify({"error": "El motivo debe tener al menos 8 caracteres"}), 400

    tipo_visita = d.get("tipo_visita") or "preventiva"
    if tipo_visita not in ("garantia","correctiva","preventiva","inspeccion"):
        tipo_visita = "preventiva"
    estado_nuevo = d.get("estado_nuevo") or "operativo"
    if estado_nuevo not in ("critico","en_mantencion","operativo"):
        estado_nuevo = "operativo"
    # Fecha por defecto: +48 horas. Solo el superadmin puede modificarla en frontend,
    # pero validamos también en backend que el campo no quede vacío.
    fecha_prog = d.get("fecha_programada") or (datetime.today().date() + timedelta(days=2)).isoformat()
    # Multi-técnico: aceptamos lista de IDs (1..10). Compatibilidad con tecnico_id único.
    tecnico_ids_raw = d.get("tecnico_ids") or []
    if not isinstance(tecnico_ids_raw, list):
        tecnico_ids_raw = []
    if not tecnico_ids_raw and d.get("tecnico_id"):
        tecnico_ids_raw = [d.get("tecnico_id")]
    tecnico_ids = []
    for x in tecnico_ids_raw:
        try:
            xi = int(x)
            if xi > 0:
                tecnico_ids.append(xi)
        except (TypeError, ValueError):
            pass
    tecnico_ids = list(dict.fromkeys(tecnico_ids))[:10]  # dedupe + máx 10

    # Resolver nombres y tarifas desde la tabla. El campo `tecnico` (texto libre)
    # se conserva por compatibilidad con visitas viejas — se usa nombre concatenado.
    tecnicos_data = []
    if tecnico_ids:
        placeholders_t = ",".join(["%s"] * len(tecnico_ids))
        rows_t = mysql_fetchall(
            f"SELECT id,nombre,tarifa_visita FROM mant_tecnicos WHERE id IN ({placeholders_t}) AND activo=1",
            tuple(tecnico_ids)
        ) or []
        tecnicos_data = [dict(r) for r in rows_t]
    tecnico_id     = tecnicos_data[0]["id"]     if tecnicos_data else None
    tecnico_nombre = (
        ", ".join(t["nombre"] for t in tecnicos_data)
        if tecnicos_data
        else ((d.get("tecnico") or "").strip()[:200] or None)
    )

    titulo_input = (d.get("titulo") or "").strip()[:200]

    # Hora inicio / fin — autocompleta ceros: '8' -> '08:00', '12' -> '12:00'
    hora_inicio = _normalize_hora(d.get("hora_inicio"))
    hora_fin    = _normalize_hora(d.get("hora_fin"))

    # Costo (CLP). Si el usuario no lo indica → cálculo automático:
    # tarifa_visita × cantidad de técnicos asignados (default $50.000 por técnico).
    costo = d.get("costo")
    try:
        costo = float(costo) if costo not in (None, "", 0, "0") else None
        if costo is not None and costo < 0:
            costo = None
    except (TypeError, ValueError):
        costo = None
    if costo is None and tecnicos_data:
        costo_auto = 0.0
        for t in tecnicos_data:
            tarifa = float(t.get("tarifa_visita") or 50000)
            costo_auto += tarifa
        costo = costo_auto if costo_auto > 0 else None

    observaciones = (d.get("observaciones") or "").strip()[:1000] or None

    # Repuestos asociados (lista de objetos con sku/descripcion/cantidad/costo_unitario/origen)
    repuestos_raw = d.get("repuestos") or []
    if not isinstance(repuestos_raw, list):
        repuestos_raw = []
    repuestos = []
    repuestos_total = 0.0
    for rp in repuestos_raw[:50]:  # max 50 por visita
        try:
            desc = (rp.get("descripcion") or "").strip()[:300]
            if not desc:
                continue
            cant = float(rp.get("cantidad") or 1) or 1.0
            cunit = float(rp.get("costo_unitario") or 0) or 0.0
            ctot = round(cant * cunit, 2)
            repuestos.append({
                "sku":            (rp.get("sku") or "").strip()[:80] or None,
                "producto_id":    int(rp["producto_id"]) if rp.get("producto_id") else None,
                "descripcion":    desc,
                "cantidad":       cant,
                "costo_unitario": cunit,
                "costo_total":    ctot,
                "origen":         rp.get("origen") if rp.get("origen") in ("manual","erp","catalogo") else "manual",
                "notas":          (rp.get("notas") or "").strip()[:500] or None,
            })
            repuestos_total += ctot
        except (TypeError, ValueError):
            continue
    # Si vinieron repuestos y el costo no fue indicado manualmente,
    # se SUMA al costo total de la visita
    if repuestos_total > 0 and costo is not None:
        costo += repuestos_total
    elif repuestos_total > 0 and costo is None:
        costo = repuestos_total

    # Generar número de OT correlativo
    numero_ot = _next_ot_number()

    # Cargar todos los equipos seleccionados (verificar que pertenecen al cliente)
    placeholders = ",".join(["%s"] * len(mids))
    rows = mysql_fetchall(
        f"SELECT id,nombre,sku,serie,cantidad,doc_origen FROM mant_maquinas "
        f"WHERE id IN ({placeholders}) AND cliente_id=%s",
        tuple(mids) + (cid,)
    )
    if not rows or len(rows) != len(mids):
        return jsonify({"error": "Algunos equipos no pertenecen a este cliente"}), 400

    # Construir título y descripción
    tipos_label = {
        "garantia":   "Cambio / Garantía",
        "correctiva": "Reparación correctiva",
        "preventiva": "Mantención preventiva",
        "inspeccion": "Inspección / Levantamiento",
    }
    titulo = titulo_input or f"{tipos_label[tipo_visita]} — {len(rows)} equipo(s)"

    detalle_equipos = "\n".join(
        f"  • {r['nombre']} (SKU {r.get('sku') or '—'}, Serie {r.get('serie') or '—'})"
        for r in rows
    )
    descripcion = (
        f"VISITA MULTI-EQUIPO — {tipos_label[tipo_visita]}\n"
        f"Equipos involucrados ({len(rows)}):\n{detalle_equipos}\n\n"
        f"Motivo:\n{motivo}\n"
    )

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # 1. Crear la visita única (con número de OT correlativo)
            cur.execute(
                """INSERT INTO mant_visitas
                   (numero_ot,cliente_id,titulo,fecha_programada,hora_inicio,hora_fin,
                    tipo,estado,descripcion,observaciones,tecnico,tecnico_id,
                    costo,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,'programada',%s,%s,%s,%s,%s,%s)""",
                (numero_ot, cid, titulo, fecha_prog, hora_inicio, hora_fin,
                 tipo_visita, descripcion, observaciones, tecnico_nombre, tecnico_id,
                 costo, current_username())
            )
            vid = cur.lastrowid

            # 1.b Insertar técnicos asignados (N:N) si hay
            if tecnicos_data:
                for t in tecnicos_data:
                    cur.execute(
                        """INSERT INTO mant_visita_tecnicos (visita_id, tecnico_id, costo)
                           VALUES (%s, %s, %s)""",
                        (vid, t["id"], float(t.get("tarifa_visita") or 50000))
                    )

            # 1.c Insertar repuestos asociados a la visita
            for rp in repuestos:
                cur.execute(
                    """INSERT INTO mant_visita_repuestos
                       (visita_id, sku, producto_id, descripcion, cantidad,
                        costo_unitario, costo_total, origen, notas)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                    (vid, rp["sku"], rp["producto_id"], rp["descripcion"],
                     rp["cantidad"], rp["costo_unitario"], rp["costo_total"],
                     rp["origen"], rp["notas"])
                )

            # 2. Cambiar estado de los equipos
            cur.execute(
                f"UPDATE mant_maquinas SET estado_op=%s WHERE id IN ({placeholders})",
                (estado_nuevo,) + tuple(mids)
            )
            # 3. Loguear en cada equipo (auditoría)
            for r in rows:
                cur.execute(
                    "INSERT INTO mant_logs (entidad,entidad_id,accion,detalle,usuario) "
                    "VALUES ('maquina',%s,'visita_multi',%s,%s)",
                    (r["id"],
                     f"Visita {vid} ({tipo_visita}) — fecha {fecha_prog}. {motivo[:120]}",
                     current_username())
                )
        conn.commit()
        _mant_log("cliente", cid, "visita_multi_creada",
                  f"Visita {vid}: {len(rows)} equipos, tipo {tipo_visita}, fecha {fecha_prog}")
        return jsonify({
            "ok": True,
            "visita_id": vid,
            "numero_ot": numero_ot,
            "equipos_afectados": len(rows),
            "tecnicos_asignados": len(tecnicos_data),
            "repuestos_count": len(repuestos),
            "repuestos_total": repuestos_total,
            "costo_calculado": costo,
            "fecha_programada": fecha_prog,
            "cliente_id": cid,
        })
    except Exception as e:
        try:
            conn.rollback()
        except Exception:
            pass
        return jsonify({"error": f"Error al crear la visita: {str(e)}"}), 500
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════
# MÓDULO: TÉCNICOS — Catálogo de técnicos asignables a visitas
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/tecnicos")
@_mant_required
def mant_tecnicos_index():
    """Listado de técnicos."""
    q       = (request.args.get("q") or "").strip()
    estado  = request.args.get("estado", "activo")  # activo | inactivo | todos
    where, params = ["1=1"], []
    if estado == "activo":
        where.append("activo=1")
    elif estado == "inactivo":
        where.append("activo=0")
    if q:
        where.append("(nombre LIKE %s OR rut LIKE %s OR especialidad LIKE %s OR email LIKE %s)")
        like = f"%{q}%"
        params.extend([like, like, like, like])
    sql = f"""SELECT t.*,
                (SELECT COUNT(*) FROM mant_visitas WHERE tecnico_id=t.id) AS visitas_total,
                (SELECT COUNT(*) FROM mant_visitas
                  WHERE tecnico_id=t.id AND estado='programada'
                    AND fecha_programada >= CURDATE())                    AS visitas_pendientes
              FROM mant_tecnicos t
              WHERE {' AND '.join(where)}
              ORDER BY t.activo DESC, t.nombre"""
    tecnicos = mysql_fetchall(sql, tuple(params)) or []
    return render_template("mantenciones/tecnicos.html",
        tecnicos = [dict(r) for r in tecnicos],
        filtros  = {"q": q, "estado": estado},
    )


@app.route("/mantenciones/api/tecnicos", methods=["GET"])
@_mant_required
def mant_tecnicos_list_api():
    """JSON ligero para dropdowns. Solo activos por defecto."""
    incluir_inactivos = request.args.get("all") == "1"
    sql = "SELECT id,nombre,especialidad,nivel,telefono,email,tarifa_visita,activo,es_externo FROM mant_tecnicos"
    if not incluir_inactivos:
        sql += " WHERE activo=1"
    sql += " ORDER BY nombre"
    rows = mysql_fetchall(sql, ()) or []
    return jsonify([dict(r) for r in rows])


@app.route("/mantenciones/api/colaboradores-search", methods=["GET"])
@_mant_required
def mant_colab_search():
    """
    Autocomplete de colaboradores (HR) para importarlos como técnicos.
    Búsqueda por nombre o RUT, devuelve datos personales + dirección.
    """
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify([])
    like = f"%{q}%"
    rows = mysql_fetchall(
        f"""SELECT c.id, c.nombre_completo, c.rut, c.email, c.telefono,
                   c.direccion, c.comuna, c.region,
                   cg.nombre AS cargo
              FROM `{HRM_COLAB_TABLE}` c
              LEFT JOIN `{HRM_CARGOS_TABLE}` cg ON cg.id = c.cargo_id
             WHERE (c.nombre_completo LIKE %s OR c.rut LIKE %s)
               AND c.estado='activo'
             ORDER BY c.nombre_completo
             LIMIT 12""",
        (like, like)
    ) or []
    return jsonify([dict(r) for r in rows])


@app.route("/mantenciones/api/tecnicos", methods=["POST"])
@_mant_required
def mant_tecnico_crear():
    d = request.get_json(silent=True) or {}
    nombre = (d.get("nombre") or "").strip()[:200]
    if len(nombre) < 3:
        return jsonify({"error": "El nombre es obligatorio (mín 3 caracteres)"}), 400
    nivel = d.get("nivel") or "junior"
    if nivel not in ("junior","senior","externo"):
        nivel = "junior"
    es_externo = 1 if d.get("es_externo") else 0
    try:
        tarifa = float(d.get("tarifa_visita")) if d.get("tarifa_visita") not in (None, "") else None
        if tarifa is not None and tarifa < 0:
            tarifa = None
    except (TypeError, ValueError):
        tarifa = None
    fecha_ingreso = d.get("fecha_ingreso") or None

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_tecnicos
                   (nombre,rut,especialidad,nivel,telefono,email,direccion,comuna,region,
                    foto_url,notas,tarifa_visita,activo,es_externo,empresa_externa,
                    fecha_ingreso,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,1,%s,%s,%s,%s)""",
                (nombre,
                 (d.get("rut") or "").strip()[:20] or None,
                 (d.get("especialidad") or "").strip()[:160] or None,
                 nivel,
                 (d.get("telefono") or "").strip()[:50] or None,
                 (d.get("email") or "").strip()[:200] or None,
                 (d.get("direccion") or "").strip()[:300] or None,
                 (d.get("comuna") or "").strip()[:100] or None,
                 (d.get("region") or "").strip()[:100] or None,
                 (d.get("foto_url") or "").strip()[:500] or None,
                 (d.get("notas") or "").strip() or None,
                 tarifa,
                 es_externo,
                 (d.get("empresa_externa") or "").strip()[:200] or None,
                 fecha_ingreso,
                 current_username())
            )
            tid = cur.lastrowid
        conn.commit()
        _mant_log("cliente", 0, "tecnico_creado", f"Técnico {tid} ({nombre}) — nivel {nivel}")
        return jsonify({"ok": True, "id": tid, "nombre": nombre})
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({"error": f"No se pudo crear el técnico: {str(e)}"}), 500
    finally:
        conn.close()


@app.route("/mantenciones/tecnicos/<int:tid>")
@_mant_required
def mant_tecnico_ficha(tid):
    tecnico = mysql_fetchone("SELECT * FROM mant_tecnicos WHERE id=%s", (tid,))
    if not tecnico:
        flash("Técnico no encontrado", "danger")
        return redirect(url_for("mant_tecnicos_index"))
    visitas = mysql_fetchall(
        """SELECT v.*, c.razon_social
             FROM mant_visitas v
             JOIN mant_clientes c ON c.id=v.cliente_id
            WHERE v.tecnico_id=%s
            ORDER BY v.fecha_programada DESC LIMIT 50""",
        (tid,)
    ) or []
    # Stats
    stats = mysql_fetchone(
        """SELECT
              COUNT(*)                                                    AS total,
              SUM(estado='programada')                                    AS programadas,
              SUM(estado='completada')                                    AS completadas,
              SUM(estado='cancelada')                                     AS canceladas,
              COALESCE(SUM(costo),0)                                      AS costo_total
            FROM mant_visitas WHERE tecnico_id=%s""",
        (tid,)
    ) or {}
    return render_template("mantenciones/tecnico_ficha.html",
        tecnico = dict(tecnico),
        visitas = [dict(r) for r in visitas],
        stats   = dict(stats),
    )


@app.route("/mantenciones/api/tecnicos/<int:tid>", methods=["PUT"])
@_mant_required
def mant_tecnico_editar(tid):
    if not mysql_fetchone("SELECT id FROM mant_tecnicos WHERE id=%s", (tid,)):
        return jsonify({"error": "Técnico no encontrado"}), 404
    d = request.get_json(silent=True) or {}

    # Solo se actualizan los campos que llegan en el payload (PATCH-like)
    campos_validos = {
        "nombre":(str,200), "rut":(str,20), "especialidad":(str,160),
        "telefono":(str,50), "email":(str,200), "direccion":(str,300),
        "comuna":(str,100), "region":(str,100), "foto_url":(str,500),
        "notas":(str,5000), "empresa_externa":(str,200),
    }
    sets, params = [], []
    for k,(typ,maxlen) in campos_validos.items():
        if k in d:
            v = (d.get(k) or "").strip()[:maxlen] or None
            sets.append(f"{k}=%s"); params.append(v)
    if "nivel" in d:
        nv = d["nivel"] if d["nivel"] in ("junior","senior","externo") else "junior"
        sets.append("nivel=%s"); params.append(nv)
    if "activo" in d:
        sets.append("activo=%s"); params.append(1 if d["activo"] else 0)
    if "es_externo" in d:
        sets.append("es_externo=%s"); params.append(1 if d["es_externo"] else 0)
    if "tarifa_visita" in d:
        try:
            tv = float(d["tarifa_visita"]) if d["tarifa_visita"] not in (None, "") else None
        except (TypeError, ValueError):
            tv = None
        sets.append("tarifa_visita=%s"); params.append(tv)
    if "fecha_ingreso" in d:
        sets.append("fecha_ingreso=%s"); params.append(d["fecha_ingreso"] or None)

    if not sets:
        return jsonify({"error": "Sin cambios"}), 400

    params.append(tid)
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_tecnicos SET {', '.join(sets)} WHERE id=%s", tuple(params))
        conn.commit()
        _mant_log("cliente", 0, "tecnico_editado", f"Técnico {tid} actualizado")
        return jsonify({"ok": True})
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/mantenciones/api/tecnicos/<int:tid>", methods=["DELETE"])
@_mant_required
def mant_tecnico_eliminar(tid):
    """Soft delete: marca activo=0 si tiene visitas asociadas; hard delete si no."""
    if not g.permissions.get("superadmin"):
        return jsonify({"error": "Solo superadmin puede eliminar técnicos"}), 403
    n_vis = mysql_fetchone("SELECT COUNT(*) AS n FROM mant_visitas WHERE tecnico_id=%s", (tid,))
    tiene_visitas = (n_vis or {}).get("n", 0) > 0

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            if tiene_visitas:
                cur.execute("UPDATE mant_tecnicos SET activo=0 WHERE id=%s", (tid,))
                accion = "desactivado (tiene visitas históricas)"
            else:
                cur.execute("DELETE FROM mant_tecnicos WHERE id=%s", (tid,))
                accion = "eliminado"
        conn.commit()
        _mant_log("cliente", 0, "tecnico_eliminado", f"Técnico {tid} {accion}")
        return jsonify({"ok": True, "soft": tiene_visitas})
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/mantenciones/api/maquinas/<int:mid>/solicitar-cambio", methods=["POST"])
@_mant_required
def mant_maquina_solicitar_cambio(mid):
    """
    Solicita cambio/reparación de un equipo dañado:
      1. Marca el equipo como 'critico' (o 'en_mantencion')
      2. Crea una visita programada (tipo=garantia o correctiva) con
         descripción que incluye el motivo del daño y la cantidad afectada
      3. Loguea la acción para auditoría

    Body JSON:
      cantidad_afectada: int (cuántas unidades del equipo están dañadas)
      motivo: str (mín 8 chars, descripción del daño)
      tipo_visita: 'garantia' | 'correctiva' (default garantia)
      fecha_programada: 'YYYY-MM-DD' (default: hoy + 3 días)
      estado_nuevo: 'critico' | 'en_mantencion' (default critico)
    """
    d = request.get_json(silent=True) or {}
    motivo = (d.get("motivo") or "").strip()
    if len(motivo) < 8:
        return jsonify({"error": "El motivo debe tener al menos 8 caracteres"}), 400

    cant_afectada = int(d.get("cantidad_afectada") or 1)
    tipo_visita   = d.get("tipo_visita") or "garantia"
    if tipo_visita not in ("garantia","correctiva","preventiva","inspeccion"):
        tipo_visita = "garantia"
    estado_nuevo  = d.get("estado_nuevo") or "critico"
    if estado_nuevo not in ("critico","en_mantencion","operativo"):
        estado_nuevo = "critico"
    fecha_prog = d.get("fecha_programada") or (
        datetime.today().date() + timedelta(days=3)
    ).isoformat()
    tecnico_asign = (d.get("tecnico") or "").strip()[:200] or None

    # Cargar el equipo + cliente
    maq = mysql_fetchone(
        "SELECT id,cliente_id,nombre,sku,serie,cantidad,doc_origen "
        "FROM mant_maquinas WHERE id=%s", (mid,)
    )
    if not maq:
        return jsonify({"error": "Equipo no encontrado"}), 404
    cid = maq["cliente_id"]

    # Construir descripción completa para la visita
    titulo = f"Cambio/garantía: {maq.get('nombre','')[:80]}"
    descripcion = (
        f"SOLICITUD DE CAMBIO/REPARACIÓN\n"
        f"Equipo: {maq.get('nombre','')}\n"
        f"SKU del modelo: {maq.get('sku','') or '—'}\n"
        f"N° serie: {maq.get('serie','') or '—'}\n"
        f"Documento origen: {maq.get('doc_origen','') or '—'}\n"
        f"Unidades afectadas: {cant_afectada} de {maq.get('cantidad',1)}\n"
        f"Motivo: {motivo}\n"
    )

    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # 1. Cambiar estado del equipo
            cur.execute(
                "UPDATE mant_maquinas SET estado_op=%s WHERE id=%s",
                (estado_nuevo, mid)
            )
            # 2. Crear visita programada
            cur.execute(
                """INSERT INTO mant_visitas
                   (cliente_id,titulo,fecha_programada,tipo,estado,descripcion,tecnico,created_by)
                   VALUES (%s,%s,%s,%s,'programada',%s,%s,%s)""",
                (cid, titulo, fecha_prog, tipo_visita, descripcion, tecnico_asign, current_username())
            )
            vid = cur.lastrowid
        conn.commit()
        _mant_log("maquina", mid, "solicitud_cambio",
                  f"{cant_afectada} unidad(es), {tipo_visita}, fecha {fecha_prog}, motivo: {motivo[:80]}")
        return jsonify({
            "ok": True,
            "visita_id": vid,
            "cliente_id": cid,
            "fecha_programada": fecha_prog,
        })
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════
# SUCURSALES — información adicional opcional del cliente
# Para clientes con múltiples ubicaciones (cadenas, holding, franquicias).
# Cada sucursal tiene encargado y opcionalmente contacto secundario.
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/clientes/<int:cid>/sucursales", methods=["GET"])
@_mant_required
def mant_sucursales_list(cid):
    rows = mysql_fetchall(
        "SELECT * FROM mant_sucursales WHERE cliente_id=%s AND activo=1 ORDER BY nombre",
        (cid,)
    )
    return jsonify({"ok": True, "sucursales": [dict(r) for r in (rows or [])]})


@app.route("/mantenciones/api/clientes/<int:cid>/sucursales", methods=["POST"])
@_mant_required
def mant_sucursal_add(cid):
    d = request.get_json(silent=True) or {}
    nombre = (d.get("nombre") or "").strip()[:200]
    if not nombre:
        return jsonify({"error": "El nombre de la sucursal es obligatorio"}), 400
    es_principal = 1 if d.get("es_principal") else 0
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Si esta nueva va a ser principal, desmarcar las anteriores
            if es_principal:
                cur.execute(
                    "UPDATE mant_sucursales SET es_principal=0 WHERE cliente_id=%s AND es_principal=1",
                    (cid,)
                )
            cur.execute(
                """INSERT INTO mant_sucursales
                   (cliente_id,nombre,direccion,comuna,ciudad,region,
                    encargado_nombre,encargado_cargo,encargado_tel,encargado_email,
                    contacto2_nombre,contacto2_cargo,contacto2_tel,contacto2_email,
                    notas,es_principal,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (cid, nombre,
                 (d.get("direccion") or "").strip()[:300] or None,
                 (d.get("comuna") or "").strip()[:100] or None,
                 (d.get("ciudad") or "").strip()[:100] or None,
                 (d.get("region") or "").strip()[:100] or None,
                 (d.get("encargado_nombre") or "").strip()[:200] or None,
                 (d.get("encargado_cargo") or "").strip()[:120] or None,
                 (d.get("encargado_tel") or "").strip()[:50] or None,
                 (d.get("encargado_email") or "").strip()[:200] or None,
                 (d.get("contacto2_nombre") or "").strip()[:200] or None,
                 (d.get("contacto2_cargo") or "").strip()[:120] or None,
                 (d.get("contacto2_tel") or "").strip()[:50] or None,
                 (d.get("contacto2_email") or "").strip()[:200] or None,
                 (d.get("notas") or "").strip() or None,
                 es_principal,
                 current_username())
            )
            sid = cur.lastrowid
        conn.commit()
        _mant_log("cliente", cid, "sucursal_agregada", nombre + (" (principal)" if es_principal else ""))
        return jsonify({"ok": True, "id": sid, "es_principal": bool(es_principal)})
    finally:
        conn.close()


@app.route("/mantenciones/api/sucursales/<int:sid>", methods=["PUT"])
@_mant_required
def mant_sucursal_update(sid):
    d = request.get_json(silent=True) or {}
    fields = ["nombre","direccion","comuna","ciudad","region",
              "encargado_nombre","encargado_cargo","encargado_tel","encargado_email",
              "contacto2_nombre","contacto2_cargo","contacto2_tel","contacto2_email",
              "notas"]
    sets, vals = [], []
    for f in fields:
        if f in d:
            v = (d[f] or "").strip() if isinstance(d[f], str) else d[f]
            sets.append(f"{f}=%s")
            vals.append(v if v else None)
    # Manejo de es_principal con lógica única (solo una principal por cliente)
    es_principal = d.get("es_principal")
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            if es_principal is not None:
                ep = 1 if es_principal else 0
                if ep:
                    # Obtener el cliente_id de esta sucursal
                    row = mysql_fetchone("SELECT cliente_id FROM mant_sucursales WHERE id=%s", (sid,))
                    if row:
                        cur.execute(
                            "UPDATE mant_sucursales SET es_principal=0 WHERE cliente_id=%s AND id<>%s",
                            (row["cliente_id"], sid)
                        )
                sets.append("es_principal=%s")
                vals.append(ep)
            if not sets:
                return jsonify({"error": "Sin cambios"}), 400
            cur.execute(f"UPDATE mant_sucursales SET {','.join(sets)} WHERE id=%s",
                        vals + [sid])
        conn.commit()
        # Log con cliente_id y nombre de la sucursal
        suc_info = mysql_fetchone("SELECT cliente_id, nombre FROM mant_sucursales WHERE id=%s", (sid,))
        if suc_info:
            campos_mod = ", ".join([f.split("=")[0] for f in sets[:6]])
            _mant_log("cliente", suc_info["cliente_id"], "sucursal_actualizada",
                      f"{suc_info.get('nombre') or 'sucursal'} — campos: {campos_mod}")
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/sucursales/<int:sid>/marcar-principal", methods=["POST"])
@_mant_required
def mant_sucursal_marcar_principal(sid):
    """Marca una sucursal como principal (desmarca cualquier otra del mismo cliente)."""
    row = mysql_fetchone(
        "SELECT cliente_id, nombre FROM mant_sucursales WHERE id=%s AND activo=1",
        (sid,)
    )
    if not row:
        return jsonify({"error": "Sucursal no encontrada"}), 404
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE mant_sucursales SET es_principal=0 WHERE cliente_id=%s",
                (row["cliente_id"],)
            )
            cur.execute(
                "UPDATE mant_sucursales SET es_principal=1 WHERE id=%s",
                (sid,)
            )
        conn.commit()
        _mant_log("cliente", row["cliente_id"], "sucursal_principal_cambiada", row["nombre"])
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/sucursales/<int:sid>", methods=["DELETE"])
@_mant_required
def mant_sucursal_del(sid):
    """Soft delete: marca activo=0 para preservar histórico."""
    suc_info = mysql_fetchone("SELECT cliente_id, nombre FROM mant_sucursales WHERE id=%s", (sid,))
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("UPDATE mant_sucursales SET activo=0 WHERE id=%s", (sid,))
        conn.commit()
        if suc_info:
            _mant_log("cliente", suc_info["cliente_id"], "sucursal_eliminada", suc_info.get("nombre") or f"sucursal #{sid}")
        return jsonify({"ok": True})
    finally:
        conn.close()


# ── CONTRATOS ─────────────────────────────────────────────────────────

MANT_UPLOADS = os.path.join(BASE_DIR, "static", "uploads", "mantenciones")
os.makedirs(MANT_UPLOADS, exist_ok=True)

ALLOWED_CONTRATO = {"pdf", "doc", "docx"}


@app.route("/mantenciones/api/clientes/<int:cid>/contratos", methods=["POST"])
@_mant_required
def mant_contrato_subir(cid):
    f = request.files.get("archivo")
    d = request.form
    if not f or not f.filename:
        return jsonify({"error": "Sin archivo"}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_CONTRATO:
        return jsonify({"error": "Tipo no permitido"}), 400

    fname  = secure_filename(f"{cid}_{int(time.time())}_{f.filename}")
    fpath  = os.path.join(MANT_UPLOADS, fname)
    f.save(fpath)

    tipo = "pdf" if ext == "pdf" else ("word" if ext in ("doc","docx") else "otro")
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_contratos
                   (cliente_id,nombre,archivo_nombre,archivo_path,archivo_tipo,
                    fecha_inicio,fecha_vencimiento,es_indefinido,
                    monto_mensual,monto_anual,frecuencia_meses,notas,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (cid, d.get("nombre","Contrato"), f.filename, fname, tipo,
                 d.get("fecha_inicio") or None, d.get("fecha_vencimiento") or None,
                 1 if d.get("es_indefinido") else 0,
                 float(d.get("monto_mensual",0) or 0),
                 float(d.get("monto_anual",0) or 0),
                 int(d.get("frecuencia_meses",0) or 0),
                 d.get("notas",""), current_username())
            )
            ctid = cur.lastrowid
        conn.commit()
        _mant_log("contrato", ctid, "subido", f.filename)
        return jsonify({"ok": True, "id": ctid})
    finally:
        conn.close()


@app.route("/mantenciones/api/contratos/<int:ctid>", methods=["PUT"])
@_mant_required
def mant_contrato_update(ctid):
    d = request.get_json(silent=True) or {}
    allowed = ["nombre","fecha_inicio","fecha_vencimiento","es_indefinido",
               "monto_mensual","monto_anual","frecuencia_meses","notas","estado"]
    sets = [f"{f}=%s" for f in allowed if f in d]
    vals = [d[f] for f in allowed if f in d]
    if not sets:
        return jsonify({"error": "Sin campos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_contratos SET {','.join(sets)} WHERE id=%s",
                        vals + [ctid])
        conn.commit()
        # Log con campos modificados
        ct_info = mysql_fetchone("SELECT cliente_id, nombre FROM mant_contratos WHERE id=%s", (ctid,))
        campos_mod = ", ".join([f for f in allowed if f in d][:6])
        _mant_log("contrato", ctid, "actualizado", f"campos: {campos_mod}")
        if ct_info:
            _mant_log("cliente", ct_info["cliente_id"], "contrato_actualizado",
                      f"{ct_info.get('nombre') or ''} — {campos_mod}")
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/contratos/<int:ctid>/archivo")
@_mant_required
def mant_contrato_archivo(ctid):
    """Sirve el archivo del contrato. Todos los usuarios mant pueden VER; solo superadmin descarga."""
    from flask import send_from_directory
    ct = mysql_fetchone("SELECT * FROM mant_contratos WHERE id=%s", (ctid,))
    if not ct or not ct.get("archivo_path"):
        return "Archivo no encontrado", 404
    as_dl = request.args.get("download") == "1"
    # Descargar: solo superadmin
    if as_dl and not g.permissions.get("superadmin"):
        return ("Acceso restringido — solo el Superadministrador puede "
                "descargar archivos de contratos."), 403
    return send_from_directory(MANT_UPLOADS, ct["archivo_path"],
                               as_attachment=as_dl,
                               download_name=ct["archivo_nombre"] if as_dl else None)


# ── ANÁLISIS IA DE CONTRATO ───────────────────────────────────────────

@app.route("/mantenciones/api/contratos/<int:ctid>/analizar", methods=["POST"])
@_mant_required
def mant_contrato_analizar(ctid):
    """
    Llama a Claude API para analizar el contrato de mantención.
    Extrae: puntos críticos, frecuencia sugerida, alertas, score.
    """
    ct = mysql_fetchone("SELECT ct.*, cl.razon_social FROM mant_contratos ct "
                        "JOIN mant_clientes cl ON cl.id=ct.cliente_id "
                        "WHERE ct.id=%s", (ctid,))
    if not ct:
        return jsonify({"error": "Contrato no encontrado"}), 404

    # Leer texto del contrato si es PDF
    texto_contrato = ""
    fpath = os.path.join(MANT_UPLOADS, ct["archivo_path"] or "")
    if os.path.exists(fpath):
        ext = ct["archivo_path"].rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            try:
                import pdfplumber
                with pdfplumber.open(fpath) as pdf:
                    texto_contrato = "\n".join(
                        p.extract_text() or "" for p in pdf.pages[:15]
                    )
            except Exception:
                pass
        elif ext in ("doc", "docx"):
            try:
                import docx
                doc = docx.Document(fpath)
                texto_contrato = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                pass

    # Datos del contrato para enriquecer el prompt
    datos_extra = (
        f"Cliente: {ct['razon_social']}\n"
        f"Fecha inicio: {ct['fecha_inicio']}\n"
        f"Fecha vencimiento: {ct['fecha_vencimiento']}\n"
        f"Monto mensual: ${ct['monto_mensual']}\n"
        f"Frecuencia declarada: {ct['frecuencia_meses']} meses\n"
    )
    if not texto_contrato:
        texto_contrato = "(Texto no extraíble — análisis basado en metadatos)"

    prompt_sistema = """Eres un experto jurídico y técnico en contratos de mantención
de equipos de fitness para gimnasios y centros deportivos en Chile (ILUS Fitness).
Analiza el contrato con criterio profesional. Responde SIEMPRE en JSON con esta estructura EXACTA:
{
  "tipo_contrato": "Preventivo|Correctivo|Full|Garantía|Inspección|Otro",
  "resumen": "2-3 oraciones resumiendo el contrato",
  "score": 0-100,
  "nivel_riesgo": "alto|medio|bajo",
  "vigencia_inicio": "YYYY-MM-DD o null",
  "vigencia_fin": "YYYY-MM-DD o null",
  "es_indefinido": true_o_false,
  "frecuencia_sugerida_meses": número_entero,
  "sla_horas": número_entero_o_null,
  "incluye_mant_gratis": true_o_false,
  "incluye_repuestos": true_o_false,
  "cobertura_descripcion": "descripción de qué cubre el contrato",
  "costo_mensual": número_o_null,
  "costo_por_mant": número_o_null,
  "costo_total": número_o_null,
  "clausulas_criticas": ["clausula1","clausula2",...],
  "puntos_criticos": ["punto1","punto2",...],
  "alertas": ["alerta1","alerta2",...],
  "mejoras_prioritarias": ["mejora1","mejora2","mejora3"]
}
Sé específico sobre equipos fitness (treadmills, bikes, elípticas, pesas, etc.).
Detecta SLA, penalidades, cláusulas de exclusión y riesgos operativos para el prestador."""

    prompt_usuario = f"""Analiza este contrato de mantención:

METADATOS:
{datos_extra}

TEXTO DEL CONTRATO:
{texto_contrato[:6000]}

Devuelve SOLO el JSON, sin texto adicional."""

    ai_key = _get_ai_key()
    if not ai_key:
        return jsonify({"error": "⚠️ API de IA no configurada. Agrega ANTHROPIC_API_KEY en Railway."}), 503
    try:
        import anthropic as _anthropic
        client = _anthropic.Anthropic(api_key=ai_key)
        msg = client.messages.create(
            model   = "claude-opus-4-5",
            max_tokens = 1500,
            system  = prompt_sistema,
            messages = [{"role": "user", "content": prompt_usuario}]
        )
        raw = msg.content[0].text.strip()
        # Limpiar posibles bloques de código
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        resultado = json.loads(raw)
    except Exception as e:
        return jsonify({"error": f"Error IA: {str(e)}"}), 500

    # Guardar resultado en DB (estructura expandida + trazabilidad de usuario)
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """UPDATE mant_contratos SET
                   ai_analizado=1, ai_fecha=%s, ai_usuario=%s,
                   ai_resumen=%s,
                   ai_puntos_criticos=%s, ai_alertas=%s, ai_mejoras=%s,
                   ai_clausulas=%s, ai_cobertura=%s, ai_tipo_contrato=%s,
                   ai_frecuencia_sug=%s, ai_score=%s,
                   ai_vigencia_inicio=%s, ai_vigencia_fin=%s,
                   nivel_riesgo=%s,
                   sla_horas=%s,
                   incluye_mant_gratis=%s, incluye_repuestos=%s,
                   costo_por_mant=%s, costo_total=%s,
                   frecuencia_meses=COALESCE(NULLIF(frecuencia_meses,0),%s),
                   es_indefinido=COALESCE(NULLIF(es_indefinido,0),%s),
                   monto_mensual=COALESCE(NULLIF(monto_mensual,0),%s)
                   WHERE id=%s""",
                (datetime.now(),
                 current_username(),
                 resultado.get("resumen",""),
                 json.dumps(resultado.get("puntos_criticos",[]),    ensure_ascii=False),
                 json.dumps(resultado.get("alertas",[]),            ensure_ascii=False),
                 json.dumps(resultado.get("mejoras_prioritarias",[]),ensure_ascii=False),
                 json.dumps(resultado.get("clausulas_criticas",[]), ensure_ascii=False),
                 resultado.get("cobertura_descripcion",""),
                 resultado.get("tipo_contrato",""),
                 resultado.get("frecuencia_sugerida_meses"),
                 resultado.get("score"),
                 resultado.get("vigencia_inicio") or None,
                 resultado.get("vigencia_fin") or None,
                 resultado.get("nivel_riesgo","medio"),
                 resultado.get("sla_horas") or None,
                 1 if resultado.get("incluye_mant_gratis") else 0,
                 1 if resultado.get("incluye_repuestos") else 0,
                 resultado.get("costo_por_mant") or None,
                 resultado.get("costo_total") or None,
                 resultado.get("frecuencia_sugerida_meses"),
                 1 if resultado.get("es_indefinido") else 0,
                 resultado.get("costo_mensual") or None,
                 ctid)
            )
        conn.commit()
        _mant_log("contrato", ctid, "analizado_ia", f"score={resultado.get('score')} riesgo={resultado.get('nivel_riesgo')}")
        return jsonify({"ok": True, "resultado": resultado})
    except Exception as db_err:
        conn.rollback()
        print(f"[MANT] ERROR guardando análisis contrato {ctid}: {db_err}")
        return jsonify({"error": f"IA OK pero error guardando: {db_err}"}), 500
    finally:
        conn.close()


# ── GESTIÓN CONTRATO — cláusulas y variables personalizadas ──────────

@app.route("/mantenciones/api/contratos/<int:ctid>/clausulas", methods=["PUT"])
@_mant_required
def mant_contrato_clausulas(ctid):
    """Guarda cláusulas personalizadas y variables adicionales del contrato."""
    d = request.get_json(silent=True) or {}
    clausulas = d.get("clausulas", [])   # lista de {titulo, texto, tipo}
    variables = d.get("variables", {})   # dict campo→valor
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            # Actualizar también campos directos si vienen
            extra_sets, extra_vals = [], []
            for campo in ("sla_horas", "frecuencia_meses", "monto_mensual",
                          "monto_anual", "notas", "nivel_riesgo"):
                if campo in variables:
                    extra_sets.append(f"{campo}=%s")
                    extra_vals.append(variables.pop(campo) or None)
            sets = ["clausulas_custom=%s", "variables_extra=%s"] + extra_sets
            vals = [
                json.dumps(clausulas, ensure_ascii=False),
                json.dumps(variables, ensure_ascii=False),
            ] + extra_vals + [ctid]
            cur.execute(
                f"UPDATE mant_contratos SET {','.join(sets)} WHERE id=%s", vals
            )
        conn.commit()
        _mant_log("contrato", ctid, "clausulas_actualizadas",
                  f"{len(clausulas)} cláusulas, {len(variables)} variables")
        return jsonify({"ok": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/mantenciones/api/contratos/<int:ctid>/clausulas", methods=["GET"])
@_mant_required
def mant_contrato_clausulas_get(ctid):
    """Devuelve cláusulas personalizadas y variables de un contrato."""
    ct = mysql_fetchone(
        "SELECT clausulas_custom, variables_extra, sla_horas, frecuencia_meses, "
        "monto_mensual, monto_anual, notas, nivel_riesgo FROM mant_contratos WHERE id=%s",
        (ctid,)
    )
    if not ct:
        return jsonify({"error": "No encontrado"}), 404
    return jsonify({
        "clausulas": json.loads(ct.get("clausulas_custom") or "[]"),
        "variables": json.loads(ct.get("variables_extra") or "{}"),
        "campos": {
            "sla_horas": ct.get("sla_horas"),
            "frecuencia_meses": ct.get("frecuencia_meses"),
            "monto_mensual": ct.get("monto_mensual"),
            "monto_anual": ct.get("monto_anual"),
            "notas": ct.get("notas",""),
            "nivel_riesgo": ct.get("nivel_riesgo","medio"),
        }
    })


# ── VISITAS / AGENDA ──────────────────────────────────────────────────

@app.route("/mantenciones/api/visitas", methods=["GET"])
@_mant_required
def mant_visitas_api():
    """
    Devuelve visitas enriquecidas para el calendario inteligente.
    Cada evento incluye: cliente, dirección, técnicos N:N, horario, costo.
    Filtros opcionales:  start, end, cliente_id, tecnico_id
    """
    desde = request.args.get("start", "")
    hasta = request.args.get("end", "")
    cid   = request.args.get("cliente_id")
    tid   = request.args.get("tecnico_id")
    where, params = [], []
    if desde: where.append("v.fecha_programada >= %s"); params.append(desde[:10])
    if hasta: where.append("v.fecha_programada <= %s"); params.append(hasta[:10])
    if cid:   where.append("v.cliente_id=%s"); params.append(int(cid))
    if tid:   where.append(
        "v.id IN (SELECT visita_id FROM mant_visita_tecnicos WHERE tecnico_id=%s)"
    ); params.append(int(tid))

    sql = ("""SELECT v.*, c.razon_social, c.direccion, c.comuna, c.region
                FROM mant_visitas v
                JOIN mant_clientes c ON c.id=v.cliente_id""")
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY v.fecha_programada, v.hora_inicio LIMIT 500"
    rows = mysql_fetchall(sql, tuple(params)) or []
    rows = [dict(r) for r in rows]

    # Cargar técnicos N:N de todas las visitas en una sola query
    tecs_por_visita = {}
    if rows:
        vids = [r["id"] for r in rows]
        ph = ",".join(["%s"] * len(vids))
        tec_rows = mysql_fetchall(
            f"""SELECT vt.visita_id, t.id, t.nombre, t.especialidad, t.nivel,
                       t.es_externo, t.foto_url
                  FROM mant_visita_tecnicos vt
                  JOIN mant_tecnicos t ON t.id=vt.tecnico_id
                 WHERE vt.visita_id IN ({ph})
                 ORDER BY t.nombre""",
            tuple(vids)
        ) or []
        for tr in tec_rows:
            tecs_por_visita.setdefault(tr["visita_id"], []).append({
                "id":           tr["id"],
                "nombre":       tr["nombre"],
                "especialidad": tr["especialidad"],
                "nivel":        tr["nivel"],
                "es_externo":   bool(tr.get("es_externo")),
                "foto_url":     tr.get("foto_url"),
            })

    # Colores por TIPO (mantención preventiva, correctiva, etc.)
    TIPO_COLOR = {
        "preventiva": "#16a34a",
        "correctiva": "#dc2626",
        "garantia":   "#2563eb",
        "inspeccion": "#f59e0b",
    }
    EST_BORDER = {
        "programada":  "#0f172a",
        "completada":  "#166534",
        "cancelada":   "#9ca3af",
        "reagendada":  "#7c3aed",
    }

    events = []
    for r in rows:
        tecs = tecs_por_visita.get(r["id"], [])
        # Si no hay técnicos N:N pero sí hay tecnico_id legacy, fallback
        if not tecs and r.get("tecnico_id"):
            t_legacy = mysql_fetchone(
                "SELECT id,nombre,especialidad,nivel,es_externo,foto_url FROM mant_tecnicos WHERE id=%s",
                (r["tecnico_id"],)
            )
            if t_legacy:
                tecs = [dict(t_legacy)]
        # Para visitas viejas con solo el campo texto `tecnico`
        elif not tecs and r.get("tecnico"):
            tecs = [{"id": None, "nombre": r["tecnico"], "especialidad": None,
                     "nivel": None, "es_externo": False, "foto_url": None}]

        # Helper para serializar hora (puede venir como timedelta o string)
        def _h(v):
            if v is None: return None
            if hasattr(v, "total_seconds"):
                t = int(v.total_seconds())
                return f"{t//3600:02d}:{(t%3600)//60:02d}"
            s = str(v)
            return s[:5] if len(s) >= 5 else s

        events.append({
            "id":          r["id"],
            "title":       r["titulo"] or (r.get("tipo","").capitalize() + " programada"),
            "fecha":       str(r["fecha_programada"]) if r["fecha_programada"] else None,
            "hora_inicio": _h(r.get("hora_inicio")),
            "hora_fin":    _h(r.get("hora_fin")),
            "tipo":        r.get("tipo"),
            "estado":      r.get("estado"),
            "descripcion": (r.get("descripcion") or "")[:300],
            "costo":       float(r.get("costo") or 0),
            "cliente": {
                "id":           r["cliente_id"],
                "razon_social": r["razon_social"],
                "direccion":    r.get("direccion") or "",
                "comuna":       r.get("comuna") or "",
                "region":       r.get("region") or "",
            },
            "tecnicos":    tecs,
            "n_tecnicos":  len(tecs),
            "color_tipo":   TIPO_COLOR.get(r.get("tipo"), "#6b7280"),
            "color_borde":  EST_BORDER.get(r.get("estado"), "#0f172a"),
        })
    return jsonify(events)


@app.route("/mantenciones/api/visitas", methods=["POST"])
@_mant_required
def mant_visita_crear():
    d = request.get_json(silent=True) or {}
    if not d.get("cliente_id") or not d.get("fecha_programada"):
        return jsonify({"error": "cliente_id y fecha_programada requeridos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_visitas
                   (cliente_id,contrato_id,titulo,fecha_programada,hora_inicio,hora_fin,
                    tecnico,tipo,estado,descripcion,costo,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (d["cliente_id"], d.get("contrato_id") or None,
                 d.get("titulo","Mantención"), d["fecha_programada"],
                 d.get("hora_inicio") or None, d.get("hora_fin") or None,
                 d.get("tecnico",""), d.get("tipo","preventiva"),
                 d.get("estado","programada"), d.get("descripcion",""),
                 float(d.get("costo",0) or 0), current_username())
            )
            vid = cur.lastrowid
        conn.commit()
        _mant_log("visita", vid, "creada", d.get("titulo",""))
        return jsonify({"ok": True, "id": vid})
    finally:
        conn.close()


@app.route("/mantenciones/api/visitas/<int:vid>", methods=["PUT"])
@_mant_required
def mant_visita_update(vid):
    d = request.get_json(silent=True) or {}
    allowed = ["titulo","fecha_programada","fecha_realizada","hora_inicio","hora_fin",
               "tecnico","tipo","estado","descripcion","observaciones","costo","contrato_id"]
    sets = [f"{f}=%s" for f in allowed if f in d]
    vals = [d[f] for f in allowed if f in d]
    if not sets:
        return jsonify({"error": "Sin campos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_visitas SET {','.join(sets)} WHERE id=%s",
                        vals + [vid])
        conn.commit()
        _mant_log("visita", vid, "actualizada")
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/visitas/<int:vid>", methods=["DELETE"])
@_mant_required
def mant_visita_del(vid):
    # Capturar info ANTES de borrar para el log
    v_info = mysql_fetchone(
        "SELECT cliente_id, numero_ot, titulo, fecha_programada FROM mant_visitas WHERE id=%s", (vid,)
    )
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM mant_visitas WHERE id=%s", (vid,))
        conn.commit()
        if v_info:
            ot = v_info.get("numero_ot") or f"V-{vid}"
            _mant_log("visita", vid, "eliminada", f"{ot} — {v_info.get('titulo') or ''}")
            _mant_log("cliente", v_info.get("cliente_id"), "visita_eliminada", f"{ot}")
        return jsonify({"ok": True})
    finally:
        conn.close()


# ═════════════════════════════════════════════════════════════════════
# OT — LISTADO GLOBAL DE ÓRDENES DE TRABAJO
# Vista pro: todas las OTs con filtros (estado, técnico, cliente, rango)
# ═════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/ots")
@_mant_required
def mant_ots_list():
    """Listado global de todas las Órdenes de Trabajo (visitas) con filtros."""
    estado     = (request.args.get("estado") or "").strip().lower()
    tipo       = (request.args.get("tipo") or "").strip().lower()
    tecnico_id = request.args.get("tecnico_id")
    cliente_id = request.args.get("cliente_id")
    q          = (request.args.get("q") or "").strip()

    where = ["1=1"]
    params = []

    if estado in ("programada", "completada", "cancelada", "reagendada", "en_curso"):
        where.append("v.estado=%s")
        params.append(estado)
    if tipo in ("preventiva", "correctiva", "garantia", "inspeccion", "levantamiento", "instalacion"):
        where.append("v.tipo=%s")
        params.append(tipo)
    if tecnico_id:
        try:
            params.append(int(tecnico_id))
            where.append("v.tecnico_id=%s")
        except (TypeError, ValueError):
            pass
    if cliente_id:
        try:
            params.append(int(cliente_id))
            where.append("v.cliente_id=%s")
        except (TypeError, ValueError):
            pass
    if q:
        where.append("(c.razon_social LIKE %s OR v.titulo LIKE %s OR v.numero_ot LIKE %s)")
        like = f"%{q}%"
        params.extend([like, like, like])

    ots = []
    kpis = {"total": 0, "programadas": 0, "completadas": 0, "atrasadas": 0}
    tecnicos = []
    fatal_error = None
    try:
        # ── PERF: en vez de 3 subqueries correlacionadas POR fila (~600 queries
        # con LIMIT 200), usamos LEFT JOIN a derivadas agregadas. Total: 1 query.
        sql = (
            "SELECT v.id, v.numero_ot, v.tipo, v.estado, v.titulo, v.descripcion, "
            "       v.fecha_programada, v.hora_inicio, v.hora_fin, v.costo, "
            "       v.cliente_id, c.razon_social, c.comuna AS cli_comuna, "
            "       v.tecnico_id, t.nombre AS tecnico_nombre, "
            "       COALESCE(tar.n_tareas, 0)    AS n_tareas, "
            "       COALESCE(tar.n_completas, 0) AS n_completas, "
            "       COALESCE(fot.n_fotos, 0)     AS n_fotos "
            "  FROM mant_visitas v "
            "  JOIN mant_clientes c ON c.id=v.cliente_id "
            "  LEFT JOIN mant_tecnicos t ON t.id=v.tecnico_id "
            "  LEFT JOIN ( "
            "       SELECT visita_id, "
            "              COUNT(*) AS n_tareas, "
            "              SUM(CASE WHEN completada=1 THEN 1 ELSE 0 END) AS n_completas "
            "         FROM mant_visita_tareas "
            "       GROUP BY visita_id"
            "  ) tar ON tar.visita_id = v.id "
            "  LEFT JOIN ( "
            "       SELECT visita_id, COUNT(*) AS n_fotos "
            "         FROM mant_visita_fotos "
            "       GROUP BY visita_id"
            "  ) fot ON fot.visita_id = v.id "
            f" WHERE {' AND '.join(where)} "
            " ORDER BY v.fecha_programada DESC, v.id DESC "
            " LIMIT 200"
        )
        ots = mysql_fetchall(sql, tuple(params)) or []
        ots = [dict(o) for o in ots]
    except Exception as e:
        # Fallback simple si subqueries fallan (ej. tablas no migradas aún)
        try:
            ots = mysql_fetchall(
                "SELECT v.id, v.numero_ot, v.tipo, v.estado, v.titulo, v.descripcion, "
                "       v.fecha_programada, v.hora_inicio, v.hora_fin, v.costo, "
                "       v.cliente_id, c.razon_social, c.comuna AS cli_comuna, "
                "       v.tecnico_id, v.tecnico AS tecnico_nombre, "
                "       0 AS n_tareas, 0 AS n_completas, 0 AS n_fotos "
                "  FROM mant_visitas v "
                "  JOIN mant_clientes c ON c.id=v.cliente_id "
                f" WHERE {' AND '.join(where)} "
                " ORDER BY v.fecha_programada DESC, v.id DESC LIMIT 200",
                tuple(params)
            ) or []
            ots = [dict(o) for o in ots]
        except Exception as e2:
            fatal_error = f"{e} / fallback: {e2}"
            ots = []

    try:
        kpis = mysql_fetchone(
            "SELECT "
            " COUNT(*) AS total, "
            " SUM(CASE WHEN estado='programada' THEN 1 ELSE 0 END) AS programadas, "
            " SUM(CASE WHEN estado='completada' THEN 1 ELSE 0 END) AS completadas, "
            " SUM(CASE WHEN estado='programada' AND fecha_programada < CURDATE() THEN 1 ELSE 0 END) AS atrasadas "
            "FROM mant_visitas"
        ) or kpis
    except Exception:
        pass

    try:
        tecnicos = mysql_fetchall(
            "SELECT id, nombre FROM mant_tecnicos WHERE estado='activo' ORDER BY nombre"
        ) or []
        tecnicos = [dict(t) for t in tecnicos]
    except Exception:
        tecnicos = []

    return render_template(
        "mantenciones/ots_list.html",
        ots=ots,
        kpis=kpis,
        tecnicos=tecnicos,
        fatal_error=fatal_error,
        filtros={
            "estado": estado, "tipo": tipo,
            "tecnico_id": tecnico_id, "cliente_id": cliente_id, "q": q,
        },
    )


# ═════════════════════════════════════════════════════════════════════
# OT — PÁGINA FICHA (ficha completa con tabs)
# ═════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/ot/<int:vid>")
@_mant_required
def mant_ot_ficha(vid):
    """Ficha completa de una OT/visita con tabs (Tareas/Fotos/Repuestos/Bitácora)."""
    visita = mysql_fetchone(
        "SELECT v.*, c.razon_social, c.direccion AS cli_direccion, "
        "       c.comuna AS cli_comuna, c.email_empresa AS cli_email, "
        "       c.tel_empresa AS cli_tel, c.contacto_nombre AS cli_contacto, "
        "       c.contacto_telefono AS cli_contacto_tel, "
        "       ct.nombre AS contrato_nombre, "
        "       t.nombre AS tecnico_principal, t.telefono AS tecnico_tel, "
        "       t.email AS tecnico_email, t.foto_url AS tecnico_foto "
        "  FROM mant_visitas v "
        "  JOIN mant_clientes c ON c.id=v.cliente_id "
        "  LEFT JOIN mant_contratos ct ON ct.id=v.contrato_id "
        "  LEFT JOIN mant_tecnicos t ON t.id=v.tecnico_id "
        " WHERE v.id=%s",
        (vid,)
    )
    if not visita:
        flash("OT no encontrada.", "danger")
        return redirect(url_for("mant_index"))

    visita = dict(visita)

    # Técnicos adicionales asignados (N:N)
    tecnicos = mysql_fetchall(
        "SELECT vt.rol, vt.horas, vt.costo, t.id, t.nombre, t.telefono, "
        "       t.email, t.foto_url, t.especialidad, t.nivel "
        "  FROM mant_visita_tecnicos vt "
        "  JOIN mant_tecnicos t ON t.id=vt.tecnico_id "
        " WHERE vt.visita_id=%s ORDER BY (vt.rol='lider') DESC, t.nombre",
        (vid,)
    ) or []
    tecnicos = [dict(t) for t in tecnicos]

    # Máquinas del cliente (para asignar a tareas)
    maquinas = mysql_fetchall(
        "SELECT id, nombre, serie, sku FROM mant_maquinas "
        " WHERE cliente_id=%s ORDER BY nombre",
        (visita["cliente_id"],)
    ) or []
    maquinas = [dict(m) for m in maquinas]

    # Tareas con stats
    tareas_stats = mysql_fetchone(
        "SELECT COUNT(*) AS total, "
        "       SUM(CASE WHEN completada=1 THEN 1 ELSE 0 END) AS completas "
        "  FROM mant_visita_tareas WHERE visita_id=%s",
        (vid,)
    ) or {"total": 0, "completas": 0}

    fotos_count = mysql_fetchone(
        "SELECT COUNT(*) AS n FROM mant_visita_fotos WHERE visita_id=%s",
        (vid,)
    ) or {"n": 0}

    return render_template(
        "mantenciones/ot_ficha.html",
        visita=visita,
        tecnicos=tecnicos,
        maquinas=maquinas,
        stats={
            "total_tareas":    tareas_stats.get("total", 0),
            "completas":       tareas_stats.get("completas", 0) or 0,
            "fotos":           fotos_count.get("n", 0),
        },
    )


# ═════════════════════════════════════════════════════════════════════
# OT — TAREAS DENTRO DE UNA VISITA
# Checklist que el técnico ejecuta paso a paso. Cada tarea puede estar
# asociada a una máquina específica (mant_maquinas) para trazabilidad.
# ═════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/visitas/<int:vid>/tareas", methods=["GET"])
@_mant_required
def mant_visita_tareas_get(vid):
    """Lista las tareas de una OT con su estado."""
    rows = mysql_fetchall(
        "SELECT t.id, t.orden, t.titulo, t.descripcion, t.tipo, t.cantidad, "
        "       t.completada, t.completada_at, t.completada_por, t.observaciones, "
        "       t.maquina_id, m.nombre AS maquina_nombre, m.serie AS maquina_serie "
        "  FROM mant_visita_tareas t "
        "  LEFT JOIN mant_maquinas m ON m.id=t.maquina_id "
        " WHERE t.visita_id=%s ORDER BY t.orden ASC, t.id ASC",
        (vid,)
    ) or []
    return jsonify([dict(r) for r in rows])


@app.route("/mantenciones/api/visitas/<int:vid>/tareas/nueva", methods=["POST"])
@_mant_required
def mant_visita_tarea_nueva(vid):
    """Crea una tarea dentro de la OT. Acepta JSON o form-data."""
    d = request.get_json(silent=True) or request.form
    titulo = (d.get("titulo") or "").strip()[:300]
    if not titulo:
        return jsonify({"ok": False, "error": "Título requerido"}), 400
    descripcion = (d.get("descripcion") or "").strip()
    tipo = (d.get("tipo") or "otro").strip().lower()
    if tipo not in ("inspeccion", "cambio", "reparacion", "limpieza",
                    "levantamiento", "instalacion", "garantia", "otro"):
        tipo = "otro"
    maquina_id = d.get("maquina_id")
    try:
        maquina_id = int(maquina_id) if maquina_id else None
    except (TypeError, ValueError):
        maquina_id = None
    cantidad = int(d.get("cantidad") or 1)
    # Determinar orden = max + 1
    row = mysql_fetchone(
        "SELECT COALESCE(MAX(orden),0)+1 AS nx FROM mant_visita_tareas WHERE visita_id=%s",
        (vid,)
    )
    orden = (row or {}).get("nx", 1)
    nid = mysql_execute(
        "INSERT INTO mant_visita_tareas "
        "(visita_id, orden, titulo, descripcion, tipo, maquina_id, cantidad, created_by) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
        (vid, orden, titulo, descripcion, tipo, maquina_id, cantidad, current_username())
    )
    try:
        mysql_execute(
            "INSERT INTO mant_logs (entidad, entidad_id, accion, detalle, usuario) "
            "VALUES ('visita', %s, 'tarea_agregada', %s, %s)",
            (vid, f"Tarea: {titulo}", current_username())
        )
    except Exception:
        pass
    return jsonify({"ok": True, "id": nid, "orden": orden})


@app.route("/mantenciones/api/visitas/<int:vid>/tareas/<int:tid>", methods=["POST", "PATCH", "DELETE"])
@_mant_required
def mant_visita_tarea_update(vid, tid):
    """Actualiza, completa o elimina una tarea."""
    if request.method == "DELETE":
        mysql_execute("DELETE FROM mant_visita_tareas WHERE id=%s AND visita_id=%s", (tid, vid))
        return jsonify({"ok": True, "deleted": True})

    d = request.get_json(silent=True) or request.form
    action = (d.get("action") or "").strip().lower()
    user = current_username()

    if action == "toggle":
        # Cambiar estado completada
        row = mysql_fetchone(
            "SELECT completada FROM mant_visita_tareas WHERE id=%s AND visita_id=%s",
            (tid, vid)
        )
        if not row:
            return jsonify({"ok": False, "error": "Tarea no encontrada"}), 404
        new_state = 0 if row["completada"] else 1
        mysql_execute(
            "UPDATE mant_visita_tareas "
            "   SET completada=%s, "
            "       completada_at=%s, "
            "       completada_por=%s "
            " WHERE id=%s AND visita_id=%s",
            (
                new_state,
                datetime.now() if new_state else None,
                user if new_state else None,
                tid, vid
            )
        )
        try:
            mysql_execute(
                "INSERT INTO mant_logs (entidad, entidad_id, accion, detalle, usuario) "
                "VALUES ('visita', %s, %s, %s, %s)",
                (vid,
                 "tarea_completada" if new_state else "tarea_reabierta",
                 f"Tarea ID {tid}",
                 user)
            )
        except Exception:
            pass
        return jsonify({"ok": True, "completada": bool(new_state)})

    # Update genérico
    fields = []
    vals = []
    for f in ("titulo", "descripcion", "tipo", "observaciones"):
        if f in d:
            fields.append(f"{f}=%s")
            vals.append((d.get(f) or "").strip()[:500])
    for f in ("cantidad", "orden", "maquina_id"):
        if f in d:
            try:
                fields.append(f"{f}=%s")
                vals.append(int(d.get(f)) if d.get(f) else None)
            except (TypeError, ValueError):
                continue
    if not fields:
        return jsonify({"ok": False, "error": "Nada para actualizar"}), 400
    vals.extend([tid, vid])
    mysql_execute(
        f"UPDATE mant_visita_tareas SET {', '.join(fields)} WHERE id=%s AND visita_id=%s",
        tuple(vals)
    )
    return jsonify({"ok": True})


# ═════════════════════════════════════════════════════════════════════
# OT — PROTOCOLOS RÁPIDOS (generadores de checklist 1-click)
# Crean en lote N tareas estándar sobre una OT existente.
#
# Protocolos disponibles:
#   · levantamiento     → 1 tarea por máquina del cliente (fotos + serie)
#   · cambio_garantia   → 1 tarea cambio por máquina seleccionada
#   · inspeccion_pm     → checklist estándar preventivo (8 ítems)
#   · instalacion_pm    → checklist de instalación nueva
# ═════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/visitas/<int:vid>/protocolo/<protocolo>", methods=["POST"])
@_mant_required
def mant_visita_protocolo(vid, protocolo):
    """Aplica un protocolo a la OT: genera N tareas estándar.

    POST JSON:
      maquinas_ids: [int]  (opcional, p/ cambio_garantia restringe a IDs específicos)
      reset:        bool   (default false; si true borra tareas existentes antes)
    """
    visita = mysql_fetchone(
        "SELECT id, cliente_id, tipo FROM mant_visitas WHERE id=%s", (vid,)
    )
    if not visita:
        return jsonify({"ok": False, "error": "OT no encontrada"}), 404

    visita = dict(visita)
    cliente_id = visita["cliente_id"]
    d = request.get_json(silent=True) or {}
    reset = bool(d.get("reset"))
    maquinas_ids = d.get("maquinas_ids") or []
    user = current_username()

    # Validar protocolo
    if protocolo not in ("levantamiento", "cambio_garantia", "inspeccion_pm", "instalacion_pm"):
        return jsonify({"ok": False, "error": "Protocolo desconocido"}), 400

    # Obtener máquinas del cliente
    if maquinas_ids:
        placeholders = ",".join(["%s"] * len(maquinas_ids))
        maquinas = mysql_fetchall(
            f"SELECT id, nombre, serie, sku FROM mant_maquinas "
            f"WHERE cliente_id=%s AND id IN ({placeholders})",
            tuple([cliente_id] + [int(x) for x in maquinas_ids])
        ) or []
    else:
        maquinas = mysql_fetchall(
            "SELECT id, nombre, serie, sku FROM mant_maquinas "
            "WHERE cliente_id=%s ORDER BY nombre",
            (cliente_id,)
        ) or []
    maquinas = [dict(m) for m in maquinas]

    # Si reset, borrar tareas existentes
    if reset:
        mysql_execute("DELETE FROM mant_visita_tareas WHERE visita_id=%s", (vid,))

    # Orden inicial = max(orden) + 1
    last_orden = mysql_fetchone(
        "SELECT COALESCE(MAX(orden),0) AS n FROM mant_visita_tareas WHERE visita_id=%s",
        (vid,)
    ) or {"n": 0}
    orden_base = int(last_orden.get("n", 0)) + 1

    creadas = 0

    # ─── PROTOCOLO 1: LEVANTAMIENTO FOTOGRÁFICO COMPLETO ──────────────
    # Para cada máquina del cliente: 1 tarea con checklist visual
    # ──────────────────────────────────────────────────────────────────
    if protocolo == "levantamiento":
        if not maquinas:
            return jsonify({
                "ok": False,
                "error": "El cliente no tiene máquinas registradas. Agrega máquinas desde la ficha del cliente."
            }), 400
        for i, m in enumerate(maquinas):
            mant_titulo = f"📸 Levantamiento: {m['nombre'] or 'Máquina'}"
            if m.get("serie"):
                mant_titulo += f" (S/N: {m['serie']})"
            descripcion = (
                "Protocolo de levantamiento fotográfico:\n"
                "  1) Foto general del equipo (frontal y lateral)\n"
                "  2) Foto de la placa / número de serie legible\n"
                "  3) Foto de la marca y modelo\n"
                "  4) Detalle de cualquier daño visible (si aplica)\n"
                "  5) Foto del equipo instalado / ubicado en sala"
            )
            mysql_execute(
                "INSERT INTO mant_visita_tareas "
                "(visita_id, orden, titulo, descripcion, tipo, maquina_id, cantidad, created_by) "
                "VALUES (%s,%s,%s,%s,'levantamiento',%s,1,%s)",
                (vid, orden_base + i, mant_titulo[:300], descripcion, m["id"], user)
            )
            creadas += 1

    # ─── PROTOCOLO 2: CAMBIO POR GARANTÍA ─────────────────────────────
    # Para cada máquina seleccionada: 1 tarea de cambio con checklist
    # ──────────────────────────────────────────────────────────────────
    elif protocolo == "cambio_garantia":
        if not maquinas:
            return jsonify({
                "ok": False,
                "error": "Selecciona al menos una máquina a cambiar."
            }), 400
        for i, m in enumerate(maquinas):
            mant_titulo = f"🔄 Cambio por garantía: {m['nombre'] or 'Máquina'}"
            if m.get("serie"):
                mant_titulo += f" (S/N saliente: {m['serie']})"
            descripcion = (
                "Protocolo de cambio por garantía:\n"
                "  1) Foto del equipo SALIENTE antes del retiro\n"
                "  2) Foto de la placa / N° serie del equipo saliente\n"
                "  3) Foto de la falla / motivo de garantía (si aplica)\n"
                "  4) Foto del equipo NUEVO (entrante) con su serie\n"
                "  5) Foto del equipo instalado y funcionando\n"
                "  6) Registrar N° de serie nuevo en el sistema (campo abajo)\n"
                "  7) Firma / conformidad del cliente"
            )
            mysql_execute(
                "INSERT INTO mant_visita_tareas "
                "(visita_id, orden, titulo, descripcion, tipo, maquina_id, cantidad, created_by) "
                "VALUES (%s,%s,%s,%s,'cambio',%s,1,%s)",
                (vid, orden_base + i, mant_titulo[:300], descripcion, m["id"], user)
            )
            creadas += 1

    # ─── PROTOCOLO 3: INSPECCIÓN PREVENTIVA ───────────────────────────
    elif protocolo == "inspeccion_pm":
        items = [
            ("inspeccion", "🔍 Inspección visual general del equipo", None),
            ("limpieza",   "🧹 Limpieza profunda de componentes", None),
            ("inspeccion", "🔋 Verificar tensión y voltaje", None),
            ("inspeccion", "⚙️ Lubricación de partes móviles", None),
            ("inspeccion", "🔌 Verificar conexiones eléctricas y aislantes", None),
            ("inspeccion", "📊 Test de funcionamiento bajo carga", None),
            ("inspeccion", "🩺 Diagnóstico de software / firmware (si aplica)", None),
            ("inspeccion", "📋 Registrar lecturas y observaciones finales", None),
        ]
        for i, (tipo, titulo, mid) in enumerate(items):
            mysql_execute(
                "INSERT INTO mant_visita_tareas "
                "(visita_id, orden, titulo, tipo, maquina_id, cantidad, created_by) "
                "VALUES (%s,%s,%s,%s,%s,1,%s)",
                (vid, orden_base + i, titulo[:300], tipo, mid, user)
            )
            creadas += 1

    # ─── PROTOCOLO 4: INSTALACIÓN NUEVA ───────────────────────────────
    elif protocolo == "instalacion_pm":
        items = [
            "📦 Verificar embalaje y accesorios completos",
            "🧰 Armado / ensamble del equipo",
            "📍 Ubicación y nivelación en sala",
            "🔌 Conexión eléctrica y verificación de toma",
            "⚡ Encendido inicial y test de funciones básicas",
            "📊 Configuración de software / firmware",
            "👤 Capacitación al cliente sobre uso básico",
            "📸 Foto del equipo instalado funcionando",
            "✍️ Firma de conformidad del cliente",
        ]
        for i, t in enumerate(items):
            mysql_execute(
                "INSERT INTO mant_visita_tareas "
                "(visita_id, orden, titulo, tipo, cantidad, created_by) "
                "VALUES (%s,%s,%s,'instalacion',1,%s)",
                (vid, orden_base + i, t[:300], user)
            )
            creadas += 1

    # Log de auditoría
    try:
        mysql_execute(
            "INSERT INTO mant_logs (entidad, entidad_id, accion, detalle, usuario) "
            "VALUES ('visita', %s, %s, %s, %s)",
            (vid, "protocolo_aplicado", f"{protocolo} → {creadas} tareas", user)
        )
    except Exception:
        pass

    return jsonify({
        "ok": True,
        "protocolo": protocolo,
        "tareas_creadas": creadas,
        "maquinas_afectadas": len(maquinas) if protocolo in ("levantamiento", "cambio_garantia") else 0,
    })


# ═════════════════════════════════════════════════════════════════════
# OT — FOTOS DE LA VISITA (galería con multi-upload)
# ═════════════════════════════════════════════════════════════════════

MANT_FOTOS_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "static", "uploads", "mantenciones"
)
os.makedirs(MANT_FOTOS_DIR, exist_ok=True)


@app.route("/mantenciones/api/visitas/<int:vid>/fotos", methods=["GET"])
@_mant_required
def mant_visita_fotos_get(vid):
    """Lista las fotos de una OT."""
    rows = mysql_fetchall(
        "SELECT id, archivo_path, archivo_nombre, tipo_foto, descripcion, "
        "       tarea_id, maquina_id, tomada_por, tomada_at, file_size_kb "
        "  FROM mant_visita_fotos "
        " WHERE visita_id=%s ORDER BY tomada_at DESC, id DESC",
        (vid,)
    ) or []
    out = []
    for r in rows:
        rd = dict(r)
        # tomada_at puede ser datetime → string ISO
        if rd.get("tomada_at"):
            rd["tomada_at_iso"] = rd["tomada_at"].isoformat() if hasattr(rd["tomada_at"], "isoformat") else str(rd["tomada_at"])
        out.append(rd)
    return jsonify(out)


@app.route("/mantenciones/api/visitas/<int:vid>/fotos/subir", methods=["POST"])
@_mant_required
def mant_visita_fotos_subir(vid):
    """Sube una o varias fotos a la OT. Soporta multi-upload."""
    files = request.files.getlist("fotos") or request.files.getlist("imagenes")
    if not files or all(not f.filename for f in files):
        return jsonify({"ok": False, "error": "No se envió ningún archivo"}), 400

    tipo_foto    = (request.form.get("tipo_foto") or "general").strip().lower()
    if tipo_foto not in ("antes", "durante", "despues", "serie", "falla",
                          "reparacion", "general", "levantamiento"):
        tipo_foto = "general"
    descripcion  = (request.form.get("descripcion") or "").strip()[:500]
    tarea_id     = request.form.get("tarea_id")
    maquina_id   = request.form.get("maquina_id")
    try:
        tarea_id = int(tarea_id) if tarea_id else None
    except (TypeError, ValueError):
        tarea_id = None
    try:
        maquina_id = int(maquina_id) if maquina_id else None
    except (TypeError, ValueError):
        maquina_id = None

    saved = 0
    errors = []
    user = current_username()
    for i, f in enumerate(files):
        if not f or not f.filename:
            continue
        ext, err = _validate_uploaded_image(f, label=f.filename)
        if err:
            errors.append(err)
            continue
        # Carpeta por visita (organización)
        visita_dir = os.path.join(MANT_FOTOS_DIR, f"v{vid}")
        os.makedirs(visita_dir, exist_ok=True)
        fname = f"v{vid}_{int(time.time())}_{i}_{secure_filename(f.filename)}"
        fpath = os.path.join(visita_dir, fname)
        try:
            f.save(fpath)
            size_kb = os.path.getsize(fpath) // 1024
        except Exception as e:
            errors.append(f"Error guardando {f.filename}: {e}")
            continue
        try:
            mysql_execute(
                "INSERT INTO mant_visita_fotos "
                "(visita_id, tarea_id, maquina_id, archivo_path, archivo_nombre, "
                " tipo_foto, descripcion, tomada_por, file_size_kb) "
                "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                (
                    vid, tarea_id, maquina_id,
                    f"uploads/mantenciones/v{vid}/{fname}",
                    f.filename[:300],
                    tipo_foto, descripcion or None, user, size_kb
                )
            )
            saved += 1
        except Exception as e:
            errors.append(f"Error BD {f.filename}: {e}")
    try:
        if saved:
            mysql_execute(
                "INSERT INTO mant_logs (entidad, entidad_id, accion, detalle, usuario) "
                "VALUES ('visita', %s, 'fotos_subidas', %s, %s)",
                (vid, f"{saved} foto(s) tipo={tipo_foto}", user)
            )
    except Exception:
        pass
    return jsonify({"ok": True, "saved": saved, "errors": errors[:3]})


@app.route("/mantenciones/api/visitas/<int:vid>/fotos/<int:fid>", methods=["DELETE"])
@_mant_required
def mant_visita_foto_delete(vid, fid):
    """Elimina una foto de la OT (archivo + registro)."""
    row = mysql_fetchone(
        "SELECT archivo_path FROM mant_visita_fotos WHERE id=%s AND visita_id=%s",
        (fid, vid)
    )
    if row and row.get("archivo_path"):
        try:
            fp = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                "static", row["archivo_path"]
            )
            if os.path.exists(fp):
                os.remove(fp)
        except Exception:
            pass
    mysql_execute("DELETE FROM mant_visita_fotos WHERE id=%s AND visita_id=%s", (fid, vid))
    return jsonify({"ok": True, "deleted": True})


# ═════════════════════════════════════════════════════════════════════
# GALERÍA POR MÁQUINA (inventario fotográfico permanente)
# ═════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/maquinas/<int:mid>/fotos", methods=["GET"])
@_mant_required
def mant_maquina_fotos_get(mid):
    rows = mysql_fetchall(
        "SELECT id, archivo_path, archivo_nombre, tipo_foto, descripcion, "
        "       es_principal, visita_origen, tomada_por, tomada_at "
        "  FROM mant_maquina_fotos WHERE maquina_id=%s "
        "ORDER BY es_principal DESC, tomada_at DESC",
        (mid,)
    ) or []
    return jsonify([dict(r) for r in rows])


@app.route("/mantenciones/api/maquinas/<int:mid>/fotos/subir", methods=["POST"])
@_mant_required
def mant_maquina_fotos_subir(mid):
    files = request.files.getlist("fotos") or request.files.getlist("imagenes")
    if not files or all(not f.filename for f in files):
        return jsonify({"ok": False, "error": "No se envió ningún archivo"}), 400
    tipo_foto = (request.form.get("tipo_foto") or "principal").strip().lower()
    if tipo_foto not in ("principal", "serie", "marca", "detalle",
                          "instalada", "daño", "antes_reparacion", "despues_reparacion"):
        tipo_foto = "principal"
    descripcion = (request.form.get("descripcion") or "").strip()[:500]
    user = current_username()

    saved = 0
    errors = []
    for i, f in enumerate(files):
        if not f or not f.filename:
            continue
        ext, err = _validate_uploaded_image(f, label=f.filename)
        if err:
            errors.append(err)
            continue
        mq_dir = os.path.join(MANT_FOTOS_DIR, f"m{mid}")
        os.makedirs(mq_dir, exist_ok=True)
        fname = f"m{mid}_{int(time.time())}_{i}_{secure_filename(f.filename)}"
        fpath = os.path.join(mq_dir, fname)
        try:
            f.save(fpath)
        except Exception as e:
            errors.append(f"Error guardando {f.filename}: {e}")
            continue
        try:
            mysql_execute(
                "INSERT INTO mant_maquina_fotos "
                "(maquina_id, archivo_path, archivo_nombre, tipo_foto, descripcion, tomada_por) "
                "VALUES (%s,%s,%s,%s,%s,%s)",
                (
                    mid,
                    f"uploads/mantenciones/m{mid}/{fname}",
                    f.filename[:300],
                    tipo_foto, descripcion or None, user
                )
            )
            saved += 1
        except Exception as e:
            errors.append(f"Error BD {f.filename}: {e}")
    return jsonify({"ok": True, "saved": saved, "errors": errors[:3]})


# ── REPUESTOS DE UNA VISITA ──────────────────────────────────────────

@app.route("/mantenciones/api/visitas/<int:vid>/repuestos", methods=["GET"])
@_mant_required
def mant_visita_repuestos_get(vid):
    rows = mysql_fetchall(
        "SELECT id, sku, producto_id, descripcion, cantidad, "
        "       costo_unitario, costo_total, origen, notas "
        "  FROM mant_visita_repuestos "
        " WHERE visita_id=%s ORDER BY id",
        (vid,)
    ) or []
    return jsonify([dict(r) for r in rows])


@app.route("/mantenciones/api/productos-search", methods=["GET"])
@_mant_required
def mant_productos_search():
    """
    Buscador de productos del catálogo local (para repuestos).
    Busca por SKU o nombre.
    """
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify([])
    like = f"%{q}%"
    # Tabla 'products' es el catálogo de etiquetas/productos local
    rows = mysql_fetchall(
        "SELECT sku, nombre, marca, precio_venta "
        "  FROM products "
        " WHERE sku LIKE %s OR nombre LIKE %s "
        " ORDER BY sku LIMIT 15",
        (like, like)
    ) or []
    return jsonify([dict(r) for r in rows])


# ── ENVÍO DE EMAIL: visita agendada (con OT) ─────────────────────────

@app.route("/mantenciones/api/visitas/<int:vid>/enviar-email", methods=["POST"])
@_mant_required
def mant_visita_enviar_email(vid):
    """
    Envía un email al cliente notificando que se agendó la visita,
    incluyendo el número de OT, fecha, técnicos y costo.
    Body JSON opcional:
      destinatario: email (si no se pasa, usa contacto_email del cliente)
    """
    d = request.get_json(silent=True) or {}
    v = mysql_fetchone(
        """SELECT v.*, c.razon_social, c.contacto_email, c.contacto_nombre,
                  c.direccion, c.comuna, c.region
             FROM mant_visitas v
             JOIN mant_clientes c ON c.id=v.cliente_id
            WHERE v.id=%s""",
        (vid,)
    )
    if not v:
        return jsonify({"error": "Visita no encontrada"}), 404

    destinatario = (d.get("destinatario") or v.get("contacto_email") or "").strip()
    if not destinatario:
        return jsonify({"error": "El cliente no tiene email registrado. Indica un destinatario."}), 400

    # Cargar técnicos asignados
    tecs = mysql_fetchall(
        """SELECT t.nombre, t.especialidad
             FROM mant_visita_tecnicos vt
             JOIN mant_tecnicos t ON t.id=vt.tecnico_id
            WHERE vt.visita_id=%s""",
        (vid,)
    ) or []
    # Repuestos
    reps = mysql_fetchall(
        "SELECT descripcion, cantidad, costo_total FROM mant_visita_repuestos WHERE visita_id=%s",
        (vid,)
    ) or []

    ot   = v.get("numero_ot") or f"V-{vid:05d}"
    fecha_str = v["fecha_programada"].strftime("%d/%m/%Y") if v.get("fecha_programada") else "—"
    horario = ""
    if v.get("hora_inicio"):
        h_ini = str(v["hora_inicio"])[:5]
        h_fin = str(v["hora_fin"])[:5] if v.get("hora_fin") else ""
        horario = f"{h_ini}{(' – ' + h_fin) if h_fin else ''}"
    tipo_label = {
        "preventiva":"Mantención preventiva",
        "correctiva":"Reparación correctiva",
        "garantia":"Cambio / Garantía",
        "inspeccion":"Inspección"
    }.get(v.get("tipo"), v.get("tipo") or "Visita")
    tecs_html = ", ".join(t["nombre"] for t in tecs) if tecs else (v.get("tecnico") or "Por asignar")
    reps_html = ""
    if reps:
        reps_html = "<h3 style='margin:18px 0 6px;font-size:14px;color:#0f172a'>Repuestos asociados</h3><ul style='font-size:13px;color:#374151'>"
        for rp in reps:
            reps_html += f"<li>{rp['descripcion']} × {rp['cantidad']} — ${int(rp['costo_total'] or 0):,}</li>".replace(",", ".")
        reps_html += "</ul>"
    costo_html = ""
    if v.get("costo"):
        costo_html = f"<p style='font-size:14px'><strong>Costo estimado:</strong> ${int(v['costo']):,} CLP</p>".replace(",", ".")

    saludo = f"Estimado/a {v.get('contacto_nombre') or v['razon_social']}"
    body = f"""
    <p style="font-size:14px;color:#374151">{saludo},</p>
    <p style="font-size:14px;color:#374151">
      Le informamos que se ha programado una visita técnica en sus dependencias:
    </p>
    <table style="width:100%;border-collapse:collapse;margin:14px 0;font-size:13.5px">
      <tr><td style="padding:7px 10px;background:#f9fafb;font-weight:600">N° de Orden</td>
          <td style="padding:7px 10px;font-family:monospace;color:#dc2626;font-weight:700">{ot}</td></tr>
      <tr><td style="padding:7px 10px;background:#f9fafb;font-weight:600">Tipo</td>
          <td style="padding:7px 10px">{tipo_label}</td></tr>
      <tr><td style="padding:7px 10px;background:#f9fafb;font-weight:600">Fecha</td>
          <td style="padding:7px 10px"><strong>{fecha_str}</strong>{(' · ' + horario) if horario else ''}</td></tr>
      <tr><td style="padding:7px 10px;background:#f9fafb;font-weight:600">Técnico(s)</td>
          <td style="padding:7px 10px">{tecs_html}</td></tr>
      {('<tr><td style="padding:7px 10px;background:#f9fafb;font-weight:600">Dirección</td><td style="padding:7px 10px">' + (v.get('direccion') or '') + ', ' + (v.get('comuna') or '') + '</td></tr>') if v.get('direccion') else ''}
    </table>
    {costo_html}
    {reps_html}
    <p style="font-size:13px;color:#6b7280;margin-top:18px">
      Si necesita reagendar o cancelar la visita, por favor responda este correo
      o comuníquese con su ejecutivo asignado.
    </p>
    """
    asunto = f"Visita técnica programada — {ot} ({fecha_str})"
    html = _comm_render_email_document(asunto, body, subtitle=f"OT {ot}")
    ok = _send_ilus_email(destinatario, asunto, html, evento="visita_agendada")

    if ok:
        _mant_log("visita", vid, "email_enviado", f"a {destinatario} — OT {ot}")
        return jsonify({"ok": True, "destinatario": destinatario})
    return jsonify({"error": "No se pudo enviar el email. Revisa la configuración SMTP."}), 500


# ── CALENDARIO ────────────────────────────────────────────────────────

@app.route("/mantenciones/calendario")
@_mant_required
def mant_calendario():
    clientes = mysql_fetchall(
        "SELECT id, razon_social FROM mant_clientes WHERE estado='activo' ORDER BY razon_social",
        ()
    ) or []
    tecnicos = mysql_fetchall(
        "SELECT id, nombre, especialidad, nivel, es_externo, foto_url "
        "FROM mant_tecnicos WHERE activo=1 ORDER BY nombre",
        ()
    ) or []
    return render_template("mantenciones/calendario.html",
        clientes = [dict(r) for r in clientes],
        tecnicos = [dict(r) for r in tecnicos],
    )


# ── ANÁLISIS ECONÓMICO ────────────────────────────────────────────────

@app.route("/mantenciones/analisis")
@_mant_required
def mant_analisis():
    # Ingresos por contrato (12 meses)
    ingresos_mes = mysql_fetchall(
        """SELECT YEAR(fecha_inicio) AS anio, MONTH(fecha_inicio) AS mes,
                  COALESCE(SUM(monto_mensual),0) AS total_mensual,
                  COUNT(*) AS n_contratos
           FROM mant_contratos
           WHERE estado IN ('vigente','indefinido')
             AND fecha_inicio >= DATE_SUB(NOW(), INTERVAL 12 MONTH)
           GROUP BY anio, mes ORDER BY anio, mes""",
        ()
    )
    # Proyección próximos 6 meses
    activos_total = mysql_fetchone(
        "SELECT COALESCE(SUM(monto_mensual),0) AS mrr FROM mant_contratos "
        "WHERE estado IN ('vigente','indefinido')", ()
    ) or {}
    mrr = float(activos_total.get("mrr", 0))

    # Visitas por mes (últimos 6)
    visitas_mes = mysql_fetchall(
        """SELECT YEAR(fecha_programada) AS anio, MONTH(fecha_programada) AS mes,
                  COUNT(*) AS total,
                  SUM(estado='completada') AS completadas,
                  COALESCE(SUM(costo),0) AS ingresos_visitas
           FROM mant_visitas
           WHERE fecha_programada >= DATE_SUB(NOW(), INTERVAL 6 MONTH)
           GROUP BY anio, mes ORDER BY anio, mes""",
        ()
    )
    # Top clientes por valor
    top_clientes = mysql_fetchall(
        """SELECT c.razon_social,
                  COALESCE(SUM(ct.monto_mensual),0) AS mrr,
                  COUNT(ct.id) AS n_contratos,
                  COUNT(v.id) AS n_visitas
           FROM mant_clientes c
           LEFT JOIN mant_contratos ct ON ct.cliente_id=c.id
                     AND ct.estado IN ('vigente','indefinido')
           LEFT JOIN mant_visitas v ON v.cliente_id=c.id
           WHERE c.estado='activo'
           GROUP BY c.id ORDER BY mrr DESC LIMIT 10""",
        ()
    )
    # Contratos por vencer
    por_vencer = mysql_fetchall(
        """SELECT ct.*, cl.razon_social,
                  DATEDIFF(ct.fecha_vencimiento, CURDATE()) AS dias_restantes
           FROM mant_contratos ct
           JOIN mant_clientes cl ON cl.id=ct.cliente_id
           WHERE ct.estado IN ('por_vencer','vigente')
             AND ct.fecha_vencimiento IS NOT NULL
             AND ct.es_indefinido=0
           ORDER BY ct.fecha_vencimiento LIMIT 15""",
        ()
    )

    return render_template("mantenciones/analisis.html",
        ingresos_mes = [dict(r) for r in ingresos_mes],
        mrr          = mrr,
        visitas_mes  = [dict(r) for r in visitas_mes],
        top_clientes = [dict(r) for r in top_clientes],
        por_vencer   = [dict(r) for r in por_vencer],
    )


# Cache global de UF — la UF cambia 1 vez al día, no tiene sentido pegarle
# a mindicador.cl en cada request (~200-600ms por hit).
_UF_CACHE = {"ts": 0, "uf": None, "fecha": None, "error": None}
_UF_CACHE_TTL = 3600  # 1 hora


@app.route("/api/uf-actual")
def api_uf_actual():
    """Devuelve el valor actual de la UF desde mindicador.cl (cacheado 1h)."""
    now = time.time()
    if _UF_CACHE["uf"] is not None and (now - _UF_CACHE["ts"]) < _UF_CACHE_TTL:
        return jsonify({"uf": _UF_CACHE["uf"], "fecha": _UF_CACHE["fecha"],
                        "ok": True, "cached": True})
    try:
        import urllib.request as _ur
        req = _ur.Request(
            "https://mindicador.cl/api/uf",
            headers={"User-Agent": "ILUSApp/1.0"}
        )
        with _ur.urlopen(req, timeout=6) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        uf_val = float(data["serie"][0]["valor"])
        fecha  = data["serie"][0]["fecha"][:10]
        _UF_CACHE.update({"ts": now, "uf": uf_val, "fecha": fecha, "error": None})
        return jsonify({"uf": uf_val, "fecha": fecha, "ok": True})
    except Exception as e:
        # Si tenemos un valor cacheado aunque sea viejo, devolverlo en lugar de error
        if _UF_CACHE["uf"] is not None:
            return jsonify({"uf": _UF_CACHE["uf"], "fecha": _UF_CACHE["fecha"],
                            "ok": True, "cached": True, "stale": True})
        return jsonify({"uf": None, "error": str(e), "ok": False})


# ── BÚSQUEDA PRODUCTOS ERP ───────────────────────────────────────────

@app.route("/mantenciones/api/productos/buscar")
@_mant_required
def mant_productos_buscar():
    """
    Typeahead de productos ERP por SKU o descripción.
    Usa la REST API (funciona desde Railway).
    Devuelve [{sku, nombre}].
    """
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify([])
    TOKEN = ERP_CONFIG.get("api_token", "")
    resultados = []
    seen = set()
    try:
        body = _erp_get(
            "/productos",
            {"search": q, "empresa": "01", "fields": "KOPR,NOKOPR,NOKOTI", "visible": "true"},
            TOKEN, timeout=8,
        )
        for p in (body.get("data") or []):
            sku    = (p.get("KOPR")   or "").strip().upper()
            nombre = (p.get("NOKOPR") or "").strip()
            tipo   = (p.get("NOKOTI") or p.get("TIPO") or "").strip()
            if not sku or sku in seen:
                continue
            resultados.append({"sku": sku, "nombre": nombre, "tipo": tipo})
            seen.add(sku)
            if len(resultados) >= 25:
                break
    except Exception:
        pass  # Si REST falla devolvemos lista vacía (mejor que 500)
    return jsonify(resultados)


# ── BÚSQUEDA DOCUMENTO ERP (como cubicador) ───────────────────────────

@app.route("/mantenciones/api/documento")
@_mant_required
def mant_documento_erp():
    """
    Busca un documento ERP por tipo + número usando la REST API (igual que cubicador).
    Devuelve header y líneas de productos.
    GET ?tido=FCV&nudo=12345
    """
    tido = request.args.get("tido", "").strip().upper()
    nudo = request.args.get("nudo", "").strip()
    if not tido or not nudo:
        return jsonify({"error": "Ingresa tipo y número de documento"}), 400

    # GDP = Guía Despacho Provisional (despachos parciales). El ERP usa GDP para
    # documentos en estado "Parc" antes de facturación final. Incluido para
    # importar equipos desde guías de despacho provisorias.
    TIDOS_VALIDOS = {"FCV","BLV","NVI","NVV","GDV","GDP","VD","WEB","FCO"}
    if tido not in TIDOS_VALIDOS:
        return jsonify({"error": f"Tipo '{tido}' no válido. Usa: {', '.join(sorted(TIDOS_VALIDOS))}"}), 400

    try:
        header, lineas = _cubicador_fetch(tido, nudo)
    except ConnectionError as ce:
        return jsonify({"error": str(ce)}), 503
    except Exception as e:
        return jsonify({"error": f"Error al consultar ERP: {e}"}), 500

    if not header:
        return jsonify({"error": f"Documento {tido} {nudo} no encontrado en el ERP"}), 404

    # Formatear líneas para el wizard
    # NOTA: _cubicador_fetch devuelve el campo "descripcion_erp" (no "descripcion")
    _ZZ_SKUS = {"ZZENVIO","ZZINGREPUESTO","ZZSERVTEC","ZZRETIRO","ZZINSTALACION","ZZINGARREQUIP"}
    items = []
    filtradas = 0
    total_lineas = len(lineas)
    for l in lineas:
        sku  = (l.get("sku") or l.get("KOPRCT") or "").strip().upper()
        nom  = (l.get("descripcion_erp") or l.get("descripcion") or l.get("nombre") or l.get("NOKOPR") or "").strip()
        qty  = int(float(l.get("cantidad") or l.get("qty") or 1))
        if not nom and not sku:
            filtradas += 1
            continue
        # Excluir SKUs de servicio/flete (ya marcados con es_zz por _cubicador_fetch)
        if l.get("es_zz") or sku.upper() in _ZZ_SKUS:
            filtradas += 1
            continue
        items.append({"sku": sku, "nombre": nom or sku, "cantidad": max(qty, 1)})

    debug = request.args.get("debug") == "1"
    resp = {
        "ok":           True,
        "tido":         header.get("tido", tido),
        "nudo":         header.get("nudo_display", nudo),
        "fecha":        header.get("fecha",""),
        "cliente":      header.get("cliente_nombre",""),
        "rut":          header.get("cliente_rut",""),
        "direccion":    header.get("direccion",""),
        "comuna":       header.get("comuna",""),
        "items":        items,
        "total_lineas": total_lineas,
        "filtradas":    filtradas,
    }
    if debug:
        resp["raw_lineas"] = [{
            "sku":   l.get("sku") or l.get("KOPRCT",""),
            "desc":  l.get("descripcion_erp") or l.get("descripcion",""),
            "qty":   l.get("cantidad",0),
            "es_zz": l.get("es_zz",False),
            "keys":  list(l.keys())[:12],
        } for l in lineas[:20]]
    return jsonify(resp)


# ── BÚSQUEDA ERP ─────────────────────────────────────────────────────

def _rut_norm(raw: str) -> str:
    """Normaliza RUT: quita puntos, espacios y guión. Devuelve solo dígitos + DV opcional."""
    if not raw:
        return ""
    return raw.replace(".", "").replace(" ", "").replace("-", "").upper()


def _ruts_equivalentes(rut_a: str, rut_b: str) -> bool:
    """True si dos RUT son el mismo (comparando sin puntos/DV)."""
    a = _rut_norm(rut_a)
    b = _rut_norm(rut_b)
    if not a or not b:
        return False
    # Comparar sin DV (último carácter)
    a_sin = a[:-1] if len(a) > 1 else a
    b_sin = b[:-1] if len(b) > 1 else b
    return a_sin == b_sin or a == b


@app.route("/mantenciones/api/clientes/<int:cid>/equipos-import-mismatch", methods=["POST"])
@_mant_required
def mant_equipos_import_mismatch(cid):
    """
    Registra en mant_logs cuando el usuario importa equipos de un documento ERP
    cuyo RUT NO coincide con el de la ficha de cliente. Pide motivo justificado.
    """
    d = (request.get_json(silent=True) or {})
    motivo  = (d.get("motivo")  or "").strip()
    rut_doc = (d.get("rut_doc") or "").strip()
    tido    = (d.get("tido")    or "").strip()
    nudo    = (d.get("nudo")    or "").strip()
    if len(motivo) < 8:
        return jsonify({"error": "El motivo debe tener al menos 8 caracteres."}), 400
    detalle = f"Documento {tido} {nudo} (RUT cliente doc: {rut_doc}) — Motivo: {motivo}"
    _mant_log("cliente", cid, "import_rut_mismatch", detalle)
    return jsonify({"ok": True})


# TIDOs válidos para listar como "documentos del cliente" (todos los G* + ventas)
# COV=Cotización, FCV=Factura, BLV=Boleta, NVI=Nota Interna, NVV=Nota Venta,
# GDV=Guía Despacho Venta, GDP=Guía Despacho Provisional, GTR=Guía Traslado,
# GRD=Guía Retiro/Devolución, FCO=Factura Compra
_RANDOM_TIDOS_VENTA = ('FCV','BLV','NVI','NVV','GDV','GDP','GTR','GRD','FCO','COV')


@app.route("/mantenciones/api/buscar-erp-sql", methods=["POST"])
@_mant_required
def mant_buscar_erp_sql():
    """
    Búsqueda inteligente de documentos en Random ERP via SQL Server directo.

    Detecta automáticamente el tipo de búsqueda según lo que escribió el usuario:
      - Solo dígitos 7-9 chars  → RUT (busca todos los docs del cliente)
      - Solo dígitos 1-6 chars  → Número de documento (busca por NUDO)
      - Texto                   → Razón social (LIKE en NOKOEN)

    Maneja prefijos VD/WEB que se almacenan como NVV con NUDO prefijado.
    """
    d = request.get_json(silent=True) or {}
    q = (d.get("q") or "").strip()
    if len(q) < 3:
        return jsonify({"error": "Mínimo 3 caracteres"}), 400

    pool = _random_sql_pool()
    if pool is None:
        return jsonify({
            "error": "Conexión a Random ERP no configurada. Pídele al admin que setee RANDOM_SQL_* en Railway.",
            "documentos": [], "sin_conexion": True
        }), 200

    # Normalizar query
    q_clean   = q.replace(".", "").replace(" ", "").replace("-", "").upper()
    is_digits = q_clean.isdigit()
    tidos_in  = "','".join(_RANDOM_TIDOS_VENTA)

    docs = []
    modo = ""
    try:
        # ── Modo 1: RUT (7-9 dígitos) ───────────────────────────
        if is_digits and 7 <= len(q_clean) <= 9:
            modo = "rut"
            rut_base = q_clean[:-1] if len(q_clean) >= 8 else q_clean
            docs = _random_sql_query(f"""
                SELECT TOP 100
                    e.IDMAEEDO, e.TIDO, e.NUDO, e.ENDO,
                    e.FEEMDO, e.VANEDO, e.VAIVDO, e.VABRDO,
                    e.ESPGDO, e.ESDO,
                    en.NOKOEN AS razon_social,
                    en.RTEN   AS rut
                FROM MAEEDO e
                LEFT JOIN MAEEN en ON LTRIM(RTRIM(en.RTEN)) =
                      LEFT(LTRIM(RTRIM(e.ENDO)),
                           CHARINDEX('-', LTRIM(RTRIM(e.ENDO)) + '-') - 1)
                WHERE (e.ENDO LIKE %s OR e.ENDO LIKE %s)
                  AND e.TIDO IN ('{tidos_in}')
                ORDER BY e.FEEMDO DESC
            """, (f"{rut_base}%", f"%{q_clean}%")) or []

        # ── Modo 2: Número de documento (1-6 dígitos) ──────────
        if not docs and is_digits and 1 <= len(q_clean) <= 7:
            modo = "numero"
            # NUDO en MAEEDO se guarda con padding de ceros (10 chars) o con prefijo VD/WEB
            nudo_padded = q_clean.zfill(10)
            nudo_vd     = f"VD{q_clean.zfill(8)}"
            nudo_web    = f"WEB{q_clean.zfill(7)}"
            docs = _random_sql_query(f"""
                SELECT TOP 50
                    e.IDMAEEDO, e.TIDO, e.NUDO, e.ENDO,
                    e.FEEMDO, e.VANEDO, e.VAIVDO, e.VABRDO,
                    e.ESPGDO, e.ESDO,
                    en.NOKOEN AS razon_social,
                    en.RTEN   AS rut
                FROM MAEEDO e
                LEFT JOIN MAEEN en ON LTRIM(RTRIM(en.RTEN)) =
                      LEFT(LTRIM(RTRIM(e.ENDO)),
                           CHARINDEX('-', LTRIM(RTRIM(e.ENDO)) + '-') - 1)
                WHERE e.NUDO IN (%s, %s, %s)
                  AND e.TIDO IN ('{tidos_in}')
                ORDER BY e.FEEMDO DESC
            """, (nudo_padded, nudo_vd, nudo_web)) or []

        # ── Modo 3: Razón social (texto) ───────────────────────
        # Estrategia 2-pasos para evitar JOIN calculado lento:
        #   3a) Buscar primero los RUTs en MAEEN (tabla chica, índice por NOKOEN)
        #   3b) Luego traer documentos por ENDO IN (rut1-N, rut2-N, ...)
        if not docs and not is_digits:
            modo = "nombre"
            q_like = f"%{q.upper()}%"
            # 3a) RUTs que coinciden con el nombre buscado
            ruts = _random_sql_query("""
                SELECT TOP 20 LTRIM(RTRIM(RTEN)) AS rut, LTRIM(RTRIM(NOKOEN)) AS razon
                FROM MAEEN
                WHERE UPPER(NOKOEN) LIKE %s AND TIEN IN ('C','A')
            """, (q_like,)) or []
            if ruts:
                # 3b) Construir IN(...) seguro (los valores vienen de MAEEN, no del usuario)
                rut_map = {r['rut']: r['razon'] for r in ruts}
                # ENDO en MAEEDO incluye DV (formato 65206047-N), generamos patrones LIKE
                like_clauses = " OR ".join(["e.ENDO LIKE %s"] * len(rut_map))
                params = tuple(f"{rut}%" for rut in rut_map.keys())
                docs_raw = _random_sql_query(f"""
                    SELECT TOP 50
                        e.IDMAEEDO, e.TIDO, e.NUDO, e.ENDO,
                        e.FEEMDO, e.VANEDO, e.VAIVDO, e.VABRDO,
                        e.ESPGDO, e.ESDO
                    FROM MAEEDO e
                    WHERE ({like_clauses})
                      AND e.TIDO IN ('{tidos_in}')
                    ORDER BY e.FEEMDO DESC
                """, params) or []
                # Enriquecer con razón social mapeada por prefijo de RUT
                for r in docs_raw:
                    endo = (r.get("ENDO") or "").strip()
                    rut_clean = endo.split("-")[0] if "-" in endo else endo
                    r["rut"] = rut_clean
                    r["razon_social"] = rut_map.get(rut_clean, "")
                docs = docs_raw
    except PermissionError as pe:
        return jsonify({"error": f"Bloqueado por seguridad: {pe}"}), 403

    # Formatear respuesta
    out = []
    for r in docs:
        nudo_raw = (r.get("NUDO") or "").strip()
        tido_raw = (r.get("TIDO") or "").strip()
        # Detectar prefijo VD/WEB en NUDO (esos se guardan como TIDO=NVV)
        tido_display = tido_raw
        if tido_raw == "NVV" and nudo_raw.startswith("VD"):
            tido_display = "VD"
            nudo_display = nudo_raw[2:].lstrip("0") or "0"
        elif tido_raw == "NVV" and nudo_raw.startswith("WEB"):
            tido_display = "WEB"
            nudo_display = nudo_raw[3:].lstrip("0") or "0"
        else:
            nudo_display = nudo_raw.lstrip("0") or "0"

        fe = r.get("FEEMDO")
        out.append({
            "idmaeedo":     r.get("IDMAEEDO"),
            "tido":         tido_raw,                  # tido real para fetch detalle
            "nudo":         nudo_raw,                  # nudo real para fetch detalle
            "tido_display": tido_display,              # para mostrar al usuario
            "nudo_display": nudo_display,
            "rut":          (r.get("rut")  or r.get("ENDO") or "").strip(),
            "razon_social": (r.get("razon_social") or "").strip().title(),
            "fecha":        fe.strftime("%d/%m/%Y") if fe else "",
            "fecha_iso":    fe.strftime("%Y-%m-%d") if fe else "",
            "valor_neto":   float(r.get("VANEDO") or 0),
            "valor_total":  float(r.get("VABRDO") or 0),
            "estado_pago":  (r.get("ESPGDO") or "").strip(),
        })

    return jsonify({
        "ok": True,
        "modo": modo,
        "documentos": out,
        "count": len(out),
        "query": q,
    })


@app.route("/mantenciones/api/buscar-erp", methods=["POST"])
@_mant_required
def mant_buscar_erp():
    """Busca cliente/documentos en el ERP por RUT o razón social."""
    d   = request.get_json(silent=True) or {}
    q   = d.get("q", "").strip()
    if not q:
        return jsonify({"error": "Término de búsqueda requerido"}), 400
    ERP_SALES = ERP_CONFIG.get("table_sales", "HEBDOC")
    erp_conn = get_erp_conn()
    if not erp_conn:
        return jsonify({
            "error": "No hay conexión directa al ERP. Usa la búsqueda por número de documento.",
            "documentos": [], "sin_conexion": True
        }), 200
    # Normalizar el query:
    #  - quitar puntos/guión/espacios para que "65.206.047-1" matchee "65206047" o "65206047-1"
    #  - upper para que el LIKE en NRAZON sea case-insensitive (NRAZON suele estar en mayúsculas)
    q_clean    = q.replace(".", "").replace(" ", "")
    q_sin_dv   = q_clean.split("-")[0] if "-" in q_clean else q_clean
    q_like     = f"%{q.upper()}%"
    q_like_clean = f"%{q_clean}%"
    q_like_sindv = f"%{q_sin_dv}%"
    try:
        with erp_conn.cursor() as cur:
            cur.execute(
                f"""SELECT DISTINCT
                       TRIM(d.NRAZON) AS razon_social,
                       TRIM(d.NRUC)   AS rut,
                       TRIM(d.TIDO)   AS tipo_doc,
                       TRIM(d.NUDO)   AS num_doc,
                       d.FEMIS        AS fecha,
                       TRIM(d.NOMBR)  AS producto,
                       TRIM(d.KOPRCT) AS sku,
                       d.CANTD        AS cantidad
                    FROM `{ERP_SALES}` d
                    WHERE (
                          UPPER(TRIM(d.NRAZON)) LIKE %s
                       OR TRIM(d.NRUC) LIKE %s
                       OR REPLACE(REPLACE(TRIM(d.NRUC),'.',''),' ','') LIKE %s
                       OR REPLACE(REPLACE(TRIM(d.NRUC),'.',''),' ','') LIKE %s
                    )
                      AND d.TIDO IN ('FCV','BLV','NVV','VD','WEB','FCO','NVI','GDV','GDP')
                    ORDER BY d.FEMIS DESC
                    LIMIT 100""",
                (q_like, q_like_clean, q_like_clean, q_like_sindv)
            )
            rows = cur.fetchall()
        erp_conn.close()
        docs = {}
        for r in rows:
            key = f"{r['tipo_doc']} {r['num_doc']}"
            if key not in docs:
                docs[key] = {
                    "razon_social": r["razon_social"],
                    "rut":          r["rut"],
                    "tipo_doc":     r["tipo_doc"],
                    "num_doc":      r["num_doc"],
                    "fecha":        str(r["fecha"]) if r["fecha"] else "",
                    "lineas":       []
                }
            docs[key]["lineas"].append({
                "sku":      r["sku"],
                "nombre":   r["producto"],
                "cantidad": int(r["cantidad"] or 1),
            })
        return jsonify({"ok": True, "documentos": list(docs.values())})
    except Exception as e:
        return jsonify({
            "error": "No hay conexión directa al ERP. Usa la búsqueda por número de documento.",
            "documentos": [], "sin_conexion": True
        }), 200


# ── DOCUMENTOS ERP POR RUT (auto-search en tab Equipos) ─────────────────
@app.route("/mantenciones/api/clientes/<int:cid>/documentos-erp")
@_mant_required
def mant_documentos_por_rut(cid):
    """
    Busca en ERP todos los documentos (FCV/BLV/GDV) del cliente identificado por su RUT.
    Estrategia: generar TODOS los formatos posibles del RUT en Python y usar IN(...)
    para que el ERP pueda aprovechar el índice de NRUC — sin funciones por fila.
    GET /mantenciones/api/clientes/<cid>/documentos-erp
    """
    cliente = mysql_fetchone("SELECT rut, razon_social FROM mant_clientes WHERE id=%s", (cid,))
    if not cliente or not cliente.get("rut"):
        return jsonify({"sin_rut": True, "documentos": []})

    rut = cliente["rut"].strip()

    # ── Normalización: extraer solo dígitos + K ─────────────────────────
    rut_digits = re.sub(r"[^0-9kK]", "", rut).upper()   # ej: "760087564"
    dv         = rut_digits[-1] if len(rut_digits) > 1 else ""
    cuerpo     = rut_digits[:-1] if len(rut_digits) > 1 else rut_digits  # sin DV

    # ── Formatos que el ERP podría estar usando (basado en observación) ──
    # El ERP típicamente almacena RUTs sin puntos, y frecuentemente sin DV.
    candidatos = set(filter(None, [
        cuerpo,                       # 76008756   ← más común en ERP
        rut_digits,                   # 760087564  (con DV pegado)
        f"{cuerpo}-{dv}" if dv else None,  # 76008756-4
        f"{cuerpo}-{dv.lower()}" if dv else None,  # 76008756-k (minúscula)
        rut.strip(),                  # tal como está guardado en nuestro sistema
    ]))
    placeholders = ",".join(["%s"] * len(candidatos))
    params_rut   = list(candidatos)

    ERP_SALES = ERP_CONFIG.get("table_sales", "HEBDOC")
    erp_conn = get_erp_conn()
    if not erp_conn:
        return jsonify({
            "sin_conexion": True,
            "rut": rut,
            "documentos": [],
            "msg": "Conexión al ERP no disponible desde Railway. "
                   "Usa la búsqueda por número de documento."
        }), 200

    try:
        with erp_conn.cursor() as cur:
            # IN con literales exactos → el ERP puede usar índice en NRUC.
            # Sin funciones sobre la columna = mucho más rápido.
            cur.execute(
                f"""SELECT
                       TRIM(d.NRAZON) AS razon_social,
                       TRIM(d.NRUC)   AS rut,
                       TRIM(d.TIDO)   AS tipo_doc,
                       TRIM(d.NUDO)   AS num_doc,
                       d.FEMIS        AS fecha,
                       TRIM(d.NOMBR)  AS producto,
                       TRIM(d.KOPRCT) AS sku,
                       d.CANTD        AS cantidad
                    FROM `{ERP_SALES}` d
                    WHERE d.TIDO IN ('FCV','BLV','GDV','VD','WEB','NVI','NVV')
                      AND TRIM(d.NRUC) IN ({placeholders})
                    ORDER BY d.FEMIS DESC
                    LIMIT 300""",
                params_rut
            )
            rows = cur.fetchall()
        erp_conn.close()

        docs = {}
        for r in rows:
            key = f"{r['tipo_doc']}|{r['num_doc']}"
            if key not in docs:
                docs[key] = {
                    "razon_social": r["razon_social"] or "",
                    "rut":          r["rut"] or "",
                    "tipo_doc":     r["tipo_doc"],
                    "num_doc":      str(r["num_doc"] or "").lstrip("0") or str(r["num_doc"]),
                    "num_doc_raw":  str(r["num_doc"] or ""),
                    "fecha":        str(r["fecha"])[:10] if r["fecha"] else "",
                    "lineas":       []
                }
            nom = (r["producto"] or "").strip()
            sku = (r["sku"] or "").strip().upper()
            if nom or sku:
                docs[key]["lineas"].append({
                    "sku":      sku,
                    "nombre":   nom or sku,
                    "cantidad": int(float(r["cantidad"] or 1)),
                })

        return jsonify({
            "ok":         True,
            "rut":        rut,
            "cliente":    cliente["razon_social"],
            "documentos": list(docs.values())
        })
    except Exception as e:
        erp_conn and erp_conn.close()
        return jsonify({
            "sin_conexion": True,
            "rut": rut,
            "documentos": [],
            "msg": "No se pudo conectar al ERP. Usa la búsqueda por número de documento."
        }), 200


# ══════════════════════════════════════════════════════════════════════
#  REPORTES DE SERVICIO — INFORME POST SERVICIO
# ══════════════════════════════════════════════════════════════════════

MANT_REPORTES_UPLOADS = os.path.join(BASE_DIR, "static", "uploads", "mantenciones", "reportes")
MANT_REPORTES_HTML    = os.path.join(BASE_DIR, "static", "uploads", "mantenciones", "reportes", "html")
os.makedirs(MANT_REPORTES_UPLOADS, exist_ok=True)
os.makedirs(MANT_REPORTES_HTML, exist_ok=True)

ALLOWED_REPORT_IMG = {"jpg","jpeg","png","gif","webp"}
ALLOWED_ADJUNTO    = {"pdf","doc","docx","jpg","jpeg","png","gif","webp","xlsx","xls"}


def _save_reporte_html_snapshot(rid):
    """Genera y persiste HTML del reporte en disco. Devuelve ruta relativa o None."""
    try:
        r = mysql_fetchone(
            "SELECT r.*, c.razon_social, c.rut FROM mant_reportes r "
            "JOIN mant_clientes c ON c.id=r.cliente_id WHERE r.id=%s", (rid,)
        )
        if not r: return None
        rep = dict(r)
        cliente = {"razon_social": rep.pop("razon_social",""), "rut": rep.pop("rut","")}
        for k in ("objetivos","trabajos","observaciones","maquinas_json","ai_acciones"):
            if rep.get(k):
                try: rep[k] = json.loads(rep[k])
                except: rep[k] = []
        html = _reporte_to_html(rep, cliente)
        fname = f"informe_{rid}_{int(time.time())}.html"
        fpath = os.path.join(MANT_REPORTES_HTML, fname)
        with open(fpath, "w", encoding="utf-8") as fh:
            fh.write(html)
        rel = f"uploads/mantenciones/reportes/html/{fname}"
        # Borrar snapshot anterior para no acumular
        try:
            old = mysql_fetchone("SELECT html_path FROM mant_reportes WHERE id=%s",(rid,))
            if old and old.get("html_path") and old["html_path"] != rel:
                op = os.path.join(BASE_DIR, "static", old["html_path"])
                if os.path.exists(op): os.remove(op)
        except Exception: pass
        mysql_execute(
            "UPDATE mant_reportes SET html_path=%s, html_generated_at=NOW() WHERE id=%s",
            (rel, rid)
        )
        return rel
    except Exception as exc:
        print(f"[REPORTE HTML SNAPSHOT] {exc}")
        return None


@app.route("/mantenciones/api/clientes/<int:cid>/reportes", methods=["GET"])
@_mant_required
def mant_reportes_list(cid):
    rows = mysql_fetchall(
        "SELECT id,tipo,estado,ticket_num,asunto,tecnico_junior,tecnico_senior,"
        "fecha_inicio,fecha_cierre,ai_diagnostico,ai_fecha,html_path,html_generated_at,"
        "created_by,created_at "
        "FROM mant_reportes WHERE cliente_id=%s ORDER BY created_at DESC", (cid,)
    )
    def _fmt(r):
        d = dict(r)
        for k in ("fecha_inicio","fecha_cierre","ai_fecha","created_at","html_generated_at"):
            if d.get(k): d[k] = str(d[k])[:16]
        if d.get("html_path"):
            d["html_url"] = f"/static/{d['html_path']}"
        return d
    return jsonify([_fmt(r) for r in rows])


@app.route("/mantenciones/api/reportes/<int:rid>/regenerar-html", methods=["POST"])
@_mant_required
def mant_reporte_regenerar_html(rid):
    """Regenera snapshot HTML del reporte y devuelve la URL."""
    rel = _save_reporte_html_snapshot(rid)
    if not rel:
        return jsonify({"error":"No se pudo generar"}), 500
    return jsonify({"ok": True, "html_url": f"/static/{rel}"})


@app.route("/mantenciones/api/clientes/<int:cid>/reportes", methods=["POST"])
@_mant_required
def mant_reporte_crear(cid):
    d = request.get_json(silent=True) or {}
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_reportes
                   (cliente_id,tipo,estado,ticket_num,asunto,
                    tecnico_junior,tecnico_senior,
                    fecha_solicitado,fecha_inicio,fecha_cierre,
                    antecedentes,objetivos,trabajos,observaciones,
                    maquinas_json,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (cid, d.get("tipo","mantencion"), d.get("estado","borrador"),
                 d.get("ticket_num",""), d.get("asunto",""),
                 d.get("tecnico_junior",""), d.get("tecnico_senior",""),
                 d.get("fecha_solicitado") or None,
                 d.get("fecha_inicio") or None, d.get("fecha_cierre") or None,
                 d.get("antecedentes",""),
                 json.dumps(d.get("objetivos",[]),    ensure_ascii=False),
                 json.dumps(d.get("trabajos",[]),     ensure_ascii=False),
                 json.dumps(d.get("observaciones",[]),ensure_ascii=False),
                 json.dumps(d.get("maquinas",[]),     ensure_ascii=False),
                 current_username())
            )
            rid = cur.lastrowid
        conn.commit()
        _mant_log("reporte", rid, "creado", d.get("asunto",""))
        # Snapshot HTML automático
        try: _save_reporte_html_snapshot(rid)
        except Exception: pass
        return jsonify({"ok": True, "id": rid})
    finally:
        conn.close()


@app.route("/mantenciones/api/reportes/<int:rid>", methods=["GET"])
@_mant_required
def mant_reporte_get(rid):
    r = mysql_fetchone("SELECT * FROM mant_reportes WHERE id=%s", (rid,))
    if not r: return jsonify({"error":"No encontrado"}), 404
    d = dict(r)
    for k in ("fecha_solicitado","fecha_inicio","fecha_cierre","ai_fecha","created_at","updated_at"):
        if d.get(k): d[k] = str(d[k])[:10]
    for k in ("objetivos","trabajos","observaciones","maquinas_json","fotos_json","ai_acciones"):
        d[k] = json.loads(d[k] or "[]")
    # Adjuntos
    fotos = mysql_fetchall(
        "SELECT id,nombre,archivo_path,tipo,created_at FROM mant_contrato_adjuntos "
        "WHERE contrato_id=%s AND tipo='imagen' ORDER BY created_at", (rid,)
    )
    d["fotos"] = [{"id":f["id"],"nombre":f["nombre"],
                   "url":f"/static/uploads/mantenciones/reportes/{f['archivo_path']}"} for f in fotos]
    return jsonify(d)


@app.route("/mantenciones/api/reportes/<int:rid>", methods=["PUT"])
@_mant_required
def mant_reporte_update(rid):
    d = request.get_json(silent=True) or {}
    allowed = ["tipo","estado","ticket_num","asunto","tecnico_junior","tecnico_senior",
               "fecha_solicitado","fecha_inicio","fecha_cierre",
               "antecedentes","objetivos","trabajos","observaciones","maquinas_json"]
    sets, vals = [], []
    for f in allowed:
        if f not in d: continue
        if f in ("objetivos","trabajos","observaciones","maquinas_json"):
            sets.append(f"{f}=%s")
            vals.append(json.dumps(d[f], ensure_ascii=False))
        else:
            sets.append(f"{f}=%s")
            vals.append(d[f] or None)
    if not sets: return jsonify({"error":"Sin campos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_reportes SET {','.join(sets)} WHERE id=%s", vals+[rid])
        conn.commit()
        # Regenerar snapshot HTML
        try: _save_reporte_html_snapshot(rid)
        except Exception: pass
        # Log
        rep_info = mysql_fetchone("SELECT cliente_id, asunto, ticket_num FROM mant_reportes WHERE id=%s", (rid,))
        campos_mod = ", ".join([f for f in allowed if f in d][:6])
        _mant_log("reporte", rid, "actualizado", f"campos: {campos_mod}")
        if rep_info:
            _mant_log("cliente", rep_info["cliente_id"], "reporte_actualizado",
                      f"#{rep_info.get('ticket_num') or rid} — {rep_info.get('asunto') or ''}")
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/reportes/<int:rid>", methods=["DELETE"])
@_mant_required
def mant_reporte_del(rid):
    rep_info = mysql_fetchone("SELECT cliente_id, asunto, ticket_num FROM mant_reportes WHERE id=%s", (rid,))
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM mant_reportes WHERE id=%s", (rid,))
        conn.commit()
        if rep_info:
            label = f"#{rep_info.get('ticket_num') or rid} — {rep_info.get('asunto') or ''}"
            _mant_log("reporte", rid, "eliminado", label)
            _mant_log("cliente", rep_info["cliente_id"], "reporte_eliminado", label)
        return jsonify({"ok": True})
    finally:
        conn.close()


@app.route("/mantenciones/api/reportes/<int:rid>/fotos", methods=["POST"])
@_mant_required
def mant_reporte_foto_subir(rid):
    """Sube foto al registro fotográfico del reporte."""
    f = request.files.get("foto")
    if not f or not f.filename:
        return jsonify({"error":"Sin archivo"}), 400
    ext = f.filename.rsplit(".",1)[-1].lower()
    if ext not in ALLOWED_REPORT_IMG:
        return jsonify({"error":"Tipo no permitido"}), 400
    fname  = secure_filename(f"rep{rid}_{int(time.time())}_{f.filename}")
    fpath  = os.path.join(MANT_REPORTES_UPLOADS, fname)
    f.save(fpath)
    size   = os.path.getsize(fpath)
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_contrato_adjuntos
                   (contrato_id,cliente_id,tipo,nombre,archivo_nombre,archivo_path,
                    mime_type,tamaño_bytes,created_by)
                   SELECT %s,cliente_id,'imagen',%s,%s,%s,%s,%s,%s
                   FROM mant_reportes WHERE id=%s""",
                (rid, f.filename, f.filename, fname,
                 f"image/{ext}", size, current_username(), rid)
            )
            aid = cur.lastrowid
        conn.commit()
        # Log
        rep_info = mysql_fetchone("SELECT cliente_id FROM mant_reportes WHERE id=%s", (rid,))
        _mant_log("reporte", rid, "foto_subida", f.filename)
        if rep_info:
            _mant_log("cliente", rep_info["cliente_id"], "reporte_foto_subida",
                      f"reporte #{rid} — {f.filename}")
        return jsonify({"ok":True,"id":aid,
                        "url":f"/static/uploads/mantenciones/reportes/{fname}",
                        "nombre":f.filename})
    finally:
        conn.close()


@app.route("/mantenciones/api/reportes/<int:rid>/analizar-ia", methods=["POST"])
@_mant_required
def mant_reporte_analizar(rid):
    """Análisis IA del reporte: diagnóstico, acciones sugeridas, alertas."""
    r = mysql_fetchone("SELECT r.*, c.razon_social FROM mant_reportes r "
                       "JOIN mant_clientes c ON c.id=r.cliente_id WHERE r.id=%s", (rid,))
    if not r: return jsonify({"error":"No encontrado"}), 404

    maquinas  = json.loads(r.get("maquinas_json") or "[]")
    trabajos  = json.loads(r.get("trabajos") or "[]")
    obs       = json.loads(r.get("observaciones") or "[]")
    ai_key    = _get_ai_key()
    if not ai_key:
        return jsonify({"error":"API IA no configurada"}), 503

    prompt = f"""Eres el sistema de análisis técnico de ILUS Sport & Health Solution SPA.
Analiza este Informe Post Servicio y genera un diagnóstico inteligente.

CLIENTE: {r['razon_social']}
TIPO: {r['tipo']}
TICKET: {r['ticket_num']}
TÉCNICO: {r['tecnico_junior']} / Senior: {r['tecnico_senior']}
FECHAS: {r['fecha_inicio']} → {r['fecha_cierre']}

ANTECEDENTES:
{r.get('antecedentes','')}

TRABAJOS REALIZADOS:
{chr(10).join(f'• {t}' for t in trabajos)}

OBSERVACIONES:
{chr(10).join(f'• {o}' for o in obs)}

EQUIPOS INTERVENIDOS:
{chr(10).join(f'• {m.get("sku","")} {m.get("descripcion","")} x{m.get("cantidad",1)}: {m.get("observacion","")}' for m in maquinas)}

Responde SOLO en JSON con esta estructura:
{{
  "diagnostico": "Diagnóstico técnico conciso (2-3 párrafos)",
  "estado_flota": "bueno|regular|critico",
  "indice_salud": 0-100,
  "acciones": [
    {{"urgencia":"alta|media|baja","tipo":"correctiva|preventiva|cotizacion|seguimiento|contrato",
      "titulo":"...", "descripcion":"...", "plazo":"...", "costo_estimado":null_o_numero}}
  ],
  "piezas_criticas": ["parte1","parte2"],
  "notificaciones_sugeridas": [
    {{"tipo":"email|sistema","titulo":"...","mensaje":"...","destinatario":"cliente|tecnico|admin"}}
  ],
  "requiere_cotizacion": true_o_false,
  "items_cotizacion": [{{"sku":"","descripcion":"","cant":1,"precio_unit":0}}]
}}"""

    try:
        import anthropic as _anthropic
        client = _anthropic.Anthropic(api_key=ai_key)
        msg = client.messages.create(
            model="claude-opus-4-5", max_tokens=2000,
            messages=[{"role":"user","content":prompt}]
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
        resultado = json.loads(raw)
    except Exception as e:
        return jsonify({"error":f"Error IA: {e}"}), 500

    # Guardar diagnóstico + crear notificaciones sugeridas automáticamente
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE mant_reportes SET ai_diagnostico=%s, ai_acciones=%s, "
                "ai_fecha=%s, ai_usuario=%s WHERE id=%s",
                (resultado.get("diagnostico",""),
                 json.dumps(resultado.get("acciones",[]), ensure_ascii=False),
                 datetime.now(), current_username(), rid)
            )
            # Crear notificaciones sugeridas automáticamente
            for n in resultado.get("notificaciones_sugeridas",[]):
                cur.execute(
                    """INSERT INTO mant_notificaciones
                       (cliente_id,entidad,entidad_id,tipo,titulo,mensaje,
                        canal,estado,created_by)
                       VALUES (%s,'reporte',%s,'ai_alerta',%s,%s,%s,'pendiente',%s)""",
                    (r["cliente_id"], rid, n.get("titulo",""),
                     n.get("mensaje",""), n.get("tipo","sistema"),
                     current_username())
                )
        conn.commit()
        _mant_log("reporte", rid, "analizado_ia",
                  f"salud={resultado.get('indice_salud')} estado={resultado.get('estado_flota')}")
        # Regenerar snapshot HTML con IA incorporada
        try: _save_reporte_html_snapshot(rid)
        except Exception: pass
        return jsonify({"ok":True, "resultado":resultado})
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════
#  REPORTES — exportación Word + envío email
# ══════════════════════════════════════════════════════════════════════

def _build_reporte_docx(rep, cliente):
    """Genera un DOCX corporativo del reporte. Devuelve bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # python-docx no instalado — construir HTML simulado descargable como .doc
        return None

    import io as _io
    doc = Document()

    # Márgenes
    for s in doc.sections:
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin   = Cm(2.0)
        s.right_margin  = Cm(2.0)

    # Estilo
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Título
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("INFORME POST-SERVICIO")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ILUS Sport & Health Solution SPA  |  RUT 76.996.964-0")
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # Tabla cabecera (cliente / ticket / técnicos / fechas)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Light Grid Accent 2'
    info = [
        ("Cliente",      cliente.get('razon_social','—')),
        ("RUT",          cliente.get('rut','—')),
        ("Ticket",       rep.get('ticket_num','—')),
        ("Asunto",       rep.get('asunto','—')),
    ]
    for i,(k,v) in enumerate(info):
        tbl.cell(i,0).text = k
        tbl.cell(i,1).text = str(v)
    for row in tbl.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Datos del servicio
    p = doc.add_paragraph()
    p.add_run("DATOS DEL SERVICIO").bold = True
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    svc = doc.add_table(rows=4, cols=2)
    svc.style = 'Light Grid Accent 2'
    svc_data = [
        ("Tipo",            rep.get('tipo','—')),
        ("Estado",          rep.get('estado','—')),
        ("Técnico Junior",  rep.get('tecnico_junior','—')),
        ("Técnico Senior",  rep.get('tecnico_senior','—')),
    ]
    for i,(k,v) in enumerate(svc_data):
        svc.cell(i,0).text = k
        svc.cell(i,1).text = str(v) if v else '—'
        svc.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Antecedentes
    if rep.get('antecedentes'):
        doc.add_paragraph().add_run("ANTECEDENTES").bold = True
        doc.add_paragraph(rep['antecedentes'])

    # Listas (objetivos, trabajos, observaciones)
    for label, key in [("OBJETIVOS","objetivos"), ("TRABAJOS REALIZADOS","trabajos"), ("OBSERVACIONES","observaciones")]:
        items = rep.get(key) or []
        if isinstance(items, str):
            try: items = json.loads(items)
            except: items = []
        if items:
            doc.add_paragraph().add_run(label).bold = True
            for it in items:
                doc.add_paragraph(str(it), style='List Bullet')

    # Equipos / máquinas
    maquinas = rep.get('maquinas_json') or []
    if isinstance(maquinas, str):
        try: maquinas = json.loads(maquinas)
        except: maquinas = []
    if maquinas:
        doc.add_paragraph().add_run("EQUIPOS ATENDIDOS").bold = True
        mt = doc.add_table(rows=1, cols=4)
        mt.style = 'Light Grid Accent 2'
        hdr = mt.rows[0].cells
        for i,h in enumerate(['Equipo','SKU/Serie','Estado','Notas']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for m in maquinas:
            row = mt.add_row().cells
            row[0].text = str(m.get('nombre',''))
            row[1].text = str(m.get('sku','') or m.get('serie',''))
            row[2].text = str(m.get('estado','OK'))
            row[3].text = str(m.get('notas',''))

    # Análisis IA
    if rep.get('ai_diagnostico'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("DIAGNÓSTICO TÉCNICO IA").bold = True
        p.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        doc.add_paragraph(rep['ai_diagnostico'])

    acciones = rep.get('ai_acciones') or []
    if isinstance(acciones, str):
        try: acciones = json.loads(acciones)
        except: acciones = []
    if acciones:
        doc.add_paragraph().add_run("ACCIONES SUGERIDAS").bold = True
        for a in acciones:
            doc.add_paragraph(str(a), style='List Bullet')

    # Pie
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')} — ILUS Sport & Health")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _reporte_to_html(rep, cliente):
    """Genera HTML del reporte para preview o email."""
    def _ls(items):
        if isinstance(items, str):
            try: items = json.loads(items)
            except: items = []
        if not items: return '<em style="color:#9ca3af">—</em>'
        return '<ul style="margin:6px 0 12px 20px">' + ''.join(f'<li>{i}</li>' for i in items) + '</ul>'

    maq = rep.get('maquinas_json') or []
    if isinstance(maq, str):
        try: maq = json.loads(maq)
        except: maq = []
    maq_html = ''
    if maq:
        maq_html = '<table style="width:100%;border-collapse:collapse;margin:8px 0 16px"><thead><tr style="background:#fafbfc"><th style="text-align:left;padding:8px;border-bottom:2px solid #eaecf0;font-size:.75rem">Equipo</th><th style="text-align:left;padding:8px;border-bottom:2px solid #eaecf0;font-size:.75rem">SKU/Serie</th><th style="text-align:left;padding:8px;border-bottom:2px solid #eaecf0;font-size:.75rem">Estado</th></tr></thead><tbody>'
        for m in maq:
            maq_html += f'<tr><td style="padding:8px;border-bottom:1px solid #f3f4f6;font-size:.82rem">{m.get("nombre","")}</td><td style="padding:8px;border-bottom:1px solid #f3f4f6;font-size:.78rem;color:#6b7280">{m.get("sku") or m.get("serie","")}</td><td style="padding:8px;border-bottom:1px solid #f3f4f6;font-size:.78rem">{m.get("estado","OK")}</td></tr>'
        maq_html += '</tbody></table>'

    ai_block = ''
    if rep.get('ai_diagnostico'):
        ai_block = f'<div style="background:linear-gradient(135deg,#faf5ff,#eff6ff);border:1px solid #ddd6fe;border-radius:10px;padding:14px 18px;margin:16px 0"><div style="font-size:.7rem;font-weight:800;color:#7c3aed;text-transform:uppercase;letter-spacing:.5px;margin-bottom:6px">🤖 Diagnóstico IA</div><div style="font-size:.86rem;line-height:1.55;color:#374151">{rep["ai_diagnostico"]}</div></div>'

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Informe — {cliente.get('razon_social','')}</title></head>
<body style="font-family:-apple-system,'Segoe UI',Arial,sans-serif;color:#0f172a;max-width:780px;margin:0 auto;padding:24px;background:#fff">
  <div style="text-align:center;border-bottom:3px solid #cc0000;padding-bottom:16px;margin-bottom:24px">
    <div style="font-size:1.6rem;font-weight:900;color:#cc0000;letter-spacing:1px">INFORME POST-SERVICIO</div>
    <div style="font-size:.78rem;color:#6b7280;margin-top:4px">ILUS Sport &amp; Health Solution SPA · RUT 76.996.964-0</div>
  </div>
  <table style="width:100%;border-collapse:collapse;margin-bottom:18px;font-size:.82rem">
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc;width:140px"><strong>Cliente</strong></td><td style="padding:6px;border:1px solid #eaecf0">{cliente.get('razon_social','—')}</td></tr>
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc"><strong>RUT</strong></td><td style="padding:6px;border:1px solid #eaecf0">{cliente.get('rut','—')}</td></tr>
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc"><strong>Ticket</strong></td><td style="padding:6px;border:1px solid #eaecf0">{rep.get('ticket_num','—')}</td></tr>
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc"><strong>Asunto</strong></td><td style="padding:6px;border:1px solid #eaecf0">{rep.get('asunto','—')}</td></tr>
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc"><strong>Tipo</strong></td><td style="padding:6px;border:1px solid #eaecf0">{rep.get('tipo','—')} <span style="color:#9ca3af">·</span> Estado: {rep.get('estado','—')}</td></tr>
    <tr><td style="padding:6px;border:1px solid #eaecf0;background:#fafbfc"><strong>Técnicos</strong></td><td style="padding:6px;border:1px solid #eaecf0">{rep.get('tecnico_senior','—')} {('· ' + rep.get('tecnico_junior')) if rep.get('tecnico_junior') else ''}</td></tr>
  </table>
  {f'<h3 style="font-size:.95rem;color:#cc0000;margin-top:20px">Antecedentes</h3><p style="font-size:.86rem;line-height:1.55;color:#374151">{rep["antecedentes"]}</p>' if rep.get('antecedentes') else ''}
  <h3 style="font-size:.95rem;color:#cc0000;margin-top:20px">Objetivos</h3>{_ls(rep.get('objetivos'))}
  <h3 style="font-size:.95rem;color:#cc0000;margin-top:20px">Trabajos realizados</h3>{_ls(rep.get('trabajos'))}
  <h3 style="font-size:.95rem;color:#cc0000;margin-top:20px">Observaciones</h3>{_ls(rep.get('observaciones'))}
  {f'<h3 style="font-size:.95rem;color:#cc0000;margin-top:20px">Equipos atendidos</h3>{maq_html}' if maq else ''}
  {ai_block}
  <div style="text-align:center;color:#9ca3af;font-size:.7rem;margin-top:32px;padding-top:16px;border-top:1px solid #eaecf0">
    Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')} · ILUS Sport &amp; Health
  </div>
</body></html>"""


@app.route("/mantenciones/api/reportes/<int:rid>/word", methods=["GET"])
@_mant_required
def mant_reporte_word(rid):
    """Descarga el reporte como DOCX."""
    r = mysql_fetchone(
        "SELECT r.*, c.razon_social, c.rut FROM mant_reportes r "
        "JOIN mant_clientes c ON c.id=r.cliente_id WHERE r.id=%s", (rid,)
    )
    if not r: return jsonify({"error":"No encontrado"}), 404
    rep = dict(r)
    cliente = {"razon_social": rep.pop("razon_social",""), "rut": rep.pop("rut","")}
    # Decode JSON fields
    for k in ("objetivos","trabajos","observaciones","maquinas_json","ai_acciones"):
        if rep.get(k):
            try: rep[k] = json.loads(rep[k])
            except: rep[k] = []
    docx_bytes = _build_reporte_docx(rep, cliente)
    if not docx_bytes:
        return jsonify({"error":"python-docx no instalado en el servidor"}), 503
    fname = f"informe_{rid}_{cliente.get('razon_social','cliente').replace(' ','_')[:40]}.docx"
    resp = make_response(docx_bytes)
    resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    resp.headers["Content-Disposition"] = f'attachment; filename="{fname}"'
    _mant_log("reporte", rid, "exportar_word")
    return resp


@app.route("/mantenciones/api/reportes/<int:rid>/html", methods=["GET"])
@_mant_required
def mant_reporte_html(rid):
    """Devuelve preview HTML del reporte."""
    r = mysql_fetchone(
        "SELECT r.*, c.razon_social, c.rut FROM mant_reportes r "
        "JOIN mant_clientes c ON c.id=r.cliente_id WHERE r.id=%s", (rid,)
    )
    if not r: return "Reporte no encontrado", 404
    rep = dict(r)
    cliente = {"razon_social": rep.pop("razon_social",""), "rut": rep.pop("rut","")}
    return _reporte_to_html(rep, cliente)


@app.route("/mantenciones/api/reportes/<int:rid>/enviar", methods=["POST"])
@_mant_required
def mant_reporte_enviar(rid):
    """Envía el reporte por email al cliente o a destinatarios custom."""
    d = request.get_json(silent=True) or {}
    r = mysql_fetchone(
        "SELECT r.*, c.razon_social, c.rut, c.contacto_email, c.contacto_nombre "
        "FROM mant_reportes r JOIN mant_clientes c ON c.id=r.cliente_id WHERE r.id=%s", (rid,)
    )
    if not r: return jsonify({"error":"No encontrado"}), 404
    rep = dict(r)
    cliente = {
        "razon_social": rep.pop("razon_social",""),
        "rut":          rep.pop("rut",""),
        "contacto_email": rep.pop("contacto_email",""),
        "contacto_nombre": rep.pop("contacto_nombre",""),
    }
    for k in ("objetivos","trabajos","observaciones","maquinas_json","ai_acciones"):
        if rep.get(k):
            try: rep[k] = json.loads(rep[k])
            except: rep[k] = []

    destinatarios = d.get("destinatarios") or [cliente["contacto_email"]] if cliente["contacto_email"] else []
    destinatarios = [e.strip() for e in destinatarios if e and "@" in e]
    if not destinatarios:
        return jsonify({"error":"Sin destinatarios válidos. Configura email del cliente o pasa 'destinatarios'."}), 400

    asunto = d.get("asunto") or f"Informe post-servicio {rep.get('ticket_num') or '#'+str(rid)} — {cliente['razon_social']}"
    mensaje_extra = d.get("mensaje","").strip()

    # Construir HTML del email — siempre con la plantilla corporativa
    html_reporte = _reporte_to_html(rep, cliente)
    body_email = (
        f"<p>Estimado/a {cliente['contacto_nombre'] or cliente['razon_social']},</p>"
        f"<p>Adjunto encontrarás el informe de servicio realizado. {mensaje_extra}</p>"
        f"<p>Saludos,<br><strong>Equipo ILUS Sport &amp; Health</strong></p>"
        f"<hr style='border:none;border-top:1px solid #eaecf0;margin:18px 0'>"
        f"{html_reporte}"
    )
    html_email = _comm_render_email_document(asunto, body_email, "Informe de servicio")

    # Adjunto Word
    attachments = []
    try:
        docx_bytes = _build_reporte_docx(rep, cliente)
        if docx_bytes:
            attachments.append({
                "filename": f"informe_{rep.get('ticket_num') or rid}.docx",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": docx_bytes,
            })
    except Exception as exc:
        print(f"[REPORTE WORD] {exc}")

    # Enviar (usa el helper existente _send_ilus_email si está disponible)
    sent = False
    err  = None
    try:
        sent = _send_ilus_email(", ".join(destinatarios), asunto, html_email,
                                attachments=attachments)
    except TypeError:
        # Si el helper no acepta attachments, mandar sin adjunto
        try: sent = _send_ilus_email(", ".join(destinatarios), asunto, html_email)
        except Exception as exc: err = str(exc)
    except Exception as exc:
        err = str(exc)

    # Log
    _mant_log("reporte", rid, "enviar_email",
              f"to={','.join(destinatarios)} ok={sent} err={err or '-'}")

    if sent:
        # Crear notificación de envío
        try:
            mysql_execute(
                "INSERT INTO mant_notificaciones (cliente_id,entidad,entidad_id,tipo,titulo,mensaje,canal,estado,destinatario,fecha_envio,created_by) "
                "VALUES (%s,'reporte',%s,'sla',%s,%s,'email','enviada',%s,NOW(),%s)",
                (r["cliente_id"], rid, f"Informe enviado al cliente",
                 f"Informe {rep.get('ticket_num') or rid} enviado a {', '.join(destinatarios)}",
                 ", ".join(destinatarios), current_username())
            )
        except Exception: pass

    return jsonify({"ok": bool(sent), "destinatarios": destinatarios, "error": err})


# ── ADJUNTOS DE CONTRATOS (multi-archivo) ─────────────────────────────

ALLOWED_ADJUNTO_TIPOS = {
    "pdf":"contrato","doc":"contrato","docx":"contrato",
    "jpg":"imagen","jpeg":"imagen","png":"imagen","gif":"imagen","webp":"imagen",
    "xlsx":"solicitud","xls":"solicitud",
}

@app.route("/mantenciones/api/contratos/<int:ctid>/adjuntos", methods=["GET"])
@_mant_required
def mant_adjuntos_list(ctid):
    rows = mysql_fetchall(
        "SELECT id,tipo,nombre,archivo_nombre,mime_type,tamaño_bytes,descripcion,created_by,created_at "
        "FROM mant_contrato_adjuntos WHERE contrato_id=%s ORDER BY created_at DESC", (ctid,)
    )
    def _fmt(r):
        d = dict(r)
        d["created_at"] = str(d["created_at"])[:16] if d.get("created_at") else ""
        d["url"] = f"/static/uploads/mantenciones/{d['archivo_nombre']}"
        return d
    return jsonify([_fmt(r) for r in rows])


@app.route("/mantenciones/api/contratos/<int:ctid>/adjuntos", methods=["POST"])
@_mant_required
def mant_adjunto_subir(ctid):
    ct = mysql_fetchone("SELECT cliente_id FROM mant_contratos WHERE id=%s", (ctid,))
    if not ct: return jsonify({"error":"Contrato no encontrado"}), 404
    f = request.files.get("archivo")
    if not f or not f.filename: return jsonify({"error":"Sin archivo"}), 400
    ext  = f.filename.rsplit(".",1)[-1].lower()
    if ext not in ALLOWED_ADJUNTO:
        return jsonify({"error":f"Tipo .{ext} no permitido"}), 400
    tipo  = ALLOWED_ADJUNTO_TIPOS.get(ext, "otro")
    fname = secure_filename(f"ct{ctid}_{int(time.time())}_{f.filename}")
    fpath = os.path.join(MANT_UPLOADS, fname)
    f.save(fpath)
    size  = os.path.getsize(fpath)
    nombre = request.form.get("nombre") or f.filename
    descripcion = request.form.get("descripcion","")
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_contrato_adjuntos
                   (contrato_id,cliente_id,tipo,nombre,archivo_nombre,archivo_path,
                    mime_type,tamaño_bytes,descripcion,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                (ctid, ct["cliente_id"], tipo, nombre, fname, fname,
                 f.content_type or f"application/{ext}", size,
                 descripcion, current_username())
            )
            aid = cur.lastrowid
        conn.commit()
        # Log
        _mant_log("contrato", ctid, "adjunto_subido", f"{tipo} · {nombre}")
        _mant_log("cliente", ct["cliente_id"], "adjunto_contrato", f"{tipo} · {nombre}")
        return jsonify({"ok":True,"id":aid,"nombre":nombre,
                        "url":f"/static/uploads/mantenciones/{fname}","tipo":tipo})
    finally:
        conn.close()


@app.route("/mantenciones/api/adjuntos/<int:aid>", methods=["DELETE"])
@_mant_required
def mant_adjunto_del(aid):
    adj = mysql_fetchone(
        "SELECT archivo_nombre, nombre, contrato_id, cliente_id, tipo FROM mant_contrato_adjuntos WHERE id=%s",
        (aid,)
    )
    if not adj: return jsonify({"error":"No encontrado"}),404
    try:
        fpath = os.path.join(MANT_UPLOADS, adj["archivo_nombre"])
        if os.path.exists(fpath): os.remove(fpath)
    except Exception: pass
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM mant_contrato_adjuntos WHERE id=%s",(aid,))
        conn.commit()
        label = f"{adj.get('tipo') or 'archivo'} · {adj.get('nombre') or aid}"
        if adj.get("contrato_id"):
            _mant_log("contrato", adj["contrato_id"], "adjunto_eliminado", label)
        if adj.get("cliente_id"):
            _mant_log("cliente", adj["cliente_id"], "adjunto_eliminado", label)
        return jsonify({"ok":True})
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════
#  REPUESTOS — gestión de repuestos por cliente / visita / reporte
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/clientes/<int:cid>/repuestos", methods=["GET"])
@_mant_required
def mant_repuestos_list(cid):
    """Lista repuestos del cliente. Filtros opcionales: ?tipo=, ?estado=, ?visita_id=, ?reporte_id="""
    where = ["cliente_id=%s"]; params = [cid]
    for k in ("tipo","estado"):
        v = request.args.get(k)
        if v: where.append(f"{k}=%s"); params.append(v)
    for k in ("visita_id","reporte_id","maquina_id"):
        v = request.args.get(k)
        if v and v.isdigit(): where.append(f"{k}=%s"); params.append(int(v))
    rows = mysql_fetchall(
        f"SELECT * FROM mant_repuestos WHERE {' AND '.join(where)} "
        f"ORDER BY created_at DESC",
        tuple(params)
    )
    def _fmt(r):
        d = dict(r)
        for k in ('cantidad','costo_unitario','precio_venta'):
            d[k] = float(d[k] or 0)
        d['margen'] = round((d['precio_venta'] - d['costo_unitario']) * d['cantidad'], 2)
        d['margen_pct'] = round(((d['precio_venta'] - d['costo_unitario']) / d['costo_unitario'] * 100), 1) if d['costo_unitario'] else None
        d['costo_total'] = round(d['costo_unitario'] * d['cantidad'], 2)
        d['venta_total'] = round(d['precio_venta'] * d['cantidad'], 2)
        d['fecha'] = d['fecha'].isoformat() if d.get('fecha') else None
        d['created_at'] = str(d.get('created_at'))[:16] if d.get('created_at') else ''
        return d
    repuestos = [_fmt(r) for r in rows]
    # Resumen
    totales = {
        'count':         len(repuestos),
        'venta_total':   round(sum(r['venta_total'] for r in repuestos if r['tipo']=='venta'), 2),
        'costo_total':   round(sum(r['costo_total'] for r in repuestos), 2),
        'margen_total':  round(sum(r['margen']      for r in repuestos if r['tipo']=='venta'), 2),
        'garantia_count':sum(1 for r in repuestos if r['tipo']=='garantia'),
        'venta_count':   sum(1 for r in repuestos if r['tipo']=='venta'),
    }
    return jsonify({"repuestos": repuestos, "totales": totales})


@app.route("/mantenciones/api/clientes/<int:cid>/repuestos", methods=["POST"])
@_mant_required
def mant_repuesto_crear(cid):
    d = request.get_json(silent=True) or {}
    nombre = (d.get("nombre") or "").strip()
    if not nombre:
        return jsonify({"error":"Nombre obligatorio"}), 400
    fields = {
        'cliente_id': cid,
        'visita_id':  d.get('visita_id') or None,
        'reporte_id': d.get('reporte_id') or None,
        'maquina_id': d.get('maquina_id') or None,
        'sku':        (d.get('sku') or '').strip()[:120],
        'nombre':     nombre[:400],
        'descripcion':(d.get('descripcion') or '').strip(),
        'cantidad':   float(d.get('cantidad') or 1),
        'costo_unitario': float(d.get('costo_unitario') or 0),
        'precio_venta':   float(d.get('precio_venta') or 0),
        'moneda':     (d.get('moneda') or 'CLP')[:8],
        'tipo':       (d.get('tipo') or 'venta'),
        'estado':     (d.get('estado') or 'cotizado'),
        'proveedor':  (d.get('proveedor') or '').strip()[:200],
        'documento':  (d.get('documento') or '').strip()[:120],
        'fecha':      d.get('fecha') or None,
        'observacion':(d.get('observacion') or '').strip(),
        'created_by': current_username(),
    }
    cols = list(fields.keys())
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                f"INSERT INTO mant_repuestos ({','.join(cols)}) "
                f"VALUES ({','.join(['%s']*len(cols))})",
                tuple(fields[c] for c in cols)
            )
            new_id = cur.lastrowid
        conn.commit()
        _mant_log("repuesto", new_id, "crear", f"{fields['nombre']} (cliente {cid})")
        return jsonify({"ok":True, "id":new_id})
    finally:
        conn.close()


@app.route("/mantenciones/api/repuestos/<int:rid>", methods=["PUT"])
@_mant_required
def mant_repuesto_update(rid):
    d = request.get_json(silent=True) or {}
    allowed = ['sku','nombre','descripcion','cantidad','costo_unitario','precio_venta',
               'moneda','tipo','estado','proveedor','documento','fecha','observacion',
               'visita_id','reporte_id','maquina_id']
    sets, vals = [], []
    for f in allowed:
        if f in d:
            sets.append(f"{f}=%s"); vals.append(d[f] if d[f] not in ('','null') else None)
    if not sets:
        return jsonify({"error":"Sin campos"}), 400
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(f"UPDATE mant_repuestos SET {','.join(sets)} WHERE id=%s", vals+[rid])
        conn.commit()
        _mant_log("repuesto", rid, "editar")
        return jsonify({"ok":True})
    finally:
        conn.close()


@app.route("/mantenciones/api/repuestos/<int:rid>", methods=["DELETE"])
@_mant_required
def mant_repuesto_del(rid):
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM mant_repuestos WHERE id=%s", (rid,))
        conn.commit()
        _mant_log("repuesto", rid, "eliminar")
        return jsonify({"ok":True})
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════
#  FINANZAS DEL CLIENTE — agregados de ingresos/costos/margen
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/clientes/<int:cid>/finanzas")
@_mant_required
def mant_cliente_finanzas(cid):
    """Devuelve agregados financieros del cliente para los últimos N meses."""
    meses = int(request.args.get("meses") or 12)
    if meses <= 0:
        fecha_corte = "1900-01-01"
    else:
        fecha_corte = (datetime.now().date() - timedelta(days=meses*31)).isoformat()

    # Repuestos
    rep_rows = mysql_fetchall(
        "SELECT tipo, cantidad, costo_unitario, precio_venta, fecha "
        "FROM mant_repuestos WHERE cliente_id=%s AND (fecha IS NULL OR fecha >= %s)",
        (cid, fecha_corte)
    )
    rep_venta_total = 0
    rep_costo_total = 0
    garantia_costo  = 0
    for r in rep_rows:
        c = float(r["cantidad"] or 0)
        cu = float(r["costo_unitario"] or 0)
        pv = float(r["precio_venta"] or 0)
        if r["tipo"] == "venta":
            rep_venta_total += pv * c
            rep_costo_total += cu * c
        elif r["tipo"] == "garantia":
            garantia_costo += cu * c
        else:  # reposicion / consumo
            rep_costo_total += cu * c

    # Visitas con costo
    visitas = mysql_fetchall(
        "SELECT tipo, estado, costo, fecha_programada FROM mant_visitas "
        "WHERE cliente_id=%s AND fecha_programada >= %s",
        (cid, fecha_corte)
    )
    visitas_costo = sum(float(v["costo"] or 0) for v in visitas if v.get("costo"))
    visitas_count = len([v for v in visitas if v.get("estado")=="completada"])

    # Contrato — estimación lineal (monto_mensual × meses_vigentes)
    contrato_estimado = 0
    contratos = mysql_fetchall(
        "SELECT monto_mensual, fecha_inicio, fecha_vencimiento, es_indefinido "
        "FROM mant_contratos WHERE cliente_id=%s AND estado='vigente'", (cid,)
    )
    hoy = datetime.now().date()
    for ct in contratos:
        m = float(ct["monto_mensual"] or 0)
        if m <= 0: continue
        fi = ct.get("fecha_inicio") or hoy
        if isinstance(fi, datetime): fi = fi.date()
        meses_vigentes = max(0, min(meses or 12, ((hoy.year - fi.year)*12 + hoy.month - fi.month)))
        contrato_estimado += m * meses_vigentes

    ingresos_total = rep_venta_total + visitas_costo + contrato_estimado
    costos_total   = rep_costo_total + garantia_costo
    margen         = ingresos_total - costos_total
    ticket_prom    = (ingresos_total / visitas_count) if visitas_count else 0

    # Por mes (últimos 12)
    months_es = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
    por_mes = []
    for i in range(11,-1,-1):
        ref_year, ref_month = hoy.year, hoy.month - i
        while ref_month <= 0: ref_month += 12; ref_year -= 1
        ingreso_mes = 0
        for r in rep_rows:
            if r.get("fecha") and r["tipo"]=="venta":
                fr = r["fecha"]
                if isinstance(fr, datetime): fr = fr.date()
                if fr.year == ref_year and fr.month == ref_month:
                    ingreso_mes += float(r["precio_venta"] or 0) * float(r["cantidad"] or 0)
        for v in visitas:
            if v.get("fecha_programada") and v.get("estado")=="completada":
                fv = v["fecha_programada"]
                if isinstance(fv, datetime): fv = fv.date()
                if fv.year == ref_year and fv.month == ref_month:
                    ingreso_mes += float(v["costo"] or 0)
        por_mes.append({"label": months_es[ref_month-1], "year": ref_year, "ingreso": round(ingreso_mes,2)})

    return jsonify({
        "ingresos_total":  round(ingresos_total, 2),
        "costos_total":    round(costos_total, 2),
        "margen":          round(margen, 2),
        "garantia_costo":  round(garantia_costo, 2),
        "visitas_count":   visitas_count,
        "ticket_promedio": round(ticket_prom, 2),
        "repuestos_venta": round(rep_venta_total, 2),
        "visitas_costo":   round(visitas_costo, 2),
        "contrato_estimado": round(contrato_estimado, 2),
        "por_mes":         por_mes,
    })


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENTOS DEL CLIENTE — listado unificado de adjuntos
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/clientes/<int:cid>/documentos")
@_mant_required
def mant_cliente_documentos(cid):
    """Devuelve TODOS los documentos asociados al cliente (multi-fuente)."""
    items = []

    # 1. Adjuntos (contratos + multi-archivo)
    rows = mysql_fetchall(
        "SELECT id,tipo,nombre,archivo_nombre,mime_type,tamaño_bytes,created_by,created_at "
        "FROM mant_contrato_adjuntos WHERE cliente_id=%s", (cid,)
    )
    for r in rows:
        items.append({
            "kind":   "adjunto",
            "id":     r["id"],
            "tipo":   r["tipo"],
            "nombre": r["nombre"] or r.get("archivo_nombre"),
            "url":    f"/static/uploads/mantenciones/{r['archivo_nombre']}",
            "size_kb": round((r.get("tamaño_bytes") or 0)/1024) if r.get("tamaño_bytes") else None,
            "created_by": r.get("created_by"),
            "created_at": str(r["created_at"])[:16] if r.get("created_at") else "",
            "fuente": "Adjunto",
            "deletable": True,
        })

    # 2. Reportes con HTML guardado
    rep_rows = mysql_fetchall(
        "SELECT id, ticket_num, asunto, html_path, html_generated_at, created_by "
        "FROM mant_reportes WHERE cliente_id=%s AND html_path IS NOT NULL", (cid,)
    )
    for r in rep_rows:
        items.append({
            "kind":   "reporte_html",
            "id":     r["id"],
            "tipo":   "reporte",
            "nombre": f"Informe {r.get('ticket_num') or r['id']} — {r.get('asunto') or 'Sin asunto'}",
            "url":    f"/static/{r['html_path']}",
            "size_kb": None,
            "created_by": r.get("created_by"),
            "created_at": str(r["html_generated_at"])[:16] if r.get("html_generated_at") else "",
            "fuente": "Reporte (HTML)",
            "deletable": False,
        })

    # 3. Contrato — archivo principal
    ct_rows = mysql_fetchall(
        "SELECT id,nombre,archivo_nombre,archivo_path,archivo_tipo,created_by,created_at "
        "FROM mant_contratos WHERE cliente_id=%s AND archivo_path IS NOT NULL", (cid,)
    )
    for r in ct_rows:
        items.append({
            "kind":   "contrato_principal",
            "id":     r["id"],
            "tipo":   "contrato",
            "nombre": r.get("nombre") or r.get("archivo_nombre") or f"Contrato #{r['id']}",
            "url":    f"/mantenciones/api/contratos/{r['id']}/archivo",
            "size_kb": None,
            "created_by": r.get("created_by"),
            "created_at": str(r["created_at"])[:16] if r.get("created_at") else "",
            "fuente": "Contrato principal",
            "deletable": False,
        })

    # Ordenar por created_at desc
    items.sort(key=lambda x: x.get("created_at",""), reverse=True)
    return jsonify(items)


@app.route("/mantenciones/api/clientes/<int:cid>/documentos", methods=["POST"])
@_mant_required
def mant_cliente_documento_subir(cid):
    """Sube un documento suelto al cliente como adjunto general."""
    f = request.files.get("archivo")
    if not f or not f.filename: return jsonify({"error":"Sin archivo"}), 400
    ext = f.filename.rsplit(".",1)[-1].lower() if "." in f.filename else ""
    if ext not in ALLOWED_ADJUNTO: return jsonify({"error":"Tipo no permitido"}), 400
    fname = secure_filename(f"cli{cid}_{int(time.time())}_{f.filename}")
    fpath = os.path.join(MANT_UPLOADS, fname)
    f.save(fpath)
    size  = os.path.getsize(fpath)
    tipo  = ALLOWED_ADJUNTO_TIPOS.get(ext, "otro")
    nombre = (request.form.get("nombre") or f.filename)[:300]

    # Crear contrato dummy si no existe (para satisfacer FK)
    ct = mysql_fetchone("SELECT id FROM mant_contratos WHERE cliente_id=%s ORDER BY id LIMIT 1", (cid,))
    if not ct:
        # Si el cliente no tiene contrato, creamos uno mínimo
        conn = get_mysql()
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO mant_contratos (cliente_id,nombre,estado,es_indefinido,created_by) "
                "VALUES (%s,'Contenedor de documentos','indefinido',1,%s)",
                (cid, current_username())
            )
            ct_id = cur.lastrowid
        conn.commit(); conn.close()
    else:
        ct_id = ct["id"]

    mysql_execute(
        "INSERT INTO mant_contrato_adjuntos "
        "(contrato_id,cliente_id,tipo,nombre,archivo_nombre,archivo_path,mime_type,tamaño_bytes,created_by) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)",
        (ct_id, cid, tipo, nombre, fname, fname, f.mimetype or '', size, current_username())
    )
    _mant_log("documento", cid, "subir", nombre)
    return jsonify({"ok":True})


@app.route("/mantenciones/api/documentos/<kind>/<int:did>", methods=["DELETE"])
@_mant_required
def mant_documento_del(kind, did):
    if kind == "adjunto":
        adj = mysql_fetchone(
            "SELECT archivo_nombre, nombre, cliente_id FROM mant_contrato_adjuntos WHERE id=%s",
            (did,)
        )
        if adj:
            try:
                fp = os.path.join(MANT_UPLOADS, adj["archivo_nombre"])
                if os.path.exists(fp): os.remove(fp)
            except Exception: pass
            mysql_execute("DELETE FROM mant_contrato_adjuntos WHERE id=%s",(did,))
            if adj.get("cliente_id"):
                _mant_log("cliente", adj["cliente_id"], "documento_eliminado",
                          adj.get("nombre") or f"adjunto #{did}")
            return jsonify({"ok":True})
    return jsonify({"error":"Tipo no eliminable o no encontrado"}), 400


# ══════════════════════════════════════════════════════════════════════
#  EMAIL MANUAL — enviar email arbitrario al cliente con trazabilidad
# ══════════════════════════════════════════════════════════════════════

@app.route("/mantenciones/api/clientes/<int:cid>/email-manual", methods=["POST"])
@_mant_required
def mant_email_manual(cid):
    d = request.get_json(silent=True) or {}
    destinatario = (d.get("destinatario") or "").strip()
    asunto       = (d.get("asunto") or "").strip()
    mensaje      = (d.get("mensaje") or "").strip()
    if not (destinatario and asunto and mensaje):
        return jsonify({"error":"Completa destinatario, asunto y mensaje"}), 400
    if "@" not in destinatario:
        return jsonify({"error":"Email no válido"}), 400

    cliente = mysql_fetchone("SELECT razon_social FROM mant_clientes WHERE id=%s",(cid,))
    if not cliente: return jsonify({"error":"Cliente no encontrado"}), 404

    # Envolver SIEMPRE con la plantilla corporativa (formato del preview)
    body = (
        f"<p>{mensaje.replace(chr(10), '<br>')}</p>"
        f"<hr style='border:none;border-top:1px solid #e5e7eb;margin:18px 0'>"
        f"<p style='font-size:.78rem;color:#9ca3af'>Enviado por {current_username()}</p>"
    )
    html = _comm_render_email_document(asunto, body, cliente.get("razon_social",""))

    sent = False; err = None
    try:
        sent = _send_ilus_email(destinatario, asunto, html)
    except Exception as exc: err = str(exc)

    # Registrar en notificaciones
    try:
        mysql_execute(
            "INSERT INTO mant_notificaciones "
            "(cliente_id,entidad,entidad_id,tipo,titulo,mensaje,canal,estado,destinatario,fecha_envio,created_by) "
            "VALUES (%s,'cliente',%s,'sla',%s,%s,'email',%s,%s,NOW(),%s)",
            (cid, cid, asunto, mensaje[:1000],
             'enviada' if sent else 'fallida', destinatario, current_username())
        )
    except Exception: pass
    _mant_log("cliente", cid, "email_manual",
              f"to={destinatario} ok={sent} err={err or '-'}")
    return jsonify({"ok": bool(sent), "error": err})


# ── NOTIFICACIONES ────────────────────────────────────────────────────

@app.route("/mantenciones/api/notificaciones")
@_mant_required
def mant_notif_list():
    """Lista notificaciones (todas o filtradas por cliente)."""
    cid    = request.args.get("cliente_id")
    estado = request.args.get("estado","")
    sql    = ("SELECT n.*,c.razon_social FROM mant_notificaciones n "
              "LEFT JOIN mant_clientes c ON c.id=n.cliente_id WHERE 1=1")
    params = []
    if cid:    sql += " AND n.cliente_id=%s"; params.append(int(cid))
    if estado: sql += " AND n.estado=%s";    params.append(estado)
    sql += " ORDER BY n.created_at DESC LIMIT 100"
    rows = mysql_fetchall(sql, tuple(params))
    def _fmt(r):
        d = dict(r)
        d["created_at"] = str(d["created_at"])[:16] if d.get("created_at") else ""
        return d
    return jsonify([_fmt(r) for r in rows])


@app.route("/mantenciones/api/notificaciones", methods=["POST"])
@_mant_required
def mant_notif_crear():
    """Crea una notificación manualmente."""
    d = request.get_json(silent=True) or {}
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """INSERT INTO mant_notificaciones
                   (cliente_id,entidad,entidad_id,tipo,titulo,mensaje,canal,
                    destinatario,estado,created_by)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'pendiente',%s)""",
                (d.get("cliente_id"), d.get("entidad","cliente"),
                 d.get("entidad_id"), d.get("tipo","otro"),
                 d.get("titulo",""), d.get("mensaje",""),
                 d.get("canal","sistema"), d.get("destinatario",""),
                 current_username())
            )
            nid = cur.lastrowid
        conn.commit()
        return jsonify({"ok":True,"id":nid})
    finally:
        conn.close()


@app.route("/mantenciones/api/notificaciones/<int:nid>", methods=["PUT"])
@_mant_required
def mant_notif_update(nid):
    """Cambia estado de una notificación (leida, ignorada, enviada)."""
    d = request.get_json(silent=True) or {}
    estado = d.get("estado","leida")
    conn = get_mysql()
    try:
        with conn.cursor() as cur:
            extra = ", fecha_lectura=%s" if estado in ("leida","ignorada") else ""
            vals  = ([datetime.now()] if extra else []) + [estado, nid]
            cur.execute(
                f"UPDATE mant_notificaciones SET estado=%s{extra} WHERE id=%s",
                [estado] + ([datetime.now()] if extra else []) + [nid]
            )
        conn.commit()
        return jsonify({"ok":True})
    finally:
        conn.close()


@app.route("/mantenciones/api/notificaciones/<int:nid>/enviar", methods=["POST"])
@_mant_required
def mant_notif_enviar(nid):
    """Envía la notificación por email usando Resend."""
    notif = mysql_fetchone(
        "SELECT n.*,c.contacto_email,c.razon_social FROM mant_notificaciones n "
        "LEFT JOIN mant_clientes c ON c.id=n.cliente_id WHERE n.id=%s", (nid,)
    )
    if not notif: return jsonify({"error":"No encontrado"}), 404

    destinatario = notif.get("destinatario") or notif.get("contacto_email","")
    if not destinatario:
        return jsonify({"error":"Sin email destinatario"}), 400

    html_body = f"""
    <div style="font-family:sans-serif;max-width:600px;margin:auto">
      <div style="background:#CC0000;padding:20px;text-align:center">
        <h2 style="color:#fff;margin:0">ILUS Sport & Health</h2>
      </div>
      <div style="padding:24px;background:#f9f9f9">
        <h3 style="color:#111">{notif['titulo']}</h3>
        <p style="color:#374151;line-height:1.6">{notif['mensaje']}</p>
        <p style="color:#6b7280;font-size:.85rem">
          Cliente: <strong>{notif.get('razon_social','')}</strong>
        </p>
      </div>
      <div style="padding:12px;text-align:center;color:#9ca3af;font-size:.78rem">
        ILUS Sport & Health Solution SPA · Sistema de Mantenciones
      </div>
    </div>"""

    try:
        resultado = _send_via_resend(
            to=destinatario,
            subject=f"[ILUS Mantenciones] {notif['titulo']}",
            html=html_body,
            from_addr="mantenciones@sphs.cl"
        )
        conn = get_mysql()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "UPDATE mant_notificaciones SET estado='enviada',fecha_envio=%s WHERE id=%s",
                    (datetime.now(), nid)
                )
            conn.commit()
        finally:
            conn.close()
        return jsonify({"ok":True})
    except Exception as e:
        return jsonify({"error":str(e)}), 500


@app.route("/mantenciones/notificaciones")
@_mant_required
def mant_notificaciones_centro():
    """Centro de notificaciones global."""
    notifs = mysql_fetchall(
        "SELECT n.*,c.razon_social FROM mant_notificaciones n "
        "LEFT JOIN mant_clientes c ON c.id=n.cliente_id "
        "ORDER BY n.created_at DESC LIMIT 200"
    )
    pendientes = sum(1 for n in notifs if n.get("estado")=="pendiente")
    return render_template("mantenciones/notificaciones.html",
                           notifs=[dict(n) for n in notifs],
                           pendientes=pendientes)


# ─────────────────────────────────────────────
#  Arranque — inicializar tablas al cargar módulo
#  (funciona con `python app.py` Y `flask run`)
# ─────────────────────────────────────────────

try:
    from pickups_module import register_pickup_routes
    register_pickup_routes(app, globals())
    print("[ILUS] Modulo de retiros registrado.")
except Exception as _pickup_reg_err:
    print(f"[ILUS][WARN] register_pickup_routes: {_pickup_reg_err}")

try:
    with app.app_context():
        init_db()
    print("[ILUS] Tablas inicializadas correctamente.")
except Exception as _init_err:
    print(f"[ILUS][WARN] init_db: {_init_err}")

# init_transporte_tables usa get_mysql() (sin app context), siempre funciona
try:
    init_transporte_tables()
    print("[ILUS] Tablas de transporte OK.")
except Exception as _tr_init_err:
    print(f"[ILUS][WARN] init_transporte_tables: {_tr_init_err}")

try:
    init_comunicaciones_tables()
    print("[ILUS] Tablas de comunicaciones OK.")
except Exception as _comm_init_err:
    print(f"[ILUS][WARN] init_comunicaciones_tables: {_comm_init_err}")

try:
    init_mantenciones_tables()
    print("[ILUS] Tablas de mantenciones OK.")
except Exception as _mant_init_err:
    print(f"[ILUS][WARN] init_mantenciones_tables: {_mant_init_err}")

# ── PLAN DE MEJORA IA ─────────────────────────────────────────────────────────

@app.route("/mantenciones/api/clientes/<int:cid>/plan-mejora", methods=["POST"])
@_mant_required
def mant_plan_mejora(cid):
    """
    Genera un plan de mejora y proyecciones para el próximo ciclo usando Claude AI.
    Considera: equipos del cliente, contratos vigentes, análisis IA previos.
    """
    cliente = mysql_fetchone("SELECT * FROM mant_clientes WHERE id=%s", (cid,))
    if not cliente:
        return jsonify({"error": "Cliente no encontrado"}), 404

    maquinas  = mysql_fetchall("SELECT * FROM mant_maquinas WHERE cliente_id=%s AND estado!='baja'", (cid,))
    contratos = mysql_fetchall("SELECT * FROM mant_contratos WHERE cliente_id=%s", (cid,))
    visitas   = mysql_fetchall(
        "SELECT * FROM mant_visitas WHERE cliente_id=%s ORDER BY fecha_programada DESC LIMIT 10", (cid,)
    )

    ai_key = _get_ai_key()
    if not ai_key:
        return jsonify({"error": "API de IA no configurada"}), 503

    # Construir contexto del cliente
    equipos_txt = "\n".join([
        f"- {m['nombre']} (SKU: {m.get('sku','N/A')}, Cant: {m.get('cantidad',1)}, "
        f"Doc: {m.get('doc_origen','')}, Fecha: {m.get('doc_fecha','')}, Estado: {m.get('estado_op','operativo')})"
        for m in maquinas
    ]) or "Sin equipos registrados"

    ct_vigentes = [c for c in contratos if c.get("estado") in ("vigente", "indefinido")]
    contratos_txt = "\n".join([
        f"- {c['nombre']} | Tipo: {c.get('ai_tipo_contrato','N/A')} | "
        f"Monto: ${c.get('monto_mensual',0):,.0f}/mes | Frecuencia: cada {c.get('frecuencia_meses','?')} meses | "
        f"SLA: {c.get('sla_horas','?')}h | Score IA: {c.get('ai_score','N/A')} | "
        f"Incluye repuestos: {'Sí' if c.get('incluye_repuestos') else 'No'} | "
        f"Vencimiento: {c.get('fecha_vencimiento','indefinido')}"
        for c in ct_vigentes
    ]) or "Sin contratos vigentes"

    ai_mejoras_txt = ""
    for c in contratos:
        if c.get("ai_mejoras"):
            try:
                mejoras = json.loads(c["ai_mejoras"])
                ai_mejoras_txt += f"\nMejoras del contrato {c['nombre']}: " + "; ".join(mejoras)
            except Exception:
                pass
        if c.get("ai_puntos_criticos"):
            try:
                puntos = json.loads(c["ai_puntos_criticos"])
                ai_mejoras_txt += f"\nPuntos críticos: " + "; ".join(puntos)
            except Exception:
                pass

    from datetime import date
    hoy = date.today()

    prompt = f"""Eres un experto en gestión de mantenimiento de equipos fitness para ILUS Fitness Chile.
Analiza la siguiente información del cliente y genera un PLAN DE MEJORA Y PROYECCIÓN para el próximo ciclo de mantención.
Responde ÚNICAMENTE con JSON estructurado, sin texto adicional.

CLIENTE: {cliente['razon_social']} | RUT: {cliente.get('rut','')} | Ciudad: {cliente.get('ciudad','')}
FECHA HOY: {hoy}

EQUIPOS ({len(maquinas)} en total):
{equipos_txt}

CONTRATOS VIGENTES:
{contratos_txt}

ANÁLISIS IA PREVIO:
{ai_mejoras_txt or 'No disponible'}

Estructura JSON requerida:
{{
  "resumen_ejecutivo": "2-3 oraciones del estado actual y oportunidades",
  "estado_flota": "bueno|regular|critico",
  "indice_salud": número_0_a_100,
  "proxima_visita": {{
    "fecha_sugerida": "YYYY-MM-DD",
    "tipo": "preventiva|correctiva|inspeccion",
    "duracion_horas": número,
    "prioridad": "alta|media|baja",
    "razon": "por qué urgente"
  }},
  "recomendaciones_equipos": [
    {{
      "equipo": "nombre del equipo",
      "estado": "ok|atencion|urgente",
      "accion": "acción específica a tomar",
      "plazo": "inmediato|30 días|60 días|90 días"
    }}
  ],
  "proyeccion_12_meses": [
    {{
      "mes": "YYYY-MM",
      "tipo_visita": "preventiva|correctiva",
      "descripcion": "qué se hará",
      "costo_estimado": número_o_null
    }}
  ],
  "propuestas_mejora": [
    {{
      "titulo": "título de la propuesta",
      "descripcion": "detalle de qué mejorar",
      "impacto": "alto|medio|bajo",
      "categoria": "contrato|equipo|proceso|costos"
    }}
  ],
  "alertas_criticas": ["alerta 1", "alerta 2"],
  "oportunidades_comerciales": ["oportunidad 1", "oportunidad 2"]
}}"""

    try:
        import anthropic as _anthropic
        ai  = _anthropic.Anthropic(api_key=ai_key)
        msg = ai.messages.create(
            model="claude-opus-4-5",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        plan = json.loads(raw)
        return jsonify({"ok": True, "plan": plan, "cliente": cliente["razon_social"]})
    except Exception as e:
        return jsonify({"error": f"Error IA: {e}"}), 500


if __name__ == "__main__":
    print("=" * 45)
    print("  ILUS - Sistema de Etiquetas")
    print("  http://localhost:5000")
    print("=" * 45)
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") != "production"
    app.run(debug=debug, host="0.0.0.0", port=port)
